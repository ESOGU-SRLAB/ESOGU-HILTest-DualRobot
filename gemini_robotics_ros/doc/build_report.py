#!/usr/bin/env python3
"""
gemini_robotics_ros geliştirme raporunu (.docx) üretir.

    python3 doc/build_report.py

Şekiller doc/figures/ altındadır. Rapor metnindeki her sayı bu depoda
ölçülmüştür; kaynakları ilgili bölümlerde belirtilmiştir.
"""

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "gemini_robotics_ros_raporu.docx")

ACCENT = RGBColor(0x1E, 0x51, 0x7B)
MUTED = RGBColor(0x55, 0x5F, 0x6B)
GOOD = RGBColor(0x1E, 0x84, 0x49)
BAD = RGBColor(0xC0, 0x39, 0x2B)


# --------------------------------------------------------------------------
# yardımcılar
# --------------------------------------------------------------------------

def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def para(doc, text, italic=False, size=10.5, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(doc, chunks, space_after=6):
    """chunks: [(metin, {'b':True,'i':True,'code':True,'color':RGBColor}), ...]"""
    p = doc.add_paragraph()
    for text, fmt in chunks:
        run = p.add_run(text)
        run.bold = fmt.get("b", False)
        run.italic = fmt.get("i", False)
        if fmt.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        if fmt.get("color") is not None:
            run.font.color.rgb = fmt["color"]
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    shade = OxmlElement("w:shd")
    shade.set(qn("w:val"), "clear")
    shade.set(qn("w:fill"), "F2F4F6")
    p._p.get_or_add_pPr().append(shade)
    return p


def figure(doc, filename, caption, width=6.2):
    path = os.path.join(FIG, filename)
    if not os.path.exists(path):
        para(doc, f"[şekil bulunamadı: {filename}]", italic=True, color=BAD)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    cap.paragraph_format.space_after = Pt(12)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, name in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(name)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)
            if str(value).startswith(("/", "_", "ur10e", "scan_", "tool_", "render_")):
                run.font.name = "Consolas"
                run.font.size = Pt(9)
    # Satırlar sayfa sınırında BÖLÜNMESİN: bölündüğünde metnin bir kısmı
    # sonraki sayfanın tepesinde başlıksız, boş görünen bir şeride düşüyor.
    for row in t.rows:
        no_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(no_split)

    if widths:
        # Sabit yerleşim ŞART: autofit açıkken Word/LibreOffice hücre
        # genişliklerini içeriğe göre yeniden dağıtıyor ve verilen oranları
        # yok sayıyor (kısa sütunlar gereksiz yere sarıyordu).
        # (autofit=False zaten w:tblLayout'u "fixed" yapar; elle ikinci bir
        # tblLayout eklemek tblPr'yi bozuyor ve LibreOffice hepsini yok sayıyor.)
        t.autofit = False
        for i, w in enumerate(widths):
            # Hem tblGrid hem her hücre: yalnız biri verilirse okuyucular
            # diğerine bakıp yine kendi dağıtımını yapabiliyor.
            t.columns[i].width = Inches(w)
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# --------------------------------------------------------------------------
# rapor
# --------------------------------------------------------------------------

def build():
    doc = Document()
    style_document(doc)
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.85)
        section.top_margin = section.bottom_margin = Inches(0.8)

    # ---- kapak ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Gemini Robotics ER 2 ile\nDille Yönlendirilen Pick and Place")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("UR10e + OnRobot VGC10 + SICK Visionary-T Mini\n"
                      "Gazebo Harmonic / ROS 2 Humble")
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("\ngemini_robotics_ros paketi — geliştirme raporu\n"
                       "ESOGÜ IFARLAB · 10–11 Ağustos 2026")
    run.font.size = Pt(11)

    doc.add_paragraph()
    para(doc,
         "Bu rapordaki her sayısal değer çalışan hücrede ya da paketin kendi "
         "mesh/URDF dosyalarında ölçülmüştür. Ölçümün nasıl yapıldığı ilgili "
         "bölümde belirtilmiştir; doğrulanmamış varsayımlar açıkça öyle "
         "işaretlenmiştir.",
         italic=True, color=MUTED)

    doc.add_page_break()

    # ---- 1. yönetici özeti ----
    h(doc, "1. Yönetici Özeti", 1)
    para(doc,
         "Doğal dille verilen bir komutu — örneğin “konveyördeki şeyi al ve "
         "toolkit'in üst sırasındaki boş gözlerden birine koy” — UR10e kolunun "
         "fiziksel hareketine çeviren uçtan uca bir hat kuruldu. Algı tarafında "
         "Google'ın Gemini Robotics ER 2 modeli, hareket tarafında MoveIt "
         "kullanılıyor. Sistem Gazebo Harmonic'te çalışıyor ve gerçek donanıma "
         "geçiş tek bir launch argümanına indirgendi.")

    para(doc, "Gün içinde çözülen başlıca problemler:")
    bullet(doc, "Düğümün MoveIt çağrılarından sonra tamamen sağır kalması "
                "(rclpy executor'dan düşme) — kök nedeni bulundu ve kalıcı olarak kapatıldı.")
    bullet(doc, "ER 2'nin konveyördeki parçayı görememesi — sorun modelde değil, "
                "ona gönderdiğimiz görüntüdeydi; kontrast 1.7/255 idi, 89.8/255'e çıkarıldı.")
    bullet(doc, "Konveyör ve toolkit'in aynı kadraja sığmaması — çok pozlu tarama eklendi.")
    bullet(doc, "Vakum ucunun yüzeye dik gelmemesi — 90°'lik hata önce 30°'ye, "
                "sonra tam sıfıra indirildi (Bölüm 7).")
    bullet(doc, "Kabın parçaya hiç ulaşmaması — uç eleman TCP'si 156 mm girilmişti, "
                "mesh'te ölçülen 84.2 mm; kap her hedefin 71.8 mm üstünde duruyordu. "
                "Sabit artık URDF'ten ölçülüyor, elle yazılmıyor (Bölüm 7.9).")
    bullet(doc, "Başarısız bir hareketin “başarılı” raporlanması — sessiz yanlış "
                "raporlama kapatıldı.")
    bullet(doc, "Parça kavrandıktan SONRA bırakma pozunun planlanamaması — hedefler "
                "artık kavramadan önce /compute_ik ile sınanıyor (Bölüm 10).")
    bullet(doc, "Gazebo'da parçanın gerçekten taşınması (DetachableJoint) ve taşınan "
                "parçanın MoveIt çarpışma sahnesine iliştirilmesi (Bölüm 11).")
    bullet(doc, "Parçanın kalınlığının elle girilmesi — boyut artık derinlikten "
                "ölçülüyor; yüzde yerine histogram modu kullanılıyor (Bölüm 11.2).")

    figure(doc, "fig_orientation_error.png",
           "Şekil 1 — Vakum ucunun diklik hatasının gün içindeki seyri. "
           "Her değer TF'ten ya da mesh geometrisinden ölçülmüştür.")

    # ---- 2. Gemini Robotics ----
    h(doc, "2. Gemini Robotics ER 2 nedir, ne yapar?", 1)
    para(doc,
         "Gemini Robotics ER 2 (embodied reasoning), Google'ın robotik için "
         "uyarladığı bir görsel-dil modelidir. Bu projede kullanılan sürüm "
         "gemini-robotics-er-2-preview'dir.")

    h(doc, "2.1 Yapabildikleri", 2)
    bullet(doc, "Pointing: “konveyördeki dikdörtgen bloğu göster” gibi serbest "
                "metin bir sorguya karşılık görüntü üzerinde piksel koordinatı "
                "döndürür. Çıktı [y, x] biçiminde ve 0–1000 aralığına normalize edilmiştir.")
    bullet(doc, "Planlama: dağınık bir kullanıcı cümlesini yapılandırılmış bir "
                "göreve ayrıştırır (pick hedefi / place hedefi / gerekçe).")
    bullet(doc, "İlişkisel akıl yürütme: “boş göz”, “üst sıra” gibi nesne sınıfı "
                "olmayan, ilişkiyle tanımlı hedefleri çözebilir.")

    h(doc, "2.2 Yapamadıkları — mimariyi belirleyen kısıt", 2)
    rich(doc, [
        ("ER 2 ", {}),
        ("aksiyon üretmez", {"b": True}),
        (". Çıktısı yalnızca metindir; eklem açısı, kartezyen hedef ya da "
         "yörünge vermez. Aksiyon üreten VLA sürümü (Gemini Robotics 1.5) "
         "yalnızca seçili partnerlere açıktır.", {}),
    ])
    para(doc,
         "Bu kısıt mimariyi zorunlu kılar: model “beyin”, MoveIt “eller”. ER 2 "
         "nereye bakılacağını ve neyin alınacağını söyler; nereye nasıl "
         "gidileceğini bizim yazdığımız geometri ve hareket katmanı hesaplar. "
         "Kamerayı da kendisi çeviremez — tarama pozları arasındaki hareketi "
         "biz sürüyoruz (Bölüm 6).")

    h(doc, "2.3 Ölçülen maliyet", 2)
    table(doc,
          ["Çağrı", "Süre", "Not"],
          [["Planlama", "7.4 – 7.6 s", "görev başına bir kez"],
           ["Pointing (pick)", "4.6 – 4.9 s", "tarama pozu başına bir kez"],
           ["Pointing (place)", "4.6 – 16.4 s", "toolkit pozunda daha uzun"],
           ["Görev başına toplam", "~13 – 25 s", "yalnızca ER; hareket hariç"]],
          widths=[2.0, 1.6, 2.8])
    para(doc,
         "Karşılaştırma için: aynı kutuyu klasik görüntü işleme ile (düzlemden "
         "yükselen bağlı bileşen) bulmak ~5 ms sürüyor ve deterministik. ER 2'nin "
         "kazandırdığı şey hız değil, eğitim verisi gerektirmeden serbest metinle "
         "ve ilişkisel ifadelerle çalışabilmesi.", italic=True)

    doc.add_page_break()

    # ---- 3. mimari ----
    h(doc, "3. Sistem Mimarisi", 1)
    para(doc, "Komuttan harekete akış:")
    code(doc,
         "/gemini/command  (std_msgs/String, serbest metin)\n"
         "      │\n"
         "      ├─ 1. tarama pozu 0'a git (konveyör)\n"
         "      ├─ 2. derinlikten ER görüntüsü üret        [depth_render.py]\n"
         "      ├─ 3. ER 2: planlama  -> {pick, place}     [er_client.py]\n"
         "      ├─ 4. ER 2: pointing  -> piksel (u, v)     [er_client.py]\n"
         "      ├─ 5. piksel -> 3B nokta (world frame)     [geometry.py]\n"
         "      ├─ 6. yüzey düzlemi + normal + denetim     [grasp.py]\n"
         "      ├─ 7. hedef bulunamazsa sonraki tarama pozu\n"
         "      └─ 8. APPROACH → TOUCH → VAKUM → LIFT → PLACE → BIRAK\n"
         "                                                  [pick_place_node.py]")

    h(doc, "3.1 Modüller", 2)
    table(doc,
          ["Dosya", "Sorumluluk"],
          [["er_client.py",
            "ER 2 istemcisi. Modalite anlatım metinleri (IMAGE_CONTEXT), pointing ve "
            "planlama prompt'ları, JSON ayrıştırma, 0–1000 normalize koordinatın "
            "piksele çevrilmesi."],
           ["depth_render.py",
            "Derinlik haritasından ER'ye verilecek görüntüyü üretir: relief, normals, "
            "turbo, gray, intensity. Baskın düzlem uydurma ve yüzey normali hesabı burada."],
           ["geometry.py",
            "Piksel → 3B nokta (nokta bulutundan ya da derinlik + CameraInfo ile), "
            "TF ile world frame'e taşıma. Normaller için yalnızca dönme uygulayan "
            "ayrı bir fonksiyon var; öteleme eklenirse normal anlamsızlaşır."],
           ["grasp.py",
            "Yüzey yaması PCA'sı, düzlemsellik/eğim denetimi, normalin dünya eksenine "
            "oturtulması, yaklaşma kuaterniyonu, kuaterniyon-matris dönüşümleri."],
           ["locator.py",
            "Sensör abonelikleri, ER görüntüsünün üretimi ve /gemini/er_image'a "
            "yayınlanması, sorgudan 3B tespite kadar olan zincir."],
           ["vacuum.py",
            "VGC10 vakum kavrayıcı. Eklem değil, Modbus komutu "
            "(vgc10_msgs/OnRobotVGOutput)."],
           ["pick_place_node.py",
            "Görev akışı, çok pozlu tarama, MoveIt hareketleri, uç eleman geometrisi, "
            "durum yayını."],
           ["perception_node.py",
            "Yalnızca algı: /gemini/query alır, /gemini/detections ve RViz "
            "işaretçileri yayınlar. Robotu hareket ettirmez."],
           ["tool_geometry.py",
            "Uç eleman geometrisini URDF + mesh'ten ÖLÇER: dönel yüzey ekseni "
            "(emme yönü) ve ağız halkasının merkezi (TCP). Hem ros2 run measure_tcp "
            "olarak, hem de düğümün açılış doğrulaması olarak çalışır."],
           ["reachability.py",
            "Hedefi kavramadan önce /compute_ik ile sınar. İki fazlı: önce "
            "çarpışma denetimiyle, sonra denetimsiz — böylece “çarpışma” ile "
            "“menzil dışı” ayrılır. Servis yoksa doğrulama sessizce devre dışı "
            "kalır, görevi hiçbir zaman bloklamaz."],
           ["payload.py",
            "Kavranan parçanın üç boyutunu ve düzlem içi yönelimini derinlikten "
            "ölçer: destek düzlemi (histogram modu), bağlantı kısıtlı segmentasyon, "
            "ayak izi için 2B PCA."]],
          widths=[1.75, 4.65])

    h(doc, "3.2 Arayüzler", 2)
    table(doc,
          ["Topic", "Tip", "Yön", "Açıklama"],
          [["/gemini/command", "String", "giriş", "serbest metin görev komutu"],
           ["/gemini/query", "String", "giriş", "yalnızca algı sorgusu"],
           ["/gemini/status", "String", "çıkış", "JSON durum akışı (state, detay, konum, normal)"],
           ["/gemini/detections", "String", "çıkış", "JSON tespit listesi"],
           ["/gemini/er_image", "Image", "çıkış", "ER 2'nin gerçekte gördüğü kare"],
           ["/gemini/markers", "MarkerArray", "çıkış", "RViz: tespit küreleri ve normal okları"]],
          widths=[1.5, 1.0, 0.7, 3.2])
    rich(doc, [
        ("/gemini/er_image", {"code": True}),
        (" teşhis açısından kritiktir: derinlikten üretilmiş bir görüntüyle "
         "çalışırken “model neden orayı gösterdi” sorusunun tek cevabı, ona "
         "gerçekte ne gönderdiğimizi görmektir. Bu yayın sorgudan bağımsız olarak "
         "sürekli çalışır.", {}),
    ])

    doc.add_page_break()

    # ---- 4. modalite ----
    h(doc, "4. Modalite Seçimi: Neden RGB Değil, Derinlik?", 1)
    para(doc,
         "Fiziksel dünyadaki kamera SICK Visionary-T Mini V3S145-1A'dır ve "
         "RGB modu yoktur — sürücü kaynağında (visionary_t_mini.cpp:221-225) "
         "yalnızca depth, points, intensity ve statemap yayıncıları vardır. "
         "Gazebo'nun RGB kamerasıyla çalışan bir algı hattı gerçek robota "
         "taşınamazdı.")
    para(doc,
         "Bu yüzden ER 2'ye verilen görüntü her zaman derinlikten üretilir. "
         "Sim 32FC1/metre, gerçek 16UC1/milimetre yayınlar; ikisi de aynı "
         "fonksiyonda metreye indirgenir, böylece iki dünyada üretilen görüntü "
         "aynı görünür.")
    para(doc,
         "ER 2'ye görüntünün fotoğraf olmadığı, rengin mesafe ya da yükseklik "
         "kodladığı prompt içinde açıkça anlatılır. Bu anlatım render moduna göre "
         "değişir; yanlış anlatım, doğru görüntünün yanlış yorumlanmasına yol açar.")

    # ---- 5. relief ----
    h(doc, "5. Relief Render: Düz Zeminde Düz Cismi Görünür Kılmak", 1)
    para(doc,
         "İlk denemelerde ER 2 konveyördeki kutuyu bulamıyordu. Kutunun görüntüde "
         "olmadığı sanıldı; ölçüm bunun yanlış olduğunu gösterdi. Kutu kadrajın "
         "içindeydi ve derinlik verisi doğruydu — üst yüzü banttan 29 mm yüksek "
         "ölçülüyordu. Sorun, ona gönderdiğimiz görüntüde kutunun görünmemesiydi.")

    h(doc, "5.1 Kök neden", 2)
    para(doc, "İki bağımsız etken vardı:")
    bullet(doc, "normals modu yüzey yönelimini boyar. Bir düzlemin üstünde duran "
                "düz bir kutunun üst yüzü, bantla AYNI normale sahiptir; dolayısıyla "
                "aynı tonda çıkar. Kutu yalnızca ince bir kenar çizgisi olarak görünür.")
    bullet(doc, "Sabit 0.3–4.0 m'lik metrik pencere, 29 mm'lik farkı 3.7 m'lik "
                "skalaya sıkıştırır: 255 seviyenin ~2'si.")
    para(doc,
         "Her iki etken de gerçek SICK'te birebir aynı şekilde ortaya çıkardı; "
         "bu sim'e özgü bir kusur değildi.", italic=True)

    figure(doc, "fig_render_contrast.png",
           "Şekil 2 — Kutu ile bant arasındaki ton farkı. Ölçüm, kutunun üst yüzünden "
           "ve hemen yanındaki bant yüzeyinden alınan aynı satırlardaki bölgelerin "
           "ortalama parlaklığı ile yapıldı.")

    figure(doc, "fig_conveyor_normals.png",
           "Şekil 3 — normals modu (eski varsayılan). Kutu, sağdaki parlak bant "
           "yüzeyinin üzerinde yalnızca ince bir dikdörtgen kenar olarak seçiliyor. "
           "Kontrast 1.7/255.", width=4.6)

    h(doc, "5.2 Çözüm", 2)
    para(doc,
         "relief adında yeni bir render modu yazıldı. Sahnedeki baskın düz yüzeyi "
         "(bant, masa) bulur ve o yüzeye GÖRE yüksekliği renklendirir. Pencere "
         "kameranın uzaklığına değil yüzeyin kendisine göre tanımlıdır; bu yüzden "
         "kol yaklaşıp uzaklaştıkça bozulmaz ve gerçek kamerada aynı şekilde hesaplanır.")

    figure(doc, "fig_conveyor_relief.png",
           "Şekil 4 — relief modu. Aynı kare, aynı derinlik verisi. Kutu artık koyu "
           "bant üzerinde açık mavi bir dikdörtgen; yandaki bant rayları da "
           "yükseklikleriyle ayrışıyor. Kontrast 89.8/255.", width=4.6)

    h(doc, "5.3 Düzlem uydurmada üç yaklaşım", 2)
    para(doc,
         "Baskın düzlemi bulmak sanıldığından zordu. Üç yöntem denendi; ilk ikisi "
         "ölçümle elendi:")
    table(doc,
          ["Yöntem", "Sonuç"],
          [["Tüm geçerli piksellere doğrudan SVD",
            "BAŞARISIZ — raylar ve arka plan uyumu çekiyor, normal 26° eğiliyor; "
            "MAD kırpması her turda daha kötüye yakınsıyor."],
           ["Derinlik histogramının modundan tohumlama",
            "KISMEN — kameraya dik bakışta mükemmel (0.6° hata), ama 25° eğimde "
            "aynı derinlik dilimine farklı yüzeyler düşüyor ve düzlem bandın "
            "yerine RAYLARA kilitleniyor."],
           ["Normal medyanından tohumla → yükseklik modundan ayır → uyum",
            "ÇALIŞIYOR — 0/10/25/40/55° eğimde ve 5 mm gürültüde doğru."]],
          widths=[2.4, 4.0])
    para(doc,
         "Üçüncü yöntemin kilit noktası şudur: bant, raylar ve kutu PARALEL "
         "yüzeylerdir. Normal seçimi üçünü birden alır ve uyum ortalarından geçer; "
         "bunları ayıran şey yükseklik histogramının modudur.", italic=True)

    h(doc, "5.4 Doğrulama", 2)
    bullet(doc, "Sentetik düzlemler: 0°, 10°, 25°, 40°, 55° eğimde ve 5 mm "
                "gürültüde kutu yüksekliği 29.0 mm, zemin 0.0 mm okundu.")
    bullet(doc, "Gerçek kare (canlı hücre): kutu +29.5 mm, bant +1.2 mm. "
                "Gerçek değer 29 mm.")
    bullet(doc, "İşaret kuralı testle kilitlendi: düzlemin kendi medyanı tanımı "
                "gereği sıfırdır ve yön belirlemede KULLANILAMAZ. Bu hata bir kez "
                "yapıldı ve tüm yükseklikleri ters çevirdi.")

    doc.add_page_break()

    # ---- 6. çok pozlu tarama ----
    h(doc, "6. Çok Pozlu Tarama", 1)
    para(doc,
         "Kamera bileğe monte olduğu için tek bir bakış pozundan hücrenin tamamı "
         "görünmüyor: konveyör bir pozdan, toolkit bambaşka bir pozdan görünüyor "
         "ve ikisi arasında 1.17 m ray hareketi var. ER'ye planlama sorulan poz "
         "konveyör pozu olduğundan, toolkit o karede hiç yok.")

    h(doc, "6.1 Neden ucuz oldu", 2)
    para(doc,
         "Tespitler kamera frame'inde değil world frame'inde saklanıyor. Bu sayede "
         "farklı pozlardan bakılarak toplanan sonuçlar tek bir haritada birikiyor. "
         "Bu olmasaydı çok pozlu tarama zor bir problem olurdu.")

    h(doc, "6.2 Poz başına render modu", 2)
    para(doc,
         "Tek bir render modu iki bakışa birden yetmedi. Konveyöre tepeden "
         "bakılıyor: parça düz bir yüzeyin üstünde bir ÇIKINTI, relief şart. "
         "Toolkit'e yandan bakılıyor: sahne 0.68–1.84 m'ye yayılıyor, relief'in "
         "yükseklik penceresine piksellerin ancak %50'si giriyor, gerisi doyuyor.")

    figure(doc, "fig_toolkit_relief_saturated.png",
           "Şekil 5 — Toolkit pozunda relief. Sahne penceresi aştığı için görüntü "
           "doyuyor; ER bu karede hiçbir şey bulamadı.", width=4.3)

    figure(doc, "fig_toolkit_normals.png",
           "Şekil 6 — Aynı kare, normals modu. Raf yapısı ve boş gözler net "
           "seçiliyor. ER bu karede boş gözü buldu.", width=4.3)

    para(doc,
         "Render modu değiştiğinde ER'nin modalite anlatımı da otomatik "
         "değişiyor; aksi halde yönelim gölgeli bir görüntüye bakarken ona "
         "“renk yüksekliktir” demiş olurduk.")

    h(doc, "6.3 Ölçülen tarama pozları", 2)
    code(doc,
         'scan_pose_names: ["conveyor", "toolkit"]\n'
         "scan_pose_joints: [\n"
         "  0.54520, -1.08288, -0.60626,  0.53539, -3.07005,  1.08265,  0.00141,\n"
         "  1.71329, -1.40881, -1.01420,  0.73822, -2.86483,  1.40865, -0.54466,\n"
         "]\n"
         'scan_pose_render_mode:  ["relief", "normals"]\n'
         "scan_pose_height_lo_m:  [-0.005, -0.080]\n"
         "scan_settle_sec: 1.0")
    para(doc,
         "Eklem sırası: base_to_robot_mount (lineer ray), shoulder_pan, "
         "shoulder_lift, elbow, wrist_1, wrist_2, wrist_3. /joint_states'in name "
         "dizisi bu sırada DEĞİLDİR; oradan kopyalanan değerler yeniden sıralanmalıdır.")

    h(doc, "6.4 İki ince nokta", 2)
    rich(doc, [("Boş göz bir çukurdur. ", {"b": True}),
               ("relief'in yükseklik penceresinin alt sınırı negatif olmazsa "
                "boş göz yüzeyle aynı tonda kalır. 45 mm derin bir göz için "
                "ölçüldü: alt sınır −0.005 m iken kontrast 28/255, −0.08 m iken "
                "50/255.", {})])
    rich(doc, [("Poz değişiminden sonra beklemek zorunludur. ", {"b": True}),
               ("Kamera ve nokta bulutu bir önceki pozun karesini taşıyor olabilir. "
                "Taze veri gelmeden sorgu sorulursa ER'ye YANLIŞ sahne gösterilir, "
                "üstelik tespit o anki TF ile çevrildiği için sessizce hatalı bir "
                "3B nokta üretilir — gürültülü bir hata değil, sessiz bir hata.", {})])

    doc.add_page_break()

    # ---- 7. vakum ucu (asıl bölüm) ----
    h(doc, "7. Vakum Ucunun Yüzeye Tam Dik Getirilmesi", 1)
    para(doc,
         "Bu bölüm, günün en çok iterasyon gerektiren problemini ayrıntısıyla "
         "anlatır. Emme kabı dönel simetriktir ama yüzeye DİK oturmak zorundadır: "
         "Ø40 mm'lik bir kapta 4°'lik eğim, kenarda ~1.6 mm boşluk demektir ve "
         "gerçek donanımda vakum sızdırır.")

    h(doc, "7.1 Belirti", 2)
    para(doc,
         "Kol parçaya doğru gidiyordu ama emme kabı yüzeye dik değil, YATAY "
         "geliyordu. Üstelik MoveIt hiçbir hata vermiyordu: istenen oryantasyona "
         "ulaştığını raporluyordu ve bu doğruydu.")

    figure(doc, "fig_wrong_orientation.png",
           "Şekil 7 — Kullanıcı ekran görüntüsü: kol kutunun (yeşil nokta bulutu) "
           "üzerinde, ancak uç eleman yatay duruyor.", width=4.0)

    h(doc, "7.2 İlk ölçüm: MoveIt suçsuz", 2)
    para(doc,
         "Kol o pozdayken TF'ten uç eleman frame'inin eksenleri okundu:")
    code(doc,
         "ur10e_suction_cup   pos = (0.745, 0.102, 1.007)\n"
         "   +X = (-0.000, +1.000, -0.011)\n"
         "   +Y = (+1.000, +0.000, +0.026)\n"
         "   +Z = (+0.026, -0.011, -1.000)     <- tam AŞAĞI")
    para(doc,
         "Frame'in +Z'si tam aşağı bakıyordu, yani komut doğru uygulanmıştı. "
         "Demek ki hata komutta değil, VARSAYIMDAYDI: kod, emme yönünün frame'in "
         "+Z'si olduğunu kabul ediyordu.")

    h(doc, "7.3 İkinci ipucu: frame orijini alet ekseninde değil", 2)
    para(doc, "URDF zinciri incelendi:")
    code(doc,
         "tool0 --(0.04267, 0.056052, -0.0195)--> vacuum_gripper\n"
         "      --(0.070588, 0.044218, +0.0195)--> suction_cup\n"
         "toplam: (0.1133, 0.1003, 0.0000)   <- TAMAMEN YANAL, eksende sıfır")
    para(doc,
         "Bir vakum kabının frame orijininin, alet ekseni boyunca sıfır ofsetle "
         "15 cm yanda olması fiziksel olarak anlamsızdır. Bu, frame'in CAD'den "
         "geldiği gibi bırakıldığının işaretiydi.")

    h(doc, "7.4 Birinci düzeltme: siluet analizi (+X) — yetersiz", 2)
    para(doc,
         "Mesh üç eksenden siluet olarak çizdirildi. X ekseninden bakınca daire, "
         "diğer iki eksenden sap ve çan profili görünüyordu.")

    figure(doc, "fig_cup_silhouettes.png",
           "Şekil 8 — suction_cup.stl mesh'inin üç ortografik silueti. Soldaki "
           "(X ekseninden bakış) dairesel; bu, kabın ekseninin kabaca X yönünde "
           "olduğunu gösterir — ama 30°'lik sapmayı ayırt etmeye yetmez.", width=6.4)

    para(doc,
         "Buna dayanarak emme ekseni +X kabul edildi ve kod genelleştirildi. "
         "Sonuç iyileşti (90° → 30°) ama kullanıcı “hâlâ tam dik değil” dedi. "
         "Yapılan ölçüm bunu doğruladı: frame'in +X'i artık tam dikeydi, yani "
         "sorun eksen seçiminde değil, eksenin GERÇEK yönündeydi.")

    h(doc, "7.5 İkinci düzeltme: dönel yüzey ekseni uydurma", 2)
    para(doc,
         "Kesin sonuç için geometrik bir uyum yapıldı. Bir dönel yüzeyde her "
         "yüzey normali ekseni keser; buradan u (eksen yönü) ve c (eksen üzerinde "
         "bir nokta) için şu kısıt yazılır:")
    code(doc, "u · ( n_i × (p_i − c) ) = 0        her i üçgeni için")
    para(doc,
         "Bu kısıt, alan ağırlıklı olarak ve dönüşümlü (önce u, sonra c) çözüldü. "
         "Sonuç:")
    code(doc,
         "DÖNEL EKSEN = (0.8661, 0.4999, 0.0002)\n"
         "            = (cos 30°, sin 30°, 0)\n"
         "frame +X ile açı = 30.00°")
    para(doc,
         "Tam 30° çıkması, bunun gürültü değil CAD'de kasıtlı bir montaj açısı "
         "olduğunu gösterir.")

    h(doc, "7.6 Çapraz doğrulama", 2)
    table(doc,
          ["Bağımsız ölçüm", "Bulunan yön", "Fark"],
          [["Dönel yüzey ekseni uydurma", "(0.8661, 0.4999, 0.0002)", "—"],
           ["Ağız düzlemi uyumu", "(0.8709, 0.4915, -0.0004)", "0.56°"],
           ["Montaj yüzü normalinin tersi", "(0.8584, 0.4956, -0.1327)", "7.64°"]],
          widths=[2.6, 2.4, 1.4])

    para(doc, "Denenip işe yaramayan yöntemler de kayda değer:")
    bullet(doc, "Yüzey normali kovaryansı — özdeğerler neredeyse izotropik "
                "(0.299 / 0.304 / 0.398), mesh temiz bir silindir olmadığı için sonuçsuz.")
    bullet(doc, "Açık kenar (mesh boundary) analizi — mesh su geçirmez, açık kenar yok.")
    bullet(doc, "En geniş düz uç yüz araması — kabın ağzını değil MONTAJ yüzünü buluyor.")

    h(doc, "7.7 Belirleyici kanıt: yarıçap profili", 2)
    para(doc,
         "Uydurulan eksenin doğruluğunun en okunaklı kanıtı, eksen boyunca alınan "
         "kesitlerin yarıçapıdır. Yanlış eksende profil tutarsızca zıplar; doğru "
         "eksende fiziksel olarak anlamlı bir parça çıkar.")

    figure(doc, "fig_radius_profile.png",
           "Şekil 9 — Kesit yarıçapının eksen boyunca değişimi. Sağdaki profil "
           "bir vakum kabının beklenen geometrisidir: ince sap, 30 mm boyunca "
           "sabit 7.1 mm şaft, sonra Ø40 mm ağza açılan çan.")

    h(doc, "7.8 Uç noktası (TCP) düzeltmesi", 2)
    rich(doc, [
        ("MoveIt hedefi uç eleman link frame'inin ORİJİNİ için verilir; parçaya "
         "değen yer ise kabın ağzıdır. Kodda tüm hareket hedefleri kap ağzı "
         "cinsinden verilir ve MoveIt'e gönderilmeden önce şu dönüşümden geçer: ", {}),
        ("frame = uç − R · tip_offset", {"code": True}),
        (". Log her iki değeri de yazar.", {}),
    ])
    para(doc,
         "Bu dönüşümde tip_offset'in hatası doğrudan mesafeye yazılır ve iki "
         "yönlüdür: offset gereğinden UZUNSA frame geride durur ve kap yüzeye "
         "hiç değmez; KISAYSA kap yüzeye bastırır. Aradaki fark oryantasyondan "
         "bağımsızdır — kap tam dik gelse bile yanlış yerde durur.")

    h(doc, "7.9 Sonradan bulunan hata: TCP 72 mm fazla girilmişti", 2)
    para(doc,
         "11 Ağustos 2026'da, kolun konveyördeki kutuya gözle görülür biçimde "
         "yaklaşmadığı fark edildi. İlk şüphe algıya gitti — ama ölçüm tersini "
         "söyledi.")
    para(doc, "Mesh doğrudan ölçüldüğünde:")
    table(doc,
          ["", "config'te yazan", "Mesh'te ölçülen"],
          [["Ağız merkezi", "(0.1351, 0.0780, 0.0000)", "(0.0730, 0.0421, 0.0000)"],
           ["Frame orijininden uzaklık", "156.0 mm", "84.2 mm"],
           ["Ağız halkası", "—", "yarıçap 19.5 mm, düzlemsellik 0.07 mm"]],
          widths=[1.8, 2.0, 2.9])
    rich(doc, [
        ("Mesh eksen boyunca 84.2 mm'de bitiyor. 156 mm'lik nokta kabın "
         "72 mm ÖTESİNDE, havadadır. ", {"b": True}),
        ("Sonuç: kap her hedefin 71.8 mm üstünde duruyor, vakum boşa çekiyordu.", {}),
    ])
    para(doc,
         "Bu hatanın sinsiliği belirtisinde: kol yanlış yerde durunca akla önce "
         "kamera geliyor, oysa ER 2 de derinlik de doğruydu. Aynı raporun 9. "
         "şeklindeki yarıçap profili zaten mesh'in 84 mm'de bittiğini gösteriyordu "
         "— sayı ile grafik çelişiyordu ve grafik haklıydı.", italic=True)

    h(doc, "7.9.1 Kalıcı çözüm: sabiti ölçen araç", 3)
    para(doc,
         "Sayıyı düzeltmek yetmez; aynı hatanın bir daha yazılamaması gerekir. "
         "Bu iki sabit tahmin edilecek şey değil, robotun kendi modelinde yazılı. "
         "Artık oradan okunuyorlar:")
    code(doc,
         "ros2 run gemini_robotics_ros measure_tcp\n"
         "  → URDF'i /robot_state_publisher'dan alır (MoveIt'in planladığı model)\n"
         "  → link'in mesh'ini, ölçeğini ve origin'ini çözer\n"
         "  → mesh'e dönel yüzey ekseni uydurur      : emme yönü\n"
         "  → eksen boyunca en uzak halkanın merkezi : TCP\n"
         "  → yapıştırılacak YAML satırlarını yazar")
    para(doc,
         "Ayrıca pick_place_node açılışta aynı ölçümü yapıp config'tekiyle "
         "karşılaştırır; 3 mm'den fazla fark varsa hata seviyesinde log basar. "
         "Doğrulama bloklamaz ve görevi düşürmez — amaç, sessizce yanlış bir "
         "sayıyla çalışmayı imkânsız kılmaktır. Kap değişirse elle güncellenecek "
         "hiçbir şey yoktur: mesh değişir, ölçüm değişir.")

    h(doc, "7.10 Üçüncü düzeltme: normalin dünya eksenine oturtulması", 2)
    para(doc,
         "Geometri tamamen doğru olduktan sonra bile ölçülen sapma sıfır değildi. "
         "Bırakma pozunda TF'ten okundu: 3.99°. Kaynak artık geometri değil, "
         "yüzey normalinin gürültülü derinlik yamasından PCA ile kestirilmesiydi.")
    para(doc,
         "Bant ve raf yüzeyleri işlenmiş, dünya eksenlerine göre düz yüzeylerdir; "
         "o birkaç derece sensör gürültüsüdür. Normal, verilen tolerans içindeyse "
         "en yakın dünya eksenine tam olarak oturtulur.")
    rich(doc, [("Sıra önemlidir: ölç → denetle → oturt. ", {"b": True}),
               ("Kavranabilirlik eğim denetimi ÖLÇÜLEN normale bakmalıdır; "
                "oturtulmuş normale bakarsa gerçekten eğik bir yüzey de “düz” "
                "görünür ve denetim anlamını yitirir.", {})])
    para(doc,
         "Bunun bir modelleme varsayımı olduğu unutulmamalıdır: yüzeylerin eksene "
         "hizalı olduğunu kabul eder. Gerçekten eğik yüzeylerle çalışılacaksa "
         "normal_snap_deg parametresi 0.0 yapılarak kapatılır.", italic=True)

    h(doc, "7.11 Sonuç", 2)
    para(doc,
         "Uç elemanın doğru gelmesi İKİ ayrı büyüklüğün birden doğru olmasını "
         "gerektiriyor ve bunlar bağımsız olarak yanlış olabiliyor: yüzeye "
         "diklik (açı) ve ağzın yüzeye ulaşması (mesafe).")
    table(doc,
          ["Aşama", "Varsayım / düzeltme", "Açı sapması", "Mesafe hatası"],
          [["Başlangıç", "emme ekseni = frame +Z", "90°", "—"],
           ["1. düzeltme", "emme ekseni = frame +X", "30°", "—"],
           ["2. düzeltme", "emme ekseni = (cos30°, sin30°, 0)", "3.99°", "71.8 mm"],
           ["3. düzeltme", "normal dünya eksenine oturtuluyor", "0.00°", "71.8 mm"],
           ["4. düzeltme", "TCP mesh'ten ölçüldü: 84.2 mm", "0.00°", "0.04 mm"]],
          widths=[1.15, 2.95, 1.1, 1.5])
    para(doc,
         "Canlı koşuda log her tespitte ölçülen açı sapmasını yazıyor: konveyörde "
         "1.23°, toolkit'te 11.13°. İkisi de 15°'lik tolerans içinde kaldığı için "
         "oturtuldu ve uç her iki yüzeye de tam dik geldi. Mesafe tarafı ise artık "
         "uçtan uca testle bağlanmış durumda: hedefi frame'e çevirip kabın ağzını "
         "geri hesaplayan test, kalan hatanın mikron mertebesinde olduğunu "
         "doğruluyor (test/test_tool_geometry.py).")

    doc.add_page_break()

    # ---- 8. altyapı hataları ----
    h(doc, "8. Bulunan ve Kapatılan Altyapı Hataları", 1)

    h(doc, "8.1 Düğümün executor'dan düşmesi", 2)
    para(doc,
         "En sinsi hata buydu. İlk MoveIt hareketinden sonra düğüm tamamen sağır "
         "kalıyordu: süreç yaşıyor, CPU normal, /gemini/er_image yayınlanmaya "
         "devam ediyor, ama komut aboneliği hiç tetiklenmiyor ve ros2 param get "
         "yanıt vermiyordu.")
    para(doc, "Kök neden rclpy'nin düğüm sahiplik kuralıdır:")
    code(doc,
         "rclpy.spin_once(node)  →  get_global_executor().add_node(node)\n"
         "Executor.add_node      →  node.executor = self\n"
         "Node.executor setter   →  ESKİ executor.remove_node(self)   ← düşüş burada\n"
         "spin_once finally      →  global_executor.remove_node(node)")
    para(doc,
         "pymoveit2 beş ayrı yerde rclpy.spin_once(self._node, ...) çağırıyor. "
         "Sonuç: ilk hareketten sonra düğüm HİÇBİR executor'a ait değil. "
         "MultiThreadedExecutor boş bir düğüm kümesiyle sonsuza kadar bekliyor.")
    para(doc, "Bağımsız bir test betiğiyle mekanizma birebir üretildi:")
    code(doc,
         "baseline delivered            : True\n"
         "node in executor._nodes       : True\n"
         "--- rclpy.spin_once(node) sonrası ---\n"
         "node in executor._nodes       : False\n"
         "delivered after spin_once     : False\n"
         "--- executor.add_node(node) sonrası ---\n"
         "delivered after repair        : True")
    para(doc,
         "Çözüm: hareket beklerken spin_once çağırmayan bir bekleme yazıldı ve her "
         "MoveIt çağrısından sonra düğüm executor'ına geri bağlanıyor.")

    h(doc, "8.2 Başarısız hareketin “başarılı” raporlanması", 2)
    para(doc,
         "Kartezyen bir plan %46.7'de kaldı, kol parçaya hiç değmedi, ama akış "
         "devam etti: vakum havada açıldı ve görev “tamamlandı” diye bitti.")
    para(doc,
         "Sebep, spin_once'ı kaldırırken pymoveit2'nin “hiç hareket istenmedi” "
         "korumasının da düşmesiydi. Planlama başarısız olduğunda hedef hiç "
         "gönderilmiyor, bayrak hiç set edilmiyor ve bekleme anında dönüp BİR "
         "ÖNCEKİ hareketten kalan başarı bayrağını okuyordu.")
    para(doc,
         "Çözüm: her komuttan önce başarı bayrağı sıfırlanıyor. Sessiz yanlış "
         "rapor, kilitlenmeden daha tehlikelidir.", italic=True)

    h(doc, "8.3 Kısaltılmış kartezyen planın tam sayılması", 2)
    rich(doc, [
        ("pymoveit2, kartezyen planı ", {}),
        ("fraction >= eşik", {"code": True}),
        (" olduğunda YÜRÜTÜR — ve kısaltılmış yörünge de yürütülür. Eşik 0.9 "
         "olduğu için 148 mm'lik APPROACH→TOUCH hareketinin 14.8 mm eksik "
         "kalması “başarı” sayılabiliyordu.", {}),
    ])
    para(doc,
         "Vakumlu kavramada eksik kalan son milimetreler doğrudan tutamamak "
         "demek. Eşik parametreye çıkarıldı ve 0.99'a çekildi "
         "(cartesian_min_fraction).")

    h(doc, "8.4 Derinlik bandının yanlış eksende filtrelemesi", 2)
    para(doc,
         "ER 2'nin gösterdiği pikselin çevresine düzlem oturtulurken, yamaya "
         "arka planın karışmaması için bir derinlik bandı uygulanıyor. Band "
         "noktanın z bileşenine bakıyordu: optik frame'de (z = ileri) doğru, "
         "ama Gazebo bulutu x-ileri konvansiyonunda geliyor — orada z YANAL "
         "bir eksendir ve filtre arka planı hiç elemiyordu.")
    para(doc,
         "Derinlik artık BAKIŞ IŞINI boyunca ölçülüyor (noktanın çapa ışınına "
         "izdüşümü). Kamera her iki konvansiyonda da orijinde olduğu için bu "
         "ölçüm ikisinde de aynı şeyi verir; sim ile gerçek arasında sessizce "
         "ayrışacak bir varsayım kalmadı.", italic=True)

    h(doc, "8.5 Düğümün çökmesi: iki executor tek düğümü paylaşıyor", 2)
    para(doc,
         "11 Ağustos'ta görev TOUCH hareketini planlarken süreç exit 1 ile "
         "öldü. Log'da üç ayrı traceback vardı; üçü de aynı kök nedenin "
         "ardılıydı ve ilk bakışta yanlış yere bakılıyordu.")
    code(doc,
         "pymoveit2 (moveit2.py:527, 639, 747, 1189, 1281):\n"
         "    while not future.done():\n"
         "        rclpy.spin_once(self._node, timeout_sec=1.0)\n\n"
         "rclpy.spin_once  ->  get_global_executor().add_node(node)\n"
         "                 ->  bizim MultiThreadedExecutor'dan SÖKÜLÜR")
    para(doc,
         "Bizim executor o sırada wait_for_ready_callbacks içinde waitable'ları "
         "dolaşıyordu. Action client'ın wait set indeksleri artık öteki "
         "executor'ın daha küçük wait set'ine göre yazılmıştı:")
    code(doc,
         "RCLError: Failed to get number of ready entities for action client:\n"
         "  wait set index for status subscription is out of bounds")
    para(doc,
         "Bu istisna executor.spin()'den dışarı çıktı; main yalnızca "
         "KeyboardInterrupt yakaladığı için finally çalışıp düğümü yok etti ve "
         "rclpy'yi kapattı. Sonrası domino: görev thread'i bir sonraki "
         "spin_once'ta ölü context'e çarpıp AttributeError verdi, hata "
         "yakalayıcısı da kapanmış publisher'a yazmaya çalışıp InvalidHandle "
         "verdi. MoveIt tarafı sapasağlamdı - aynı saniyede kartezyen yolu "
         "%100 hesapladığını loglamıştı.")
    rich(doc, [
        ("Test yazarken beklenmedik bir ayrıntı çıktı: ", {}),
        ("spin_once'ın finally'sindeki remove_node, düğümün KENDİ executor "
         "referansına hiç dokunmuyor.", {"b": True}),
        (" Çağrıdan sonra düğüm ne bizim ne global executor'ın listesinde, ama "
         "node.executor hâlâ global'i gösteriyor. Sahiplik kalıcı olarak el "
         "değiştiriyor.", {}),
    ])
    para(doc,
         "Çözüm, süreç içinde rclpy.spin_once'ı sarmalamak: düğümün zaten bir "
         "sahibi varsa global executor'a hiç verilmiyor, kısa bir uyku ile "
         "bekleniyor. Callback'leri bizim executor çalıştırmaya devam ettiği "
         "için future yine tamamlanıyor.")
    para(doc,
         "pymoveit2'nin kendisi DEĞİŞTİRİLMEDİ: o fork'u viewpoint_planner ve "
         "multirobot_viewpoint_planner da kullanıyor ve orada aynı değişiklik, "
         "düğümü kendi executor'ında spin etmeyen bir çağırana kilitlenme "
         "yaşatırdı. Yama yalnızca bu düğümün sürecinde geçerli.", italic=True)

    h(doc, "8.6 Diğerleri", 2)
    table(doc,
          ["Sorun", "Kök neden", "Çözüm"],
          [["Konveyör Gazebo'ya spawn olmuyor",
            "Dünya adı 'cem', spawn komutları '-world ifarlab' kullanıyordu",
            "Üç spawn çağrısı düzeltildi"],
           ["/sim/depth/image rgb8 taşıyordu",
            "Köprü, derinlik yerine RGB topic'ine bağlıydı",
            "Kaynak .../depth_image olarak düzeltildi"],
           ["Tüm ROS timer'ları donuk",
            "gz iki clock yayınlar; köprü SESSİZ olana bağlıydı",
            "/world/cem/clock'a bağlandı; 8 s'de 801 mesaj doğrulandı"],
           ["Normals render'ı neredeyse siyah",
            "Işık yönü ters; kameraya bakan yüzeyler lambert=0 alıyordu",
            "Işık Z bileşeni pozitife çevrildi"],
           ["Normaller gürültülü",
            "Tek piksel farkı, 5 mm derinlik gürültüsünü ~25° eğime çeviriyor",
            "Bilateral filtre + geniş stencil (std 48.5 → 5.1)"]],
          widths=[1.9, 2.4, 2.1])

    doc.add_page_break()

    # ---- 9. kullanım ----
    h(doc, "9. Kurulum ve Kullanım", 1)
    para(doc, "Hücreyi başlatın, sonra paketi çalıştırın:")
    code(doc,
         "export GEMINI_API_KEY=...        # anahtar config dosyasına YAZILMAZ\n"
         "\n"
         "ros2 launch my_robot_cell_control hil_test_whole_unified.launch.py \\\n"
         "     use_vacuum_gripper:=true\n"
         "\n"
         "ros2 launch gemini_robotics_ros gemini_pick_place.launch.py             # sim\n"
         "ros2 launch gemini_robotics_ros gemini_pick_place.launch.py mode:=real  # gerçek\n"
         "\n"
         'ros2 topic pub --once /gemini/command std_msgs/msg/String \\\n'
         '     "{data: \'konveyördeki şeyi al, toolkitin üst sırasındaki boş göze koy\'}"')

    para(doc,
         "Launch'ın TEK argümanı vardır: mode. Sim ile gerçek arasındaki tek fark "
         "kamera ve saattir; hareket tarafı değişmez.")
    table(doc,
          ["", "mode:=sim (varsayılan)", "mode:=real"],
          [["bulut / derinlik", "/sim/pointcloud, /sim/depth/image", "/points, /depth"],
           ["camera_info", "/sim/camera_info", "/camera_info"],
           ["kodlama", "32FC1, metre", "16UC1, milimetre"],
           ["use_sim_time", "true", "false"],
           ["MoveIt", "/move_action", "/move_action — aynı"]],
          widths=[1.6, 2.6, 2.2])
    para(doc,
         "MoveIt her iki modda da gerçek /move_action'a komut gönderir; Gazebo'daki "
         "robot real_to_sim_bridge sayesinde bunu aynalar. Gerçek robota geçerken "
         "hareket tarafında değişecek hiçbir şey yoktur — bu bilinçli bir tasarım "
         "tercihidir.")

    h(doc, "9.1 Ölçülmüş kritik parametreler", 2)
    table(doc,
          ["Parametre", "Değer", "Kaynağı"],
          [["tool_approach_vector", "[0.8660254, 0.5, 0.0]", "measure_tcp: dönel eksen (30.01°)"],
           ["tool_tip_offset", "[0.072920, 0.042100, 0.0]", "measure_tcp: ağız halkası, 84.2 mm"],
           ["cartesian_min_fraction", "0.99", "0.9'da TOUCH 14.8 mm eksik kalabiliyordu"],
           ["normal_snap_deg", "15.0", "ölçülen sapmalar 1.2–11.1° arasında"],
           ["render_mode", "poz başına: relief / normals", "kontrast ölçümleri"],
           ["descend_before_release", "false", "raf gözü gövdeyi almıyor"],
           ["deprojection", "cloud", "depth yolu bu hücrede 90° dönük çıkıyor"],
           ["validate_targets", "true", "kavramadan önce IK sınaması (Bölüm 10)"],
           ["approach_distance_candidates", "[0.15 … 0.06]", "ilk ulaşılabilir yaklaşma yüksekliği seçilir"],
           ["payload_measure", "true", "boyut derinlikten ölçülür, elle girilmez"],
           ["payload_margin_m", "0.005", "ölçüm belirsizliği için kenar payı"]],
          widths=[2.0, 2.2, 2.2])

    # ---- 10. açık maddeler ----
    doc.add_page_break()

    # ---- 10. hedef doğrulama ----
    h(doc, "10. Hedefin Ulaşılabilirliğini Kavramadan Önce Sınamak", 1)
    para(doc,
         "ER 2'nin gösterdiği nokta DOĞRU olabilir ama kolun oraya gitmesi "
         "imkânsız olabilir. Bunu hareket sırasında öğrenmek pahalıdır: "
         "11 Ağustos'ta parça kavrandıktan sonra bırakma pozu planlanamadı ve "
         "kol havada parçayla asılı kaldı.")

    h(doc, "10.1 Toolkit gerçekte ne", 2)
    para(doc,
         "stackable_bin mesh'ine dikey ışın atıldı. Toolkit tek bir raf değil, "
         "beş katlı bir ünite:")
    table(doc,
          ["Ölçüm", "Değer"],
          [["Raf tabanları (z)", "0.002 | 0.259 | 0.517 | 0.727 | 0.938 m"],
           ["Yapının tepesi", "1.095 m"],
           ["Göz iç yüksekliği", "211 mm"],
           ["Göz ağzı genişliği", "172 mm"],
           ["Gözün açıldığı yön", "−X (Y'de iki yan da duvar)"],
           ["vacuum_gripper gövde çapı", "167 mm"]],
          widths=[2.2, 4.2])
    rich(doc, [
        ("Sonuç: üst sıra dışındaki hiçbir göze bu uç elemanla girilemez. ",
         {"b": True}),
        ("Yukarıdan dik inmek tavan yüzünden imkânsız, yandan girmek de 172 mm "
         "ağza 167 mm gövde sığmadığı için. En üstteki göz (taban 0.938) o XY'de "
         "üstünde hiçbir yüzey olmadığı için yukarı açıktır ve oraya "
         "bırakılabilir.", {}),
    ])

    h(doc, "10.2 MoveIt'e sorulan soru", 2)
    para(doc, "Tahmin yerine ölçüm - /compute_ik iki kez çağrıldı:")
    table(doc,
          ["Poz", "IK (çarpışmalı)", "IK (çarpışmasız)", "Sonuç"],
          [["ER'nin seçtiği göz", "ÇÖZÜLEMEDİ (−31)", "ULAŞILABİLİR",
            "kap ↔ bin, 9.5 mm"],
           ["Üst sıra, taban +30 mm", "ULAŞILABİLİR", "ULAŞILABİLİR", "geçerli"],
           ["Üst sıra, taban +73 mm", "ULAŞILABİLİR", "ULAŞILABİLİR", "geçerli"],
           ["Üst sıra, taban +150 mm", "ULAŞILABİLİR", "ULAŞILABİLİR",
            "önkol ↔ kablo kanalı"],
           ["Konveyör yaklaşma", "ULAŞILABİLİR", "ULAŞILABİLİR", "geçerli"]],
          widths=[1.9, 1.6, 1.6, 1.6])
    para(doc,
         "İki aşamalı olmasının sebebi, “olmuyor”un iki farklı anlamı olması: "
         "çarpışmada yaklaşma mesafesini küçültmek işe yarayabilir, "
         "erişimsizlikte yaramaz - hedefin kendisi bırakılmalıdır.")

    h(doc, "10.3 Eklenen üç davranış", 2)
    bullet(doc, "Her aday hedef, kavramadan önce /compute_ik ile sınanıyor; "
                "geçmeyen aday atlanıp ER'nin verdiği bir sonraki deneniyor.")
    bullet(doc, "Yaklaşma mesafesi uyarlanıyor: 0.15 → 0.12 → 0.10 → 0.08 → "
                "0.06. Tek sabit değer bantla raf gözüne aynı anda uymuyor.")
    bullet(doc, "Pointing prompt'u artık EŞLEŞEN BÜTÜN örnekleri istiyor. Tek "
                "aday gelirse doğrulama sadece daha erken başarısız olur, "
                "görevi kurtarmaz.")
    para(doc,
         "/compute_ik yoksa doğrulama kendini kapatır ve hedefi REDDETMEZ: "
         "doğrulanamayan bir hedefi reddetmek, çalışan bir görevi altyapı "
         "eksikliği yüzünden durdurmak olurdu.", italic=True)

    h(doc, "10.4 Komutun kendisi de bir parametredir", 2)
    para(doc,
         "İlk komut “...in the toolkits top row” diyordu ve ER üstten İKİNCİ "
         "sırayı gösterdi - o da tavanlı olan. Komut şöyle değiştirildi:")
    code(doc,
         "Pick up the object on the conveyor belt and place it into an empty\n"
         "bin on the topmost level of the toolkit rack - the level that is\n"
         "open from above, with no shelf above it.")
    rich(doc, [
        ("Kritik ekleme “open from above, with no shelf above it”. ", {"b": True}),
        ("Bu bir süsleme değil, ayırt edici özelliğin kendisi: beş kattan "
         "yalnızca birinin üstü açık. ER 2 ilişkisel akıl yürütebildiği için "
         "“üstünde raf olmayan” doğrulanabilir bir ipucu, “en üstteki”nden çok "
         "daha güçlü. Bu komutla görev uçtan uca tamamlandı ve ER dört aday "
         "birden döndürdü.", {}),
    ])

    doc.add_page_break()

    # ---- 11. parçanın taşınması ----
    h(doc, "11. Parçanın Gerçekten Taşınması", 1)
    para(doc,
         "Kavranan parça artık robotun bir parçasıdır ve iki ayrı yerde temsil "
         "edilmesi gerekir. İkisi farklı sistemlerdir ve farklı sebeplerle "
         "gereklidir.")
    table(doc,
          ["Nerede", "Ne için", "Nasıl"],
          [["Gazebo", "parça kolla birlikte gitsin",
            "DetachableJoint eklentisi, /vacuum/attach ve /vacuum/detach"],
           ["MoveIt", "planlayıcı parçayı görsün",
            "sahneye kutu eklenip ur10e_suction_cup'a iliştirilir"]],
          widths=[1.1, 2.2, 3.1])
    rich(doc, [
        ("MoveIt tarafı GERÇEK ROBOTTA DA gereklidir", {"b": True}),
        (": bilinmezse kol parçayı rafa, konveyöre ya da kendine sürter - "
         "planlayıcı için orada hiçbir şey yoktur.", {}),
    ])

    h(doc, "11.1 Gazebo: ölçülen üç davranış", 2)
    para(doc,
         "gz-sim 8.14'te vakum/suction sistemi YOK; tek yol koparılabilir "
         "eklem. Davranışı ayrı bir test dünyasında ölçüldü:")
    table(doc,
          ["Adım", "Yükün z'si", "Sonuç"],
          [["Başlangıç", "0.800 m", "eklenti BAĞLI başlıyor"],
           ["detach", "0.025 m", "yere düştü - çalışıyor"],
           ["attach + tutucuyu 1 m taşı", "1.025 m",
            "yükü peşinden getirdi - yeniden bağlanma çalışıyor"]],
          widths=[2.0, 1.4, 3.0])
    para(doc,
         "Birincisi bir tuzak: eklenti, child_model dünyada belirir belirmez "
         "eklemi kurar. RedCube 22. saniyede spawn edildiği için launch 26. "
         "saniyede bir kez detach yayınlar; yoksa kutu ilk hareketten itibaren "
         "kola kaynaklı olur ve sürüklenir. Düğüm de her görev başında ayırır.")
    para(doc,
         "Üçüncüsü tasarımı belirledi: bağ, mesajın ULAŞTIĞI ANDAKİ göreli "
         "pozla kurulur. Bu yüzden attach, vakum komutuyla aynı anda - kap "
         "parçaya değerken - gönderilir.")
    para(doc,
         "Köprü tek yönlüdür (ROS → Gazebo) ve yalnızca sim launch'ında vardır; "
         "gerçek hücrede bu topic'leri kimse dinlemez, dolayısıyla aynı "
         "yayınlar zararsızdır. Mode overlay'lerine yeni anahtar eklemek "
         "gerekmedi.", italic=True)

    h(doc, "11.2 Parçanın boyutu derinlikten ölçülüyor", 2)
    rich(doc, [
        ("Yaygın sezginin tersi: derinlik kamerasında YÜKSEKLİK doğrudan "
         "ölçülen büyüklüktür.", {"b": True}),
        (" Zor olan yanal boyutlardır, çünkü onlar “cisim nerede bitiyor” "
         "kararını, yani bir segmentasyonu gerektirir.", {}),
    ])
    para(doc, "Aynı kareden dördü birden çıkarılıyor:")
    bullet(doc, "Destek düzlemi: üst yüz zaten uydurulmuş (fit_surface); parça "
                "düz bir yüzeyin üstünde durduğuna göre destek ona paraleldir "
                "ve tek bilinmeyen uzaklığıdır. Noktaların normal boyunca "
                "izdüşümü 1B histograma dönüşür, destek alttaki baskın tepedir.")
    bullet(doc, "Segment: destek düzleminin üstünde kalan ve ER'nin gösterdiği "
                "piksele BAĞLI pikseller. Bağlılık şart - yan yana duran ikinci "
                "bir cisim yoksa tek bir dev kutuya birleşir.")
    bullet(doc, "Kutu: o noktaların düzlem içindeki 2B PCA'sı uzunluk, genişlik "
                "ve düzlem içi YÖNELİMİ verir.")
    para(doc,
         "Yükseklik en güvenilir olanıdır çünkü iki gürültülü nokta arasındaki "
         "fark değil, her biri yüzlerce pikselden çıkan iki SEVİYE arasındaki "
         "farktır.")

    h(doc, "11.3 Yüzdelik tuzağı", 2)
    para(doc,
         "İlk sürümde üst seviye 98. yüzdelikten alınıyordu. Test bunu yakaladı:")
    table(doc,
          ["Piksel gürültüsü", "p98 ile ölçülen", "Tepe (mode) ile"],
          [["0 mm", "30 mm", "30 mm"],
           ["2 mm", "35 mm  (%17 fazla)", "30 mm"],
           ["5 mm", "daha da fazla", "±3 mm içinde"]],
          widths=[2.0, 2.2, 2.2])
    para(doc,
         "Sebep basit: p98 yaklaşık 2σ'ya denk gelir, yani yüzdelik gürültüyü "
         "SİSTEMATİK olarak yukarı yanlar ve hata sensör kötüleştikçe büyür. "
         "Histogram tepesi yansızdır.")
    para(doc,
         "Bedeli: düz olmayan bir üst yüzde baskın yüzeyin üstündeki çıkıntılar "
         "kutuya girmez. Bu, hattın zaten kabul ettiği bir varsayım - "
         "is_graspable düzlemsellik artığı 4 mm'yi aşan yüzeyi reddediyor, yani "
         "kavradığımız her şeyin üstü tanım gereği düz.", italic=True)

    h(doc, "11.4 Ölçümün göremedikleri", 2)
    para(doc, "Hepsi SESSİZCE eksik sonuç verir; sırayla en tehlikelisi:")
    table(doc,
          ["Sınır", "Sonucu"],
          [["Malzeme: mat siyah, parlak metal, şeffaf",
            "ToF ya veri vermez ya saçma verir - hassasiyet değil, veri sorunu"],
           ["Destek düzleminin ALTI görünmez",
            "ayaklı / oyuk tabanlı parçada siluet ölçülür, gövde değil"],
           ["“Düz yüzeyin üstünde duruyor” varsayımı",
            "parça başka bir şeyin üstündeyse referans düzlem yanlış"],
           ["Oklüzyon",
            "arka yüz görünmez, ayak izi EKSİK çıkar (yükseklik etkilenmez)"]],
          widths=[2.6, 3.8])
    para(doc,
         "Bu yüzden ölçüm her kenara payla (varsayılan 5 mm) şişirilir ve "
         "başarısız olursa yapılandırılmış payload_size'a düşülür; log "
         "hangisinin kullanıldığını yazar. Ölçemediğinde UYDURMAZ - uydurulmuş "
         "bir kutu, elle verilen yedekten daha tehlikelidir çünkü kimse yanlış "
         "olduğunu fark etmez.")

    doc.add_page_break()

    # ---- 12. arayüz ve anahtar ----
    h(doc, "12. Arayüz Entegrasyonu ve API Anahtarı", 1)

    h(doc, "12.1 Dashboard'dan tek tuşla görev", 2)
    para(doc, "user_interface paketindeki Pick & Place senaryosu yeniden bağlandı:")
    code(doc,
         "Pick & Place butonu\n"
         "  -> hil_test_whole_unified.launch.py use_vacuum_gripper:=true\n"
         "     (+ toggle açıksa use_fake_hardware / use_mock_hardware ...)\n"
         "Robot is Ready butonu\n"
         "  -> ros2 launch gemini_robotics_ros gemini_pick_place.launch.py \\\n"
         "         mode:=sim   (Use Fake Hardware AÇIK)\n"
         "         mode:=real  (Use Fake Hardware KAPALI)\n"
         "  -> ardından komut penceresi açılır")
    para(doc,
         "Pencereye yazılan serbest metin /gemini/command'a yayınlanır. İki "
         "ayrıntı kasıtlı: metindeki tek tırnaklar YAML için ikiye katlanıp tüm "
         "yük kabuk için shlex.quote'tan geçirilir (aksi halde “toolkit's” yazan "
         "bir komut satırı bozar), ve “ros2 topic pub --once” Humble'da "
         "varsayılan olarak bir abone bulunana kadar bekler - düğümler henüz "
         "kalkmamışken Confirm'e basılsa bile mesaj kaybolmaz.")

    h(doc, "12.2 API anahtarı artık kabuğa bağlı değil", 2)
    para(doc,
         "Anahtar yalnızca ortam değişkeninden okunuyordu. Ortam değişkeni "
         "sadece onu export eden kabuktan başlatılan süreçlere geçer; düğüm "
         "masaüstü kısayolundan, bir IDE terminalinden, systemd'den ya da "
         ".bashrc düzenlenmeden önce açılmış bir terminalden başlatılırsa "
         "anahtar yoktur - ve belirti “her açılışta yeniden girmek gerekiyor” "
         "gibi görünür.")
    code(doc,
         "arama sırası:\n"
         "  1. GEMINI_API_KEY        (ortam)\n"
         "  2. GOOGLE_API_KEY        (ortam)\n"
         "  3. ~/.config/gemini/api_key   (chmod 600)\n"
         "  4. /etc/gemini/api_key")
    para(doc,
         "Anahtar hâlâ depoya girmez: bu yollar ev dizinindedir, paketin içinde "
         "değil. .bashrc'de de anahtarın kendisi değil, dosyayı okuyan tek bir "
         "satır durur.")

    doc.add_page_break()

    h(doc, "13. Devir Teslim: Ne Doğrulandı, Ne Doğrulanmadı", 1)
    para(doc,
         "Bu bölüm çalışmaya ara verildiği andaki durumu kaydeder. Ayrım "
         "kasıtlıdır: “yazıldı ve testleri geçti” ile “çalışan hücrede "
         "görüldü” aynı şey değildir ve bu paketteki en pahalı hatalar "
         "(156 mm'lik TCP, p98 yükseklik sapması) tam olarak bu ikisinin "
         "karıştırılmasından doğdu.")

    h(doc, "13.1 Çalışan hücrede doğrulandı", 2)
    bullet(doc, "Uçtan uca görev: komut → planlama → çok pozlu tarama → kavrama "
                "→ toolkit'e bırakma, yeni prompt ile başarıyla tamamlandı.")
    bullet(doc, "TCP düzeltmesi: TF'te frame ile uç arasındaki fark 84 mm "
                "olarak görüldü; kap artık yüzeye değiyor.")
    bullet(doc, "ER 2, place sorgusunda dört ayrı boş göz adayı döndürdü "
                "(prompt “tüm örnekler” isteyecek şekilde düzeltildikten sonra).")
    bullet(doc, "Arayüzden görev başlatma ve serbest metin komut yayınlama "
                "(Bölüm 12).")
    bullet(doc, "Executor çakışmasının kapatılması: düğüm artık MoveIt "
                "çağrılarından sonra sağır kalmıyor, RCLError çökmesi görülmedi.")

    h(doc, "13.2 Yazıldı, testleri geçti — HÜCREDE HENÜZ ÇALIŞTIRILMADI", 2)
    para(doc,
         "Aşağıdakiler oturum sonunda hücre kapalıyken tamamlandı. Sonraki "
         "oturumun ilk işi bunları gerçek/simüle hücrede koşturmak olmalıdır.",
         italic=True, color=MUTED)
    bullet(doc, "DetachableJoint attach/detach: yalnızca ayrı bir test dünyasında "
                "ölçüldü, tam hücre launch'ında koşulmadı. Kontrol edilecek: "
                "spawn sonrası ilk detach'ın (t=26 s) gerçekten düşmesi, "
                "kavrama anında parçanın kola bağlanması, bırakmada düşmesi.")
    bullet(doc, "Kavranan parçanın MoveIt sahnesine iliştirilmesi "
                "(attach_collision_object + touch_links). Taşırken planlamanın "
                "parça yüzünden gereksiz başarısız olup olmadığına bakılmalı.")
    bullet(doc, "Derinlikten yük ölçümü (payload.py): yalnızca sentetik sahnelerde "
                "sınandı. Gerçek RedCube'de ölçülen boyut 140×83×30 mm ile "
                "karşılaştırılmalı; log hangi kaynağın kullanıldığını yazıyor.")
    bullet(doc, "measure_tcp'nin düğüm açılışındaki doğrulama olarak koşması.")
    bullet(doc, "Gazebo konveyörünün konumu düzeltildi. IFRA paketinin SDF'i kendi "
                "STL'inin koordinatlarında duruyordu; gerçek dünyada ölçülmüş konum "
                "my_robot_cell_macro.xacro'daki conveyor_belt linkidir. İki mesh'in "
                "bant yüzeyi dikey ışınla ölçülüp hizalandı: x −92.0 mm, y −52.0 mm, "
                "z +8.4 mm. Bant uzunluğu iki mesh'te de 956.0 mm çıktı — aynı "
                "konveyör. Sim'de URDF konveyör linki yorumda olduğu için MoveIt bu "
                "engeli GÖRMÜYOR; kayma yalnızca fizikte etkiliydi. Konveyör tarama "
                "pozu bu kaymadan sonra kontrol edilmeli: parça kadrajda ~105 mm "
                "yana kaydı.")
    bullet(doc, "test_reachability_live.py hâlâ ATLANIYOR — move_group ayakta "
                "değilken çalışmıyor. Hücre açıkken bir kez koşturulmalı.")

    h(doc, "13.3 Bilinen sınırlar", 2)
    bullet(doc, "Toolkit'in alt sıraları bu gripper ile kullanılamaz: göz açıklığı "
                "172 mm, vacuum_gripper gövdesi Ø167 mm. Yalnızca üst sıra "
                "yerleştirilebilir; bırakma da yaklaşma pozundan (~15 cm) yapılıyor.")
    bullet(doc, "Yük ölçümü destek düzleminin ALTINI göremez: ayaklı ya da "
                "girintili tabanlı bir parça silüeti kadar okunur. Kapanma "
                "(occlusion) ayak izini KISA verir. Mat siyah, parlak metal ve "
                "saydam yüzeyleri ToF zaten okuyamaz. Hepsi sessizce başarısız olur.")
    bullet(doc, "DetachableJoint'te child_model sabit bir SDF adıdır (RedCube). "
                "İkinci bir parça için ikinci bir eklenti bloğu gerekir; "
                "çalışma anında hedef değiştirilemez.")
    bullet(doc, "Gerçek kamera topic adları sürücü kaynağından ve hücre launch "
                "dosyasından çıkarıldı ancak GERÇEK DONANIMLA DOĞRULANMADI.")
    bullet(doc, "Place araması, pick'in bulunduğu pozdan başlıyor; görev başına "
                "yaklaşık 5 saniyelik bir ER çağrısı boşa gidiyor.")

    h(doc, "13.4 Güvenlik: bekleyen iş", 2)
    rich(doc, [
        ("Gemini API anahtarı ", {}),
        ("DÖNDÜRÜLMELİ", {"b": True}),
        (". Bir hata ayıklama adımında anahtar açık metin olarak terminal "
         "çıktısına düştü. aistudio.google.com/apikey adresinden eski anahtar "
         "iptal edilip yenisi ~/.config/gemini/api_key dosyasına (chmod 600) "
         "yazılmalıdır. Bu satır, iş yapılana kadar raporda kalır.", {}),
    ])

    h(doc, "13.5 Değerlendirilebilecek iyileştirmeler", 2)
    bullet(doc, "Klasik görüntü işleme ile hibrit çalışma: bant üstündeki tek "
                "parçayı bulmak için bağlı bileşen analizi ~5 ms sürüyor ve "
                "deterministik; ER'yi ilişkisel hedeflere (boş göz, üst sıra) "
                "ayırmak maliyeti düşürür.")
    bullet(doc, "Pick ve place sorgularının tek ER çağrısında birleştirilmesi.")
    bullet(doc, "Gözün içine inebilecek daha ince bir uç eleman ya da toolkit'in "
                "yeniden konumlandırılması — alt sıraları kullanılabilir kılar.")

    doc.add_paragraph()
    para(doc,
         "Rapor sonu. Şekiller doc/figures/ altındadır; bu belge "
         "doc/build_report.py ile yeniden üretilebilir.",
         italic=True, color=MUTED)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print("yazıldı:", path, os.path.getsize(path), "bayt")
