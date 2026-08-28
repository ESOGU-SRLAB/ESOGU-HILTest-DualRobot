#!/usr/bin/env python3
"""
makale_uret_tr.py
=================
ESOGÜ Mühendislik ve Mimarlık Fakültesi Dergisi (TR Dizin) için TÜRKÇE makaleyi
üretir -> `ESOGU_MMF_Gemini_Makale_TR.docx`.

    python3 makale_figurler.py     # önce şekiller
    python3 makale_uret_tr.py      # sonra makale

Makale ELLE DÜZENLENMEZ. Yerleşim `makale_duzen.py` içindedir (derginin kendi
`Dosya.docx` şablonundan gelen stiller); içerik burada. Bir sayı değişecekse
kaynağı bu dosyadır.

SAYILARIN KAYNAĞI
-----------------
Bu makaledeki her ölçülen değer üç yerden birinden gelir ve hiçbiri elle
girilmemiştir:

  * `gemini_robotics_ros` paketinin kodu ve `config/gemini_params.yaml`
    (parametreler, eşikler, uç eleman geometrisi),
  * kaynak koda 17 Ağustos 2026'da yazılmış ölçüm blokları (yama kirlenmesi -
    düzlem eğimi tablosu),
  * gerçek hücrede 20 Ağustos 2026'da alınan kayıt
    (`figures/gemini_pick_place.mp4`) içindeki düğüm loglarının epoch
    damgaları (aşama süreleri, varış sapmaları, vakum yüzdesi, ER 2 gecikmesi).

Kayıttaki damgalar görev saatine 1787219765,7 çıkarılarak çevrildi; bu sabit
kayıttaki en geç "son satır" damgasından belirlendiği için mutlak sıfır
±0,5 s belirsiz, aşama SÜRELERİ ise damga farkı olduğundan tamdır.
"""

from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

import makale_duzen as duzen
from makale_figurler import _KARELER
from makale_duzen import (
    BODY_PT, Builder, FIG, PAGE_W, SMALL_PT, abstract_block, open_template,
    style_run,
)

OUT = duzen.OUT


class BuilderTR(Builder):
    FIG_LABEL = "Şekil"
    TAB_LABEL = "Tablo"

    def p(self, text):
        """Gövde paragrafı; `**kalın**` ve `*italik*` işaretlerini uygular."""
        para = self.para("", style="Paragraf",
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=6, space_after=0)
        self.rich(para, text)
        return para

    def photo_grid(self, caption, paneller, panel_cm=8.55, cols=2,
                   sayfa_basi=False):
        """Fotoğraf ızgarası: her kare AYRI bir resim olarak gömülür.

        Dört kareyi matplotlib'de tek bir PNG'ye birleştirmek yerine
        kenarlıksız bir tabloya tek tek koymanın tek sebebi düzenlenebilirlik:
        yazar Word'de bir kareyi seçip kırpabilir, değiştirebilir ya da
        yerini oynatabilir. Birleşik resimde bunların hiçbiri yapılamaz.

        `paneller`: [(dosya_adı, altyazı), ...]
        """
        self.fig_no += 1
        self._begin_wide()

        ncol = cols
        nrow = -(-len(paneller) // ncol)
        hucre = PAGE_W / ncol
        tbl = self.doc.add_table(rows=nrow * 2, cols=ncol)
        self.body.remove(tbl._tbl)
        self._add(tbl._tbl)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        for k, (dosya, alt) in enumerate(paneller):
            r, c = divmod(k, ncol)
            yol = FIG / dosya
            hucre_resim = tbl.cell(r * 2, c)
            par = hucre_resim.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.keep_with_next = True
            if k == 0 and sayfa_basi:
                # Tam sayfa şekil: ızgara her hâlükârda kendi sayfasında
                # başlasın, yoksa satırlar iki sayfaya bölünüyor.
                par.paragraph_format.page_break_before = True
            if yol.exists():
                par.add_run().add_picture(str(yol), width=Cm(panel_cm))
            else:
                print(f"  ! eksik kare {dosya}")
            alt_par = tbl.cell(r * 2 + 1, c).paragraphs[0]
            alt_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            alt_par.paragraph_format.space_before = Pt(1)
            alt_par.paragraph_format.space_after = Pt(4)
            self.rich(alt_par, alt, size=SMALL_PT)

        self._fix_layout(tbl, [hucre] * ncol)
        self._kenarliksiz(tbl)
        self._bolunmez(tbl)

        cap = self.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=2, space_after=6, keep_wide=True)
        ad, _, ek = caption.partition("|")
        self.rich(cap, f"**{self.FIG_LABEL} {self.fig_no}.** {ad.strip()}",
                  size=SMALL_PT)
        if ek.strip():
            self.rich(cap, "  " + ek.strip(), size=SMALL_PT - 0.5, colour=duzen.GREY)
        self._end_wide(cap)
        self.last = cap
        return self.fig_no

    @staticmethod
    def _bolunmez(tbl):
        """Satırlar sayfa sonunda ikiye ayrılmasın."""
        for row in tbl.rows:
            trpr = row._tr.get_or_add_trPr()
            for eski in trpr.findall(qn("w:cantSplit")):
                trpr.remove(eski)
            trpr.append(OxmlElement("w:cantSplit"))

    @staticmethod
    def _kenarliksiz(tbl):
        """Izgara tablosunun her kenarını kapatır; çizgi görünmesin."""
        for row in tbl.rows:
            for cell in row.cells:
                tcpr = cell._tc.get_or_add_tcPr()
                for eski in tcpr.findall(qn("w:tcBorders")):
                    tcpr.remove(eski)
                kenarlar = OxmlElement("w:tcBorders")
                for yon in ("top", "left", "bottom", "right"):
                    el = OxmlElement(f"w:{yon}")
                    el.set(qn("w:val"), "nil")
                    kenarlar.append(el)
                tcpr.append(kenarlar)



# ═════════════════════════════════════════════════════════════════════
# ön sayfa
# ═════════════════════════════════════════════════════════════════════
TITLE_TR = ("GÖRÜ-DİL MODELİNİN İKİ BOYUTLU İŞARETİNDEN VAKUMLA KAVRAMAYA: "
            "BİR UR10e HÜCRESİNDE DERİNLİK TABANLI GEOMETRİ ZİNCİRİ")
TITLE_EN = ("FROM A VISION-LANGUAGE MODEL'S TWO-DIMENSIONAL POINT TO A VACUUM "
            "GRASP: A DEPTH-BASED GEOMETRY CHAIN IN A UR10e CELL")

KEYWORDS_TR = ["Görü-dil modeli", "Vakumlu kavrama", "Derinlik algılama",
               "Yüzey normali kestirimi", "ROS 2"]
KEYWORDS_EN = ["Vision-language model", "Vacuum grasping", "Depth sensing",
               "Surface normal estimation", "ROS 2"]

ABSTRACT_TR = (
    "Görü-dil modelleri doğal dille verilen bir komutu görüntü üzerinde "
    "konumlandırabilmekte, ancak yaygın olarak erişilebilen gömülü akıl "
    "yürütme modelleri robot eylemi üretmemektedir; çıktıları iki boyutlu "
    "bir noktadır. Vakumlu bir kavrayıcı ise noktayı değil, yüzeye dik "
    "oturan tam bir poz gerektirir; emme kabı eğik oturursa sızdırır ve "
    "parça hiç tutulamaz. Bu çalışmada, bir görü-dil modelinin işaret "
    "ettiği pikselden uygulanabilir bir vakum kavrama pozuna giden ölçüm "
    "zinciri kurulmuş ve UR10e kolu, OnRobot VGC10 vakumlu kavrayıcı ve "
    "uçuş zamanı ilkeli SICK Visionary-T Mini kamerasından oluşan gerçek "
    "bir hücrede çalıştırılmıştır. Kamera renkli görüntü üretmediği için "
    "modele verilen girdi derinlikten türetilen bir kabartma render'dır. "
    "İşaret edilen pikselin çevresindeki nokta yaması, bakış ışını "
    "boyunca dar bir banda indirgenerek arka plandan ayrılmakta; banda "
    "tekil değer ayrışımıyla düzlem oturtulmakta; en küçük tekil vektör "
    "yüzey normalini, artığın karekök ortalaması ise kavramayı reddetme "
    "kapısını vermektedir. Normal, tolerans içindeyse en yakın dünya "
    "eksenine oturtulmakta; yaklaşma, iniş ve kaldırma dünya düşeyi "
    "boyunca değil bu normal boyunca yapılmaktadır. Uç eleman geometrisi "
    "elle girilmek yerine robot modelinden ölçülmektedir. Sistem gerçek "
    "robotta on kez çalıştırılmış, sekizinde görev tamamlanmıştır; on "
    "koşu bir başarı oranı bildirmeye yetmediğinden bulgu bir "
    "yapılabilirlik kanıtı olarak sunulmakta ve başarısızlıkların "
    "dağılımı vurgulanmaktadır. Kavrama tarafında hiç başarısızlık "
    "görülmemiş, iki başarısızlığın ikisi de modelin bırakma hedefini "
    "yanlış konumlandırmasından kaynaklanmıştır."
)

ABSTRACT_EN = (
    "Vision-language models can ground a natural-language command in an "
    "image, yet the available embodied-reasoning models emit no action — "
    "only a two-dimensional point. A vacuum gripper needs a full pose "
    "perpendicular to the surface: a tilted cup leaks and the part is "
    "never picked. This study builds the measurement chain that turns "
    "that pixel into an executable vacuum grasp pose, and runs it on a "
    "real cell: a UR10e arm, an OnRobot VGC10 vacuum gripper and a SICK "
    "Visionary-T Mini time-of-flight camera. The camera produces no "
    "colour image, so the model is fed a relief rendering derived from "
    "depth. The point patch around the indicated pixel is reduced to a "
    "narrow band along the viewing ray, which removes the background; a "
    "plane is fitted to the band by singular value decomposition, whose "
    "smallest singular vector gives the surface normal and whose "
    "root-mean-square residual gates the grasp. The normal is snapped to "
    "the nearest world axis when within tolerance, and approach, descent "
    "and lift follow that normal rather than world vertical. End-effector "
    "geometry is measured from the robot model, not entered by hand. The "
    "system was run ten times on the real robot and completed the task in "
    "eight of them; ten runs cannot support a success rate, so the result "
    "is reported as a feasibility study and the distribution of the "
    "failures is what carries information. No grasp failure occurred, and "
    "both failures came from the model mislocating the release target."
)


def front_matter(b: Builder) -> None:
    def title(text):
        p = b.para("", style="Makale Başlığı",
                   align=WD_ALIGN_PARAGRAPH.CENTER, front=True)
        style_run(p.add_run(text), BODY_PT + 2, bold=True)

    title(TITLE_TR)
    p = b.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, front=True)
    style_run(p.add_run("Adı SOYADI"), BODY_PT)
    style_run(p.add_run("1*"), BODY_PT - 3.5).font.superscript = True
    style_run(p.add_run(",  Adı SOYADI"), BODY_PT)
    style_run(p.add_run("2"), BODY_PT - 3.5).font.superscript = True
    style_run(p.add_run("  (Dergi editörlüğü tarafından basım aşamasında "
                        "yazılacaktır.)"), BODY_PT)
    for i in (1, 2):
        p = b.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER,
                   front=True)
        style_run(p.add_run(f"{i} "), SMALL_PT - 1)
        style_run(p.add_run(f"Yazar {i} Adresi, ORCID No : "
                            "https://orcid.org/"), SMALL_PT - 1)
    b.para("", style="Normal", front=True)

    abstract_block(b, "Anahtar Kelimeler", "Öz", KEYWORDS_TR, ABSTRACT_TR,
                   "Araştırma Makalesi")
    b.para("", style="Normal", front=True)
    title(TITLE_EN)
    b.para("", style="Normal", front=True)
    abstract_block(b, "Keywords", "Abstract", KEYWORDS_EN, ABSTRACT_EN,
                   "Research Article")
    b.para("", style="Normal", front=True)


def _madde(b, metin):
    p = b.para("", style="Paragraf", align=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=3, space_after=0)
    p.paragraph_format.left_indent = Cm(0.4)
    b.rich(p, "•  " + metin)


# ═════════════════════════════════════════════════════════════════════
# 1. Giriş
# ═════════════════════════════════════════════════════════════════════
def _giris(b: Builder) -> None:
    b.h1("1. Giriş")
    b.p(
        "Bir endüstriyel hücreye doğal dille iş tarif edebilmek, robot "
        "programlamanın en pahalı kısmını, yani her yeni parça ve her yeni "
        "yerleşim için yazılan hareket kodunu ortadan kaldırma vaadini taşır. "
        "Görü-dil modelleri bu vaadin görme ve anlama tarafını büyük ölçüde "
        "karşılamaktadır: bir görüntü ile bir cümle verildiğinde, cümlenin "
        "işaret ettiği nesneyi görüntü üzerinde konumlandırabilmektedirler. "
        "Ancak vaadin robot tarafı aynı hızda gelmemiştir. Eylem üreten "
        "görü-dil-eylem modelleri ya kapalıdır ya da belirli gövdeler üzerinde "
        "toplanmış gösterim verisiyle eğitilmiştir; yaygın olarak erişilebilen "
        "gömülü akıl yürütme modelleri ise robot aksiyonu üretmez. Çıktıları "
        "metin ve görüntü düzleminde bir koordinattır."
    )
    b.p(
        "Bu ayrım, uygulamada göründüğünden daha belirleyicidir. Vakumlu bir "
        "kavrayıcı bir noktaya değil bir *poza* gider ve bu pozun yöneliminde "
        "hata payı çok dardır: emme kabı yüzeye dik oturmazsa aradaki açıklık "
        "vakumu sızdırır ve parça hiç tutulamaz. Bu çalışmada kullanılan Ø39 mm "
        "çapındaki kapta, dört derecelik bir eğim kabın bir kenarında yaklaşık "
        "1,6 mm'lik bir boşluk demektir; gerçek donanımda bu boşluk vakumun hiç "
        "kurulamaması ile sonuçlanmaktadır. Dolayısıyla modelin verdiği iki "
        "boyutlu nokta ile robotun ihtiyaç duyduğu altı serbestlik dereceli poz "
        "arasında kapatılması gereken bir açık kalmaktadır ve bu açık, modelin "
        "daha iyi olmasıyla değil, ölçümle kapatılmaktadır."
    )
    b.p(
        "Bu makale tam olarak o ölçüm zincirini konu edinmektedir. Katkı, yeni "
        "bir model ya da yeni bir öğrenme yöntemi değildir; modelin işaret "
        "ettiği pikselden başlayıp uygulanabilir bir vakum kavrama pozuna varan, "
        "her adımı gerçek hücrede ölçülmüş bir geometri zinciridir. Zincir, "
        "renkli görüntü üretmeyen bir uçuş zamanı kamerasıyla çalışacak biçimde "
        "kurulmuş, Evrensel Robotlar (Universal Robots) UR10e kolu, OnRobot "
        "VGC10 vakumlu kavrayıcı ve SICK Visionary-T Mini kamerasından oluşan "
        "bir hücrede Robot İşletim Sistemi 2 (Robot Operating System 2, ROS 2) "
        "üzerinde gerçeklenmiş ve gerçek robotta çalıştırılmıştır."
    )
    b.p(
        "Makalenin devamı şöyle düzenlenmiştir. İkinci bölüm, görü-dil "
        "modellerinin robot manipülasyonunda kullanımına ve vakumlu kavrama "
        "planlamasına ilişkin yazını, bu çalışmanın doldurduğu boşluğu "
        "belirterek özetlemektedir. Üçüncü bölüm hücreyi, modele verilen girdinin "
        "neden derinlikten üretildiğini ve geometri zincirinin her adımını "
        "eşitlikleriyle vermektedir. Dördüncü bölüm gerçek robotta ölçülen "
        "sonuçları, beşinci bölüm bu sonuçların ne anlama geldiğini ve "
        "çalışmanın kısıtlarını tartışmakta, altıncı bölüm ise genel sonuçları "
        "vurgulamaktadır."
    )


# ═════════════════════════════════════════════════════════════════════
# 2. Bilimsel Yazın Taraması
# ═════════════════════════════════════════════════════════════════════
def _literatur(b: Builder) -> None:
    b.h1("2. Bilimsel Yazın Taraması")
    b.p(
        "Dil ile robot davranışını birbirine bağlama çabası üç ayrı hatta "
        "ilerlemiştir. Birinci hat, dil modelini bir görev planlayıcısı olarak "
        "kullanıp yürütmeyi önceden yazılmış becerilere bırakır; SayCan "
        "çalışmasında dil modelinin önerdiği her beceri, robotun o beceriyi o "
        "anda gerçekleştirebilme olasılığıyla ağırlıklandırılmakta ve böylece "
        "dilin ürettiği plan robotun yetenekleriyle sınırlandırılmaktadır "
        "(Ahn vd., 2022). Bu hatta beceri kümesi elle yazıldığı için yeni bir "
        "nesne ya da yeni bir kavrayıcı, yeni beceri yazmayı gerektirir."
    )
    b.p(
        "İkinci hat, eylemi doğrudan modele ürettirir. RT-2, görü-dil modelinin "
        "ağdan öğrendiği bilgiyi robot kontrolüne aktarmak üzere eylemleri "
        "belirteçlere (token) dönüştürerek aynı çıkış uzayında üretmektedir "
        "(Brohan vd., 2023); OpenVLA ise 970 bin gerçek robot gösteriminden "
        "eğitilmiş açık kaynaklı bir görü-dil-eylem modeli sunmakta ve daha az "
        "parametreyle kapalı modellerin üzerine çıktığını raporlamaktadır (Kim "
        "vd., 2024). Bu hattın gücü uçtan uca olmasıdır; bedeli ise gövdeye "
        "bağımlılıktır. Modelin eğitildiği gövde, kavrayıcı ve kamera ile "
        "hedeflenen hücreninki uyuşmadığında üretilen eylem doğrudan "
        "kullanılamaz. Bu çalışmadaki hücrede kavrayıcı, eklemi olmayan bir "
        "vakum kabıdır ve kamera renkli görüntü üretmemektedir; her iki özellik "
        "de bu modellerin eğitim dağılımının dışındadır."
    )
    b.p(
        "Üçüncü hat, modeli eylem üreticisi olarak değil bir *işaretleyici* olarak "
        "kullanır ve eylemi klasik geometriye bırakır. MOKA, manipülasyon "
        "problemini görüntü üzerine işaretler koyup modele görsel soru-cevap "
        "sorarak nokta temelli bir gösterime indirger (Liu, Fang, Abbeel ve "
        "Levine, 2024). RoboPoint, dil komutuna karşılık gelen uzamsal "
        "sağlayışları (affordance) doğrudan görüntü anahtar noktaları olarak "
        "kestiren ve bunun için sentetik veriyle ince ayarlanmış bir görü-dil "
        "modelidir; görsel istem yöntemlerine göre belirgin bir üstünlük "
        "raporlamaktadır (Yuan vd., 2024). Google DeepMind'ın gömülü akıl "
        "yürütme modelleri de aynı arayüzü sunmakta, iki boyutlu noktalar, "
        "kutular ve yörüngeler üretmektedir (Gemini Robotics Ekibi, 2025). Bu "
        "çalışmanın kullandığı model de bu ailedendir."
    )
    b.p(
        "Üçüncü hattın açık bıraktığı yer, işaretten poza geçiştir. Anılan "
        "çalışmaların çoğu paralel çeneli bir kavrayıcı ve masa üstü bir sahne "
        "varsaymakta, kavrama yönelimini ya görüntü düzleminde bir açı olarak "
        "ya da ayrı bir kavrama ağıyla üretmektedir. Vakumlu kavramada ise "
        "belirleyici büyüklük yüzey normalidir ve bu, yazında ağırlıklı olarak "
        "analitik ve öğrenme temelli yöntemlerin birlikte kullanıldığı ayrı bir "
        "problem olarak ele alınmıştır: Dex-Net 3.0, emme kabı ile yüzey "
        "arasındaki sızdırmazlığın kurulup kurulamayacağını ve temasın dış "
        "kuvvetlere direncini modelleyerek nokta bulutlarından sağlam vakum "
        "hedefleri hesaplamaktadır (Mahler vd., 2018). Bu yöntemler kavrama "
        "hedefini kendileri seçer; dille verilen bir komutu dinlemezler."
    )
    b.p(
        "Bu çalışmanın doldurduğu boşluk bu iki tarafın arasındadır. Eylem "
        "üretmeyen bir görü-dil modelinin işaret ettiği tek bir pikselden yola "
        "çıkıp, hedefi model seçmiş olmasına karşın kavrama pozunu ölçümle "
        "kuran, renkli görüntünün hiç bulunmadığı bir uçuş zamanı hücresinde "
        "çalışan ve gerçek robotta doğrulanmış bir zincir, bilinen yazında bir "
        "arada verilmemiştir. Türkçe yazında ise görü-dil modellerinin "
        "vakumlu kavrama ile birlikte ele alındığı bir çalışmaya "
        "rastlanmamıştır. Kullanılan hareket planlama ve donanım soyutlama "
        "altyapısı bakımından çalışma, ROS 2 (Macenski, Foote, Gerkey, "
        "Lalancette ve Woodall, 2022) ve MoveIt (Coleman, Şucan, Chitta ve "
        "Correll, 2014) üzerine oturmaktadır."
    )


# ═════════════════════════════════════════════════════════════════════
# 3. Yöntem
# ═════════════════════════════════════════════════════════════════════
def _yontem(b: Builder) -> None:
    b.h1("3. Yöntem")
    b.p(
        "Bu bölümde önce hücre ve modele verilen girdi, sonra işaretten poza "
        "giden zincirin her adımı verilmektedir. Zincirin bütünü Şekil 1'de "
        "özetlenmiştir. Şeklin anlatmak istediği ayrım şudur: modelden çıkan "
        "bilgi görüntü düzleminde bir noktadır ve poz, o noktadan *sonra*, "
        "ölçümle kurulmaktadır."
    )
    b.figure("sekil1_mimari.png",
             "Doğal dil komutundan vakum kavrama pozuna giden işlem hattı. | "
             "Üst sıra modelin gördüğü ve ürettiği bilgiyi, alt sıra pozu "
             "kuran ölçüm zincirini göstermektedir. Model yalnız bir piksel "
             "verir; yönelim bilgisi hattın hiçbir yerinde modelden gelmez.")

    # ── 3.1 ──
    b.h2("3.1. Hücre ve Donanım")
    b.p(
        "Çalışma, Eskişehir Osmangazi Üniversitesi bünyesindeki esnek üretim "
        "araştırma laboratuvarında kurulu bir hücrede yürütülmüştür. Hücrede "
        "doğrusal bir eksen üzerine yerleştirilmiş bir UR10e işbirlikçi kolu, "
        "kolun ucunda OnRobot VGC10 vakumlu kavrayıcı, bileğe bağlı bir SICK "
        "Visionary-T Mini V3S145-1A uçuş zamanı kamerası, bir taşıma bandı ve "
        "gözlere bölünmüş bir takım rafı bulunmaktadır. Görev, bant üzerindeki "
        "parçayı alıp rafın belirtilen gözüne bırakmaktır. Hücrenin bileşenleri "
        "ve yazılım yığını Tablo 1'de verilmiştir."
    )
    b.table(
        "Hücrenin Donanım ve Yazılım Bileşenleri",
        ["Bileşen", "Model / sürüm", "Çalışmadaki rolü"],
        [["Robot kolu", "Universal Robots UR10e",
          "6 eksen; doğrusal eksen üzerinde"],
         ["Kavrayıcı", "OnRobot VGC10 vakum",
          "Eklemi yok; Modbus ile komut, bağıl vakum ile geri besleme"],
         ["Kamera", "SICK Visionary-T Mini V3S145-1A",
          "Uçuş zamanı; derinlik ve yakın kızılötesi genlik, RGB yok"],
         ["Model", "gemini-robotics-er-2-preview",
          "Gömülü akıl yürütme; 2B nokta üretir, eylem üretmez"],
         ["Ara katman", "ROS 2 Humble, ros2_control",
          "Sürücüler, dönüşüm ağacı (TF), düğümler arası iletişim"],
         ["Hareket planlama", "MoveIt 2",
          "Çarpışmasız plan, ters kinematik, taşınan parçanın sahnede "
          "temsili"]],
        widths=[2.0, 2.6, 3.75])
    b.p(
        "Kavrayıcının eklemi olmadığı için robot modelinde bir planlama grubu "
        "olarak tanımlanmamıştır; kavrama, Modbus üzerinden gönderilen bir mod "
        "ve hedef vakum yüzdesi ile yapılmakta, tutup tutmadığı ise kavrayıcının "
        "bildirdiği bağıl vakum değerinden anlaşılmaktadır. Bu ayrıntı yöntem "
        "açısından önemlidir: kavramanın başarısı planlayıcıdan değil "
        "donanımdan öğrenilmektedir."
    )

    # ── 3.2 ──
    b.h2("3.2. Modele Verilen Girdi: Neden Derinlik")
    b.p(
        "Hücredeki kameranın renkli görüntü kipi yoktur; yayımladığı veriler "
        "derinlik, yakın kızılötesi genlik, durum haritası ve nokta bulutudur. "
        "Bu nedenle modele verilen görüntü, benzetim ve gerçek tarafın ikisinde "
        "de var olan tek büyüklükten, yani derinlikten üretilmektedir. Her iki "
        "taraf da önce metreye çevrilmekte, ardından *aynı* sabit metrik aralıkla "
        "(0,3-4,0 m) normalize edilip aynı biçimde render edilmektedir; "
        "benzetim ile gerçek arasındaki görüntü paritesi buradan gelmektedir."
    )
    b.p(
        "Render biçiminin seçimi, modele giden görüntünün kontrastını "
        "ölçülebilir biçimde değiştirmektedir. "
        "Yüzey normaline dayalı Lambert gölgelemesi sahneyi mat gri bir üç "
        "boyutlu render gibi göstermekte ve kenarları, eğrilikleri doğru "
        "okutmaktadır; ancak düz bir yüzeyin üstünde duran *düz* bir parçanın üst "
        "yüzü ile altındaki yüzeyin normali aynı yöne baktığı için ikisi de aynı "
        "griye boyanmaktadır. Bant üzerindeki 29 mm yüksekliğindeki bir kutu "
        "için bu, yüzün 255 seviyeden yalnızca 1,7'lik bir farkla görünmesi "
        "demektir. Bu nedenle, yüksekliği doğrudan renge taşıyan bir kabartma "
        "(relief) render kullanılmıştır: her piksel, oturduğu referans düzleme "
        "olan dik uzaklığına göre renklendirilmektedir. Aynı sahnenin iki "
        "render'ı Şekil 2'de karşılaştırılmıştır. Burada ölçülen büyüklük "
        "girdi kontrastıdır; render biçiminin görev başarısına etkisi bir "
        "bileşen çıkarma (ablation) deneyiyle ayrıca ölçülmemiştir "
        "(bkz. Bölüm 5.4)."
    )
    b.figure("sekil2_render.png",
             "Aynı sahnenin normal gölgelemeli ve kabartma render'ı. | "
             "Turuncu çerçeve bant üzerindeki kutuyu göstermektedir. Normal "
             "gölgelemede (a) yalnız kenar görünür, yüz bandın yüzüyle aynı "
             "griye düşer; kabartmada (b) yükseklik renge taşındığı için yüz "
             "ayrı bir blok hâline gelir.")
    b.p(
        "Modele ayrıca görüntünün ne olduğunu açıklayan bir bağlam bloğu "
        "verilmektedir. Bu blok olmadan model, kendisine verilen render'ı "
        "olağan bir fotoğraf sanıp renk ve doku üzerinden akıl yürütmeye "
        "çalışmaktadır. Dürüst bir sınır olarak belirtmek gerekir ki hiçbir "
        "derinlik render'ı bu modeller için dağılım-içi bir girdi değildir; "
        "nesne *kategorisi* tanıma başarımının renkli görüntüye göre düşük olması "
        "beklenmelidir. Buna karşılık geometrik sorgular – düz panel, kutunun "
        "üstü, en yakın nesne, üstü açık göz – iyi çalışmaktadır ve vakumlu "
        "kavrama zaten geometrik bir problemdir."
    )

    # ── 3.3 ──
    b.h2("3.3. Modelin Çıktısı ve Üç Boyuta Geri İzdüşüm")
    b.p(
        "Model, komutta tarif edilen nesne için görüntü düzleminde 0-1000 "
        "aralığına normalize edilmiş bir nokta döndürmektedir. Bu nokta önce "
        "görüntü çözünürlüğüne ölçeklenmekte, sonra üç boyuta taşınmaktadır. "
        "Düzenli (organized) bir nokta bulutu varsa karşılık gelen hücre "
        "doğrudan okunmakta; yoksa derinlik görüntüsü ve kamera iç "
        "parametreleri kullanılarak Eşitlik 1 ile geri izdüşüm yapılmaktadır."
    )
    b.equation("P_kam  =  ( (u − c_x)·d / f_x ,  (v − c_y)·d / f_y ,  d )")
    b.p(
        "Burada (u, v) piksel koordinatları, d o pikseldeki derinlik, "
        "(f_x, f_y) odak uzaklıkları ve (c_x, c_y) asal noktadır. Elde edilen "
        "nokta kamera çerçevesindedir ve dönüşüm ağacı üzerinden dünya "
        "çerçevesine taşınmaktadır. Konum ile *yön* ayrı ayrı taşınmaktadır: "
        "konum için tam dönüşüm, yüzey normali için yalnız dönme bileşeni "
        "uygulanmaktadır."
    )

    # ── 3.4 ──
    b.h2("3.4. Yamanın Bakış Işını Boyunca Ayıklanması")
    b.p(
        "Yönelim, işaret edilen pikselin çevresindeki nokta yamasından "
        "kestirilmektedir; bu çalışmada yama yarıçapı 12 pikseldir. Yamanın ham "
        "hâli kullanılamaz, çünkü model nesnenin kenarına yakın bir yere işaret "
        "ettiğinde yamaya arka plan karışmakta ve oturtulan düzlem iki yüzeyin "
        "arasında bir yere bakmaktadır. Bu nedenle yama, önce bakış ışını "
        "boyunca dar bir banda indirgenmektedir (Eşitlik 2)."
    )
    b.equation("B  =  { p ∈ Y  :  | p·r̂ − ‖a‖ |  ≤  δ } ,"
               "    r̂ = a / ‖a‖")
    b.p(
        "Burada Y yama noktaları kümesi, a modelin işaret ettiği pikselin üç "
        "boyutlu karşılığı, r̂ o noktaya giden bakış ışınının birim vektörü ve "
        "δ bant yarı genişliğidir (bu çalışmada 50 mm). İki tasarım kararı "
        "burada saklıdır. Birincisi, bandın bir koordinat ekseni boyunca değil "
        "*bakış ışını* boyunca ölçülmesidir; kamera her iki eksen "
        "düzenlemesinde de başlangıç noktasında bulunduğu için ışına izdüşüm, "
        "eksen düzeni değişse de aynı büyüklüğü ölçmektedir. İkincisi, bandın "
        "yamanın medyanına değil modelin işaret ettiği noktanın *kendi* derinliğine "
        "çapalanmasıdır: yama iki yüzeye yaklaşık eşit bölündüğünde – yani tam "
        "da bandın var olma nedeni olan kenar durumunda – medyan iki yüzeyin "
        "arasındaki boşluğa düşmekte ve her iki yüzey birden elenmektedir. "
        "Bandın etkisi Şekil 3'te gösterilmiştir."
    )
    b.figure("sekil3_yama.png",
             "Işın bandının düzlem oturtmaya etkisi (temsilî). | "
             "(a) Bant uygulanmadığında yamaya giren birkaç arka plan noktası "
             "düzlemi eğmekte, normal iki yüzeyin arasına bakmaktadır. "
             "(b) Bant uygulandığında arka plan noktaları düşmekte ve normal "
             "yüzeyin kendi normali olmaktadır. Ölçülen değerler Tablo 5'tedir.",
             width_cm=13.0)

    # ── 3.5 ──
    b.h2("3.5. Düzlem Oturtma, Yüzey Normali ve Düzlemsellik Kapısı")
    b.p(
        "Banda kalan noktalara en küçük kareler anlamında düzlem, noktaların "
        "ağırlık merkezine göre merkezlenmiş matrisin tekil değer ayrışımıyla "
        "oturtulmaktadır (Eşitlik 3). En küçük tekil değere karşılık gelen "
        "tekil vektör, düzlemin normalidir."
    )
    b.equation("B − c  =  U Σ Vᵀ ,     n  =  v₃ ,     "
               "c = (1/|B|) Σ p")
    b.p(
        "Normal, kameraya bakacak biçimde yönlendirilmektedir; optik çerçevede "
        "kamera başlangıçta olduğu için bu, n·c < 0 koşuluyla sağlanmaktadır. "
        "Kavramanın yapılıp yapılmayacağına ise noktaların düzleme uzaklığının "
        "karekök ortalaması ile karar verilmektedir (Eşitlik 4)."
    )
    b.equation("ρ  =  sqrt( (1/|B|) · Σ ( (p − c)·n )² )  ≤  ρ_maks")
    b.p(
        "Bu çalışmada ρ_maks = 4 mm alınmıştır; eşiği aşan bir yama kenar, köşe "
        "ya da delik sayılmakta ve kavrama *hiç* denenmemektedir. Bu ölçütün "
        "yazılışında bir dönem bulunan ve düzeltilen bir hata, kapının "
        "davranışını doğrudan etkilediği için burada belirtilmelidir: artık "
        "başlangıçta karekök ortalamadan *önce* alınarak hesaplanıyordu, yani "
        "ölçülen büyüklük karekök ortalama kare değil mutlak sapmanın "
        "ortalamasıydı. Mutlak sapmanın ortalaması her zaman karekök ortalama "
        "kareden küçük olduğundan kapı, yazıldığından daha gevşek çalışmaktaydı "
        "ve aradaki fark tam da kapının yakalaması gereken durumda – yama iki "
        "yüzeye bölündüğünde, yani artık dağılımı iki tepeli olduğunda – en "
        "büyük değeri almaktaydı. Bu farkın ölçülen sonuçları Tablo 5'te "
        "verilmiştir."
    )
    b.p(
        "Düzlemsellik kapısının yanında iki denetim daha uygulanmaktadır: banda "
        "kalan nokta sayısının bir alt sınırın (25 nokta) üzerinde olması ve "
        "normalin dünya düşeyiyle yaptığı açının bir üst sınırı aşmaması. "
        "Reddedilen her yama için gerekçe metin olarak günlüğe yazılmaktadır, "
        "böylece 'neden kavramadı' sorusunun yanıtı koşu sonrasında doğrudan "
        "okunabilmektedir."
    )

    # ── 3.6 ──
    b.h2("3.6. Normalin Dünya Eksenine Oturtulması")
    b.p(
        "Hücredeki bant ve raf yüzeyleri işlenmiş, dünya eksenlerine göre düz "
        "yüzeylerdir. Buna karşılık normal, derinlik yamasından kestirildiği "
        "için üzerine birkaç derecelik gürültü binmektedir. Bu nedenle, ölçülen "
        "normal ile en yakın dünya ekseni arasındaki açı bir toleransın "
        "altındaysa normal o eksene oturtulmaktadır (Eşitlik 5)."
    )
    b.equation("k* = argmax_k ( n · e_k ) ,     θ = arccos( n · e_{k*} ) ,\n"
               "n ← e_{k*}   eğer   θ ≤ θ_otur")
    b.p(
        "Burada e_k, altı işaretli dünya ekseninden biridir ve bu çalışmada "
        "θ_otur = 15° alınmıştır. Oturtmanın gerekçesi ölçülmüştür: bir bırakma "
        "pozunda kap düşeyden 3,99 derece sapmış durumdaydı; Ø39 mm'lik bir "
        "kapta dört derece, bir kenarda yaklaşık 1,6 mm'lik boşluk ve gerçek "
        "donanımda sızdıran bir vakum demektir. Buna karşılık oturtma bir "
        "*modelleme varsayımıdır*: yüzeyin gerçekten eğik olabileceği işlerde "
        "tolerans sıfıra çekilerek kapatılmalıdır. Bu ayrımın gözden kaçmaması "
        "için, tolerans açıkken oturtulamayan her normal için sapma ve bunun "
        "yaklaşma pozunda yol açtığı yanal kayma milimetre cinsinden uyarı "
        "olarak yazılmaktadır."
    )

    # ── 3.7 ──
    b.h2("3.7. Kavrama Pozunun Kurulması ve Uç Eleman Geometrisi")
    b.p(
        "Yönelim, emme eksenini yüzeyin *içine* bakacak biçimde hizalayan dönme "
        "olarak kurulmaktadır. Uç eleman çerçevesindeki emme ekseni t, dünya "
        "çerçevesinde −n ile çakıştırılmakta; her iki vektörden birer sağ el "
        "ortonormal taban üretilerek dönme matrisi Eşitlik 6 ile elde "
        "edilmektedir. Emme kabı dönel simetrik olduğundan sapma (yaw) açısı "
        "serbesttir ve yalnız tekrarlanabilir bir ters kinematik çözümü seçmek "
        "üzere bir referans vektörle sabitlenmektedir."
    )
    b.equation("R  =  W · Tᵀ ,     R · t  =  −n")
    b.p(
        "Bu adımda hücreye özgü ve gözden kaçması pahalı olan bir ayrıntı "
        "vardır: emme ekseni uç eleman çerçevesinin bir koordinat ekseniyle "
        "çakışmamaktadır. Kap, çerçevenin X ekseninden 30,00 derece eğik "
        "durmakta, yani (cos30°, sin30°, 0) yönünü göstermektedir. Kabın ağzı "
        "ise çerçeve başlangıcından 84,2 mm ötededir. Bu iki büyüklük elle "
        "girilmemekte, robot modelinden okunmaktadır; daha önce elle girilen "
        "156 mm'lik bir uç uzaklığı yaklaşık 7 cm'lik erişim kaybına yol "
        "açmıştı. Düğüm, açılışta aynı ölçümü yeniden yapıp yapılandırmadaki "
        "değerle karşılaştırmakta ve 3 mm'den büyük bir fark görürse hata "
        "seviyesinde uyarı basmaktadır."
    )
    b.p(
        "Yaklaşma, iniş ve kaldırma dünya düşeyi boyunca değil *ölçülen normal* "
        "boyunca yapılmaktadır; dikey bir yüzeyden parça çekmek ancak böyle "
        "doğru olmaktadır. Yaklaşma pozu, temas noktasından normal boyunca a "
        "kadar geride kurulmaktadır (Eşitlik 7)."
    )
    b.equation("p_yaklaşma  =  p_temas  +  a · n")
    b.p(
        "Buradan, normalin eğik kalmasının bedeli doğrudan çıkmaktadır: normal "
        "θ kadar eğikse yaklaşma pozu yanal olarak a·tanθ kadar kaymaktadır. "
        "Temas noktası kaymadığı için belirti yalnız inişte görünür; kol "
        "nesnenin yanından eğik inip nesneyi süpürür. Şekil 4 bu bağıntıyı "
        "kullanılan yaklaşma uzaklıkları için vermektedir. Yaklaşma uzaklığı "
        "tek bir değer değil, azalan bir aday listesidir (15, 12, 10, 8, 6 cm); "
        "ilk aday çarpışma verirse bir sonraki denenmektedir. Bırakma tarafında "
        "raf gözlerinin üstü dar olduğu için aday listesi daha uzaktan "
        "başlamaktadır (23, 21, 19 cm)."
    )
    b.figure("sekil4_egim.png",
             "Düzlem normalindeki eğimin yaklaşma pozunda yol açtığı yanal "
             "kayma. | Kayma a·tanθ ile büyümekte, yaklaşma uzaklığı a ile "
             "doğru orantılı olmaktadır. Kırmızı noktalar, yamaya karışan arka "
             "planın ölçülen iki durumunu göstermektedir (Tablo 5).",
             width_cm=12.0)
    b.p(
        "Zincirin bu noktada ürettiği büyüklükler koşu sırasında RViz'e de "
        "çizilmektedir: temas noktasına bir küre, ölçülen normal boyunca bir "
        "ok ve noktanın üstüne etiketle birlikte düzlemsellik kalıntısını "
        "milimetre cinsinden veren bir yazı. Bu çizim yalnız gösterim için "
        "değildir; kolun GERÇEKTE gideceği nokta çizildiği için, bırakma "
        "noktası çarpışma yüzünden kaydırılmışsa bu kayma da doğrudan "
        "görünür hâle gelmektedir. Şekil 5, kayıtlı koşuda alma hedefi için "
        "üretilen marker takımını göstermektedir."
    )
    b.figure("sekil5_markerlar.png",
             "Geometri zincirinin çıktısının gerçek koşudaki görüntüsü. | "
             "Kayıtlı koşudan (t = 13 s) alınmış RViz görüntüsü. Nokta bulutu "
             "derinliğe göre renklendirilmiştir; bandın üstündeki parçanın üst "
             "yüzeyi turuncu ada olarak ayrışmaktadır. Etiketin sonu RViz "
             "görünüm alanının dışında kaldığı için kalıntı değeri buradan "
             "okunmamaktadır.",
             wide=False)

    # ── 3.8 ──
    b.h2("3.8. Taşınan Parçanın Derinlikten Ölçülmesi")
    b.p(
        "Kavranan parça artık kolun bir parçasıdır ve planlayıcı onu bilmezse "
        "kol parçayı rafa, banda ya da kendine sürtmektedir. Bu nedenle parça, "
        "kavramadan sonra hareket planlama sahnesine bir kutu olarak eklenip uç "
        "elemana iliştirilmektedir. Kutunun boyutu varsayılmak yerine "
        "derinlikten ölçülmektedir. En güvenilir ölçülen büyüklük yüksekliktir: "
        "üst yüzün ve altındaki destek düzleminin baskın seviyeleri arasındaki "
        "fark, yüzlerce pikselden hesaplanmaktadır (Eşitlik 8)."
    )
    b.equation("h  =  mod( z_üst )  −  mod( z_destek )")
    b.p(
        "Baskın seviye için histogram tepesi (mod) kullanılmakta, yüksek bir "
        "yüzdelik kullanılmamaktadır; bir yüzdelik yaklaşık iki standart sapma "
        "kadar yukarı kaymakta ve yüksekliği sistematik olarak fazla "
        "göstermektedir. Ölçüm iki denetimden geçirilmektedir: yükseklik makul "
        "bir aralıkta olmalı ve parçanın ayak izi bir yatay açıklık sınırını "
        "aşmamalıdır. İkinci denetim, ölçümün parça yerine bandın kendisini "
        "ölçtüğü bir durumdan sonra eklenmiştir; orada ölçülen ayak izi "
        "881 × 347 mm, yani görüntünün yaklaşık üçte biri çıkmıştı. Ölçüm "
        "başarısız olursa yapılandırmadaki varsayılan boyuta düşülmekte ve "
        "hangisinin kullanıldığı günlüğe yazılmaktadır. Sonuç, ölçülemeyen "
        "kısımları – destek düzleminin altı, örtülü yüzler, uçuş zamanı "
        "kamerasının okuyamadığı mat siyah ya da parlak yüzeyler – telafi etmek "
        "üzere 5 mm şişirilmektedir."
    )

    # ── 3.9 ──
    b.h2("3.9. Görev Durum Makinesi ve Çok Pozlu Tarama")
    b.p(
        "Görev, aşamaları arasında geçişleri günlüğe yazan bir durum makinesi "
        "olarak yürütülmektedir: tarama pozuna gidiş, alma hedefinin bulunması, "
        "bırakma hedefinin bulunması, yaklaşma, iniş ve temas, vakumun "
        "kurulması, kaldırma, bırakma pozuna yaklaşma, bırakma, geri çekilme ve "
        "başlangıç pozuna dönüş. Her hareketin sonunda ulaşılan kap ucu konumu "
        "ile hedeflenen konum karşılaştırılıp sapma milimetre cinsinden "
        "yazılmaktadır; bu, hatanın planlamada mı, doğrusal eksende mi yoksa "
        "algıda mı olduğunu ayırmaya yarayan tek doğrudan ölçüttür."
    )
    b.p(
        "Bu sapmanın ne olduğu, ne olmadığı kadar açık yazılmalıdır. Ölçüm "
        "ROS 2'nin dönüşüm ağacından (TF) alınmaktadır: oturma beklemesinden "
        "sonra dünya çerçevesi ile uç eleman çerçevesi arasındaki dönüşüm "
        "okunmakta, üzerine robot modelinden okunan emme kabı ağzı ötelemesi "
        "uygulanmakta ve elde edilen kap ucu konumu, o harekete komut olarak "
        "verilen konumdan çıkarılmaktadır. Dolayısıyla bildirilen büyüklük "
        "kolun kendisine verilen poza ne kadar oturduğudur, yani *izleme "
        "doğruluğudur*; zincirin ürettiği hedefin dünyadaki gerçek nesneye ne "
        "kadar denk düştüğü, yani MUTLAK doğruluk değildir. Mutlak doğruluk "
        "ancak hücreden bağımsız bir optik izleme ya da metroloji sistemiyle "
        "belirlenebilir; bu çalışmada böyle bir sistem kullanılmamıştır. "
        "Kabın parça üzerinde oturduğu yerin gözlenen sapması, koşular boyunca "
        "yaklaşık 5 mm'nin içinde kalmıştır ve zincirin uçtan uca doğruluğu "
        "için makul üst sınır bu mertebedir."
    )
    b.p(
        "Kamera kolun bileğine bağlı olduğu için sorgudan önce sabit bir bakış "
        "pozuna gidilmesi zorunludur. Tek bir bakış pozu ise yetmemektedir: "
        "bant ile raf aynı kareye sığmamakta, sığsa bile ikisi için gereken "
        "yükseklik pencereleri ve render biçimleri farklı olmaktadır. Bu "
        "nedenle her hedef için ayrı bir tarama pozu tanımlanmakta ve model o "
        "pozda sorgulanmaktadır. Aşağıda gösterileceği gibi bu tasarım, "
        "kaydedilen koşuda gerçekten devreye girmiştir: bırakma hedefi bant "
        "pozundan görülememiş, raf pozuna geçildikten sonra bulunmuştur."
    )
    b.p(
        "Kavramanın gerçekten olup olmadığı benzetimden değil donanımdan "
        "öğrenilmektedir. Vakum komutu gönderildikten sonra kavrayıcının "
        "bildirdiği bağıl vakum değeri beklenmekte ve bir eşiği (%15) aşarsa "
        "kavrama başarılı sayılmaktadır. Eşik aşılmazsa kap, ölçülen normal "
        "boyunca 4 mm daha bastırılıp en çok iki kez yinelenmektedir. Bu "
        "yineleme açıkça bir *telafidir*; her kavramada devreye giriyorsa yüzey "
        "yüksekliğinin yanlış okunduğu anlamına gelir ve kalibrasyonla "
        "giderilmesi gerekir."
    )
    b.p(
        "Bırakma tarafında bir ayrıntı, ilerideki başarısızlık çözümlemesi için "
        "belirleyicidir: bırakma temas noktası, modelin işaret ettiği noktanın "
        "kendisi değildir. Nokta önce ölçülen parça boyu kadar, sonra bir "
        "bırakma açıklığı (5 cm) kadar normal boyunca yukarı taşınmakta ve kol "
        "o yükseklikte parçayı serbest bırakmaktadır. Bunun nedeni, parçanın "
        "yüzeye sürtülmeden bırakılmasıdır; ancak yan etkisi, hedef nokta "
        "gözün biraz dışına düşse bile ortaya çıkan pozun boş uzayda, "
        "erişilebilir ve çarpışmasız kalmasıdır."
    )
    b.table(
        "Geometri Zincirinin Parametreleri ve Kullanılan Değerler",
        ["Parametre", "Değer", "Ne belirler"],
        [["Yama yarıçapı", "12 piksel",
          "Küçükse normal gürültülü, büyükse komşu yüzeyler karışır"],
         ["Bant yarı genişliği δ", "50 mm",
          "Yamadan arka planın ayıklanması"],
         ["En az nokta sayısı", "25",
          "Düzlem oturtmanın anlamlı olması"],
         ["Düzlemsellik kapısı ρ_maks", "4 mm",
          "Kenar/köşe/delik yamalarının reddi"],
         ["Oturtma toleransı θ_otur", "15°",
          "Normalin dünya eksenine oturtulması"],
         ["Yaklaşma uzaklığı a", "15 / 12 / 10 / 8 / 6 cm",
          "Alma tarafı aday listesi"],
         ["Bırakma yaklaşması", "23 / 21 / 19 cm",
          "Raf gözünün üstü dar olduğu için daha uzaktan"],
         ["Bastırma yinelemesi", "4 mm × 2",
          "Vakum kurulamazsa kabın normal boyunca bastırılması"],
         ["Render aralığı", "0,3 – 4,0 m",
          "Benzetim ile gerçek arasındaki görüntü paritesi"],
         ["Vakum hedefi / onay eşiği", "%60 / %15",
          "Komut değeri ve kavramanın doğrulanma eşiği"],
         ["Parça payı", "5 mm",
          "Ölçülemeyen yüzler için kutunun şişirilmesi"]],
        widths=[2.55, 2.15, 3.65], wide=False)

    # ── 3.10 ──
    b.h2("3.10. Deney Tasarımı")
    b.p(
        "Sistem gerçek robotta, benzetim kullanılmadan, on kez uçtan uca "
        "çalıştırılmıştır. Bütün koşularda operatörün verdiği doğal dil komutu "
        "aynı tutulmuş, buna karşılık parçanın bant üzerindeki *konumu* ve "
        "parçanın *kendisi* değiştirilmiştir; koşular boyunca farklı boyutlarda "
        "kutular kullanılmıştır. Böylece değişkenlik komutta değil sahnede "
        "olmakta ve ölçülen şey, aynı komutun farklı geometrilere karşı "
        "dayanıklılığı hâline gelmektedir. Bir koşu, aşağıda ayrıntılı olarak "
        "çözümlenmek üzere görüntü ve düğüm günlükleriyle birlikte "
        "kaydedilmiştir (20 Ağustos 2026)."
    )
    b.p(
        "Başarı ölçütü ikili olarak tanımlanmıştır: parçanın banttan alınıp "
        "hedeflenen raf gözüne bırakılması başarı, bunun dışındaki her sonuç "
        "başarısızlıktır. Ara ölçütler – varış sapmaları, vakumun kurulma "
        "süresi ve değeri, model çağrısının gecikmesi, aşama süreleri – kayıtlı "
        "koşunun düğüm günlüklerinden okunmuştur."
    )
    b.p(
        "Bu kümenin niteliği baştan belirtilmelidir. On koşu, bir başarım "
        "karakterizasyonu değil bir *yapılabilirlik kanıtıdır*: zincirin gerçek "
        "donanımda uçtan uca çalıştığını ve başarısızlıkların nerede "
        "toplandığını göstermeye yeter, bir başarı ORANI bildirmeye yetmez. "
        "Hücre üretim araştırma laboratuvarında paylaşımlı kullanıldığı için "
        "koşu sayısı bu sürümde artırılamamıştır. Bölüm 4.1'de sayılarla "
        "birlikte bunlara karşılık gelen güven aralığı da verilmekte ve "
        "aralığın genişliğinden ne çıkarılabileceği açıkça yazılmaktadır."
    )


# ═════════════════════════════════════════════════════════════════════
# 4. Bulgular
# ═════════════════════════════════════════════════════════════════════
def _bulgular(b: Builder) -> None:
    b.h1("4. Bulgular")

    b.h2("4.1. Uçtan Uca Başarım")
    b.p(
        "On koşunun sekizinde görev tamamlanmış, parça banttan alınıp "
        "hedeflenen raf gözüne bırakılmıştır. İki başarısız koşuda hata, "
        "aşağıda ayrıştırıldığı üzere geometri zincirinde değil, bırakma "
        "hedefinin konumlandırılmasındadır. Sonuçlar Tablo 3'te "
        "özetlenmiştir."
    )
    b.p(
        "Bu sayıdan bir başarı oranı çıkarmak doğru olmaz ve bunun nedeni "
        "sayının kendisinde görünür: 10 koşuda 8 başarı için Clopper-Pearson "
        "%95 güven aralığı [%44, %98]'dir. Aralık, sistemin yazı tura "
        "atmaktan farksız olduğu ile neredeyse kusursuz çalıştığı arasındaki "
        "her değeri kapsamaktadır; dolayısıyla bu küme bir başarım değeri "
        "BİLDİRMEZ. Kümeden çıkarılabilecek olan, başarısızlıkların "
        "DAĞILIMIDIR: on koşunun onunda da kavrama tarafı çalışmış, iki "
        "başarısızlığın ikisi de aynı adımda, bırakma hedefinin "
        "konumlandırılmasında ortaya çıkmıştır. Rastgele dağılmış iki "
        "başarısızlık bu kadar küçük bir kümede bir şey söylemezdi; tek bir "
        "adımda toplanmış olmaları, hangi adımın kırılgan olduğunu gösteren "
        "niteliksel bir bulgudur. Oranın kendisi ancak 30-50 koşuluk bir "
        "kümeyle anlamlı hâle gelir ve bu, çalışmanın açık bir devamıdır "
        "(Bölüm 5.4)."
    )
    b.table(
        "On Gerçek Robot Koşusunun Sonucu",
        ["Sonuç", "Adet", "Açıklama"],
        [["Görev tamamlandı", "8",
          "Parça banttan alındı ve hedef göze bırakıldı"],
         ["Bırakma hedefi yanlış konumlandı", "2",
          "Model bırakma noktasını gözün dışına atadı; parça havada "
          "bırakıldı"],
         ["Kavrama başarısız (vakum kurulamadı)", "0",
          "Düzlemsellik kapısını geçen hiçbir yamada sızdırma görülmedi"],
         ["Planlama tümüyle başarısız", "0",
          "Yeniden deneme sonrası bütün hedefler planlanabildi"],
         ["Toplam", "10", "Sekiz tamamlanmış görev"]],
        widths=[3.05, 1.3, 4.0], wide=False,
        note="Bütün koşularda doğal dil komutu aynı; parçanın konumu ve "
             "boyutu koşular arasında değiştirilmiştir. Sayılar bir "
             "yapılabilirlik kümesindendir; 8/10 için Clopper-Pearson %95 "
             "güven aralığı [%44, %98] olduğundan buradan bir başarı oranı "
             "çıkarılmamalıdır.")
    b.p(
        "Başarısızlıkların ikisinin de aynı nedene dayanması, hata kaynağının "
        "rastgele değil sistematik olduğunu göstermektedir. Bu noktanın altı "
        "çizilmelidir: alma tarafında, yani bu makalenin konusu olan geometri "
        "zincirinin çalıştığı tarafta, on koşunun hiçbirinde kavrama "
        "başarısızlığı gözlenmemiştir. Başarısızlıkların tamamı, modelin "
        "bırakma hedefi için ürettiği noktanın hedef gözün dışına düşmesinden "
        "kaynaklanmıştır; bu durumda kol, kendisine verilen noktaya doğru "
        "hareket etmekte ve parçayı gözün üstünde değil yanında, havada "
        "bırakmaktadır."
    )
    b.p(
        "Hareket planlayıcısının bu iki koşuda neden itiraz etmediği ayrıca "
        "belirtilmelidir, çünkü sistemde ayrıntılı bir çarpışma sahnesi "
        "bulunmaktadır. İtiraz etmemiştir, çünkü itiraz edecek bir şey "
        "yoktur: bırakma pozu, modelin verdiği noktanın kendisi değil, o "
        "noktanın ölçülen parça boyu ve 5 cm'lik bırakma açıklığı kadar "
        "normal boyunca yukarısıdır (Bölüm 3.9). Nokta gözün birkaç santim "
        "dışına düştüğünde bile ortaya çıkan poz boş uzayda kalmakta, "
        "erişilebilir ve çarpışmasız olmaktadır; kol oraya sorunsuz gitmekte "
        "ve parçayı bırakmaktadır. Yani denetim mekanizması poz üzerinde "
        "çalışmakta, oysa hata pozda değil ANLAMDA, yani noktanın bir gözün "
        "içine düşüp düşmediğindedir. Çarpışma denetimi bu türden bir hatayı "
        "yapısal olarak yakalayamaz; yakalayabilecek olan, Bölüm 5.2'de "
        "tartışılan içerik doğrulamasıdır."
    )

    b.h2("4.2. Kayıtlı Koşunun Aşama Çözümlemesi")
    b.p(
        "Kaydedilen koşu, aşamalarının tamamı düğüm günlüklerinin epoch "
        "damgalarından okunarak çözümlenmiştir. Görev, ilk tespitten parçanın "
        "bırakılmasına kadar 100 s, başlangıç pozuna dönüşle birlikte toplam "
        "yaklaşık 117 s sürmüştür. Aşama süreleri Şekil 6 ve Şekil 7'de "
        "verilmiştir."
    )
    b.photo_grid(
        "Kayıtlı koşudan dört an. | (a) Raf tarama pozunda, model çağrısı "
        "sürerken; (b) emme kabı, bant üzerindeki parçanın üst yüzeyine "
        "oturmuş durumda; (c) vakum kurulduktan 1,4 s sonra parça banttan "
        "kalkmış durumda; (d) parça, rafın üstten açık gözünün üstüne "
        "taşınmış, bırakma öncesinde. Kareler kaydın kendi saatinden "
        "seçilmiş ve her biri ayrı ayrı kırpılmıştır; saatler Şekil 7'deki "
        "aşama sınırlarıyla aynı kaynaktandır.",
        [(dosya, alt) for dosya, _t, _kutu, alt in _KARELER],
        panel_cm=8.55, sayfa_basi=True)
    b.figure("sekil7_zaman.png",
             "Kayıtlı koşunun aşama zaman çizelgesi. | Sıfır noktası ilk "
             "loglanan olaydır (alma hedefinin bulunması). Model çağrısı ve kol "
             "hareketleri sürenin tamamına yakınını almakta, vakumun kurulması "
             "0,2 s ile ölçülebilir sürelerin en küçüğü olmaktadır.")
    b.p(
        "Karelerin arasındaki iki fark, zincirin çıktısının doğrudan "
        "gözlenebilir sonucudur. (b) ile (c) arasında yalnız 2 s vardır; kap, "
        "ölçülen normal boyunca indiği için parçanın üst yüzeyine düz "
        "oturmakta ve vakum ilk denemede kurulmaktadır. (d) ise parçanın "
        "gözün üstüne, kolun kendi gövdesinin rafa değmediği bir duruşla "
        "taşındığını göstermektedir; bu duruşu belirleyen, bırakma tarafında "
        "daha uzaktan başlayan yaklaşma aday listesidir."
    )
    b.p(
        "Çizelgeden okunan ilk sonuç, sürenin nerede harcandığıdır. Bırakma "
        "hedefi için yapılan model çağrısı, tarama pozuna varıştan sonucun "
        "işlenmesine kadar 14,8 s sürmüştür; bu süre render, ağ üzerinden model "
        "çağrısı ve geri izdüşümün toplamıdır ve tek başına görevin yaklaşık "
        "yedide biridir. Kol hareketleri ise daha da uzundur: yaklaşma ve iniş "
        "31,9 s, bırakma pozuna yaklaşma 27,0 s almıştır. Buna karşılık "
        "vakumun kurulması 0,2 s sürmüştür. Aşamalar arasına konulan 2,0 "
        "saniyelik oturma beklemeleri beş kez devreye girmiş, yani toplam "
        "sürenin 10 saniyesini oluşturmuştur."
    )

    b.h2("4.3. Çok Pozlu Taramanın Devreye Girmesi")
    b.p(
        "Kayıtlı koşuda alma hedefi, bant tarama pozunda ilk sorguda "
        "bulunmuştur. Bırakma hedefi ise aynı pozda *bulunamamıştır*: model, "
        "rafın üstten açık gözüne ilişkin sorguya hiçbir sonuç döndürmemiştir. "
        "Sistem bunun üzerine raf tarama pozuna geçmiş ve sorguyu orada "
        "yinelemiş, hedef bu kez tek seferde bulunmuştur. Bu geçiş, tasarımın "
        "gerekliliğini doğrudan göstermektedir: tek bakış pozlu bir kurulumda "
        "bu koşu, geometri zinciri kusursuz çalışsa bile başarısız olurdu."
    )
    b.p(
        "Poz değişimi sırasında eklem sarmalarının açılması da ölçülmüştür. "
        "Hedefe giderken iki eklemin tam turluk gereksiz dönüş yapacağı "
        "belirlenmiş ve hedef açılar eşdeğerlerine kaydırılarak toplam 681 "
        "derecelik dönüş elenmiştir (omuz ekseninde 360°, birinci bilek "
        "ekseninde 321°)."
    )

    b.h2("4.4. İzleme Doğruluğu ve Vakumun Kurulması")
    b.p(
        "Kayıtlı koşuda her hareketin sonunda TF'ten okunan kap ucu konumu ile "
        "o harekete komut olarak verilen konum karşılaştırılmıştır; bu "
        "bölümdeki bütün sapmalar Bölüm 3.9'da tanımlandığı anlamda izleme "
        "doğruluğudur. Temas noktasına varışta sapma üç "
        "eksende de milimetre altında kalmış, bırakma pozuna varışta iki eksende "
        "1 mm ölçülmüş, geri çekilme pozunda yeniden milimetre altına inmiştir. "
        "Yüzey normali, dünya düşeyinden 9,26 derece sapmayla ölçülmüş ve 15 "
        "derecelik tolerans içinde kaldığı için eksene oturtulmuştur. Vakum, "
        "komut gönderildikten yaklaşık 200 ms sonra %30,6 bağıl değerde "
        "kurulmuş, yani %15'lik onay eşiğinin iki katına ulaşmıştır. Bu "
        "değerler Tablo 4'te toplanmıştır."
    )
    b.p(
        "Bu sayıların ne anlama geldiği konusunda ihtiyatlı olmak gerekir. "
        "Milimetre altı değerler, kolun kendisine verilen poza oturma "
        "doğruluğunu gösterir; zincirin ürettiği hedefin nesnenin gerçek "
        "yerine ne kadar denk düştüğünü göstermez. İkincisi bağımsız bir "
        "metroloji sistemi gerektirir ve bu çalışmada ölçülmemiştir; kabın "
        "parça üzerinde oturduğu yerden yapılan gözlem, uçtan uca sapmanın "
        "yaklaşık 5 mm'nin içinde kaldığına işaret etmektedir. Bu ayrım "
        "önemsiz değildir: vakumlu kavramada başarıyı belirleyen, hedefin "
        "mutlak konumundan çok kabın yüzeye DÜZ oturmasıdır; on koşunun "
        "hiçbirinde sızdırma görülmemesi de bu yüzden anlamlıdır."
    )
    b.table(
        "Kayıtlı Koşuda Ölçülen Değerler",
        ["Ölçülen büyüklük", "Değer", "Ne gösterir"],
        [["Temas pozuna izleme sapması", "< 1 mm (3 eksen)",
          "Kolun komut edilen poza oturması (TF)"],
         ["Bırakma pozuna izleme sapması", "1 mm (2 eksen)",
          "Uzun taşıma hareketinin sonunda korunan izleme"],
         ["Geri çekilme pozuna izleme sapması", "< 1 mm (3 eksen)",
          "Parça bırakıldıktan sonra tekrarlanabilirlik"],
         ["Ölçülen normal sapması", "9,26°",
          "Tolerans içinde; eksene oturtuldu"],
         ["Vakumun kurulma süresi", "≈ 200 ms",
          "Kavrama, görev süresinde ihmal edilebilir"],
         ["Kurulan bağıl vakum", "%30,6",
          "Onay eşiğinin (%15) iki katı"],
         ["Model çağrısı gecikmesi", "14,8 s",
          "Render + ağ + geri izdüşüm; kapalı döngüye engel"],
         ["Toplam görev süresi", "≈ 117 s",
          "İlk tespitten başlangıç pozuna dönüşe"]],
        widths=[3.05, 1.9, 3.4], wide=False,
        note="Sapmalar TF'ten okunan kap ucu konumu ile komut edilen konumun "
             "farkıdır; izleme doğruluğunu verir, bağımsız bir metroloji "
             "sistemine karşı mutlak doğruluğu değil.")

    b.h2("4.5. Yaklaşma Uzaklığının ve Planlamanın Uyarlanması")
    b.p(
        "Bırakma tarafında ilk yaklaşma adayı (15 cm) çarpışma verdiği için "
        "sistem daha uzak bir adaya, 23 cm'ye geçmiştir; raf gözünün üstündeki "
        "açıklığın dar olması bu davranışı zorunlu kılmaktadır. Ayrıca bırakma "
        "pozuna yapılan ilk planlama denemesi geçersiz plan hatasıyla "
        "başarısız olmuş, aynı hedefe yapılan ikinci deneme 0,4 s sonra "
        "başarılı olmuştur. Bu iki gözlem, aday listesi ve yeniden deneme "
        "politikasının süs olmadığını, tek koşuda dahi devreye girdiğini "
        "göstermektedir."
    )

    b.h2("4.6. Düzlemsellik Kapısının Ölçülen Etkisi")
    b.p(
        "Düzlemsellik ölçütünün doğru yazılmasının ne değiştirdiği, bant "
        "üzerindeki 29 mm yüksekliğindeki bir kutuya kenarına yakın işaret "
        "edilerek ölçülmüştür. Yamaya karışan arka plan oranı arttıkça, mutlak "
        "sapmanın ortalaması ile karekök ortalama kare arasındaki fark "
        "büyümekte ve eski ölçüt kapıyı geçirirken doğru ölçüt "
        "reddetmektedir. Geçen yamanın normali ise on derece mertebesinde "
        "eğiktir ve bu, 15 cm'lik yaklaşmada 21–29 mm'lik yanal kaymaya karşılık "
        "gelmektedir (Tablo 5)."
    )
    b.table(
        "Yamaya Karışan Arka Planın Düzlemsellik Ölçütüne ve Poza Etkisi",
        ["Arka plan", "MAD (eski)", "RMS (doğru)", "Eğim",
         "Kayma (15 cm)"],
        [["%4", "2,9 mm", "5,3 mm", "8,0°", "21 mm"],
         ["%6", "4,0 mm", "6,2 mm", "10,9°", "29 mm"]],
        widths=[1.45, 1.75, 1.75, 1.25, 2.15], wide=False,
        note="MAD: mutlak sapmanın ortalaması (eski, hatalı ölçüt); RMS: "
             "karekök ortalama kare (doğru ölçüt). Kapı eşiği 4 mm'dir. Eski "
             "ölçütle %4 ve %6 kirlenmenin ikisi de kapıyı geçmekte, doğru "
             "ölçütle ikisi de reddedilmektedir.")


# ═════════════════════════════════════════════════════════════════════
# 5. Tartışma
# ═════════════════════════════════════════════════════════════════════
def _tartisma(b: Builder) -> None:
    b.h1("5. Tartışma")

    b.h2("5.1. Hatanın Yeri: Model mi, Geometri mi")
    b.p(
        "Bu çalışmanın en belirgin bulgusu, on koşuluk kümede başarısızlığın "
        "*tamamının* modelin konumlandırmasından, hiçbirinin geometri zincirinden "
        "kaynaklanmasıdır. Alma tarafında düzlemsellik kapısını geçen hiçbir "
        "yamada vakum sızdırması görülmemiş, kavrama başarısızlığı hiç "
        "yaşanmamıştır. Buna karşılık bırakma tarafında model, iki koşuda hedef "
        "gözün dışına düşen bir nokta üretmiş ve kol parçayı havada "
        "bırakmıştır."
    )
    b.p(
        "Bu asimetrinin nedeni, iki sorgunun doğası gereği farklı olmasıdır. "
        "Alma sorgusu somut ve geometriktir: bant üzerindeki nesne, "
        "çevresinden yükseklikle ayrılan tek bloktur ve kabartma render'da "
        "doğrudan görünmektedir. Bırakma sorgusu ise *ilişkiseldir* – üstü açık "
        "seviyedeki boş gözlerden kameraya en yakın olanı – ve modelin "
        "birbirine benzeyen çok sayıda göz arasında hem doluluk hem de "
        "erişilebilirlik üzerine akıl yürütmesini gerektirmektedir. Derinlik "
        "render'ının dağılım dışı bir girdi olduğu göz önüne alındığında, "
        "başarısızlığın bu tür sorgularda yoğunlaşması beklenen bir sonuçtur."
    )
    b.p(
        "Pratik çıkarım, hatanın yerine göre farklı bir karşı önlem "
        "gerektirmesidir. Geometri zincirinin hataları ölçümle kapatılabilir; "
        "modelin konumlandırma hatası ise ancak *doğrulama* ile yakalanabilir. "
        "Bu çalışmada bırakma noktası için böyle bir doğrulama "
        "bulunmamaktadır: bırakma pozu erişilebilirlik ve çarpışma bakımından "
        "denetlenmekte, ancak noktanın gerçekten bir gözün *içine* düşüp "
        "düşmediği denetlenmemektedir. Havada bırakma tam olarak bu boşluktan "
        "geçmektedir."
    )

    b.h2("5.2. Ölçüm ile Telafi Arasındaki Ayrım")
    b.p(
        "Zincirin birkaç yerinde, bir hatayı kaynağında düzeltmek ile sonucunu "
        "telafi etmek arasında seçim yapılmıştır ve bu ayrımın açıkça "
        "korunması, sistemin bakımını doğrudan etkilemektedir. Uç eleman "
        "geometrisi ölçüm tarafındadır: kabın ekseni ve ağzının uzaklığı robot "
        "modelinden okunmakta, elle girilmemektedir; daha önce elle girilen bir "
        "değer yaklaşık 7 cm erişim kaybına yol açmıştı. Buna karşılık vakum "
        "kurulamadığında kabın normal boyunca bastırılması bir *telafidir* ve "
        "kodda da böyle işaretlenmiştir; her kavramada devreye girmesi, yüzey "
        "yüksekliğinin yanlış okunduğunun göstergesidir."
    )
    b.p(
        "Aynı ayrım normalin dünya eksenine oturtulmasında da geçerlidir. "
        "Oturtma, gürültüyü bastırdığı için görünürde bir iyileştirmedir; ancak "
        "özünde 'bu hücredeki yüzeyler eksenlere diktir' varsayımıdır. Varsayım "
        "doğru olduğu sürece kazandırmakta, yanlış olduğu bir işte ise ölçülen "
        "gerçeği silmektedir. Bu nedenle tolerans bir parametre olarak açıkta "
        "bırakılmış ve oturtulamayan normaller için sapmanın yaklaşma pozunda "
        "yol açtığı yanal kayma uyarı olarak yazdırılmıştır."
    )

    b.h2("5.3. Gecikme ve Kullanım Zarfı")
    b.p(
        "Model çağrısının 14,8 saniyelik gecikmesi, sistemin kullanım zarfını "
        "belirlemektedir. Bu gecikmeyle görsel geri beslemeli, kapalı döngü bir "
        "denetim mümkün değildir; sistem yalnız 'bak, planla, hareket et' "
        "döngüsü için uygundur. Öte yandan aynı ölçüm, gecikmenin görev "
        "süresindeki payının sanıldığı kadar baskın olmadığını da "
        "göstermektedir: kol hareketleri iki model çağrısının toplamından daha "
        "uzun sürmüştür. Görev süresini kısaltmak isteyen bir çalışmanın önce "
        "hareket profillerine ve aşamalar arasındaki oturma beklemelerine "
        "bakması, model gecikmesini azaltmaya çalışmasından daha verimli "
        "olacaktır."
    )
    b.p(
        "Bununla birlikte 117 saniyelik bir çevrim süresi endüstriyel bir "
        "uygulama için yavaştır ve bu, sistemin bugünkü hâlinin açık bir "
        "sınırıdır. Ölçüm, sürenin nerede geri kazanılabileceğini de "
        "göstermektedir. Birincisi, aşamalar arasına konulan 2,0 saniyelik "
        "oturma beklemeleri toplam 10 saniye tutmaktadır; bunlar bir çözüm "
        "değil, doğrusal eksen sürücüsünün hedefe varmadan 'vardım' "
        "raporlamasına karşı konulmuş bir önlemdir ve denetleyicinin hedef "
        "toleransı ile durma hızı toleransı düzeltildiğinde tümüyle "
        "kaldırılabilir. İkincisi, bırakma hedefi için yapılan 14,8 saniyelik "
        "model çağrısı, alma sonucuna bağlı olmadığı için parçanın taşınması "
        "sırasında örtüşük yürütülebilir. Üçüncüsü, hareket profilleri bu "
        "koşularda güvenli tarafta bırakılmış hız ve ivme ölçekleriyle "
        "çalıştırılmıştır. Bu üçü birlikte çevrim süresinin yarıdan fazlasını "
        "hedeflemektedir ve sonraki çalışmanın konusudur."
    )

    b.h2("5.4. Kısıtlar")
    b.p(
        "Çalışmanın kısıtları açıkça belirtilmelidir."
    )
    _madde(b, "**Örneklem bir başarım değeri bildirmeye yetmez.** On koşuda "
              "8 başarı için %95 güven aralığı [%44, %98]'dir; çalışma bir "
              "yapılabilirlik kanıtı olarak okunmalıdır. Kümeden çıkan "
              "niteliksel bulgu, başarısızlıkların tek bir adımda "
              "toplanmasıdır. Oranın kendisi için 30-50 koşuluk, parça "
              "türü ve konumu bakımından dengelenmiş bir küme gerekir; bu, "
              "planlanan devam çalışmasıdır.")
    _madde(b, "**Tek hücre, tek kamera, tek kavrayıcı.** Sonuçlar bu hücrenin "
              "geometrisine ve bu uçuş zamanı kamerasının gürültü karakterine "
              "bağlıdır; başka bir kamerada düzlemsellik kapısının eşiği "
              "yeniden ölçülmelidir.")
    _madde(b, "**Bildirilen sapmalar izleme doğruluğudur.** Milimetre altı "
              "değerler, TF'ten okunan kap ucunun komut edilen poza oturma "
              "doğruluğunu verir. Zincirin uçtan uca MUTLAK doğruluğu "
              "bağımsız bir optik izleme ya da metroloji sistemiyle "
              "doğrulanmamıştır; gözleme dayalı üst sınır yaklaşık 5 mm "
              "mertebesindedir.")
    _madde(b, "**Bileşen çıkarma ve karşılaştırma deneyi yapılmamıştır.** "
              "Kabartma render'ın normal gölgelemeye üstünlüğü girdi "
              "kontrastı üzerinden ölçülmüş, görev başarısı üzerinden "
              "ölçülmemiştir. Sistem ayrıca klasik bir nokta bulutu "
              "bölütleme hattıyla karşılaştırılmamıştır; böyle bir hat dille "
              "verilen komutu zaten dinlemediğinden doğrudan karşılaştırma "
              "eşdeğer olmaz. Zincir model-bağımsız kurulduğu için adil "
              "karşılaştırma, AYNI zincire ER 2 yerine elle yazılmış bir "
              "bölge seçicinin bağlanmasıdır; bu düzenek doğrudan bir sonraki "
              "deney olarak kurulabilir.")
    _madde(b, "**Bırakma noktası için içerik doğrulaması yoktur.** Poz, "
              "erişilebilirlik ve çarpışma bakımından denetlenmekte; noktanın "
              "bir gözün içine düşüp düşmediği denetlenmemektedir. İki "
              "başarısız koşu bu boşluktan geçmiştir.")
    _madde(b, "**Kategori tanıma zayıftır.** Derinlik render'ı bu modeller için "
              "dağılım dışı bir girdidir; sistem geometrik sorgularda güçlü, "
              "'şu markanın kutusu' türü kategorik sorgularda zayıf "
              "olacaktır.")
    _madde(b, "**Ölçüm sessizce eksik kalabilir.** Parça yüksekliği destek "
              "düzleminin *üstünden* ölçüldüğü için ayaklı ya da oyuk tabanlı bir "
              "parçada siluet ölçülmekte, uçuş zamanı kamerasının okuyamadığı "
              "mat siyah, parlak metal ve şeffaf yüzeylerde ise ölçüm "
              "eksilmektedir.")


# ═════════════════════════════════════════════════════════════════════
# 6. Sonuçlar
# ═════════════════════════════════════════════════════════════════════
def _sonuclar(b: Builder) -> None:
    b.h1("6. Sonuçlar")
    b.p(
        "Bu çalışmada, eylem üretmeyen bir görü-dil modelinin işaret ettiği tek "
        "bir pikselden uygulanabilir bir vakum kavrama pozuna giden ölçüm "
        "zinciri kurulmuş ve gerçek bir UR10e hücresinde çalıştırılmıştır. "
        "Zincirin taşıyıcı fikri, yönelim bilgisinin modelden değil "
        "ölçümden gelmesidir: işaret edilen pikselin çevresindeki nokta yaması "
        "bakış ışını boyunca dar bir banda indirgenerek arka plandan ayrılmakta, "
        "banda tekil değer ayrışımıyla düzlem oturtulmakta, en küçük tekil "
        "vektör yüzey normalini ve artığın karekök ortalaması kavramayı "
        "reddetme kapısını vermektedir."
    )
    b.p(
        "Gerçek robotta on koşunun sekizi tamamlanmıştır. Bu, bir başarı "
        "oranı değil bir yapılabilirlik kanıtıdır; 8/10 için %95 güven "
        "aralığı [%44, %98] olduğundan orandan çok başarısızlıkların dağılımı "
        "anlamlıdır. Kayıtlı bir koşuda kolun komut edilen poza izleme "
        "sapması temas noktasında milimetre altında, bırakma pozunda 1 mm "
        "ölçülmüş; vakum yaklaşık 200 ms'de %30,6 bağıl değerde "
        "kurulmuştur. Başarısız koşuların ikisinde de hata geometri zincirinde "
        "değil, modelin bırakma hedefini yanlış konumlandırmasındadır. Bu, "
        "çalışmanın en aktarılabilir sonucudur: modelin kaba işaretinden poz "
        "kuran ölçüm zinciri, bu hücrede kavrama başarısızlığını sıfıra "
        "indirmiş; kalan hata ise modelin ilişkisel bir sorguyu "
        "konumlandırmasında yoğunlaşmıştır."
    )
    b.p(
        "Bundan çıkan araştırma sorusu doğrudandır: modelin ürettiği hedef "
        "noktası, robotun onu izlemesinden *önce* nasıl doğrulanır? Bu çalışmada "
        "bırakma noktası yalnız erişilebilirlik ve çarpışma bakımından "
        "denetlenmekte, noktanın hedeflenen boşluğun içine düşüp düşmediği "
        "denetlenmemektedir. Aynı derinlik verisinden – gözün duvarlarının ve "
        "taban düzleminin çıkarılmasıyla – böyle bir içerik denetimi kurulması "
        "ve bunun başarı oranına etkisinin daha büyük bir koşu kümesiyle "
        "ölçülmesi, doğal yönde bir devam çalışmasıdır. İkinci bir yön, "
        "yönelim kestiriminin analitik vakum kavrama modelleriyle "
        "karşılaştırılmasıdır: bu çalışmada kavramayı reddeden ölçüt tek bir "
        "düzlemsellik büyüklüğüdür ve sızdırmazlığı doğrudan modellememektedir. "
        "Üçüncü yön, zincirin model-bağımsız olmasından yararlanarak ER 2'nin "
        "yerine elle yazılmış bir bölge seçicinin ve kabartma render yerine "
        "normal gölgelemenin konulduğu bileşen çıkarma deneyleridir; bunlar, "
        "bu makalede yalnız girdi kontrastı üzerinden gerekçelendirilen "
        "tasarım kararlarını başarı üzerinden ölçecektir. Dördüncüsü, "
        "endüstriyel geçerlilik için çevrim süresidir: oturma beklemelerinin "
        "denetleyici toleranslarıyla kaldırılması, bırakma sorgusunun taşıma "
        "hareketiyle örtüştürülmesi ve hareket profillerinin hızlandırılması "
        "birlikte 117 saniyelik sürenin yarıdan fazlasını hedeflemektedir."
    )


# ═════════════════════════════════════════════════════════════════════
# arka bölümler
# ═════════════════════════════════════════════════════════════════════
REFERENCES = [
    "Ahn, M., Brohan, A., Brown, N., Chebotar, Y., Cortes, O., David, B., ... "
    "Zeng, A. (2022). Do as I can, not as I say: Grounding language in robotic "
    "affordances. arXiv:2204.01691.",

    "Brohan, A., Brown, N., Carbajal, J., Chebotar, Y., Chen, X., "
    "Choromanski, K., ... Zitkovich, B. (2023). RT-2: Vision-language-action "
    "models transfer web knowledge to robotic control. arXiv:2307.15818.",

    "Coleman, D., Şucan, I. A., Chitta, S. ve Correll, N. (2014). Reducing the "
    "barrier to entry of complex robotic software: A MoveIt! case study. "
    "Journal of Software Engineering for Robotics, 5(1), 3-16.",

    "Gemini Robotics Ekibi. (2025). Gemini Robotics 1.5: Pushing the frontier "
    "of generalist robots with advanced embodied reasoning, thinking, and "
    "motion transfer. arXiv:2510.03342.",

    "Kim, M. J., Pertsch, K., Karamcheti, S., Xiao, T., Balakrishna, A., "
    "Nair, S., ... Finn, C. (2024). OpenVLA: An open-source "
    "vision-language-action model. arXiv:2406.09246.",

    "Liu, F., Fang, K., Abbeel, P. ve Levine, S. (2024). MOKA: Open-world "
    "robotic manipulation through mark-based visual prompting. "
    "arXiv:2403.03174.",

    "Macenski, S., Foote, T., Gerkey, B., Lalancette, C. ve Woodall, W. "
    "(2022). Robot Operating System 2: Design, architecture, and uses in the "
    "wild. Science Robotics, 7(66), eabm6074.",

    "Mahler, J., Matl, M., Liu, X., Li, A., Gealy, D. ve Goldberg, K. (2018). "
    "Dex-Net 3.0: Computing robust vacuum suction grasp targets in point "
    "clouds using a new analytic model and deep learning. IEEE International "
    "Conference on Robotics and Automation (ICRA)'da sunulmuş bildiri, "
    "Brisbane, Avustralya.",

    "Yuan, W., Duan, J., Blukis, V., Pumacay, W., Krishna, R., Murali, A., ... "
    "Fox, D. (2024). RoboPoint: A vision-language model for spatial affordance "
    "prediction for robotics. arXiv:2406.10721.",
]


def _arka(b: Builder) -> None:
    b.h1("Teşekkür")
    b.p("⟨YAZAR TARAFINDAN DOLDURULACAK⟩")

    b.h1("Araştırmacıların Katkısı")
    b.p("⟨YAZAR TARAFINDAN DOLDURULACAK: her yazarın katkısı, şablondaki "
        "biçime uygun olarak yazılmalıdır.⟩")

    b.h1("Çıkar Çatışması")
    b.p("Yazarlar tarafından herhangi bir çıkar çatışması beyan edilmemiştir.")

    b.h1("Araştırma ve Yayın Etiği")
    b.p("Bu çalışmada Araştırma ve Yayın Etiğine uyulmuştur. Çalışma insan ya "
        "da hayvan denek içermediğinden etik kurul onayı gerektirmemektedir; "
        "bütün deneyler yazarların kurumunda bulunan araştırma hücresinde "
        "yürütülmüştür.")

    b.h1("Kaynaklar")
    for ref in REFERENCES:
        p = b.para("", style="Kaynaklar", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        b.rich(p, ref, size=BODY_PT)


def body(b: Builder) -> None:
    _giris(b)
    _literatur(b)
    _yontem(b)
    _bulgular(b)
    _tartisma(b)
    _sonuclar(b)
    _arka(b)


def main() -> None:
    doc, anchor = open_template()
    b = BuilderTR(doc, anchor)
    front_matter(b)
    body(b)
    b._leave_wide()
    doc.save(str(OUT))
    print(f"yazıldı: {OUT}")
    print(f"  şekil: {b.fig_no}   tablo: {b.tab_no}   eşitlik: {b.eq_no}")


if __name__ == "__main__":
    main()
