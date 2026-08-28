"""
makale_duzen.py
===============
Derginin (ESOGÜ MMF Dergisi) şablonundan makale üreten YERLEŞİM MAKİNESİ.

anomaly_detection/docs/makale_uret.py içindeki denenmiş yerleşim kodunun
birebir kopyasıdır; yalnız TEMPLATE/FIG/OUT sabitleri bu pakete çevrilmiştir.
İçerik burada DEĞİLDİR - içerik makale_uret_tr.py dosyasındadır.

Burada tutulan dizgi bilgisi:
  * sütun genişliği 8,35 cm (17,9 cm metin - 1,25 cm oluk)
  * tablolar tblLayout=fixed + tblGrid olmadan LibreOffice'te sütunu taşırır
  * tam sayfa genişliğinde öğeler gövdeye 1 sütunluk sürekli bölüm sokarak
    yapılır; ARDIŞIK geniş öğeler tek blok paylaşmalıdır (pending_wide)
"""
import copy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
PKG = HERE.parent
TEMPLATE = HERE / "Dosya.docx"
FIG = HERE / "makale_figurleri"
OUT = HERE / "ESOGU_MMF_Gemini_Makale_TR.docx"

FONT = "Cambria"
BODY_PT = 10
SMALL_PT = 9
COL_W = 8.35      # cm — one column of the two-column body
                  # (210 − 2·15 mm text width, less the 1.25 cm gutter)
PAGE_W = 17.9     # cm — full text width

INK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x44, 0x44, 0x44)


# ═════════════════════════════════════════════════════════════════════
# document skeleton
# ═════════════════════════════════════════════════════════════════════
def open_template():
    """Load the journal template and strip its body, keeping both sectPr."""
    doc = Document(str(TEMPLATE))
    body = doc.element.body
    keep_break = None
    for child in list(body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        if child.find(".//" + qn("w:sectPr")) is not None:
            keep_break = child          # paragraph that closes section 0
            continue
        body.remove(child)
    return doc, keep_break


def style_run(run, size=BODY_PT, bold=False, italic=False, colour=INK,
              name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = colour
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    return run


def _ref_kaynagi(body):
    """Şablondaki İLK sectPr; üstbilgi/altbilgi referanslarını oradan alırız."""
    for pr in body.iter(qn("w:sectPr")):
        return pr
    return None


def _basliklari_esitle(sect, body):
    """Klonlanan bölüme şablonun ilk bölümündeki üst/alt bilgiyi ver.

    Gövdenin SON sectPr'ı sayfa numarası altbilgisine (footer2) bakar; oysa
    makalenin her sayfasında CC bloğu (footer1) ve koşu başlığı olmalıdır.
    Klonu olduğu gibi bırakırsak tam sayfa kaplayan bir şekil kendi bölümüne
    düşer düşmez o sayfanın altbilgisi sessizce değişir.
    """
    kaynak = _ref_kaynagi(body)
    if kaynak is None:
        return
    for eski in [e for e in sect if e.tag.endswith("Reference")]:
        sect.remove(eski)
    ilk = sect[0] if len(sect) else None
    for ref in [e for e in kaynak if e.tag.endswith("Reference")][::-1]:
        kopya = copy.deepcopy(ref)
        if ilk is None:
            sect.append(kopya)
        else:
            ilk.addprevious(kopya)


def set_cols(paragraph, num):
    """Close the section at `paragraph` with a `num`-column layout."""
    body = paragraph.part.document.element.body
    src = body.find(qn("w:sectPr"))
    sect = copy.deepcopy(src)
    _basliklari_esitle(sect, body)
    cols = sect.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect.append(cols)
    if num == 1:
        for attr in ("w:num", "w:equalWidth"):
            if cols.get(qn(attr)) is not None:
                del cols.attrib[qn(attr)]
        cols.set(qn("w:space"), "708")
    else:
        cols.set(qn("w:num"), str(num))
        cols.set(qn("w:equalWidth"), "1")
        cols.set(qn("w:space"), "709")
    ppr = paragraph._p.get_or_add_pPr()
    for old in ppr.findall(qn("w:sectPr")):
        ppr.remove(old)
    ppr.append(sect)


# ═════════════════════════════════════════════════════════════════════
# building blocks
# ═════════════════════════════════════════════════════════════════════
class Builder:
    # Şekil ve tablo başlıklarının dil etiketi. Türkçe sürüm bunları
    # "Şekil"/"Tablo" olarak geçersiz kılar; böylece yerleşim kodu tek yerde
    # kalır ve iki sürüm arasında biçim farkı oluşmaz.
    FIG_LABEL = "Figure"
    TAB_LABEL = "Table"

    def __init__(self, doc, anchor):
        self.doc = doc
        self.anchor = anchor          # section-0 closing paragraph
        self.body = doc.element.body
        self.fig_no = 0
        self.tab_no = 0
        self.eq_no = 0
        self.last = None
        self.pending_wide = None      # last paragraph of an open 1-column block

    # -- column-span bookkeeping --------------------------------------
    def _begin_wide(self):
        """Open a full-width (single-column) block, unless one is open."""
        if self.pending_wide is not None:
            return                    # consecutive wide elements share a block
        if self.last is not None:
            set_cols(self.last, 2)

    def _end_wide(self, paragraph):
        self.pending_wide = paragraph

    def _leave_wide(self):
        if self.pending_wide is not None:
            set_cols(self.pending_wide, 1)
            self.pending_wide = None

    # -- placement -----------------------------------------------------
    def _add(self, element, front=False):
        if front:
            self.anchor.addprevious(element)
        else:
            self.body.append(element)
            # keep the final sectPr last
            sect = self.body.find(qn("w:sectPr"))
            if sect is not None:
                self.body.remove(sect)
                self.body.append(sect)

    def para(self, text="", style="Paragraf", align=None, front=False,
             size=BODY_PT, bold=False, italic=False, colour=INK,
             space_before=None, space_after=None, keep_with_next=False,
             keep_wide=False):
        if not keep_wide and not front:
            self._leave_wide()
        p = self.doc.add_paragraph()
        self.body.remove(p._p)
        self._add(p._p, front)
        try:
            p.style = self.doc.styles[style]
        except KeyError:
            p.style = self.doc.styles["Normal"]
        if align is not None:
            p.alignment = align
        if space_before is not None:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        if keep_with_next:
            p.paragraph_format.keep_with_next = True
        if text:
            self.rich(p, text, size=size, bold=bold, italic=italic,
                      colour=colour)
        self.last = p
        return p

    def rich(self, paragraph, text, size=BODY_PT, bold=False, italic=False,
             colour=INK):
        """`**bold**` and `*italic*` inline markers."""
        buf = ""
        i = 0
        while i < len(text):
            if text.startswith("**", i):
                j = text.find("**", i + 2)
                if j > 0:
                    if buf:
                        style_run(paragraph.add_run(buf), size, bold, italic, colour)
                        buf = ""
                    style_run(paragraph.add_run(text[i + 2:j]), size, True,
                              italic, colour)
                    i = j + 2
                    continue
            if text[i] == "*":
                j = text.find("*", i + 1)
                if j > 0:
                    if buf:
                        style_run(paragraph.add_run(buf), size, bold, italic, colour)
                        buf = ""
                    style_run(paragraph.add_run(text[i + 1:j]), size, bold,
                              True, colour)
                    i = j + 1
                    continue
            buf += text[i]
            i += 1
        if buf:
            style_run(paragraph.add_run(buf), size, bold, italic, colour)
        return paragraph

    # -- headings ------------------------------------------------------
    def h1(self, text):
        return self.para(text, style="Başlık-1", size=BODY_PT, bold=True,
                         space_before=12, space_after=2, keep_with_next=True)

    def h2(self, text):
        return self.para(text, style="Başlık-1", size=BODY_PT, bold=True,
                         space_before=8, space_after=2, keep_with_next=True)

    def p(self, text):
        return self.para(text, style="Paragraf",
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=6, space_after=0)

    # -- equations -----------------------------------------------------
    def equation(self, text, label=None):
        """
        `label` verilirse numara sayaçtan ALINMAZ ve sayaç ilerlemez.

        Ana denklemler arasına ek bir denklem sokmak gerektiğinde (ör. kalıntı
        tanımının sürtünmeli hâli) "2a"/"2b" gibi bir etiket kullanılır; aksi
        hâlde sonraki bütün denklem numaraları kayar ve metindeki her atıf
        sessizce yanlışa döner.
        """
        if label is None:
            self.eq_no += 1
        self.para("", style="Normal", space_before=0, space_after=0)
        p = self.para(style="Eşitlik", space_before=0, space_after=0)
        pf = p.paragraph_format
        for existing in list(pf.tab_stops):
            pf.tab_stops.clear_all()
            break
        pf.tab_stops.add_tab_stop(Cm(COL_W), WD_TAB_ALIGNMENT.RIGHT)
        style_run(p.add_run("\t" if text.startswith("\t") else ""), BODY_PT)
        self.rich(p, text, size=BODY_PT, italic=True)
        style_run(p.add_run("\t"), BODY_PT)
        style_run(p.add_run(f"({label or self.eq_no})"), BODY_PT)
        self.para("", style="Normal", space_before=0, space_after=0)
        return label or self.eq_no

    # -- figures -------------------------------------------------------
    def figure(self, filename, caption, wide=True, width_cm=None):
        path = FIG / filename
        if not path.exists():
            print(f"  ! missing figure {filename}")
            return None
        self.fig_no += 1
        if wide:
            self._begin_wide()
        else:
            self._leave_wide()
        width = width_cm or (PAGE_W if wide else COL_W)
        pic = self.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=6, space_after=2, keep_with_next=True,
                        keep_wide=True)
        pic.add_run().add_picture(str(path), width=Cm(width))
        cap = self.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=6, keep_wide=True)
        name, _, extra = caption.partition("|")
        self.rich(cap, f"**{self.FIG_LABEL} {self.fig_no}.** {name.strip()}",
                  size=SMALL_PT)
        if extra.strip():
            self.rich(cap, "  " + extra.strip(), size=SMALL_PT - 0.5,
                      colour=GREY)
        if wide:
            self._end_wide(cap)
        self.last = cap
        return self.fig_no

    # -- tables --------------------------------------------------------
    def table(self, caption, header, rows, widths=None, wide=False,
              align_right=None, note=None):
        self.tab_no += 1
        if wide:
            self._begin_wide()
        else:
            self._leave_wide()
        cap = self.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=8, space_after=3, keep_with_next=True,
                        keep_wide=True)
        self.rich(cap, f"**{self.TAB_LABEL} {self.tab_no}.** {caption}",
                  size=SMALL_PT)

        ncol = len(header)
        tbl = self.doc.add_table(rows=1 + len(rows), cols=ncol)
        self.body.remove(tbl._tbl)
        self._add(tbl._tbl)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        total = PAGE_W if wide else COL_W
        if widths is None:
            widths = [total / ncol] * ncol
        scale = total / sum(widths)
        widths = [w * scale for w in widths]

        align_right = align_right or []
        data = [header] + list(rows)
        for r, row in enumerate(data):
            for c, val in enumerate(row):
                cell = tbl.cell(r, c)
                cell.width = Cm(widths[c])
                par = cell.paragraphs[0]
                par.paragraph_format.space_before = Pt(1)
                par.paragraph_format.space_after = Pt(1)
                par.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if c in align_right
                                 and r > 0 else WD_ALIGN_PARAGRAPH.LEFT)
                self.rich(par, str(val), size=SMALL_PT - 1.0, bold=(r == 0))
        self._fix_layout(tbl, widths)
        self._rule_table(tbl)
        if note:
            n = self.para("", style="Normal", space_before=2, space_after=6,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, keep_wide=True)
            self.rich(n, note, size=SMALL_PT - 1, italic=True, colour=GREY)
            self.last = n
        else:
            self.last = self.para("", style="Normal", space_before=0,
                                  space_after=4, keep_wide=True)
        if wide:
            self._end_wide(self.last)
        return self.tab_no

    @staticmethod
    def _fix_layout(tbl, widths_cm):
        """Pin the column grid so Word and LibreOffice stop autofitting."""
        twips = [int(round(w * 567)) for w in widths_cm]
        tblpr = tbl._tbl.tblPr
        for tag in ("w:tblW", "w:tblLayout"):
            for old in tblpr.findall(qn(tag)):
                tblpr.remove(old)
        tw = OxmlElement("w:tblW")
        tw.set(qn("w:w"), str(sum(twips)))
        tw.set(qn("w:type"), "dxa")
        tblpr.append(tw)
        lay = OxmlElement("w:tblLayout")
        lay.set(qn("w:type"), "fixed")
        tblpr.append(lay)
        for old in tbl._tbl.findall(qn("w:tblGrid")):
            tbl._tbl.remove(old)
        grid = OxmlElement("w:tblGrid")
        for t in twips:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(t))
            grid.append(gc)
        tblpr.addnext(grid)
        for row in tbl.rows:
            for cell, t in zip(row.cells, twips):
                tcpr = cell._tc.get_or_add_tcPr()
                for old in tcpr.findall(qn("w:tcW")):
                    tcpr.remove(old)
                el = OxmlElement("w:tcW")
                el.set(qn("w:w"), str(t))
                el.set(qn("w:type"), "dxa")
                tcpr.append(el)

    @staticmethod
    def _rule_table(tbl):
        """Horizontal rules only: above and below the header, below the last row."""
        rows = tbl.rows
        for idx, row in enumerate(rows):
            for cell in row.cells:
                tcpr = cell._tc.get_or_add_tcPr()
                for old in tcpr.findall(qn("w:tcBorders")):
                    tcpr.remove(old)
                borders = OxmlElement("w:tcBorders")
                for edge in ("top", "left", "bottom", "right"):
                    el = OxmlElement(f"w:{edge}")
                    show = ((edge == "top" and idx <= 1)
                            or (edge == "bottom" and idx == len(rows) - 1))
                    el.set(qn("w:val"), "single" if show else "nil")
                    if show:
                        el.set(qn("w:sz"), "8" if idx == 0 or
                               idx == len(rows) - 1 else "4")
                        el.set(qn("w:space"), "0")
                        el.set(qn("w:color"), "000000")
                    borders.append(el)
                tcpr.append(borders)


# ═════════════════════════════════════════════════════════════════════
# front matter
# ═════════════════════════════════════════════════════════════════════


def abstract_block(b, kw_head, ab_head, keywords, abstract, kind):
    tbl = b.doc.add_table(rows=3, cols=2)
    b.body.remove(tbl._tbl)
    b.anchor.addprevious(tbl._tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [4.6, 13.3]
    heads = [kw_head, ab_head]
    for c in range(2):
        cell = tbl.cell(0, c)
        cell.width = Cm(widths[c])
        par = cell.paragraphs[0]
        par.paragraph_format.space_before = Pt(2)
        par.paragraph_format.space_after = Pt(2)
        style_run(par.add_run(heads[c]), BODY_PT, bold=True)

    cell = tbl.cell(1, 0)
    cell.width = Cm(widths[0])
    for i, kw in enumerate(keywords):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_run(par.add_run(kw), SMALL_PT, italic=True)

    cell = tbl.cell(1, 1)
    cell.width = Cm(widths[1])
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(2)
    style_run(par.add_run(abstract), BODY_PT, italic=True)

    cell = tbl.cell(2, 0)
    par = cell.paragraphs[0]
    style_run(par.add_run(kind), SMALL_PT, bold=True)
    par = tbl.cell(2, 1).paragraphs[0]
    style_run(par.add_run("Received / Geliş :               "
                          "Accepted / Kabul :"), SMALL_PT, colour=GREY)
    Builder._fix_layout(tbl, widths)
    Builder._rule_table(tbl)


# ═════════════════════════════════════════════════════════════════════
# body
# ═════════════════════════════════════════════════════════════════════
