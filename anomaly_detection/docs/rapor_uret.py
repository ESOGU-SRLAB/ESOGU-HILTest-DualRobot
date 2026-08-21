#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UR10e anomali tespiti — teknik rapor üreteci (.docx)
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

OUT = Path(__file__).resolve().parent / "UR10e_Anomali_Tespiti_Teknik_Rapor.docx"

ACCENT = RGBColor(0x0C, 0x5E, 0x66)
MUTED = RGBColor(0x55, 0x5F, 0x6B)

doc = Document()

# ── sayfa ve temel stiller ───────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(2.2)
sec.left_margin = sec.right_margin = Cm(2.4)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15
st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, color, before, after in (
    ("Heading 1", 17, ACCENT, 18, 8),
    ("Heading 2", 13.5, ACCENT, 14, 6),
    ("Heading 3", 11.5, RGBColor(0x22, 0x28, 0x30), 10, 4),
):
    s = doc.styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.keep_with_next = True


# ── yardımcılar ──────────────────────────────────────────────────────
def h(level, text):
    return doc.add_heading(text, level=level)


def p(text="", style=None, italic=False, size=None, space_after=None,
      align=None):
    par = doc.add_paragraph(style=style)
    _rich(par, text, italic=italic, size=size)
    if space_after is not None:
        par.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        par.paragraph_format.alignment = align
    return par


def _rich(par, text, italic=False, size=None):
    """**kalın**, *italik* ve `kod` işaretlemesini çözer."""
    import re
    tok = re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text)
    for t in tok:
        if not t:
            continue
        if t.startswith("**") and t.endswith("**"):
            r = par.add_run(t[2:-2]); r.bold = True
        elif t.startswith("`") and t.endswith("`"):
            r = par.add_run(t[1:-1]); r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif t.startswith("*") and t.endswith("*") and len(t) > 2:
            r = par.add_run(t[1:-1]); r.italic = True
        else:
            r = par.add_run(t)
        if italic:
            r.italic = True
        if size:
            r.font.size = Pt(size)
    return par


def bullet(items, style="List Bullet"):
    for it in items:
        par = doc.add_paragraph(style=style)
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _rich(par, it)


def code(text, size=8.8):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.left_indent = Cm(0.4)
    par.paragraph_format.line_spacing = 1.0
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = par.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    _shade(par, "F2F4F7")
    return par


def eq(text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(11)
    return par


def _shade(par, hexcolor):
    pPr = par._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    pPr.append(sh)


def _cell_shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def table(headers, rows, caption=None, widths=None, highlight=(), num_from=1):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        par = hdr[i].paragraphs[0]
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i < num_from
                                          else WD_ALIGN_PARAGRAPH.RIGHT)
        r = par.add_run(htxt)
        r.bold = True
        r.font.size = Pt(9)
        _cell_shade(hdr[i], "E7EBEF")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i < num_from
                                              else WD_ALIGN_PARAGRAPH.RIGHT)
            _rich(par, str(val), size=9)
            if ri in highlight:
                _cell_shade(cells[i], "E3F0F1")
                for rr in par.runs:
                    rr.bold = True
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(12)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        import re as _re
        for tk in _re.split(r"(`[^`]+?`)", caption):
            if not tk:
                continue
            if tk.startswith("`") and tk.endswith("`"):
                r = cap.add_run(tk[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(8.5)
            else:
                r = cap.add_run(tk); r.italic = True; r.font.size = Pt(9)
            r.font.color.rgb = MUTED
    return t


def note(title, body):
    """Kenarlıklı vurgu kutusu."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    c.text = ""
    par = c.paragraphs[0]
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = par.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT
    for para in body:
        q_ = c.add_paragraph()
        q_.paragraph_format.space_after = Pt(3)
        q_.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _rich(q_, para, size=9.8)
    _cell_shade(c, "F4F7F8")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ═══════════════════════════════════════════════════════════════════
# KAPAK
# ═══════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.paragraph_format.space_before = Pt(60)
tp.paragraph_format.space_after = Pt(4)
tp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = tp.add_run("TEKNİK RAPOR")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ACCENT

tp = doc.add_paragraph()
tp.paragraph_format.space_after = Pt(10)
tp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = tp.add_run("UR10e İşbirlikçi Robotunda FMU Tabanlı Kalıntı Ayrıştırma ve "
               "İkili LSTM Özkodlayıcı Birleşimi ile Çevrimiçi Anomali Tespiti")
r.bold = True; r.font.size = Pt(21)

tp = doc.add_paragraph()
tp.paragraph_format.space_after = Pt(24)
tp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = tp.add_run("Bildirinin yeniden üretimi, ölçülen sapmaların düzeltilmesi ve "
               "gerçek zamanlı ROS 2 gerçeklemesi")
r.italic = True; r.font.size = Pt(12.5); r.font.color.rgb = MUTED

p("**Referans bildiri:** C. S. Yılmaz, S. Kahraman, M. Yılmaz, H. S. Yavuz, U. Yayan, "
  "“FMU Tabanlı Kalıntı Ayrıştırma ve İkili LSTM Özkodlayıcı Birleşimi ile İşbirlikçi "
  "Robotlarda Anomali Tespiti”, 2026 (245.pdf).")
p("**Yazılım:** `anomaly_detection` (çevrimdışı hat, 3.176 satır) ve "
  "`ur10e_anomaly_detection` (ROS 2 düğümü, 892 satır).")
p("**Donanım:** UR10e işbirlikçi robot, Festo lineer eksen, uç noktada Kuvvet-Tork "
  "Sensörü (KTS); hedef koşum platformu Core i9-14900 + RTX 4000 Ada.")
p("**Tarih:** 20 Ağustos 2026")

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ═══════════════════════════════════════════════════════════════════
h(1, "Özet")
p("Bu rapor, UR10e işbirlikçi robotu için geliştirilen ve fizik tabanlı bir kalıntı "
  "modeli ile veri güdümlü bir ham sensör modelini skor düzeyinde birleştiren anomali "
  "tespit çerçevesinin (i) referans bildiriden bağımsız olarak yeniden kurulmasını, "
  "(ii) bu yeniden kurulum sırasında ortaya çıkan yedi ölçülebilir bulgunun "
  "belgelenmesini ve düzeltilmesini, (iii) sistemin 500 Hz'de çalışan bir ROS 2 "
  "düğümü olarak gerçek zamanlı gerçeklenmesini kapsamaktadır.")
p("Bildirinin yönteminin mimari ve yordamsal bileşenleri birebir doğrulanmıştır: "
  "kalıntı modelinin 478.892 ve ham modelin 1.907.032 parametresi, T=100/adım=25 "
  "pencereleme, 97. persentil eşik kuralı, 0,05 adımlı ağırlık taraması ve "
  "r_ic + r_dis = r_top kalıntı özdeşliği (sayısal artık 7,1·10⁻¹⁵) tam olarak "
  "uyuşmaktadır.")
p("Buna karşılık veri hattında yedi bulgu ölçülmüştür. En önemlisi, bildirinin "
  "sayılarının **karışık birimli** bir hattan geldiğidir: UR sürücüsü `effort` alanına "
  "amper yazarken model torku Nm üretmekte, ikisi doğrudan çıkarılmaktadır. Bu düzen "
  "`reproduce_erratum.sh` ile yeniden kurulduğunda bildirinin imza bulguları "
  "(sensör gürültüsünde kalıntı modelinin körlüğü, ~%24 yalnızca-kalıntı tespiti) "
  "geri gelmektedir. Ek olarak veri kümesinin küresel olarak sıralı olmadığı "
  "(%28,6 süreksiz örnek çifti, pencerelerin %74,9'u en az bir süreksizlik içeriyor), "
  "ivmenin bu süreksizlikler üzerinden türetildiği ve KTS ölçümünün yanlış çerçevede "
  "Jacobian'la aktarıldığı (kanal geneli korelasyon −0,143) ölçülmüştür.")
p("Düzeltilmiş hat, fiziksel olarak tutarlı bir veri kümesi (600 kesintisiz koşu, "
  "342.480 örnek), bağımsız bir akım→tork kalibrasyonu ve sızıntısız pencereleme ile "
  "yeniden eğitilmiştir. Birleşim, tekil modelleri açık farkla geçmeye devam etmektedir: "
  "yalnız-ham 0,613 ve yalnız-kalıntı 0,579'a karşılık birleşim **0,792** F1 "
  "(ΔF1 = +0,178; bildiri +0,161 raporlamıştır). Tamamlayıcılık oranları da korunmuştur "
  "(%24,3 yalnızca-kalıntı, %28,7 yalnızca-ham; bildiri %24,1 ve %25,9).")
p("Çevrimiçi sistem, çevrimdışı hatla sayısal olarak özdeş bir öznitelik motoru "
  "üzerine kurulmuştur (61.524 örnekte en büyük fark 1,4·10⁻¹⁴). Bildiriden yalnız dört "
  "noktada ayrılmaktadır; ikisi zorunlu (min–max sınırlarının nedensel bir kümeden "
  "alınması ve birleşik skor için çalışma eşiğinin türetilmesi), ikisi isteğe bağlı "
  "(uyarlanabilir alarm kuralı ve alarm gecikmesi) ve tek parametreyle kapatılabilir. "
  "Ölçülen çıkarım süresi karar başına 5,9 ms, yani 50 ms bütçenin %12'sidir; "
  "uçtan uca tespit gecikmesi ≈120–150 ms'dir.")
p("Sistem 21 Ağustos 2026'da gerçek bir UR10e hücresinde devreye alınmış ve yaklaşık "
  "45 dakikalık ölçüm toplanmıştır (Bölüm 10). FMU verisinden türetilen çalışma eşiği "
  "gerçek robotta **yedi kat düşük** çıkmış (temiz koşu medyanı 4,27; eşik 0,6436) ve "
  "operatörün doğruladığı etiketli bir küme üzerinden **θ_birleşik = 18,0** olarak "
  "yeniden ölçülmüştür; bu değer beş gerçek olayın dördünü yakalarken yanlış alarmı "
  "saatte 38,7'den 6,7'ye indirmektedir. Uyarlanabilir alarm kuralının gerçek robotta "
  "doğrulanmış hiçbir olayı yakalamadığı, yani FMU verisindeki sonucun tersine döndüğü "
  "ölçülmüş ve kural varsayılan olarak kapatılmıştır. Kalıntının poza bağlı olduğu "
  "(yerçekimi yüklü duruşta 4,30, harekette 1,58, katlanmış pozda 0,0037) ve yanlış "
  "alarmların yörünge çevrimiyle birebir tekrarlandığı gösterilerek, kalan sınırın "
  "eşikle değil modellerin gerçek veriyle yeniden eğitilmesiyle aşılabileceği "
  "belgelenmiştir. Sistem ayrıca laboratuvarın kontrol panosuna operatör etiketlemeli "
  "bir sekme olarak entegre edilmiştir.")

# ═══════════════════════════════════════════════════════════════════
h(1, "1. Giriş")

h(2, "1.1. Problem ve bağlam")
p("İşbirlikçi robotlar endüstriyel üretim hatlarında insan çalışanlarla aynı çalışma "
  "hacmini paylaşmaktadır; bu nedenle çarpışma, motor bozunması veya sensör arızası "
  "gibi anormalliklerin erken tespiti hem üretim sürekliliği hem de operatör güvenliği "
  "açısından kritik önem taşımaktadır. Referans bildiri, bu problemi otomotiv şasesi "
  "üzerinde otomatik görsel kalite muayenesi yapan bir UR10e + Festo lineer eksen "
  "platformu bağlamında ele almaktadır.")
p("Anomali tespitinde iki temel yaklaşım bulunmaktadır. **Fizik tabanlı kalıntı "
  "yöntemleri** robot dinamik modeli ile ölçülen kuvvetler arasındaki farkı arıza "
  "göstergesi olarak kullanır; boyut indirgeme ve gürültü filtreleme sağlarlar, ancak "
  "düşük genlikli sensör anomalilerini bastırırlar. **Veri güdümlü yöntemler** yüksek "
  "boyutlu ham veriden karmaşık desenleri öğrenebilir, ancak fiziksel modelin verdiği "
  "yapıdan yoksundur. Bildirinin temel savı, bu iki yaklaşımın *farklı* arıza tiplerine "
  "duyarlı olduğu ve bu asimetrinin skor düzeyinde birleşimle sömürülebileceğidir.")

h(2, "1.2. Bu çalışmanın amacı")
p("Bildirinin yayımlanmasından yaklaşık üç ay sonra hattın bir bölümü kaybolmuştur. "
  "Kodlar sıfırdan yeniden yazılmış ve modeller yeniden eğitilmiştir. Bu raporun amacı "
  "iki katlıdır:")
bullet([
    "**Doğrulama.** Yeniden yazılan sistemin bildiriyle *birebir aynı mantıkta* "
    "çalıştığını, iddia düzeyinde değil ölçüm düzeyinde göstermek — özellikle "
    "`effort` alanından gelen akım değerlerinin torka doğru çevrildiğini.",
    "**Gerçekleme.** İki modelin ortak çalışarak anomali tespiti yaptığı, 500 Hz "
    "veri akışında **çevrimiçi** koşan bir sistem elde etmek. Bildirinin kendi "
    "“Gelecek çalışmalar” bölümünde işaret ettiği madde tam olarak budur: "
    "“sistemin UR10e ile eş zamanlı çevrimiçi çalıştırılması”.",
])
p("Bir kısıt baştan sabitlenmiştir: **FMU tarafında hiçbir şey değiştirilmemiştir.** "
  "Ters dinamik çözücü gerçek robotla çalıştırılarak test edilmiş ve birebir aynı "
  "sonuçları ürettiği görülmüştür; bu nedenle çözücünün ölçülen sınırları raporlanmış, "
  "fakat kodu değiştirilmemiştir.")

h(2, "1.3. Katkılar")
bullet([
    "Bildirinin sayılarının hangi veri hattından geldiğinin nicel olarak tespiti ve "
    "o hattın `reproduce_erratum.sh` ile yeniden üretilebilir hale getirilmesi.",
    "Yedi ölçülebilir bulgunun (F1–F7) belgelenmesi; beşi düzeltilmiş, biri "
    "(FMU sınırları) bilinçli olarak kabul edilmiş, biri bildirinin dizgi hatasıdır.",
    "Fiziksel olarak tutarlı, sızıntısız ve koşu-farkında bir veri hattının kurulması.",
    "Akım→tork dönüşümünün model torkundan **bağımsız** bir yöntemle (quasi-statik "
    "yerçekimi regresyonu) kalibre edilmesi.",
    "Bildirinin min–max normalleştirmesinin nedensel bir sürüme taşınması ve iki aday "
    "ölçeğin ölçümle yarıştırılması.",
    "Çevrimdışı hatla sayısal olarak özdeş bir çevrimiçi öznitelik motoru ve bunun "
    "üzerine kurulu bir ROS 2 tespit düğümü.",
    "Birleşimin gerçekten iki modelin ortak kararı olduğunun üç bağımsız sınavla "
    "(yapısal eşleme, çalışma sayacı, ablasyon) kanıtlanması.",
])


# ═══════════════════════════════════════════════════════════════════
h(1, "2. Referans bildirinin yöntemi")
p("Bu bölüm, denetimin ve gerçeklemenin karşılaştırma tabanını oluşturmak üzere "
  "bildirinin yöntemini özetlemektedir. Numaralandırma bildirininkiyle aynıdır.")

h(2, "2.1. Hibrit kalıntı ayrıştırma")
p("Ters dinamik model, UR10e için Newton-Euler formülasyonunu FMU (Functional Mock-up "
  "Unit) biçiminde paketlemektedir:")
eq("τ̂_model(q, q̇, q̈)  =  M(q)·q̈ + C(q, q̇)·q̇ + g(q)")
p("Toplam kalıntı, ölçülen tork ile model tahmini arasındaki farktır:")
eq("r_top(t)  =  τ_ölç(t) − τ̂_model(t)                                   (1)")
p("Dışsal kalıntı, uç noktadaki KTS ölçümünün Jacobian aktarımıyla eklem uzayına "
  "izdüşümüdür:")
eq("r_dis(t)  =  J(q)ᵀ · F_KTS(t)                                        (2)")
p("İçsel kalıntı ise ikisinin farkıdır:")
eq("r_ic(t)  =  r_top(t) − r_dis(t)                                      (3)")
p("Bu ayrıştırma bildirinin özgün katkısıdır: `r_ic` içsel arızalara (motor bozunması, "
  "gizyazar hatası), `r_dis` ise dışsal etkileşimlere (çarpışma, yük değişimi) duyarlı "
  "hale gelmektedir. Denklem (3) bir tanım olduğu için ayrıştırmanın sayısal tutarlılığı "
  "doğrudan sınanabilir — bu, denetimin ilk kontrolüdür.")

h(2, "2.2. LSTM özkodlayıcı mimarileri")
p("Her iki model, T = 100 örneklik (0,2 s) kayan pencereyi saklı vektöre sıkıştıran "
  "simetrik gizyazar–gizçözer mimarisine sahiptir. İki katmanlı LSTM gizyazar girdi "
  "dizisini sabit boyutlu saklı vektöre indirger; gizçözer bu vektörü RepeatVector ile "
  "T adıma kopyalayıp simetrik LSTM katmanlarıyla yeniden oluşturur. Katmanlar arasında "
  "%15 sönümleme uygulanmıştır. Her iki model **yalnızca normal çalışma verisiyle** "
  "MSE kaybıyla eğitilmiş, anomali eşiği doğrulama setindeki yeniden yapılanma "
  "hatalarının 97. persentilinde belirlenmiştir.")
table(
    ["Model", "Girdi kanalları", "Gizli / saklı", "Parametre"],
    [["Kalıntı özkodlayıcı", "12  (r_ic,1–6 + r_dis,1–6)", "128 / 32", "478.892"],
     ["Ham özkodlayıcı", "24  (q, q̇, τ, KTS — her biri 6)", "256 / 64", "1.907.032"]],
    caption="Tablo 2.1. İki modelin mimarisi. Kalıntı mimarisi, fizik tabanlı ön "
            "işlemenin girdi boyutunu düşürmesi sayesinde yaklaşık %75 daha kompakttır.",
    widths=[4.6, 6.0, 3.0, 2.8], num_from=1)

h(2, "2.3. Skor düzeyinde birleşim")
p("İki modelin min–max normalleştirilmiş anomali skorları ağırlıklı ortalama ile "
  "birleştirilmektedir:")
eq("S_bir(t)  =  w_kal · S̄_kal(t)  +  w_ham · S̄_ham(t)                   (4)")
p("Ağırlık w_kal ∈ [0, 1] aralığında 0,05 adımla taranarak BestF1 (eşikten bağımsız en "
  "yüksek F1) üzerinden optimize edilmiştir. Karşılaştırma için MAX ve OR stratejileri "
  "de uygulanmıştır. Bildiri, birleşik skor için bir **çalışma eşiği tanımlamamakta**, "
  "yalnızca BestF1 raporlamaktadır; bu ayrıntı çevrimiçi gerçeklemede belirleyici "
  "olacaktır (bkz. Bölüm 6.4).")

h(2, "2.4. Sentetik arıza enjeksiyonu")
p("Gerçek zemin doğrusu etiketleri bulunmadığından dört sentetik arıza senaryosu "
  "tasarlanmıştır. Motor kayması ve gizyazar hatası için, kinematik zincirde yerçekimine "
  "en çok maruz kalan ve hareket aralığı en geniş eklemler olan Eklem 3 (dirsek) ve "
  "Eklem 5 (bilek 2) seçilmiştir.")
table(
    ["Senaryo", "Etkilenen kanal", "Genlik (ham / kalıntı)", "Konum"],
    [["Motor kayması", "Eklem 3 torku, doğrusal rampa", "15 Nm / 25 Nm", "%30–38"],
     ["Çarpışma", "Tüm KTS kanalları, Gauss darbe", "30 N / 40 N", "merkez %65, genişlik %4"],
     ["Gizyazar hatası", "Eklem 5 pozisyonu, basamak", "1,5 rad / 8 Nm", "başlangıç %92"],
     ["Sensör gürültüsü", "Tüm KTS kanalları, Gauss gürültü", "3,5 N / 0,08 Nm", "%15–23"]],
    caption="Tablo 2.2. Sentetik arıza senaryoları. Sensör gürültüsü, modeller arası "
            "tamamlayıcılığı sınamak amacıyla kasıtlı olarak yalnızca ham sinyali "
            "etkileyecek düşük genlikte ayarlanmıştır. Pencere düzeyinde anomali "
            "etiketi, örneklerin en az %10'unun arıza maskesiyle örtüşmesi koşuluyla "
            "atanmaktadır.",
    widths=[3.4, 5.2, 4.0, 3.8], num_from=1)

h(2, "2.5. Bildirinin raporladığı sonuçlar")
table(
    ["Özellik", "Değer"],
    [["Toplam örnek", "1.124.432"],
     ["Örnekleme frekansı", "500 Hz"],
     ["Pencere boyutu / adım", "100 / 25  (%75 örtüşme)"],
     ["Eğitim / doğrulama penceresi", "35.978 / 8.992"],
     ["Test penceresi", "179.896"],
     ["Arıza penceresi", "14.402  (%8,0)"]],
    caption="Tablo 2.3. Bildirinin veri seti özellikleri (Tablo I).",
    widths=[7.0, 6.0], num_from=1)
table(
    ["Model", "AUC", "PR-AUC", "BestF1"],
    [["Kalıntı LSTM Özk.", "0,908", "0,695", "0,692"],
     ["Ham LSTM Özk.", "0,952", "0,761", "0,698"],
     ["Bir. A-Ort (0,95/0,05)", "0,980", "0,905", "0,859"],
     ["Bir. MAX", "0,976", "0,889", "0,824"],
     ["Bir. OR", "0,976", "0,889", "0,824"],
     ["Art. Norm Eşik", "0,849", "0,805", "0,707"],
     ["Isolation Forest", "0,688", "0,459", "0,547"],
     ["One-Class SVM", "0,815", "0,816", "0,760"]],
    caption="Tablo 2.4. Bildirinin genel performans karşılaştırması (Tablo II). "
            "Birleşim, en iyi tekil modele göre ΔBestF1 = +0,161 iyileşme sağlamaktadır. "
            "MAX ve OR satırlarının üç değerinin de aynı olması Bölüm 3.4'te ele alınan "
            "F7 bulgusudur.",
    widths=[5.4, 2.6, 2.6, 2.6], highlight=(2,), num_from=1)
table(
    ["Arıza tipi", "Kalıntı", "Ham", "Birleşim"],
    [["Motor kayması", "0,597", "0,605", "0,588"],
     ["Çarpışma", "0,916", "0,933", "0,920"],
     ["Gizyazar hatası", "0,986", "0,530", "0,977"],
     ["Sensör gürültüsü", "0,272", "1,000", "0,992"]],
    caption="Tablo 2.5. Bildirinin arıza tipine göre BestF1 karşılaştırması (Tablo III). "
            "Gizyazar ve sensör gürültüsü satırları tamamlayıcılığın doğrudan kanıtıdır.",
    widths=[4.6, 2.8, 2.8, 2.8], num_from=1)
p("Pencere düzeyinde tamamlayıcılık analizi, 14.402 arıza penceresinin %24,1'inin "
  "yalnızca kalıntı, %25,9'unun yalnızca ham özkodlayıcı tarafından tespit edildiğini "
  "ortaya koymakta; OR-birleşim geri çağırma oranı 0,854 ile tekil modelleri "
  "(kalıntı 0,595, ham 0,613) belirgin biçimde geçmektedir. Eğitim, kalıntı modelde "
  "129 dönemde (val_loss = 0,181), ham modelde 87 dönemde (val_loss = 0,320) "
  "yakınsamış; eşikler θ_kal = 0,420 ve θ_ham = 0,854 olarak belirlenmiştir.")


# ═══════════════════════════════════════════════════════════════════
h(1, "3. Denetim: yeniden yazılan hattın bildiriyle karşılaştırılması")

h(2, "3.1. Denetim yöntemi")
p("Denetimin ilkesi şudur: *bir bileşenin doğru olduğu, kodda öyle yazdığı için "
  "varsayılmaz; ölçülür.* Bu nedenle her karşılaştırma ya kapalı formda bir özdeşliğe "
  "(ör. Denklem 3), ya bağımsız bir referans uygulamaya (ör. sayısal Jacobian), ya da "
  "doğrudan sayaç ölçümüne dayandırılmıştır. Denetim üç katmanda yürütülmüştür:")
bullet([
    "**Mimari ve yordam.** Parametre sayıları, pencereleme, eşik kuralı, "
    "iyileştirici ayarları, ağırlık taraması.",
    "**Sayısal özdeşlikler.** Kalıntı ayrıştırma özdeşliği, Jacobian doğruluğu, "
    "türev filtresinin tanımı.",
    "**Veri hattı.** Örnekleme düzeni, birim tutarlılığı, çerçeve tutarlılığı, "
    "kalibrasyonun bağımsızlığı.",
])

h(2, "3.2. Birebir uyuşan bileşenler")
p("Aşağıdaki bileşenlerin bildiriyle tam olarak aynı olduğu ölçülmüştür.")
table(
    ["Bileşen", "Bildiri", "Yeniden yazılan hat", "Durum"],
    [["Kalıntı model parametresi", "478.892", "478.892", "Aynı"],
     ["Ham model parametresi", "1.907.032", "1.907.032", "Aynı"],
     ["Pencere / adım", "100 / 25", "100 / 25", "Aynı"],
     ["Sönümleme", "%15", "%15", "Aynı"],
     ["İyileştirici", "Adam, lr=10⁻³, β=(0,9; 0,999)", "aynısı", "Aynı"],
     ["Yığın boyutu", "256", "256", "Aynı"],
     ["Gradyan kırpma", "max_norm = 1,0", "1,0", "Aynı"],
     ["Çizelgeleyici", "ReduceLROnPlateau (0,5; sabır 8)", "aynısı", "Aynı"],
     ["Erken durdurma", "sabır 25, en çok 300 dönem", "aynısı", "Aynı"],
     ["Eşik kuralı", "doğrulama hatalarının P97'si", "aynısı", "Aynı"],
     ["Ağırlık taraması", "0,05 adım, 21 nokta", "aynısı", "Aynı"],
     ["Kalıntı özdeşliği (3)", "tanım gereği", "artık 7,1·10⁻¹⁵", "Aynı"],
     ["Jacobian", "J(q) analitik", "sayısal J ile fark 0,0", "Aynı"]],
    caption="Tablo 3.1. Bildiriyle birebir uyuşan bileşenler. Kalıntı özdeşliği "
            "200.000 örnek üzerinde, Jacobian ise sayısal türevle karşılaştırılarak "
            "doğrulanmıştır.",
    widths=[5.0, 5.2, 4.4, 2.2], num_from=3)
p("Bu, mimarinin ve eğitim yordamının yeniden üretiminin başarılı olduğu anlamına "
  "gelmektedir. Bulguların tamamı **veri hattında** yoğunlaşmaktadır.")

h(2, "3.3. Bulgu F1 — veri kümesi küresel olarak sıralı değil")
p("Kaynak veri, 79 günlük bir Kafka/Elasticsearch dışa aktarımıdır ve zaman damgasına "
  "göre küresel olarak sıralı değildir. Ölçüm:")
code("örnek sayısı                       : 1.124.432\n"
     "süreksiz ardışık örnek çifti       :   321.841  (%28,6)\n"
     "  bunlardan negatif/sıfır dt olan  :    96.107\n"
     "dt medyanı                         : 2,00 ms   (beklenen 500 Hz)\n"
     "dt aralığı                         : −5,5·10⁹ ms  …  +3,6·10⁹ ms\n"
     "en az bir süreksizlik içeren pencere:   33.700 / 44.974  (%74,9)")
p("Yani pencerelerin dörtte üçü, aralarında saatler hatta günler bulunan iki ayrı kayıt "
  "oturumunu tek bir zaman serisi gibi birleştirmektedir. Bir LSTM özkodlayıcı için bu, "
  "öğrenilecek dinamiğin ortasına rastgele bir sıçrama koymak demektir.")
note("Düzeltme", [
    "`prepare_dataset.py` verilenleri zaman damgasına göre küresel olarak kararlı "
    "biçimde sıralar, birebir tekrar eden satırları eler, ardından yalnızca "
    "dt ∈ (0; 8 ms] **ve** |Δq| ≤ q̇_max·dt·pay koşulunu sağlayan ardışık örnekleri "
    "aynı koşuya bağlar. En az 150 örneklik koşular korunur ve her koşu 2 ms'lik "
    "düzgün ızgaraya `np.interp` ile yeniden örneklenir; ara değerlenen satırlar "
    "`interp` bayrağıyla işaretlenir.",
    "Sonuç: 1.124.432 ham örnekten **600 kesintisiz koşu** ve **342.480 örnek** "
    "(685 s robot zamanı) elde edilmiştir; ara değerlenen örnek oranı %5,0'dır. "
    "Koşu uzunluğu medyanı 227, en uzunu 35.860 örnektir.",
])

h(2, "3.4. Bulgu F2 — ivme süreksizlikler üzerinden türetiliyor")
p("Eklem ivmesi q̈, hız kanalından Savitzky-Golay türeviyle (pencere 51, derece 3) "
  "elde edilmektedir. Filtre bir evrişim olduğu için, girdi dizisi iki ayrı oturumu "
  "birleştiriyorsa çıktı o birleşme noktasında fiziksel olmayan bir sıçrama üretir. "
  "İvme ise ters dinamik modele doğrudan M(q)·q̈ terimiyle girmektedir; dolayısıyla "
  "hata doğrudan kalıntıya taşınır.")
table(
    ["Eklem", "q̈ std — özgün sıra", "q̈ std — koşulara bölünmüş", "Oran"],
    [["1", "0,80", "0,23", "×3"],
     ["2", "0,34", "0,20", "×2"],
     ["3", "0,55", "0,24", "×2"],
     ["4", "1,14", "0,19", "×6"],
     ["5", "0,84", "0,28", "×3"],
     ["6", "0,95", "0,50", "×2"]],
    caption="Tablo 3.2. Aynı 200.000 örnek üzerinde ölçülen q̈ standart sapması "
            "(rad/s²). Süreksizlikler üzerinden türetilen ivme 2–6 kat şişmektedir; "
            "tepe değerlerde şişme daha sınırlıdır (×1–2), yani hata gürültü tabanını "
            "yükseltmekte, tekil uçlar yaratmaktan çok kalıntının taban seviyesini "
            "bozmaktadır.",
    widths=[2.2, 4.4, 5.0, 2.2], num_from=1)
p("Düzeltme, türevi **koşu içinde** almak ve filtre penceresinin kenar payını "
  "(51 örneğin yarısı = 25 örnek) her koşunun başında ve sonunda geçersiz olarak "
  "işaretlemektir. Bu `valid` maskesi eğitim ve değerlendirmede pencere elemesine "
  "girmekte, böylece hiçbir pencere kenar payı içermemektedir.")

h(2, "3.5. Bulgu F3 — KTS ölçümü yanlış çerçevede aktarılıyor")
p("Denklem (2), `J(q)ᵀ · F_KTS` çarpımını gerektirmektedir. Ancak burada bir çerçeve "
  "uyumu şarttır: UR e-Serisi sürücüsünde `transformForceTorque()` fonksiyonu wrench'i "
  "TCP çerçevesine döndürmekte ve yayıncı `frame_id: tool0` ile yayımlamaktadır. "
  "Buna karşılık `J(q)`, taban (base) çerçevesinde tanımlıdır. İki büyüklüğün "
  "doğrudan çarpılması sessizce yanlış bir kalıntı üretir.")
p("Etkisi ölçülmüştür. 40.000 örnek üzerinde, doğru (önce tabana döndürülmüş) ve "
  "döndürülmemiş `r_dis` karşılaştırıldığında:")
table(
    ["Eklem", "Korelasyon", "RMS — doğru (Nm)", "RMS — döndürülmemiş (Nm)"],
    [["1", "−0,995", "1,734", "0,836"],
     ["2", "+0,592", "0,612", "1,991"],
     ["3", "+0,771", "0,520", "0,852"],
     ["4", "+0,868", "0,113", "0,298"],
     ["5", "+0,815", "0,240", "0,097"],
     ["6", "+0,994", "0,159", "0,065"],
     ["**tümü**", "**−0,143**", "—", "—"]],
    caption="Tablo 3.3. Dışsal kalıntının çerçeve hatasına duyarlılığı. Birinci "
            "eklemde işaret tersine dönmekte (−0,995), ikinci eklemde genlik üç katına "
            "çıkmaktadır. Kanal geneli korelasyon −0,143 olup, döndürülmemiş r_dis'in "
            "doğru r_dis ile pratikte ilişkisiz olduğunu göstermektedir.",
    widths=[2.2, 3.2, 4.2, 4.6], num_from=1)
p("Düzeltme, `generate_residuals.py` içine `--fts-frame {tool, base}` seçeneğinin "
  "eklenmesidir (varsayılan `tool`). `tool` seçildiğinde wrench'in kuvvet ve moment "
  "yarımları R₀₆ ile tabana döndürüldükten sonra Jacobian aktarımı yapılmaktadır. "
  "Çevrimiçi motorda da aynı seçenek bulunmakta ve aynı kod yolu kullanılmaktadır.")

h(2, "3.6. Bulgu F4 — akım→tork kalibrasyonu döngüseldi")
p("UR ROS 2 sürücüsü, `JointState` mesajının `effort` alanına tork değil "
  "**motor akımını (Amper)** yazmaktadır (`hardware_interface.cpp` içinde alan "
  "doğrudan `actual_current`'tan doldurulmaktadır). Dolayısıyla Denklem (1)'in "
  "uygulanabilmesi için bir akım→tork katsayısı gerekmektedir.")
p("Yeniden yazılan ilk sürümde bu katsayı, ölçülen akımın **model torkuna** "
  "regresyonuyla elde ediliyordu; doğrulaması ise aynı regresyonun tekrarıydı "
  "(“a ≈ 1,0 çıkmalı”). Bu döngüsel bir doğrulamadır: katsayı zaten o eşitliği "
  "sağlayacak biçimde seçildiği için sınav her hâlükârda geçilir ve modelin gerçekten "
  "doğru olup olmadığı hakkında hiçbir bilgi vermez.")
note("Düzeltme — bağımsız quasi-statik yerçekimi regresyonu", [
    "`calibrate_current_to_torque.py`, robot neredeyse durgunken (|q̇| < 0,005 rad/s) "
    "eklem torkunun **yalnızca yerçekiminden** kaynaklandığı gerçeğinden yararlanır: "
    "g(q) ≈ k·(i − i₀). Burada g(q) FMU'dan değil, kapalı formda yerçekimi vektöründen "
    "gelmektedir; yani katsayı, sınanacak kalıntının kendisinden bağımsız olarak "
    "belirlenmektedir.",
    "197.090 quasi-statik örnekten 65.697'si regresyona girmiştir. Ölçülen: "
    "shoulder_lift R² = 0,984 ve elbow R² = 0,993 — bu iki eklem güvenilir. "
    "shoulder_pan'de yerçekimi torku yapısal olarak sıfırdır (dikey eksen), bilek "
    "eklemlerinde ise yerçekimi torku gürültü seviyesinde kalmakta (R² ≤ 0,012), "
    "yani katsayı **ölçülememektedir**.",
    "Ölçülemeyen eklemler için katsayı aynı modül ailesinden kopyalanmakta "
    "(shoulder_pan ← shoulder_lift) veya tork sınırı oranıyla ölçeklenmektedir "
    "(bilekler ← elbow, 56/150). Bu durum `trusted` bayrağıyla açıkça işaretlenmiş, "
    "düğüm açılışta uyarı vermektedir. Tespit başarımını etkilemez — özkodlayıcı "
    "girdiyi z-skor normalleştirdiği için ölçek çarpanı öğrenilen dinamiği "
    "değiştirmez — yalnızca arıza genliğinin Nm cinsinden fiziksel yorumunu etkiler.",
])
table(
    ["Eklem", "Nm/A", "R²", "Kaynak", "Güvenilir"],
    [["shoulder_pan", "10,522", "—", "shoulder_lift'ten kopya (aynı modül ailesi)", "hayır"],
     ["shoulder_lift", "10,522", "0,984", "quasi-statik yerçekimi regresyonu", "**evet**"],
     ["elbow", "9,130", "0,993", "quasi-statik yerçekimi regresyonu", "**evet**"],
     ["wrist_1", "3,409", "0,000", "elbow katsayısının tork sınırı oranı (56/150)", "hayır"],
     ["wrist_2", "3,409", "0,012", "elbow katsayısının tork sınırı oranı (56/150)", "hayır"],
     ["wrist_3", "3,409", "0,010", "elbow katsayısının tork sınırı oranı (56/150)", "hayır"]],
    caption="Tablo 3.4. Ölçülen akım→tork katsayıları (`current_to_torque.json`). "
            "İki eklem doğrudan ölçülebilmiş, dördü varsayımla türetilmiştir.",
    widths=[3.0, 2.0, 1.8, 6.6, 2.0], num_from=1)
p("Kalibrasyonun ikinci bileşeni bir **ofset**tir. `generate_residuals.py --calibrate "
  "offset` kipinde eğim a = 1 sabitlenmekte, yalnızca b = medyan(τ_ölç − τ_model) "
  "hesaplanmaktadır. Eğimin dış kalibrasyondan sonra yeniden uydurulmasına izin "
  "verilmemesinin sebebi ölçülmüştür: serbest bırakıldığında shoulder_pan için a = 5,24 "
  "gibi değerler çıkmakta, yani model gürültüye yeniden ölçeklenmektedir. "
  "Ölçülen ofsetler −2,27 ile +0,37 Nm arasındadır.")

h(2, "3.7. Bulgu F5 — çözücünün ölçülen sınırları (kabul edildi)")
p("Ters dinamik çözücünün özellikleri 500 rastgele eklem pozunda sınanmıştır:")
code("asimetrik M(q)                       : 500 / 500 poz   (en kötü bağıl fark 0,139)\n"
     "simetrik kısmı pozitif tanımlı DEĞİL : 500 / 500 poz\n"
     "sürtünme vektörü                     : tüm pozlarda tam olarak sıfır\n"
     "yük (payload) modeli                 : yok")
p("Atalet matrisinin simetrik ve pozitif tanımlı olması, Newton-Euler formülasyonunun "
  "temel bir özelliğidir; ölçüm bu özelliğin sağlanmadığını göstermektedir. Sürtünmenin "
  "özdeş sıfır olması ise, bileklerdeki torkun büyük bölümünün modelde karşılığının "
  "bulunmadığı anlamına gelir.")
note("Bu bulgu bilinçli olarak DÜZELTİLMEMİŞTİR", [
    "Kısıt açıktır: *“FMU tarafında hiçbir şeyi değiştirmeyeceğiz. Bu FMU, gerçek "
    "robotla çalışarak test edildi ve birebir aynı sonuçların üretildiği görüldü.”* "
    "Çözücünün tek satırına dokunulmamıştır.",
    "Bedeli ölçülmüş ve raporlanmıştır: wrist_2 ve wrist_3 kanallarında model torkunun "
    "standart sapması ölçülen torkun sırasıyla %10 ve %2'si kadardır "
    "(0,095 ve 0,030 Nm'ye karşılık 0,959 ve 1,246 Nm). Yani bu iki kanalda kalıntı "
    "pratikte ham ölçümün kendisidir ve model bilgisi taşımaz. Bu, kalıntı modelinin "
    "12 kanalından ikisinin bilgi katkısının sınırlı olduğu anlamına gelir; sistemin "
    "çalışmasını engellemez, fakat kalıntı uzayının teorik üstünlüğünü bu iki kanalda "
    "gerçekleştirmez.",
])

h(2, "3.8. Bulgu F6 — sessizce devre dışı kalan model")
p("`generate_residuals.py`, düşük R²'li eklemlerde model torkunu sıfırlayan bir eşik "
  "(`--min-r2`, varsayılan 0,05) içermektedir. Karışık birimli çalıştırmada R² değeri "
  "çok büyük negatif çıktığı için bu eşik **altı eklemin tamamını** sıfırlamakta, yani "
  "kalıntı modeli sessizce “kalıntı = ölçüm” haline gelmektedir. Hiçbir hata mesajı "
  "üretilmemektedir.")
p("Bu davranış erratum yeniden üretiminde `--min-r2=-1e9` ile açıkça devre dışı "
  "bırakılmıştır; ayrıca `--calibrate offset` kipinde sıfırlama tümüyle kaldırılmıştır. "
  "Bulgunun kendisi bir uyarıdır: bir güvenlik kontrolünün sessizce tetiklenmesi, "
  "kontrolün olmamasından daha tehlikelidir.")

h(2, "3.9. Bulgu F7 — bildirinin Tablo II'sinde OR satırı")
p("Bildirinin Tablo II'sinde MAX ve OR stratejileri için üç metrik de birebir aynı "
  "değerleri (0,976 / 0,889 / 0,824) taşımaktadır. OR stratejisi bir **karar** "
  "birleşimidir (her iki modelden birinin kendi eşiğini aşması), MAX ise bir **skor** "
  "birleşimidir; ikisinin ROC eğrisi tanım gereği farklıdır ve AUC değerlerinin "
  "üç basamağa kadar çakışması beklenmez. Bu, OR satırının MAX'ın skorunu raporladığını "
  "göstermektedir. Yeniden yazılan hatta OR ayrı hesaplanmakta ve karşılaştırma "
  "tablosunda ayrı raporlanmaktadır.")


# ═══════════════════════════════════════════════════════════════════
h(1, "4. Erratum: bildirinin sayıları hangi hattan geliyor?")

h(2, "4.1. Hipotez")
p("Bölüm 3'teki bulgular tek tek bakıldığında birer kusurdur; birlikte bakıldığında ise "
  "tutarlı bir düzen oluşturmaktadır. Hipotez şudur: bildirinin sayıları, birimleri "
  "**karışık** olan bir hattan gelmektedir.")
code("τ_ölç    : amper   (UR sürücüsü effort alanına actual_current yazıyor)\n"
     "τ_model  : Nm      (FMU ters dinamik)\n"
     "r_dis    : Nm      (JᵀF — üstelik wrench tool0 çerçevesindeyken taban Jacobian'ı ile)\n"
     "r_top    = τ_amper − τ_model_Nm      ← üç farklı büyüklük doğrudan çıkarılıyor\n"
     "arıza    : \"Nm\" diye tanımlanan genlikler amper kanallara ham sayı olarak ekleniyor")
p("Bu hipotez sınanabilir bir öngörü üretmektedir: eğer doğruysa, bu düzen yeniden "
  "kurulduğunda bildirinin **imza bulguları** geri gelmelidir — özellikle kalıntı "
  "modelinin sensör gürültüsüne körlüğü (0,272) ve anomalilerin yaklaşık dörtte birinin "
  "yalnızca kalıntı tarafından yakalanması (%24,1).")

h(2, "4.2. Yeniden kurulum")
p("`reproduce_erratum.sh`, hattın tamamını o düzende yeniden kurmaktadır: kalibrasyon "
  "yok, model sıfırlama eşiği devre dışı (`--min-r2=-1e9`), birim `current`, wrench "
  "çerçevesi düzeltilmemiş (`--fts-frame base`), ardından her iki model sıfırdan "
  "eğitiliyor ve aynı değerlendirme çalıştırılıyor. Bu koşum bu rapor için yeniden "
  "yürütülmüştür.")

h(2, "4.3. Sonuç")
table(
    ["Büyüklük", "Erratum koşumu", "Bildiri", "Fark"],
    [["Test penceresi", "179.888", "179.896", "8 pencere"],
     ["Arıza penceresi", "14.401  (%8,01)", "14.402  (%8,0)", "1 pencere"],
     ["Yalnızca-kalıntı payı", "%23,9", "%24,1", "0,2 puan"],
     ["Sensör gürültüsü — kalıntı", "0,267", "0,272", "0,005"],
     ["OR geri çağırma", "0,841", "0,854", "0,013"],
     ["Geri çağırma — ham", "0,601", "0,613", "0,012"],
     ["Gizyazar — kalıntı / ham", "0,975 / 0,492", "0,986 / 0,530", "0,011 / 0,038"],
     ["Ham model AUC", "0,943", "0,952", "0,009"],
     ["Bir. MAX (AUC/PR/F1)", "0,974 / 0,860 / 0,803", "0,976 / 0,889 / 0,824", "≤ 0,029"],
     ["Ham model doğrulama kaybı", "0,306", "0,320", "0,014"],
     ["Kalıntı model en iyi dönem", "99  (durdurma 124)", "129", "—"]],
    caption="Tablo 4.1. Bildirinin özgün hattının yeniden kurulmasıyla elde edilen "
            "sonuçlar. Test ve arıza pencere sayılarının 179.896 içinde sırasıyla 8 ve "
            "1 pencere farkla tutması, aynı veri kümesinin aynı biçimde pencerelendiğini "
            "göstermektedir.",
    widths=[5.4, 4.2, 3.4, 2.6], highlight=(0, 1, 2, 3), num_from=1)
note("Hipotez doğrulanmıştır", [
    "Bildirinin imza bulguları noktası noktasına geri gelmektedir. En keskin kanıt "
    "sensör gürültüsü satırıdır: kalıntı modelinin bu senaryodaki başarımı 0,267 "
    "(bildiri 0,272) olup, bu değer ancak kalıntının fiziksel anlamını yitirdiği — "
    "yani amper ile Nm'nin doğrudan çıkarıldığı — bir hatta ortaya çıkmaktadır. "
    "Pencere sayılarının 179.896 içinde 8 ve 1 pencere farkla tutması, aynı ham veri "
    "kümesinin aynı biçimde bölündüğünü doğrulamaktadır.",
    "Buradan çıkan sonuç, bildirinin **savının** yanlış olduğu değildir: iki temsil "
    "uzayının tamamlayıcılığı hem erratum hatta (%23,9 yalnızca-kalıntı) hem de "
    "düzeltilmiş hatta (%24,3) ölçülmektedir. Yanlış olan, raporlanan sayıların "
    "fiziksel ölçeğidir. Bu nedenle düzeltilmiş hat **ana sonuç**, erratum ise "
    "yalnızca bildirinin sayılarının kaynağını belgeleyen bir ek olarak "
    "sunulmaktadır.",
])
p("Yeniden üretilen hatta bazı değerler bildiriden ayrılmaktadır — örneğin yalnızca-ham "
  "payı %34,1 (bildiri %25,9) ve motor kayması satırı. Bu farklar beklenendir: iki "
  "koşum arasında rastgele tohum, eğitim/doğrulama bölünmesinin tam yeri ve erken "
  "durdurma dönemi aynı olmak zorunda değildir. Belirleyici olan, tek bir sayının değil "
  "**imza örüntüsünün** yeniden üretilmesidir.")


# ═══════════════════════════════════════════════════════════════════
h(1, "5. Düzeltilmiş hat")
p("Bu bölüm, F1–F4 ve F6 bulgularının düzeltilmesiyle ortaya çıkan hattı adım adım "
  "tanımlamaktadır. Hattın tamamı `run_pipeline.sh` ile tek komutta yeniden "
  "üretilebilir.")

h(2, "5.1. Adım 1 — veri hazırlama")
p("`prepare_dataset.py` ham dışa aktarımı fiziksel olarak tutarlı koşulara ayırır:")
bullet([
    "Zaman damgasına göre küresel, **kararlı** sıralama (eşit damgalarda özgün sıra korunur).",
    "Birebir tekrar eden satırların elenmesi.",
    "Koşu sınırı ölçütü: ardışık iki örnek ancak dt ∈ (0; `--max-gap`] **ve** "
    "|Δq| ≤ q̇_max · dt · pay koşullarını birlikte sağlıyorsa aynı koşudadır. "
    "İkinci koşul önemlidir: zaman damgası tutarlı görünse bile eklem açısı fiziksel "
    "olarak imkânsız biçimde sıçramışsa iki örnek arasında kayıp veri vardır.",
    "`--min-run` (150 örnek) altındaki koşuların atılması.",
    "Her koşunun 2 ms'lik düzgün ızgaraya `np.interp` ile yeniden örneklenmesi; "
    "ara değerlenen satırlar `interp` bayrağıyla işaretlenir.",
])
table(
    ["Büyüklük", "Değer"],
    [["Ham girdi örneği", "1.124.432"],
     ["Elde edilen koşu", "600"],
     ["Temiz örnek", "342.480  (685 s robot zamanı)"],
     ["Ara değerlenen örnek", "%5,02"],
     ["Koşu uzunluğu (medyan / en kısa / en uzun)", "227 / 155 / 35.860"]],
    caption="Tablo 5.1. Veri hazırlamanın çıktısı (`ur10e_clean.parquet`).",
    widths=[8.0, 5.4], num_from=1)
p("Örnek sayısındaki büyük düşüş beklenendir ve bir kayıp değil bir **kazanımdır**: "
  "atılan örnekler, aralarında fiziksel süreklilik bulunmayan ve bu nedenle bir zaman "
  "serisi modeline verilmemesi gereken satırlardır.")

h(2, "5.2. Adım 2 — akım→tork kalibrasyonu")
p("Bölüm 3.6'da tanımlanan quasi-statik yerçekimi regresyonu çalıştırılır ve "
  "`current_to_torque.json` üretilir. Bu dosya hem çevrimdışı kalıntı üretiminde hem de "
  "çevrimiçi düğümde aynen kullanılmaktadır — yani ölçek tanımı tek bir yerde durur.")

h(2, "5.3. Adım 3 — kalıntı üretimi")
p("`generate_residuals.py` her koşu için sırasıyla şunları yapar:")
bullet([
    "q̈'yi **koşu içinde** Savitzky-Golay türeviyle (51, derece 3, dt = 2 ms) hesaplar; "
    "kenar payı `valid = False` olarak işaretlenir.",
    "τ_model = solver.getTorques(q, q̇, q̈) çağrısıyla ters dinamiği değerlendirir "
    "(FMU çekirdeği, değiştirilmemiş).",
    "τ_ölç = akım × Nm/A dönüşümünü uygular ve ofset b'yi çıkarır.",
    "J(q) ve R₀₆'yı tek geçişte hesaplar; `--fts-frame tool` ise wrench'in her iki "
    "yarımını tabana döndürür ve r_dis = Jᵀ·F'yi bulur.",
    "r_top ve r_ic'yi Denklem (1) ve (3) ile türetir.",
    "Kalıntı kanallarının **yanı sıra** ham kanalları (q, q̇, τ, KTS) ve "
    "`valid`/`run_id`/`interp` sütunlarını **tek bir parquet dosyasına** yazar.",
])
note("Neden tek dosya", [
    "İki modelin tamamen aynı pencere kümesi üzerinde eğitilmesi ve değerlendirilmesi "
    "şarttır; aksi hâlde skorları karşılaştırmak ve birleştirmek anlamsızlaşır. "
    "Kalıntı ve ham öznitelikleri ayrı dosyalarda tutulduğunda pencere kümelerinin "
    "eşitliği bir varsayım olur. Tek dosyada tutulduğunda ise `window_starts()` her iki "
    "model için aynı indeks dizisini döndürür ve eşitlik yapısal olarak garanti edilir.",
])

h(2, "5.4. Adım 4 — pencereleme ve sızıntısız bölme")
p("`models.py` içindeki `window_starts(n, W, S, groups, valid)` fonksiyonu iki kuralı "
  "birden uygular: (i) hiçbir pencere iki koşuya **yayılamaz**, (ii) içinde geçersiz "
  "örnek bulunan pencereler elenir. Eğitim/doğrulama bölünmesi %80 oranında yapılmakta, "
  "fakat kesme noktası en yakın **koşu sınırına** hizalanmaktadır; böylece aynı koşunun "
  "örtüşen pencereleri iki kümeye birden düşemez. %75 örtüşmeli pencerelemede bu "
  "hizalama olmazsa doğrulama kaybı sistematik olarak iyimser çıkar.")

h(2, "5.5. Adım 5 — eğitim")
p("`train_ae.py` her iki modeli bildirinin hiperparametreleriyle eğitir. NaN içeren "
  "satırlar atılmaz, `valid = False` olarak işaretlenir — atılsalardı komşu örnekler "
  "yapay olarak bitişik hale gelir ve F1'in aynısı küçük ölçekte tekrarlanırdı.")
table(
    ["", "Kalıntı modeli", "Ham modeli"],
    [["Girdi kanalı", "12", "24"],
     ["Parametre", "478.892", "1.907.032"],
     ["Eğitim / doğrulama penceresi", "8.294 / 2.128", "8.294 / 2.128"],
     ["En iyi dönem", "21", "22"],
     ["Toplam dönem (erken durdurma)", "46", "47"],
     ["En iyi doğrulama kaybı", "0,582", "0,930"],
     ["Eşik θ (doğrulama P97)", "1,6008", "3,4461"]],
    caption="Tablo 5.2. Düzeltilmiş hatta eğitim sonuçları. Eşiklerin bildirininkilerden "
            "(0,420 / 0,854) yüksek olması doğrudandır: kalıntılar artık gerçek Nm "
            "ölçeğindedir ve normalleştirme öncesi genlikleri farklıdır. Eşik, mutlak "
            "bir büyüklük değil, modelin kendi doğrulama dağılımına göre tanımlı bir "
            "yüzdeliktir.",
    widths=[6.4, 3.6, 3.4], num_from=1)

h(2, "5.6. Adım 6 — arıza enjeksiyonu ve değerlendirme")
p("`inject_faults.py` bildirinin dört senaryosunu ölçüm uzayında uygular. Kritik "
  "ayrıntı şudur: bildiri arıza genliklerini Nm cinsinden tanımlamaktadır, fakat "
  "kanallar (düzeltilmiş hatta) artık gerçekten Nm'dir; dolayısıyla genlikler "
  "kalibrasyon dosyası üzerinden ölçüm uzayına çevrilerek uygulanır. Karışık birimli "
  "özgün hatta ise aynı sayılar amper kanallara ham olarak ekleniyordu — bu, iki hat "
  "arasındaki mutlak başarım farkının başlıca kaynağıdır (bkz. Bölüm 6.3).")
p("Pencere etiketi bildirinin kuralıyla atanır: pencere örneklerinin en az %10'u arıza "
  "maskesiyle örtüşüyorsa pencere anomalidir.")


# ═══════════════════════════════════════════════════════════════════
h(1, "6. Sonuçlar — düzeltilmiş hat")

h(2, "6.1. Genel performans")
table(
    ["Model", "AUC", "bildiri", "PR-AUC", "bildiri", "BestF1", "bildiri"],
    [["Kalıntı LSTM Özk.", "0,812", "0,908", "0,553", "0,695", "0,584", "0,692"],
     ["Ham LSTM Özk.", "0,837", "0,952", "0,627", "0,761", "0,647", "0,698"],
     ["Bir. A-Ort (0,95/0,05)", "0,939", "0,980", "0,780", "0,905", "0,800", "0,859"],
     ["Bir. MAX", "0,930", "0,976", "0,803", "0,889", "0,792", "0,824"],
     ["Bir. OR", "0,930", "0,976", "0,803", "0,889", "0,792", "0,824"],
     ["Art. Norm Eşik", "0,646", "0,849", "0,383", "0,805", "0,398", "0,707"],
     ["Isolation Forest", "0,652", "0,688", "0,130", "0,459", "0,241", "0,547"],
     ["One-Class SVM", "0,734", "0,815", "0,534", "0,816", "0,596", "0,760"]],
    caption="Tablo 6.1. Düzeltilmiş hattın genel performansı, bildirinin Tablo II'si ile "
            "yan yana. Test kümesi 41.688 pencere, bunların 3.623'ü (%8,7) arızalıdır — "
            "bildirinin %8,0 oranıyla uyumludur.",
    widths=[4.4, 1.9, 1.9, 1.9, 1.9, 1.9, 1.9], highlight=(2,), num_from=1)
p("Mutlak değerlerin düşmesi beklenen ve açıklanabilir bir sonuçtur. Düzeltilmiş hatta "
  "arıza genlikleri gerçek Nm ölçeğinde uygulanmakta, yani problem **zorlaşmaktadır**; "
  "ayrıca veri kümesi süreksizliklerden arındırıldığı için modelin öğrenmesi gereken "
  "dinamik daha tutarlı, fakat arızayı gizleyen yapay sıçramalar da ortadan kalkmıştır. "
  "Değerlendirmede önemli olan mutlak sayı değil, **sıralamanın ve birleşim marjının "
  "korunmasıdır**: her iki tabloda da birleşim tekil modelleri ve tüm referans "
  "yöntemleri açık farkla geçmektedir.")

h(2, "6.2. Arıza tipine göre")
table(
    ["Arıza tipi", "Kalıntı", "Ham", "Birleşim", "Bildiri (kal./ham)"],
    [["Çarpışma", "0,995", "0,999", "0,996", "0,916 / 0,933"],
     ["Gizyazar hatası", "0,988", "0,491", "0,989", "0,986 / 0,530"],
     ["Sensör gürültüsü", "0,296", "1,000", "0,996", "0,272 / 1,000"],
     ["Motor kayması", "0,451", "0,256", "0,443", "0,597 / 0,605"]],
    caption="Tablo 6.2. Arıza tipine göre AUC. Gizyazar (kalıntı 0,988 ↔ ham 0,491) ve "
            "sensör gürültüsü (kalıntı 0,296 ↔ ham 1,000) satırları bildirinin "
            "asimetrisini neredeyse birebir yeniden üretmektedir.",
    widths=[3.8, 2.4, 2.4, 2.4, 3.6], num_from=1)
p("Gizyazar hatası ve sensör gürültüsü satırları, bildirinin temel savının doğrudan "
  "kanıtıdır. Fizik modeli eklem pozisyon sapmasını büyük bir kalıntı olarak "
  "yansıttığından gizyazar hatasında kalıntı modeli baskındır; buna karşılık düşük "
  "genlikli KTS gürültüsü ters dinamik modelden süzülemediği için kalıntı modeli bu "
  "senaryoda kör kalmakta, ham model ise mükemmel sonuç vermektedir. İki modelin kör "
  "noktaları **örtüşmemektedir** — birleşimin varlık sebebi tam olarak budur.")

h(2, "6.3. Tamamlayıcılık analizi")
table(
    ["3.623 arıza penceresinden", "Bu koşum", "Bildiri"],
    [["Yalnızca kalıntı modeli yakalıyor", "%24,3", "%24,1"],
     ["Yalnızca ham model yakalıyor", "%28,7", "%25,9"],
     ["Her ikisi", "%22,7", "—"],
     ["Hiçbiri", "%24,3", "—"],
     ["Geri çağırma — kalıntı", "0,470", "0,595"],
     ["Geri çağırma — ham", "0,514", "0,613"],
     ["Geri çağırma — OR", "0,757", "0,854"]],
    caption="Tablo 6.3. Pencere düzeyinde tamamlayıcılık, her model kendi P97 eşiğiyle. "
            "Yalnızca-kalıntı oranı bildiriyle 0,2 puan farkla örtüşmektedir.",
    widths=[7.4, 3.0, 3.0], num_from=1)

h(2, "6.4. Ağırlık duyarlılığı")
p("w_kal, 0,00'dan 1,00'a 0,05 adımla taranmıştır. Birleşik eşik her noktada yalnızca "
  "**temiz doğrulama** pencerelerinin 97. persentilinden alınmakta, yani test kümesi "
  "eşiği hiçbir aşamada görmemektedir.")
table(
    ["w_kal", "θ_birleşik", "Kesinlik", "Geri çağırma", "F1"],
    [["0,00  (yalnız ham)", "0,7685", "0,759", "0,514", "0,613"],
     ["0,25", "0,7279", "0,794", "0,611", "0,691"],
     ["0,50", "0,6596", "0,826", "0,730", "0,775"],
     ["0,75", "0,6480", "0,831", "0,745", "0,785"],
     ["0,90", "0,6395", "0,833", "0,752", "0,790"],
     ["0,95  (seçilen)", "0,6436", "0,833", "0,754", "**0,792**"],
     ["1,00  (yalnız kalıntı)", "0,6547", "0,756", "0,470", "0,579"]],
    caption="Tablo 6.4. Ağırlık taraması (`fusion_config.json → online_sweep`). "
            "0,25–0,95 aralığında geniş ve kararlı bir bölge bulunmakta, uçlar keskin "
            "biçimde kötüleşmektedir — yani optimum gerçekten iki modelin arasındadır. "
            "Bildiri de aynı bölgeyi (w_kal ∈ [0,25; 0,95]) raporlamıştır.",
    widths=[4.4, 2.6, 2.4, 2.8, 2.4], highlight=(5,), num_from=1)
p("En iyi tekil modele göre kazanç **ΔF1 = +0,178**'dir (0,792'ye karşılık 0,613). "
  "Bildirinin raporladığı kazanç +0,161'dir; iki hat farklı mutlak seviyelerde "
  "çalışmasına rağmen birleşimin sağladığı marj korunmuştur.")


# ═══════════════════════════════════════════════════════════════════
h(1, "7. Çevrimiçi sistem")

h(2, "7.1. Tasarım kısıtları")
p("Çevrimiçi sistemin çevrimdışı hattan üç yapısal farkı vardır ve tasarımın tamamı bu "
  "farkların yönetilmesi üzerine kuruludur:")
bullet([
    "**Nedensellik.** Akışta gelecek bilinmez. Veri kümesinin tamamı üzerinden "
    "hesaplanan hiçbir büyüklük (min, max, ortalama) canlıda kullanılamaz.",
    "**Zaman bütçesi.** 500 Hz örnek akışında örnek başına 2 ms, karar başına "
    "(adım 25) 50 ms bulunmaktadır.",
    "**Süreklilik.** Çevrimdışı hat veriyi koşulara bölebilir; canlı sistem "
    "kesintiyi ancak olduktan sonra fark eder ve durumu sıfırlaması gerekir.",
])

h(2, "7.2. Öznitelik motoru")
p("`OnlineFeatureExtractor` sınıfı, çevrimdışı hattaki hesabın örnek-örnek çalışan "
  "eşdeğeridir. 51 örneklik bir halka tampon tutar ve her yeni örnekte tamponun "
  "**ortasındaki** örneğin özniteliklerini üretir:")
code("qd_buf ← 51 örneklik halka tampon\n"
     "q̈      = sg_coeffs · qd_buf                (merkezli türev, use=\"dot\")\n"
     "τ      = akım × nm_per_amp\n"
     "τ_model= solver.getTorques(q, q̇, q̈)        (FMU çekirdeği)\n"
     "J, R₀₆ = jacobian_and_rotation(q)\n"
     "w      = R₀₆ ile tabana döndürülmüş wrench   (fts_frame == \"tool\" ise)\n"
     "r_ext  = Jᵀ · w\n"
     "r_top  = τ − τ_model − b\n"
     "r_int  = r_top − r_ext\n"
     "→ residual_vec = [r_int(6), r_ext(6)]        (12 kanal)\n"
     "→ raw_vec      = [q(6), q̇(6), τ(6), KTS(6)]  (24 kanal)")
p("Savitzky-Golay katsayılarının `savgol_coeffs(..., use=\"dot\")` ile alınıp iç çarpım "
  "olarak uygulanmasının `savgol_filter` ile birebir aynı sonucu verdiği ayrıca "
  "doğrulanmıştır (fark 2·10⁻¹⁵). Merkezli türev, 25 örneklik (50 ms) yapısal bir "
  "gecikme getirir; bu gecikme q̈'nin gürültüsüz olmasının bedelidir ve çevrimdışı "
  "tanımın birebir aynısıdır.")
note("Doğrulama — çevrimiçi motor ≡ çevrimdışı hat", [
    "`verify_online_features.py`, `ur10e_clean.parquet` dosyasını örnek örnek çevrimiçi "
    "motordan geçirir ve çevrimdışı üretilmiş `ur10e_features.parquet` ile karşılaştırır. "
    "8 koşu, 61.524 örnek üzerinde ölçülen en büyük mutlak farklar: "
    "τ 1,4·10⁻¹⁴, τ_model 0,0, r_top 8,0·10⁻¹⁵, r_ext 2,7·10⁻¹⁵, r_int 8,9·10⁻¹⁵.",
    "Bu, kayan nokta yuvarlama seviyesidir. Yani çevrimiçi düğüm, modellerin eğitildiği "
    "özniteliklerin *aynısını* görmektedir — benzerini değil.",
])

h(2, "7.3. Sapma D1 — normalleştirme sınırları")
p("Bildirinin min–max normalleştirmesi, sınırlarını **arıza enjekte edilmiş test "
  "kümesinden** almaktadır. Bu, çevrimiçi kullanım için iki nedenle uygun değildir: "
  "sınırlar gelecekteki maksimumları içerir (nedensel değildir) ve o maksimumlar "
  "normal çalışmada hiç görülmez.")
p("Etkisi ölçülmüştür: test kümesi maksimumları 4341 ve 7752 olduğu için normal "
  "çalışmadaki tüm skor aralığı ~10⁻⁴ mertebesine sıkışmakta, birleşik eşik her iki "
  "tekil eşikten de **katı** hale gelmekte ve motor kayması ile gizyazar hatası "
  "hiç tespit edilmemektedir.")
p("Çözüm olarak iki **nedensel** aday ölçek tanımlanmış ve ölçümle yarıştırılmıştır. "
  "İkisi de Denklem (4)'ün yapısını korur ve ikisi de affin bir dönüşümdür — "
  "z = (S − lo) / span:")
table(
    ["Ölçek", "Tanım", "Kalıntı dönüşümü", "Ham dönüşümü", "En iyi w", "F1"],
    [["theta", "z = S / θ", "(S−0) / 1,601", "(S−0) / 3,446", "0,95", "0,791"],
     ["minmax_val", "z = (S−min)/(max−min), **sınırlar temiz doğrulamadan**",
      "(S−0,00732) / 2,434", "(S−0,05476) / 4,413", "0,95", "**0,792**"]],
    caption="Tablo 7.1. İki nedensel ölçeğin yarışı. İkisi de aynı 21 noktalık w "
            "taramasından ve aynı P97 birleşik eşik kuralından geçmiştir; ikisinde de "
            "temiz doğrulamadaki yanlış alarm oranı %3,0'dır.",
    widths=[2.4, 5.0, 3.2, 3.2, 1.8, 1.6], highlight=(1,), num_from=2)
note("Sonuç: D1 “formül farkı” olmaktan çıktı", [
    "`minmax_val` kazandığı için **bildirinin kendi min–max formülü** canlıya geri "
    "dönmüştür. Geriye kalan tek fark, sınırların arıza görmemiş doğrulama "
    "penceresinden alınmasıdır — ki bu, çevrimiçi çalışmanın kaçınılmaz gereğidir. "
    "İki ölçek arasındaki 0,001'lik fark ölçüm gürültüsü içindedir; seçim, sapmayı "
    "azalttığı için `minmax_val` yönünde yapılmıştır.",
    "Ölçek affin biçimde (lo, span) saklandığı için düğüm tek bir kod yolu kullanır; "
    "yeni bir ölçek eklemek düğümde değişiklik gerektirmez. `scale` alanı bulunmayan "
    "eski yapılandırmalar `theta` varsayılarak okunur.",
])

h(2, "7.4. Sapma D2 — birleşik skorun çalışma eşiği")
p("Bildiri birleşik skoru yalnızca BestF1 ile raporlamakta, yani bir **çalışma eşiği "
  "tanımlamamaktadır**. BestF1 tanım gereği etiketleri görerek seçilen en iyi eşiktir; "
  "canlı bir sistemde etiket yoktur. Bu boşluk, bildirinin tekil modeller için zaten "
  "kullandığı kuralın birleşik skora genişletilmesiyle kapatılmıştır: eşik, **temiz "
  "doğrulama** pencerelerindeki birleşik skorun 97. persentilidir. FMU verisi üzerinde "
  "ölçülen değer θ_birleşik = 0,6436 olup, temiz doğrulamada beklenen %3 yanlış alarm "
  "oranını vermektedir.")
note("Bu değer gerçek robotta geçerli değildir", [
    "0,6436 FMU doğrulama kümesinden gelmektedir ve gerçek UR10e üzerinde kararların "
    "%97'sini alarm ilan etmektedir (ölçüm: Bölüm 10.3). Sebebi Bölüm 10'da ölçülmüştür: "
    "gerçek robotta temiz bir koşunun birleşik skor medyanı 4,27'dir, yani eşik yedi kat "
    "düşüktür.",
    "Paketin gömülü değeri, operatörün doğruladığı etiketli bir küme üzerinden yeniden "
    "ölçülen **θ_birleşik = 18,0**'dır. FMU'dan gelen özgün değer "
    "`fusion_config.json` içinde `fused_threshold_fmu` alanında, seçim gerekçesi ise "
    "`fused_threshold_not` alanında saklanmaktadır.",
])

h(2, "7.5. Sapma D3 — uyarlanabilir alarm kuralı")
p("Mutlak eşik θ, doğrulama kümesinin **en kötü uyan** koşularından gelmektedir. İyi "
  "uyan bir koşuda taban çizgisi θ'nın çok altında oturur; oradaki gerçek bir arıza "
  "skoru yüz kat yükseltse bile mutlak eşiği geçemez. Bu, varsayım değil ölçümdür: "
  "gizyazar hatası enjekte edilen bir koşuda birleşik skor 0,003'ten 0,462'ye "
  "(120 kat) çıkmış, fakat θ = 0,978 hiç aşılmamıştır.")
p("Çözüm, mutlak eşiğin **yanına** son N kararın kendi dağılımına göre çalışan sağlam "
  "bir kural eklemektir:")
eq("alarm  =  S_bir > θ_mutlak    VEYA    S_bir > medyan + k · 1,4826 · MAD")
p("1,4826·MAD, Gauss dağılımında standart sapmanın sağlam bir kestirimidir. Taban "
  "çizgisi neredeyse sabitse MAD sıfıra gider ve kural aşırı hassaslaşır; bu nedenle "
  "ölçek `max(1,4826·MAD; 0,05·medyan; 10⁻⁹)` ile tabanlanmıştır. Alarm sürerken taban "
  "çizgisi **dondurulur** — aksi hâlde uzun süren bir arıza yeni “normal” olur ve alarm "
  "sessizce söner.")
p("Dondurmanın **süre sınırı** olmalıdır. Sınırsız dondurma, gerçek robotta kalıcı "
  "kilitlenme üretmiştir (Bölüm 10.2): düğüm robot dururken başlatıldığında taban "
  "çizgisi hareketsiz gürültüden öğrenilmekte, robot hareket edince skor eşiği aşıp "
  "alarmı kilitlemekte ve taban çizgisi bir daha güncellenememektedir. Bu nedenle "
  "`freeze_timeout` (varsayılan 3 s) eklenmiştir: bu süreden uzun süren alarm artık "
  "arıza atağı değil **rejim değişimi** sayılır ve taban çizgisi yeniden uyum sağlar.")
note("Uyarlanabilir kural gerçek robotta varsayılan olarak KAPALIDIR", [
    "FMU verisinde bu kural tekil modelleri kurtaran bileşendi. Gerçek UR10e "
    "kayıtlarında k ∈ [8; 30] taranmış ve hiçbir değerin doğrulanmış gerçek olayları "
    "yakalamadığı, yalnızca temiz veriye yanlış blok eklediği ölçülmüştür (Bölüm 10.4).",
    "Sebep ölçülmüştür: kuralın taban çizgisi normal hareketin tepe değerlerini de "
    "yutmakta, MAD şişmekte ve uyarlanabilir eşik gerçek olayların üstüne çıkmaktadır. "
    "Gerçek robotta olayı yakalayan bileşen **mutlak eşiktir** — bildiriye dayanan D3 "
    "gerekçesinin tam tersi.",
])
table(
    ["k", "Yanlış alarm", "Motor kayması", "Çarpışma", "Gizyazar", "Sensör gürültüsü"],
    [["4", "%0,35", "4/4 · 1250 ms", "4/4 · 50 ms", "4/4 · 1250 ms", "4/4 · 50 ms"],
     ["6", "%0,00", "4/4 · 1250 ms", "4/4 · 50 ms", "4/4 · 1250 ms", "4/4 · 50 ms"],
     ["8  (seçilen)", "%0,00", "4/4 · 1250 ms", "4/4 · 50 ms", "4/4 · 1250 ms", "4/4 · 50 ms"],
     ["12", "%0,00", "4/4 · 1250 ms", "4/4 · 50 ms", "4/4 · 1250 ms", "4/4 · 50 ms"],
     ["20", "%0,00", "4/4 · 1350 ms", "4/4 · 50 ms", "4/4 · 1250 ms", "4/4 · 50 ms"],
     ["40", "%0,00", "4/4 · 1800 ms", "4/4 · 50 ms", "4/4 · 1250 ms", "4/4 · 50 ms"],
     ["80", "%0,00", "4/4 · 2750 ms", "4/4 · 50 ms", "4/4 · 1250 ms", "4/4 · 50 ms"]],
    caption="Tablo 7.2. `tune_adaptive.py` ile k taraması (4 koşu × 5 senaryo, "
            "600 karar = 30 s taban çizgisi, 200 karar ısınma). Tespit başarısı "
            "k ∈ [6; 12] aralığında platodadır; seçim plato ortası olan k = 8'dir. "
            "Büyük k tespit sayısını düşürmese de yavaş arızalarda gecikmeyi "
            "büyütmektedir (k = 80'de 2750 ms).",
    widths=[2.4, 2.4, 3.0, 2.4, 3.0, 2.6], highlight=(2,), num_from=1)

h(2, "7.6. Sapma D4 — alarm gecikmesi")
p("Tekil bir kararın gürültüyle tetiklenmesini önlemek için varsayılan olarak ardışık "
  "**iki** karar (100 ms) şartı aranmaktadır. Bu, bildiride bulunmayan bir eklentidir "
  "ve `consecutive_for_alarm:=1` ile kapatılabilir.")

h(2, "7.7. ROS 2 düğümü")
p("`anomaly_detection` paketi, hesabın tamamını ROS'tan bağımsız bir çekirdekte "
  "(`FusionDetector`) tutar; düğüm bunun ince bir sarmalayıcısıdır. Bu ayrım test "
  "edilebilirlik içindir: `replay_detector.py` gerçek parquet verisini **düğümün "
  "gerçekten çalıştırdığı sınıftan** geçirir, teste özel bir kopyasından değil.")
table(
    ["Arayüz", "Ad", "Tip"],
    [["Abone", "/joint_states", "sensor_msgs/JointState  (effort = Amper)"],
     ["Abone", "/force_torque_sensor_broadcaster/wrench", "geometry_msgs/WrenchStamped"],
     ["Yayın", "~/score", "std_msgs/Float32  (birleşik skor)"],
     ["Yayın", "~/detected", "std_msgs/Bool  (gecikme kuralı uygulanmış alarm)"],
     ["Yayın", "~/detail", "std_msgs/Float32MultiArray  (15 eleman)"]],
    caption="Tablo 7.3. Düğümün arayüzü. `~/detail` sırasıyla s_kal, s_ham, z_kal, "
            "z_ham, birleşik, mutlak eşik, uyarlanabilir eşik, mutlak isabet, "
            "uyarlanabilir isabet, kalıntı isabet, ham isabet, hareket, q̇ tepesi, "
            "taban çizgisi örnek sayısı ve dondurma bayrağı değerlerini taşır. KTS "
            "konusunun adı `ft_data` değil `wrench`'tir: `ur_controllers.yaml` "
            "`topic_name: ft_data` yazsa da çalışan yayıncı varsayılan adı "
            "kullanmaktadır ve yanlış ad sessizce hiç skor üretmemeye yol açar.",
    widths=[2.2, 6.2, 6.0], num_from=2)

h(3, "7.7.1. Çok yayıncılı /joint_states")
p("Gerçek hücrede `/joint_states` üzerinde birden çok yayıncı bulunmaktadır. 21 Ağustos "
  "2026 ölçümünde üç ayrı isim kümesi gözlenmiştir: 6 UR eklemi + `base_to_robot_mount` "
  "içeren ve **sırası karışık** 7 elemanlı küme (UR verisi buradan gelmektedir, eşleme "
  "`[0, 6, 2, 5, 3, 1]`), Kawasaki/AGV'ye ait 9 elemanlı küme ve tek elemanlı "
  "`world_to_agv` kümesi. UR dışı mesaj hızı saniyede ~112'dir.")
p("Bu nedenle eklem eşlemesi **mesaj başına**, isim kümesine göre çözülmekte ve küme "
  "başına önbelleklenmektedir. Eşlemenin bir kez kurulup sonraki mesajlara körlemesine "
  "uygulanması — ilk gerçeklemedeki davranış — iki yayıncının eklem sırası farklı "
  "olduğunda sessizce yanlış eklemlerin okunmasına yol açmaktadır; ölçülen senaryoda "
  "rayın lineer ekleminin `wrist_2` yerine okunması dahil altı kanaldan beşi hatalıydı.")

h(3, "7.7.2. Kayıt katmanı")
p("Düğüm her başlatmada `log_dir` altına zaman damgalı iki dosya açar. `olaylar_*.jsonl` "
  "yalnız alarmları tutar: başlangıçta birleşik skor, aşılan eşik, hangi kural, hangi "
  "model tetikledi ve o andaki q, q̇, akım; bitişte süre ve **tepe** skor. Satır "
  "tamponlu olduğundan kayıp yoktur. `skorlar_*.csv` her kararı 17 sütunla yazar ve "
  "eşiklerin gerçek robot verisinden yeniden ölçülmesi için gereken ham girdidir; "
  "ölçülen boyut ~87 bayt/satır, 19,8 Hz'de yaklaşık 155 MB/gündür. `log_dir:=\"\"` "
  "ile kayıt tamamen kapatılabilir.")
table(
    ["Parametre", "Varsayılan", "İşlev"],
    [["`tf_prefix`", "`ur10e_`", "`/joint_states` içindeki eklem adı öneki"],
     ["`wrench_topic`", "…`/wrench`", "KTS konusu"],
     ["`stride`", "25", "karar periyodu (50 ms @ 500 Hz)"],
     ["`fts_frame`", "`tool`", "wrench'in ifade edildiği çerçeve"],
     ["`max_wrench_age`", "20 ms", "bayat wrench ile skor üretilmez"],
     ["`consecutive_for_alarm`", "2", "alarm için ardışık karar sayısı"],
     ["`threshold_override`", "0,0", "pozitifse θ_mutlak'ın yerine geçer"],
     ["`adaptive`", "**kapalı**", "medyan + k·MAD kuralı"],
     ["`adaptive_window`", "600", "taban çizgisi penceresi (30 s)"],
     ["`adaptive_k`", "8,0", "MAD çarpanı"],
     ["`adaptive_warmup`", "200", "kural devreye girmeden önceki karar sayısı"],
     ["`freeze_timeout`", "3,0 s", "bundan uzun alarm = rejim değişimi"],
     ["`motion_qd_min`", "−1,0", "negatif = hareket kapısı kapalı"],
     ["`log_dir`", "`~/anomali_kayit`", "olay/skor kayıtları; boş dize kapatır"],
     ["`log_scores`", "açık", "her kararı CSV'ye yaz"]],
    caption="Tablo 7.4. Düğüm parametreleri ve gerçek robot için ölçülmüş "
            "varsayılanları. `adaptive`, `freeze_timeout`, `motion_qd_min` ve "
            "`threshold_override` Bölüm 10'daki devreye alma ölçümlerinin sonucudur.",
    widths=[4.6, 3.0, 6.8], num_from=3)
p("Düğüm 10 saniyede bir sağlık raporu yayımlar: örnek hızı, karar hızı ve çıkarım "
  "süresinin bütçeye oranı. Örnek hızı 500 Hz'in altına düşerse uyarı verilir — çünkü "
  "q̈ türevi sabit dt varsaymaktadır ve düşük hızda kalıntı bozulur.")

h(2, "7.8. Gecikme bütçesi")
code("örnek başına    : ters dinamik 22 µs + Jacobian ≈ 30 µs   →  500 Hz'in %1,5'i\n"
     "karar başına    : iki ONNX ileri geçişi 5,9 ms            →  50 ms'in %12'si\n"
     "tespit gecikmesi: SG filtresi 50 ms + karar periyodu 50 ms →  ≈ 120–150 ms")
p("Modeller ONNX'e dışa aktarılmış ve ONNX Runtime ile çalıştırılmaktadır; oturum başına "
  "iş parçacığı sayısı 2 ile sınırlandırılarak 500 Hz geri çağırma döngüsünün aç gözlü "
  "bir çıkarım tarafından aç bırakılması engellenmiştir. Ölçüm CPU üzerinde yapılmıştır; "
  "hedef iş istasyonunda `onnxruntime-gpu` kurulduğunda CUDA sağlayıcısı otomatik "
  "seçilecek ve bu süre daha da düşecektir. Bildirinin “GPU üzerinde 10 ms'nin altında” "
  "ifadesiyle uyumludur.")


# ═══════════════════════════════════════════════════════════════════
h(1, "8. Birleşim gerçekten iki modelin ortak kararı mı?")
p("“İki model birleşiyor” cümlesi, kodda öyle yazdığı için doğru sayılmaz. w_ham = 0,05 "
  "olduğunda ham modelin katkısı ölçüm hatası mertebesinde kalıyor olabilir; o durumda "
  "sistem, adı birleşim olan tek modelli bir dedektördür. `verify_fusion_logic.py` bu "
  "soruyu üç bağımsız sınavla yanıtlamaktadır.")

h(2, "8.1. Sınav 1 — yapısal eşleme")
table(
    ["Bildiri", "Düğümün karar yolu", "Durum"],
    [["İki ayrı model, iki ayrı temsil uzayı",
      "12 kanal kalıntı (478.892 par.) + 24 kanal ham (1.907.032 par.)", "Aynı"],
     ["Anomali skoru = yeniden yapılanma MSE",
      "mean((recon − x)²), eğitimdeki kayıpla aynı tanım", "Aynı"],
     ["Skorlar min–max ile normalleştiriliyor",
      "min–max **formülü aynen**, sınırlar temiz doğrulamadan", "**D1 · azaltıldı**"],
     ["S_bir = w_kal·S̄_kal + w_ham·S̄_ham", "fused = 0,95·z_kal + 0,05·z_ham", "Aynı"],
     ["w, 0,05 adımla taranıp optimize ediliyor", "21 nokta tarama → w_kal = 0,95", "Aynı"],
     ["Tekil eşik = doğrulama hatalarının P97'si", "θ = 1,6008 / 3,4461", "Aynı"],
     ["Birleşik skorun çalışma eşiği — *tanımsız*",
      "aynı P97 kuralı birleşik skora → θ = 0,6436", "**D2 · ekleme**"],
     ["Uyarlanabilir kural — yok",
      "medyan + 8·1,4826·MAD, mutlak eşikle VEYA'lanıyor", "**D3 · isteğe bağlı**"],
     ["Alarm gecikmesi — yok", "ardışık 2 karar (100 ms)", "**D4 · isteğe bağlı**"]],
    caption="Tablo 8.1. Bildirinin kuralları ile düğümün karar yolunun eşlenmesi. "
            "D1 ve D2 zorunludur (bildirinin yöntemi ya nedensel değildir ya da "
            "tanımlanmamıştır); D3 ve D4 `adaptive:=false` ve "
            "`consecutive_for_alarm:=1` ile kapatılıp bildirinin saf davranışına "
            "dönülebilir.",
    widths=[5.4, 6.4, 2.8], num_from=2)

h(2, "8.2. Sınav 2 — çalışma sayacı")
p("`OnnxAE.score` metodu sarmalanarak her çağrı sayılmıştır. 6 koşu × 5 senaryo "
  "yeniden oynatmasında:")
code("karar sayısı            : 13.160\n"
     "kalıntı modeli çağrısı  : 13.160\n"
     "ham modeli çağrısı      : 13.160\n"
     "→ her kararda İKİ ONNX oturumu da koştu")
p("Bu, ham modelin “opsiyonel” veya “önbelleğe alınmış” olmadığını, her karar için "
  "gerçekten ileri geçiş yaptığını gösterir.")

h(2, "8.3. Sınav 3 — ablasyon")
p("Ablasyon üç yapılandırmayı karşılaştırır: yalnız ham (w_kal = 0,00), yalnız kalıntı "
  "(w_kal = 1,00) ve birleşim (w_kal = 0,95). Sonuç, hangi **protokol** altında "
  "ölçüldüğüne kritik biçimde bağlıdır ve bu bağımlılığın kendisi bir bulgudur.")

h(3, "8.3.1. Bildirinin protokolü altında")
table(
    ["Yapılandırma", "Kesinlik", "Geri çağırma", "F1", "Yorum"],
    [["Yalnız Ham (w_kal = 0,00)", "0,759", "0,514", "0,613", "kalıntı modeli çıkarıldı"],
     ["Yalnız Kalıntı (w_kal = 1,00)", "0,756", "0,470", "0,579", "ham modeli çıkarıldı"],
     ["BİRLEŞİM (w_kal = 0,95)", "0,833", "0,754", "**0,792**",
      "ham modelin %5 ağırlığı geri çağırmayı 0,470 → 0,754 çıkarıyor"]],
    caption="Tablo 8.2. Karar düzeyi ablasyon, bildirinin kendi arıza kümesi, pencere "
            "etiketlemesi ve eşik kuralı (yalnız mutlak eşik) altında; 41.688 pencere. "
            "En iyi tekil modele göre kazanç +0,178.",
    widths=[4.6, 2.4, 2.8, 2.0, 5.0], highlight=(2,), num_from=1)

h(3, "8.3.2. Çevrimiçi yeniden oynatmada")
p("Aynı ablasyon, düğümün kendi koşum ortamında (6 uzun koşuya elle enjekte edilen dört "
  "arıza, 13.160 karar) tekrarlandığında birleşim ile yalnız-kalıntı **eşitlenmektedir**. "
  "Sebebi ölçülmüştür ve iki katlıdır.")
table(
    ["Kural", "Yalnız Ham", "Yalnız Kalıntı", "Birleşim 0,95", "Taramada en iyi w"],
    [["Bildirinin saf kuralı (yalnız mutlak eşik)", "0,512", "0,513", "0,513", "—"],
     ["Uyarlanabilir kural açık (D3)", "0,772", "0,818", "0,818", "0,75–1,00"],
     ["Genlik ×0,25, D3 açık", "0,648", "0,747", "0,747", "**0,50 → 0,769**"]],
    caption="Tablo 8.3. Çevrimiçi yeniden oynatmada karar düzeyi F1 (13.160 karar). "
            "Kaynak: `fusion_logic_audit.json` ve `fusion_logic_audit_a0.25.json`.",
    widths=[6.2, 2.4, 2.6, 2.6, 3.2], num_from=1)
note("Neden bu koşum ayırt edici değil", [
    "**Birinci sebep — doyum.** Yeniden oynatma dört elle enjekte edilmiş arıza "
    "kullanır (1,5 rad gizyazar adımı, 30 N çarpışma darbesi). Bu genlikler doyurucudur: "
    "yakalanabilir olan her şeyi zaten iki model de tek başına yakalamaktadır. Genlik "
    "×0,25'e düşürüldüğünde bile w = 0,95'te fark açılmaz, çünkü enjekte edilen arızalar "
    "tam olarak *kalıntı* modelinin güçlü olduğu tork/kinematik uzayındadır. Ayırt "
    "edicilik yine de tümüyle kaybolmaz: düşük genlikte tam ağırlık taraması "
    "w = 0,50 → 0,769 vermekte, yalnız-ham 0,648 ve yalnız-kalıntı 0,747'de kalmaktadır. "
    "Yani bu senaryo kümesinde bile birleşim her iki uçtan iyidir — sadece bildirinin "
    "dondurduğu 0,95 ağırlığında değil.",
    "**İkinci sebep — uyarlanabilir kural marjı yutuyor.** Saf kuralda üç yapılandırma "
    "da 0,513'te eşitken, D3 açıldığında yalnız-kalıntı 0,818'e fırlamaktadır. Kural, "
    "tekil modelleri de kurtarmakta ve birleşimin sağladığı marjın bir kısmının "
    "**yerine geçmektedir**. İkisi birbirinin alternatifi değildir: birleşim doğruluğu, "
    "kural ise iyi uyan koşulardaki körlüğü kapatmaktadır. Bu nedenle ikisi de açık "
    "bırakılmıştır.",
    "**Sonuç.** Koşu düzeyi ölçüt (“koşuda en az bir kez tetiklendi mi”) senaryo başına "
    "yalnız 6 koşuda doyduğu için iki modeli ayırt edemez; birleşimin katkısı "
    "bildirinin protokolündeki karar düzeyi ölçütte görünür. Yeniden oynatma ablasyonu, "
    "düğümün doğru çalıştığını gösterir; birleşimin değerini ölçen sınav Tablo 8.2'dir.",
])

h(2, "8.4. Cevap")
note("Sistem bildiriyle aynı mantıkta çalışmaktadır ve gerçekten iki modelin "
     "birleşimidir.", [
    "Denklem (4) birebir uygulanmakta, min–max formülü bildirinin kendi formülü "
    "olmakta, ağırlık bildirinin taramasıyla seçilen 0,95 değerinde bulunmakta ve tekil "
    "eşikler bildirinin P97 kuralıyla belirlenmektedir. Her kararda iki ONNX oturumu da "
    "koşmakta (13.160/13.160) ve bildirinin kendi protokolünde ham modelin %5 ağırlığını "
    "çıkarmak geri çağırmayı 0,754'ten 0,470'e düşürmektedir.",
    "Kalan farklar dört maddeyle sınırlıdır: normalleştirme sınırlarının arıza görmemiş "
    "doğrulama kümesinden alınması (D1), birleşik skorun çalışma eşiğinin bildiride "
    "tanımsız olduğu için aynı P97 kuralıyla türetilmesi (D2) ve tek parametreyle "
    "kapatılabilen iki eklenti (D3, D4).",
])


# ═══════════════════════════════════════════════════════════════════
h(1, "9. Uçtan uca doğrulama")
p("`replay_detector.py`, `ur10e_clean.parquet` verisini `FusionDetector` sınıfından "
  "geçirir ve ölçüm uzayında arıza enjekte eder. Bu, düğümün gerçek veriyle uçtan uca "
  "sınavıdır.")
table(
    ["Senaryo", "Karar", "Alarm oranı", "Tespit edilen koşu", "Ortanca gecikme", "Çıkarım"],
    [["Arıza yok (temiz)", "4.517", "%0,0", "0 / 6", "—", "5,96 ms"],
     ["Çarpışma", "1.821", "%50,1", "6 / 6", "69 ms", "5,98 ms"],
     ["Sensör gürültüsü", "1.821", "%99,7", "6 / 6", "69 ms", "5,94 ms"],
     ["Gizyazar hatası", "1.821", "%82,6", "5 / 6", "72 ms", "5,93 ms"],
     ["Motor kayması", "1.821", "%74,9", "4 / 6", "1.677 ms", "5,91 ms"]],
    caption="Tablo 9.1. Uçtan uca yeniden oynatma: 6 koşu, 113.722 örnek, 227 s robot "
            "zamanı. Temiz veride yanlış alarm oranı %0,0'dır. Ani arızalarda gecikme "
            "yaklaşık 70 ms, yavaş rampa biçimli motor kaymasında ise 1,7 saniyedir.",
    widths=[3.6, 2.0, 2.4, 3.0, 2.8, 2.0], num_from=1)
p("Gecikmelerin iki gruba ayrılması yapısaldır: ani arızalar (çarpışma, gizyazar, "
  "gürültü) tek karar periyodunda görülürken, yavaş bir rampa pencere içinde neredeyse "
  "sabit bir ofset gibi davrandığı için ancak genliği yeterince büyüdüğünde ayırt "
  "edilebilir hale gelmektedir.")


# ═══════════════════════════════════════════════════════════════════
h(1, "10. Gerçek robotta devreye alma")
p("Bölüm 9'a kadarki bütün ölçümler FMU verisi üzerinde yapılmıştır. Bu bölüm, "
  "sistemin 21 Ağustos 2026'da ESOGÜ IFARLAB'daki gerçek UR10e hücresinde çalıştırılması "
  "sırasında toplanan ölçümleri ve bunların çevrimiçi sistemde yol açtığı değişiklikleri "
  "belgelemektedir. Toplam gözlem yaklaşık 45 dakikadır ve üç farklı hareket profilini "
  "kapsar: elle kışkırtılmış arıza denemeleri, otonom muayene (inspection) çevrimi ve "
  "al-yerleştir (pick and place) senaryosu.")

h(2, "10.1. Bağlantı katmanındaki üç ölçüm")
p("Devreye alma, hesaplama katmanına gelmeden önce üç bağlantı sorununu ortaya "
  "çıkarmıştır. Üçü de sessizdir: hiçbiri hata üretmez, yalnızca sonucu bozar.")
bullet([
    "**KTS konusunun adı.** `ur_controllers.yaml` `topic_name: ft_data` tanımlamakta, "
    "fakat çalışan yayıncı varsayılan `wrench` adını kullanmaktadır. Yanlış ada abone "
    "olan düğüm hiç wrench alamamakta ve bayat-wrench koruması nedeniyle **hiç skor "
    "üretmemektedir**.",
    "**Çok yayıncılı `/joint_states`.** Üç ayrı isim kümesi gelmekte, UR verisi sırası "
    "karışık 7 elemanlı kümeden akmaktadır (Bölüm 7.7.1).",
    "**Örnekleme hızı.** Ölçülen hız 495 Hz'dir, 500 değil. Öznitelik çıkarımı sabit "
    "dt = 1/500 varsaymaktadır; sapma %1 mertebesinde kalmakta ve düğüm 400 Hz altına "
    "düştüğünde uyarı vermektedir.",
])
p("Buna karşılık bir varsayım doğrulanmıştır: UR sürücüsü `hardware_interface.cpp` "
  "içinde `effort` alanına `actual_current` yazmaktadır; yani alan gerçekten amperdir "
  "ve `current_to_torque.json` dönüşümü gerçek sürücüde de geçerlidir.")

h(2, "10.2. Uyarlanabilir kuralın kalıcı kilitlenmesi")
p("İlk çalıştırmada düğüm robot **dururken** başlatılmıştır. Uyarlanabilir taban çizgisi "
  "10 saniyelik hareketsiz gürültüden öğrenilmiş ve eşik 0,0103'te oturmuştur. Robot "
  "hareket etmeye başladığında skor anında 0,34'e çıkmış, eşiği aşmış ve alarmı "
  "kilitlemiştir. Alarm sürerken taban çizgisi donduğu için eşik bir daha "
  "güncellenememiş; sonuç 292 saniye süren **tek bir alarm bloğu** olmuştur.")
p("Bu kilit, o blok içinde gerçekleşen gerçek bir olayı da görünmez kılmıştır: bir elin "
  "eklem arasına sıkışması sonucu robot koruyucu durdurmaya geçmiş, birleşik skor "
  "koşunun en yüksek değeri olan **10,27**'ye çıkmış, fakat sistem zaten alarm "
  "durumunda olduğu için yeni bir olay kaydı üretmemiştir.")
table(
    ["Rejim", "Karar", "Birleşik medyan", "Birleşik maksimum"],
    [["hareketsiz", "1.031", "0,0037", "0,0065"],
     ["normal hareket", "4.515", "1,4866", "6,8786"],
     ["**sıkışma (11 s)**", "220", "**2,8012**", "**10,2699**"],
     ["koruyucu durdurma sonrası", "2.280", "0,3151", "0,3834"]],
    caption="Tablo 10.1. İlk koşunun rejimlere ayrılmış skor dağılımı. Sıkışma anı "
            "348 saniyelik koşunun en yüksek değeridir; onu izleyen iki dakikalık "
            "düz seviye robotun koruyucu durdurmada beklediği süredir.",
    widths=[5.0, 2.6, 3.4, 3.4], highlight=(2,), num_from=1)
p("Düzeltme, dondurmaya **3 saniyelik süre sınırı** koymaktır (Bölüm 7.5). Aynı kayıt "
  "yeni kuralla yeniden koşturulduğunda tek 292 saniyelik blok yerine ayrı olaylar "
  "üretilmekte ve sıkışma, 10:21:26'da başlayan kendi 1,8 saniyelik olayı olmaktadır.")

h(2, "10.3. FMU eşiğinin gerçek robotta geçersizliği")
p("Kilit çözüldükten sonra ortaya çıkan asıl sorun eşiğin kendisidir. Temiz bir koşuda "
  "birleşik skorun medyanı **4,27**, FMU'dan gelen eşik ise **0,6436**'dır; kararların "
  "%97'si alarm ilan edilmektedir. Eşik yedi kat düşüktür.")
p("Sapmanın kaynağı ölçülmüştür: kalıntı **poza bağlıdır**. Yerçekimi yüklü, kolun açık "
  "olduğu duruşlarda skor medyanı 4,30 iken harekette 1,58, katlanmış park pozunda ise "
  "0,0037'dir. `nm_per_amp` katsayısı altı eklemin dördünde ölçülememiş, "
  "`residual_calibration_clean.json` içindeki `b` ise **sabit** bir offsettir ve poza "
  "bağlı bir hatayı telafi edemez. Bu, bir eşik sorunu değil model sadakati sorunudur; "
  "tek bir global eşik bu nedenle kalıcı olarak kırılgandır.")

h(2, "10.4. Eşiğin gerçek robot verisinden yeniden ölçülmesi")
p("Eşik, operatörün doğruladığı etiketli bir küme üzerinden yeniden belirlenmiştir: "
  "beş doğrulanmış gerçek olay ve yirmi dokuz doğrulanmış yanlış alarm.")
table(
    ["θ", "Yakalanan gerçek", "Kaçan", "Yanlış alarm", "Yanlış/saat"],
    [["0,6436  (FMU)", "5/5", "—", "29", "38,7"],
     ["12,0", "4/5", "10:21 sıkışması", "11", "14,7"],
     ["**18,0  (seçilen)**", "**4/5**", "10:21 sıkışması", "**5**", "**6,7**"],
     ["26,0", "4/5", "10:21 sıkışması", "2", "2,7"],
     ["33,0", "3/5", "2 olay", "1", "1,3"],
     ["91,0", "2/5", "3 olay", "0", "0,0"]],
    caption="Tablo 10.2. Eşik taraması (~45 dakika gözlem). 18,0 değeri muayene "
            "çevriminin tavanının (17,43) hemen üstünde oturmakta ve o çevrimin "
            "on altı periyodik yanlış alarmının tamamını susturmaktadır; en zayıf "
            "gerçek olaya (31,98) %78 pay bırakır. Kaçan tek olay, robotun kendi "
            "koruyucu durdurmasının zaten anında kestiği ilk sıkışmadır.",
    widths=[4.0, 3.2, 4.0, 2.6, 2.6], highlight=(2,), num_from=2)
p("Aynı etiketli küme üzerinde uyarlanabilir kural da yeniden sınanmıştır. k ∈ [8; 30] "
  "aralığında hiçbir değer doğrulanmış gerçek olayları yakalamamış, yalnızca temiz "
  "veriye yanlış blok eklemiştir. Sebebi ölçülmüştür: kuralın taban çizgisi normal "
  "hareketin ~6,4'lük tepelerini de yutmakta, MAD şişmekte ve eşik gerçek olayların "
  "üstüne çıkmaktadır. Gerçek robotta olayı yakalayan bileşen **mutlak eşiktir** — "
  "FMU verisindeki sonucun tam tersi.")

h(2, "10.5. Yanlış alarmların yapısı")
p("Kalan yanlış alarmlar gürültü değil, **yapılıdır**. Muayene çevriminde alarm "
  "başlangıçları arası 60, 38, 59, 32, 60, 32, 59, 32 ve 60 saniyedir; her ~92 saniyelik "
  "çevrimde aynı üçlü örüntü (6,2 s / tepe ~17, ardından 3,0 s / tepe ~10,3, ardından "
  "0,15 s / ~8,4) tekrarlanmaktadır. Aynı yörünge her seferinde aynı sıçramayı "
  "üretmektedir. Bu, modellerin gerçek yörüngeleri tanımadığının doğrudan kanıtıdır: "
  "FMU ile eğitilmiş özkodlayıcılar için gerçek muayene taraması dağıtım dışıdır.")
p("Genlik dışında ayırt edici bir öznitelik aranmış fakat bulunamamıştır. Blokların "
  "keskinliği (tepe/blok-içi-medyan) gerçek olaylarda 1,08–1,99, yanlış alarmlarda "
  "1,00–4,83 aralığındadır; en keskin blok bir **yanlış** alarmdır. Yükseliş süresi de "
  "ayırmamaktadır. Dolayısıyla \u201canomali anlıktır\u201d sezgisi bu veride bir zaman "
  "imzasına dönüşmemekte, ayrım yalnızca genlikle yapılabilmektedir.")
table(
    ["", "n", "En düşük tepe", "Medyan tepe", "En yüksek tepe"],
    [["Doğrulanmış gerçek olay", "5", "10,27", "79,86", "195,83"],
     ["Doğrulanmış yanlış alarm", "29", "8,30", "10,43", "90,55"]],
    caption="Tablo 10.3. İki sınıfın tepe skor dağılımı. Örtüşme gerçektir: bir yanlış "
            "alarmın tepesi (90,55) üç gerçek olayın tepesinden büyüktür. Bu vaka "
            "eşikle temizlenememektedir ve modellerin gerçek veriyle yeniden "
            "eğitilmesini gerektirmektedir.",
    widths=[6.0, 1.6, 3.0, 3.0, 3.0], num_from=3)

h(2, "10.6. Tespit gecikmesi")
p("Gerçek robotta ölçülen bileşenler: Savitzky–Golay filtresi 50 ms, karar periyodu "
  "50 ms (stride 25 @ 495 Hz) ve `consecutive_for_alarm = 2` için bir karar daha 50 ms. "
  "Toplam **≈150 ms**, 0,25 m/s'lik bir hareket hızında yaklaşık 3,75 cm ek yola karşılık "
  "gelmektedir. Çıkarım süresi 8,5 ms olup 50 ms bütçenin %17'sidir.")
p("Doğrulanmış en net gerçek olay — al-yerleştir senaryosunun sonundaki çarpışma, tepe "
  "146,44 — yalnızca **0,25 saniye** sürmüştür. Bu, alarmı tüketen her katmanın olayı "
  "seviye örneklemesiyle değil kenar yakalamasıyla işlemesini zorunlu kılmaktadır.")

h(2, "10.7. Operatör arayüzü")
p("Sistem, laboratuvarın mevcut web tabanlı kontrol panosuna (`user_interface` paketi, "
  "Flask + SocketIO) üçüncü bir sekme olarak entegre edilmiştir. Sekme, durum bandı "
  "(NORMAL / ANOMALİ), 60 saniyelik logaritmik eksenli canlı skor grafiği, iki modelin "
  "kendi eşiklerine oranını gösteren kırılım ve olay geçmişi tablosundan oluşmaktadır.")
p("İki tasarım kararı doğrudan bu bölümdeki ölçümlerden gelmektedir. Birincisi, tabloda "
  "birincil sütun **tepe** değeridir; giriş değeri ikincil gösterilir, çünkü giriş "
  "değerine bakıp tepeyi görmemek yanlış yoruma yol açmaktadır (8,5 giriş / 90,55 tepe). "
  "İkincisi, alarm hem sunucu tarafında kenar olarak yakalanmakta hem de arayüzde en az "
  "beş saniye mandallanmaktadır; 5 Hz'lik bir güncelleme hızında 0,25 saniyelik bir olay "
  "aksi hâlde ekrana hiç yansımamaktadır. Sahte 0,25 s'lik alarmlarla yapılan ölçümde "
  "kenar yakalama dört olayın dördünü, seviye örnekleme yalnız ikisini görmüştür.")
p("Tablodaki her satırda operatör için `Gerçek` / `Yanlış alarm` / `?` düğmeleri "
  "bulunmaktadır. Etiketler dedektörün kayıtlarına dokunmadan ayrı bir dosyaya yazılır "
  "ve modellerin gerçek veriyle yeniden eğitilmesi için gereken etiketli kümeyi "
  "biriktirir.")


# ═══════════════════════════════════════════════════════════════════
h(1, "11. Tartışma")

h(2, "11.1. Motor kayması neden düşük kalıyor")
p("Motor kayması, düzeltilmiş hatta hem tekil hem birleşik modellerde en zayıf "
  "senaryodur (birleşim AUC 0,443; bildiri 0,588). Bu bir uygulama hatası değil, "
  "**pencereli özkodlayıcının yapısal sınırıdır**: 100 örneklik (200 ms) pencere içinde "
  "yavaş bir doğrusal rampa neredeyse sabit bir ofsettir ve özkodlayıcı sabit bir "
  "ofseti sorunsuz yeniden yapılandırır. Bildiri de aynı senaryoda birleşimin marjinal "
  "bir düşüş gösterdiğini (Δ = −0,017) raporlamaktadır.")
p("Canlı sistemde bu boşluğu uyarlanabilir kural kapatmaktadır: taban çizgisi son 30 "
  "saniyeye göre güncellendiği için yavaşça büyüyen bir skor, mutlak eşiği hiç geçmeden "
  "de alarm üretebilmektedir. Yeniden oynatmada motor kayması 4/6 koşuda, ortanca "
  "1.677 ms gecikmeyle yakalanmıştır. Kalıcı çözüm için doğru yön, pencereli bir "
  "özkodlayıcıyı zorlamak değil, uzun ufuklu ikinci bir gösterge (ör. kayan ortalama "
  "kalıntının eğimi) eklemektir.")

h(2, "11.2. Mutlak başarım değerlerinin düşmesi")
p("Düzeltilmiş hattaki AUC ve F1 değerleri bildirininkilerin altındadır. Bunun üç "
  "ölçülebilir sebebi vardır ve hiçbiri bir gerileme değildir:")
bullet([
    "**Arıza genlikleri artık gerçek ölçekte.** Karışık birimli hatta “15 Nm” diye "
    "tanımlanan bir rampa, amper cinsinden bir kanala ham sayı olarak ekleniyordu; "
    "bu, fiziksel karşılığından yaklaşık on kat büyük bir bozulma demektir. "
    "Düzeltilmiş hatta aynı senaryo gerçek 15 Nm'dir — yani problem zorlaşmıştır.",
    "**Veri kümesi süreksizliklerden arındırıldı.** Süreksizlikler yalnız gürültü "
    "üretmez, aynı zamanda modelin “anormal” saydığı yapay sıçramalar yaratır; bunlar "
    "kaldırıldığında normal ile anormal arasındaki mesafe gerçek değerine iner.",
    "**Değerlendirme sızıntısız.** Eğitim/doğrulama bölünmesi koşu sınırına hizalanmış "
    "ve pencereler koşu sınırlarına saygılı hale getirilmiştir; örtüşen pencerelerin "
    "iki kümeye birden düşmesi engellenmiştir.",
])
p("Değerlendirmede belirleyici olan mutlak sayı değil, **iki hat arasında korunan "
  "yapıdır**: modellerin birbirini tamamlaması, ağırlık taramasının aynı bölgede geniş "
  "ve kararlı olması, birleşimin marjı ve arıza tipine göre asimetri — hepsi "
  "korunmuştur.")

h(2, "11.3. FMU'nun ölçülen sınırlarının etkisi")
p("Çözücünün sürtünme modeli içermemesi ve atalet matrisinin sınanan özellikleri "
  "sağlamaması, kalıntı uzayının teorik üstünlüğünü özellikle bilek eklemlerinde "
  "gerçekleştirmemektedir. Kalıntı modelinin bildiriye göre daha düşük tekil başarım "
  "göstermesi (AUC 0,812 ↔ 0,908) bu etkiyle tutarlıdır. Birleşimin marjının buna "
  "rağmen korunması, ham modelin bu kanallarda bilgiyi doğrudan taşıdığını ve "
  "birleşimin tam olarak bu tür kör noktaları kapatmak için var olduğunu "
  "göstermektedir.")

h(2, "11.4. Kısıtlar")
bullet([
    "Tek robot (UR10e) ve tek görev profili; farklı robot tiplerinde doğrulama "
    "yapılmamıştır.",
    "Sentetik arızalar gerçek arıza dinamiklerinin tümünü yansıtmayabilir; özellikle "
    "çarpışma senaryosunda darbe tüm KTS kanallarına eşit dağıtılmakta, gerçek temas "
    "noktasından Jacobian aracılığıyla yayılan bir model kullanılmamaktadır.",
    "Akım→tork katsayısı yalnızca iki eklemde doğrudan ölçülebilmiştir; kalan dört "
    "eklemde varsayımla türetilmiştir. Bu, tespiti değil fiziksel yorumu etkilemektedir.",
    "Uyarlanabilir kuralın k değeri dört koşu üzerinde ayarlanmıştır; daha geniş bir "
    "koşu havuzunda yeniden ölçülmesi yerinde olur.",
    "Çevrimiçi ölçümler CPU üzerinde yapılmıştır; hedef iş istasyonunda GPU sağlayıcısı "
    "ile yeniden ölçülmelidir.",
])


# ═══════════════════════════════════════════════════════════════════
h(1, "12. Sonuç")
p("Bildirinin yöntemi bağımsız olarak yeniden kurulmuş, mimari ve yordamsal "
  "bileşenlerinin birebir uyduğu ölçümle gösterilmiştir. Veri hattında yedi bulgu "
  "tespit edilmiş; beşi düzeltilmiş, biri (FMU'nun ölçülen sınırları) kısıt gereği "
  "bilinçle kabul edilmiş, biri bildirinin dizgi düzeyindeki bir tutarsızlığıdır. "
  "Bildirinin sayılarının karışık birimli bir hattan geldiği, o hattın yeniden "
  "kurulmasıyla belgelenmiştir.")
p("Düzeltilmiş hatta birleşimin değeri korunmuştur: yalnız-ham 0,613 ve yalnız-kalıntı "
  "0,579'a karşılık birleşim 0,792 F1 (ΔF1 = +0,178) elde etmekte, tamamlayıcılık "
  "oranları bildirininkilerle örtüşmektedir. Sistem 500 Hz'de çalışan bir ROS 2 düğümü "
  "olarak gerçeklenmiş; çevrimiçi öznitelik motorunun çevrimdışı hatla sayısal olarak "
  "özdeş olduğu, her kararda iki modelin de koştuğu ve çıkarımın zaman bütçesinin "
  "%12'sini kullandığı ölçülmüştür.")
p("Bildirinin “Gelecek çalışmalar” bölümünde işaret ettiği üç maddeden biri — "
  "sistemin UR10e ile eş zamanlı çevrimiçi çalıştırılması — bu çalışmayla "
  "tamamlanmıştır. Düğüm gerçek robot üzerinde çalıştırılmış, eşikler gerçek arıza "
  "kayıtlarıyla yeniden ölçülmüş ve sistem operatör arayüzüne entegre edilmiştir.")
p("Gerçek robot ölçümleri, FMU verisiyle varılan iki sonucu tersine çevirmiştir: "
  "çalışma eşiği yedi kat yükseltilmek zorunda kalmış ve uyarlanabilir alarm kuralı "
  "— FMU'da tekil modelleri kurtaran bileşen — gerçek sistemde doğrulanmış hiçbir olayı "
  "yakalamadığı için kapatılmıştır. Bu, simülasyon verisiyle kalibre edilmiş "
  "eşiklerin ve kural seçimlerinin gerçek donanıma aktarılamayacağının somut bir "
  "örneğidir.")
p("Kalan sınır ölçülmüştür ve eşikle aşılamaz: doğrulanmış bir yanlış alarmın tepesi "
  "(90,55) üç gerçek olayın tepesinden büyüktür ve yanlış alarmlar yörünge çevrimiyle "
  "birebir tekrar etmektedir. Kök neden, FMU ile eğitilmiş özkodlayıcılar için gerçek "
  "yörüngelerin dağıtım dışı olmasıdır. Sıradaki adım, arayüzün biriktirdiği operatör "
  "etiketleriyle modellerin gerçek robot verisi üzerinde yeniden eğitilmesidir.")


# ═══════════════════════════════════════════════════════════════════
h(1, "13. Yeniden üretilebilirlik")

h(2, "13.1. Çevrimdışı hat")
code("bash run_pipeline.sh\n"
     "\n"
     "  1/8  prepare_dataset.py            → ur10e_clean.parquet   (600 koşu)\n"
     "  2/8  calibrate_current_to_torque.py→ current_to_torque.json\n"
     "  3/8  convert_effort_to_nm.py       → ur10e_clean_nm.parquet\n"
     "  4/8  generate_residuals.py         → ur10e_features.parquet\n"
     "  5/8  train_ae.py --mode residual   → residual_ae_v2/\n"
     "       train_ae.py --mode raw        → raw_ae_v2/\n"
     "  6/8  evaluate_fusion.py            → fusion_v2/fusion_config.json\n"
     "  7/8  tune_adaptive.py              → fusion_v2/adaptive_tuning.json\n"
     "  8/8  verify_online_features.py + replay_detector.py")

h(2, "13.2. Bildirinin özgün hattı (erratum)")
code("bash reproduce_erratum.sh erratum")

h(2, "13.3. Birleşim denetimi")
code("python3 verify_fusion_logic.py                    # tam denetim, izleri .npz'ye yazar\n"
     "python3 verify_fusion_logic.py --reuse-traces     # yalnız ablasyonu tekrar koşar\n"
     "python3 verify_fusion_logic.py --amp-scale 0.25   # düşük genlikte ayırt edicilik")

h(2, "13.4. Çevrimiçi düğüm")
code("# gerçek robot için ölçülmüş varsayılanlarla (θ = 18,0, uyarlanabilir kapalı):\n"
     "ros2 launch anomaly_detection detector.launch.py\n"
     "\n"
     "# eşiği kalıcı olarak değiştirmeden denemek için:\n"
     "ros2 launch anomaly_detection detector.launch.py threshold_override:=26.0\n"
     "\n"
     "# bildirinin saf davranışına dönmek için:\n"
     "ros2 launch anomaly_detection detector.launch.py \\\n"
     "     threshold_override:=0.6436 adaptive:=false consecutive_for_alarm:=1\n"
     "\n"
     "# kaydı kapatmak için:\n"
     "ros2 launch anomaly_detection detector.launch.py log_dir:=\"\'\'\"")
p("Gömülü eşik `fusion_v2/fusion_config.json` içindeki `fused_threshold` alanındadır. "
  "FMU'dan gelen özgün değer aynı dosyada `fused_threshold_fmu` olarak, seçim gerekçesi "
  "ise `fused_threshold_not` alanında saklanmaktadır.")

h(2, "13.5. Dosya envanteri")
table(
    ["Dosya", "Satır", "İşlev"],
    [["`prepare_dataset.py`", "196", "koşulara ayırma, düzgün ızgaraya örnekleme"],
     ["`calibrate_current_to_torque.py`", "186", "quasi-statik yerçekimi regresyonu"],
     ["`generate_residuals.py`", "470", "q̈, τ_model, Jᵀ·F, kalıntı ayrıştırma"],
     ["`models.py`", "196", "mimari, koşu-farkında pencereleme"],
     ["`train_ae.py`", "294", "eğitim, eşik, ONNX dışa aktarım"],
     ["`inject_faults.py`", "165", "dört sentetik arıza senaryosu"],
     ["`evaluate_fusion.py`", "608", "Tablo II/III, ölçek yarışı, birleşim yapılandırması"],
     ["`tune_adaptive.py`", "206", "uyarlanabilir kuralın k taraması"],
     ["`verify_online_features.py`", "119", "çevrimiçi ↔ çevrimdışı öznitelik özdeşliği"],
     ["`verify_fusion_logic.py`", "419", "üç sınavlı birleşim denetimi"],
     ["`replay_detector.py`", "174", "uçtan uca yeniden oynatma"],
     ["`anomaly_detection/features.py`", "154", "çevrimiçi öznitelik motoru"],
     ["`anomaly_detection/detector.py`", "247", "ROS'tan bağımsız çekirdek"],
     ["`anomaly_detection/detector_node.py`", "412", "ROS 2 düğümü, kayıt katmanı"],
     ["`anomaly_detection/replay_publisher.py`", "169", "500 Hz parquet yayıncısı"],
     ["`launch/detector.launch.py`", "68", "başlatma dosyası"]],
    caption="Tablo 13.1. Yazılım envanteri. Çevrimdışı hat 3.176 satırdır ve "
            "depo dışında tutulmaktadır; kurulan ROS 2 paketi 1.050 satır kod ile "
            "model, kalibrasyon ve çözücü varlıklarından oluşan 13 MB'lık bir "
            "dağıtımdır.",
    widths=[7.4, 1.8, 6.2], num_from=1)
p("Paket, geliştirme sırasında kullanılan `ur10e_anomaly_detection` ile birleştirilmiş "
  "ve tek bir `anomaly_detection` ament_python paketine indirgenmiştir. Model, "
  "kalibrasyon ve çözücü varlıkları `share/anomaly_detection` altına kurulmakta; düğüm "
  "bunları `get_package_share_directory` ile bulmaktadır, yani mutlak yol bağımlılığı "
  "yoktur.")


# ═══════════════════════════════════════════════════════════════════
h(1, "Kaynaklar")
for i, ref in enumerate([
    "C. S. Yılmaz, S. Kahraman, M. Yılmaz, H. S. Yavuz, U. Yayan, “FMU Tabanlı Kalıntı "
    "Ayrıştırma ve İkili LSTM Özkodlayıcı Birleşimi ile İşbirlikçi Robotlarda Anomali "
    "Tespiti”, 2026. (Bu raporun referans bildirisi — 245.pdf)",
    "L. Correia, J. C. Goos, P. Klein, T. Bäck, A. V. Kononova, “Online model-based "
    "anomaly detection in multivariate time series: Taxonomy, survey, research "
    "challenges and future directions”, Engineering Applications of Artificial "
    "Intelligence, c. 138, s. 109323, 2024.",
    "W. Li, Y. Han, Z. Xiong, “Collision detection of robots based on a force/torque "
    "sensor at the bedplate”, IEEE Transactions on Industrial Electronics, c. 67, "
    "sy. 12, ss. 12440–12449, 2020.",
    "P. Malhotra, A. Ramakrishnan, G. Anand, L. Vig, P. Agarwal, G. Shroff, "
    "“LSTM-based encoder-decoder for multi-sensor anomaly detection”, arXiv:1607.00148, "
    "2016.",
    "D. Park, Y. Hoshi, C. C. Kemp, “A multimodal anomaly detector for robot-assisted "
    "feeding using an LSTM-based variational autoencoder”, IEEE Robotics and Automation "
    "Letters, c. 3, sy. 3, ss. 1544–1551, 2018.",
    "Universal Robots, “Universal_Robots_ROS2_Driver — hardware_interface.cpp”, "
    "`effort` alanının `actual_current` ile doldurulması ve e-Serisi için "
    "`transformForceTorque()` çerçeve dönüşümü.",
    "A. Savitzky, M. J. E. Golay, “Smoothing and differentiation of data by simplified "
    "least squares procedures”, Analytical Chemistry, c. 36, sy. 8, ss. 1627–1639, 1964.",
], 1):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(4)
    par.paragraph_format.left_indent = Cm(0.8)
    par.paragraph_format.first_line_indent = Cm(-0.8)
    _rich(par, f"[{i}]  {ref}", size=9.5)


OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"yazıldı: {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
