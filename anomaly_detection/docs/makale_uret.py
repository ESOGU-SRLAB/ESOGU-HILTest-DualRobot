#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Builds the ESOGÜ Engineering and Architecture Faculty Journal submission.

The document is generated from the journal's own template
(backup_anomaly_detection/Dosya.docx) so that page setup, the two-column body
section and every named style (Makale Başlığı, Başlık-1, Paragraf, Eşitlik,
Kaynaklar) are the journal's, not ours.

    python3 makale_uret.py

Output: docs/ESOGU_MMF_Makale.docx
Figures come from docs/figures/ and are produced by makale_figurler.py.
"""
import copy
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
PKG = HERE.parent


def find_backup():
    """`backup_anomaly_detection` dizinini bul.

    Yedek 28.08.2026'da paketin dışına alındı; paketin çalışma zamanı ona
    bağlı DEĞİL, ama makale üreteçleri bağlı (dergi şablonu, çevrimdışı
    skorlar, bildiri PDF'i). Konumu sabitlemek yerine aranıyor; taşınırsa
    `AD_BACKUP` ile gösterilebilir.
    """
    env = os.environ.get("AD_BACKUP")
    adaylar = ([Path(env).expanduser()] if env else []) + [
        PKG / "backup_anomaly_detection",              # eski yer (paket içi)
        Path.home() / "Desktop" / "backup_anomaly_detection",
        Path.home() / "backup_anomaly_detection",
        PKG.parent / "backup_anomaly_detection",       # colcon_ws/src/ altında
    ]
    for c in adaylar:
        if (c / "Dosya.docx").is_file():
            return c
    raise SystemExit(
        "backup_anomaly_detection bulunamadı. Bakılan yerler:\n  "
        + "\n  ".join(str(c) for c in adaylar)
        + "\nDoğru yeri AD_BACKUP ile verin, ör.:\n"
          "  AD_BACKUP=~/Desktop/backup_anomaly_detection python3 makale_uret.py")


BACKUP = find_backup()
TEMPLATE = BACKUP / "Dosya.docx"
FIG = HERE / "figures"
OUT = HERE / "ESOGU_MMF_Makale.docx"

FONT = "Cambria"
BODY_PT = 10
SMALL_PT = 9
COL_W = 8.35      # cm — one column of the two-column body
                  # (210 − 2·15 mm text width, less the 1.25 cm gutter)
PAGE_W = 17.9     # cm — full text width

INK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x44, 0x44, 0x44)


# ═════════════════════════════════════════════════════════════════════
# document skeleton
# ═════════════════════════════════════════════════════════════════════
def open_template():
    """Load the journal template and strip its body, keeping both sectPr."""
    doc = Document(str(TEMPLATE))
    body = doc.element.body
    keep_break = None
    for child in list(body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        if child.find(".//" + qn("w:sectPr")) is not None:
            keep_break = child          # paragraph that closes section 0
            continue
        body.remove(child)
    return doc, keep_break


def style_run(run, size=BODY_PT, bold=False, italic=False, colour=INK,
              name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = colour
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    return run


def set_cols(paragraph, num):
    """Close the section at `paragraph` with a `num`-column layout."""
    src = paragraph.part.document.element.body.find(qn("w:sectPr"))
    sect = copy.deepcopy(src)
    cols = sect.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect.append(cols)
    if num == 1:
        for attr in ("w:num", "w:equalWidth"):
            if cols.get(qn(attr)) is not None:
                del cols.attrib[qn(attr)]
        cols.set(qn("w:space"), "708")
    else:
        cols.set(qn("w:num"), str(num))
        cols.set(qn("w:equalWidth"), "1")
        cols.set(qn("w:space"), "709")
    ppr = paragraph._p.get_or_add_pPr()
    for old in ppr.findall(qn("w:sectPr")):
        ppr.remove(old)
    ppr.append(sect)


# ═════════════════════════════════════════════════════════════════════
# building blocks
# ═════════════════════════════════════════════════════════════════════
class Builder:
    # Şekil ve tablo başlıklarının dil etiketi. Türkçe sürüm bunları
    # "Şekil"/"Tablo" olarak geçersiz kılar; böylece yerleşim kodu tek yerde
    # kalır ve iki sürüm arasında biçim farkı oluşmaz.
    FIG_LABEL = "Figure"
    TAB_LABEL = "Table"

    def __init__(self, doc, anchor):
        self.doc = doc
        self.anchor = anchor          # section-0 closing paragraph
        self.body = doc.element.body
        self.fig_no = 0
        self.tab_no = 0
        self.eq_no = 0
        self.last = None
        self.pending_wide = None      # last paragraph of an open 1-column block

    # -- column-span bookkeeping --------------------------------------
    def _begin_wide(self):
        """Open a full-width (single-column) block, unless one is open."""
        if self.pending_wide is not None:
            return                    # consecutive wide elements share a block
        if self.last is not None:
            set_cols(self.last, 2)

    def _end_wide(self, paragraph):
        self.pending_wide = paragraph

    def _leave_wide(self):
        if self.pending_wide is not None:
            set_cols(self.pending_wide, 1)
            self.pending_wide = None

    # -- placement -----------------------------------------------------
    def _add(self, element, front=False):
        if front:
            self.anchor.addprevious(element)
        else:
            self.body.append(element)
            # keep the final sectPr last
            sect = self.body.find(qn("w:sectPr"))
            if sect is not None:
                self.body.remove(sect)
                self.body.append(sect)

    def para(self, text="", style="Paragraf", align=None, front=False,
             size=BODY_PT, bold=False, italic=False, colour=INK,
             space_before=None, space_after=None, keep_with_next=False,
             keep_wide=False):
        if not keep_wide and not front:
            self._leave_wide()
        p = self.doc.add_paragraph()
        self.body.remove(p._p)
        self._add(p._p, front)
        try:
            p.style = self.doc.styles[style]
        except KeyError:
            p.style = self.doc.styles["Normal"]
        if align is not None:
            p.alignment = align
        if space_before is not None:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        if keep_with_next:
            p.paragraph_format.keep_with_next = True
        if text:
            self.rich(p, text, size=size, bold=bold, italic=italic,
                      colour=colour)
        self.last = p
        return p

    def rich(self, paragraph, text, size=BODY_PT, bold=False, italic=False,
             colour=INK):
        """`**bold**` and `*italic*` inline markers."""
        buf = ""
        i = 0
        while i < len(text):
            if text.startswith("**", i):
                j = text.find("**", i + 2)
                if j > 0:
                    if buf:
                        style_run(paragraph.add_run(buf), size, bold, italic, colour)
                        buf = ""
                    style_run(paragraph.add_run(text[i + 2:j]), size, True,
                              italic, colour)
                    i = j + 2
                    continue
            if text[i] == "*":
                j = text.find("*", i + 1)
                if j > 0:
                    if buf:
                        style_run(paragraph.add_run(buf), size, bold, italic, colour)
                        buf = ""
                    style_run(paragraph.add_run(text[i + 1:j]), size, bold,
                              True, colour)
                    i = j + 1
                    continue
            buf += text[i]
            i += 1
        if buf:
            style_run(paragraph.add_run(buf), size, bold, italic, colour)
        return paragraph

    # -- headings ------------------------------------------------------
    def h1(self, text):
        return self.para(text, style="Başlık-1", size=BODY_PT, bold=True,
                         space_before=12, space_after=2, keep_with_next=True)

    def h2(self, text):
        return self.para(text, style="Başlık-1", size=BODY_PT, bold=True,
                         space_before=8, space_after=2, keep_with_next=True)

    def p(self, text):
        return self.para(text, style="Paragraf",
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=6, space_after=0)

    # -- equations -----------------------------------------------------
    def equation(self, text, label=None):
        """
        `label` verilirse numara sayaçtan ALINMAZ ve sayaç ilerlemez.

        Ana denklemler arasına ek bir denklem sokmak gerektiğinde (ör. kalıntı
        tanımının sürtünmeli hâli) "2a"/"2b" gibi bir etiket kullanılır; aksi
        hâlde sonraki bütün denklem numaraları kayar ve metindeki her atıf
        sessizce yanlışa döner.
        """
        if label is None:
            self.eq_no += 1
        self.para("", style="Normal", space_before=0, space_after=0)
        p = self.para(style="Eşitlik", space_before=0, space_after=0)
        pf = p.paragraph_format
        for existing in list(pf.tab_stops):
            pf.tab_stops.clear_all()
            break
        pf.tab_stops.add_tab_stop(Cm(COL_W), WD_TAB_ALIGNMENT.RIGHT)
        style_run(p.add_run("\t" if text.startswith("\t") else ""), BODY_PT)
        self.rich(p, text, size=BODY_PT, italic=True)
        style_run(p.add_run("\t"), BODY_PT)
        style_run(p.add_run(f"({label or self.eq_no})"), BODY_PT)
        self.para("", style="Normal", space_before=0, space_after=0)
        return label or self.eq_no

    # -- figures -------------------------------------------------------
    def figure(self, filename, caption, wide=True, width_cm=None):
        path = FIG / filename
        if not path.exists():
            print(f"  ! missing figure {filename}")
            return None
        self.fig_no += 1
        if wide:
            self._begin_wide()
        else:
            self._leave_wide()
        width = width_cm or (PAGE_W if wide else COL_W)
        pic = self.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=6, space_after=2, keep_with_next=True,
                        keep_wide=True)
        pic.add_run().add_picture(str(path), width=Cm(width))
        cap = self.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=6, keep_wide=True)
        name, _, extra = caption.partition("|")
        self.rich(cap, f"**{self.FIG_LABEL} {self.fig_no}.** {name.strip()}",
                  size=SMALL_PT)
        if extra.strip():
            self.rich(cap, "  " + extra.strip(), size=SMALL_PT - 0.5,
                      colour=GREY)
        if wide:
            self._end_wide(cap)
        self.last = cap
        return self.fig_no

    # -- tables --------------------------------------------------------
    def table(self, caption, header, rows, widths=None, wide=False,
              align_right=None, note=None):
        self.tab_no += 1
        if wide:
            self._begin_wide()
        else:
            self._leave_wide()
        cap = self.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=8, space_after=3, keep_with_next=True,
                        keep_wide=True)
        self.rich(cap, f"**{self.TAB_LABEL} {self.tab_no}.** {caption}",
                  size=SMALL_PT)

        ncol = len(header)
        tbl = self.doc.add_table(rows=1 + len(rows), cols=ncol)
        self.body.remove(tbl._tbl)
        self._add(tbl._tbl)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        total = PAGE_W if wide else COL_W
        if widths is None:
            widths = [total / ncol] * ncol
        scale = total / sum(widths)
        widths = [w * scale for w in widths]

        align_right = align_right or []
        data = [header] + list(rows)
        for r, row in enumerate(data):
            for c, val in enumerate(row):
                cell = tbl.cell(r, c)
                cell.width = Cm(widths[c])
                par = cell.paragraphs[0]
                par.paragraph_format.space_before = Pt(1)
                par.paragraph_format.space_after = Pt(1)
                par.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if c in align_right
                                 and r > 0 else WD_ALIGN_PARAGRAPH.LEFT)
                self.rich(par, str(val), size=SMALL_PT - 1.0, bold=(r == 0))
        self._fix_layout(tbl, widths)
        self._rule_table(tbl)
        if note:
            n = self.para("", style="Normal", space_before=2, space_after=6,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, keep_wide=True)
            self.rich(n, note, size=SMALL_PT - 1, italic=True, colour=GREY)
            self.last = n
        else:
            self.last = self.para("", style="Normal", space_before=0,
                                  space_after=4, keep_wide=True)
        if wide:
            self._end_wide(self.last)
        return self.tab_no

    @staticmethod
    def _fix_layout(tbl, widths_cm):
        """Pin the column grid so Word and LibreOffice stop autofitting."""
        twips = [int(round(w * 567)) for w in widths_cm]
        tblpr = tbl._tbl.tblPr
        for tag in ("w:tblW", "w:tblLayout"):
            for old in tblpr.findall(qn(tag)):
                tblpr.remove(old)
        tw = OxmlElement("w:tblW")
        tw.set(qn("w:w"), str(sum(twips)))
        tw.set(qn("w:type"), "dxa")
        tblpr.append(tw)
        lay = OxmlElement("w:tblLayout")
        lay.set(qn("w:type"), "fixed")
        tblpr.append(lay)
        for old in tbl._tbl.findall(qn("w:tblGrid")):
            tbl._tbl.remove(old)
        grid = OxmlElement("w:tblGrid")
        for t in twips:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(t))
            grid.append(gc)
        tblpr.addnext(grid)
        for row in tbl.rows:
            for cell, t in zip(row.cells, twips):
                tcpr = cell._tc.get_or_add_tcPr()
                for old in tcpr.findall(qn("w:tcW")):
                    tcpr.remove(old)
                el = OxmlElement("w:tcW")
                el.set(qn("w:w"), str(t))
                el.set(qn("w:type"), "dxa")
                tcpr.append(el)

    @staticmethod
    def _rule_table(tbl):
        """Horizontal rules only: above and below the header, below the last row."""
        rows = tbl.rows
        for idx, row in enumerate(rows):
            for cell in row.cells:
                tcpr = cell._tc.get_or_add_tcPr()
                for old in tcpr.findall(qn("w:tcBorders")):
                    tcpr.remove(old)
                borders = OxmlElement("w:tcBorders")
                for edge in ("top", "left", "bottom", "right"):
                    el = OxmlElement(f"w:{edge}")
                    show = ((edge == "top" and idx <= 1)
                            or (edge == "bottom" and idx == len(rows) - 1))
                    el.set(qn("w:val"), "single" if show else "nil")
                    if show:
                        el.set(qn("w:sz"), "8" if idx == 0 or
                               idx == len(rows) - 1 else "4")
                        el.set(qn("w:space"), "0")
                        el.set(qn("w:color"), "000000")
                    borders.append(el)
                tcpr.append(borders)


# ═════════════════════════════════════════════════════════════════════
# front matter
# ═════════════════════════════════════════════════════════════════════
TITLE_EN = "FROM OFFLINE FUSION TO ONLINE DEPLOYMENT: ANOMALY DETECTION ON A UR10e COBOT"
TITLE_TR = "ÇEVRİMDIŞI BİRLEŞİMDEN ÇEVRİMİÇİ GERÇEKLEMEYE UR10e ANOMALİ TESPİTİ"

KEYWORDS_EN = ["Anomaly detection", "Collaborative robots", "LSTM autoencoder",
               "Score-level fusion", "Real-time deployment"]
KEYWORDS_TR = ["Anomali tespiti", "İşbirlikçi robotlar", "LSTM özkodlayıcı",
               "Skor birleşimi", "Gerçek zamanlı sistem"]

ABSTRACT_EN = (
    "Early detection of anomalies in collaborative robots matters for operator "
    "safety and production continuity. Previous work introduced a score-level "
    "fusion of a physics-informed residual autoencoder and a data-driven raw "
    "signal autoencoder for a UR10e cobot, evaluated entirely offline. This "
    "paper takes that framework as its starting point, extends it in four "
    "directions and delivers it as a detector running on the physical cell. The "
    "residual definition is completed with a friction term fitted outside the "
    "validated inverse-dynamics solver, removing 87 % and 92 % of the residual "
    "spread at the wrist joints. The evaluation is placed on a run-disjoint "
    "footing with five training seeds, on which the fusion retains its "
    "advantage over both single models (F1 0.791 against 0.627 and 0.614). A "
    "physically consistent injection protocol, in which faults perturb only the "
    "measured channels and the residual is recomputed, reduces the fusion "
    "margin from +0.189 to -0.003 PR-AUC in every seed; the complementarity is "
    "traced to one injected amplitude fifty-one times smaller than its physical "
    "propagation. The detector was implemented as a 500 Hz ROS 2 node whose "
    "feature engine matches the offline pipeline to floating-point rounding. On "
    "hardware the offline-derived threshold did not transfer, because both "
    "models operate far outside their training distribution and a five-per- "
    "cent-weighted model supplied 29 % of the fused score. After recalibration "
    "on the cell the detector caught three operator-confirmed collisions with "
    "no false alarm, while a low-amplitude contact remained below normal "
    "motion."
)

ABSTRACT_TR = (
    "İşbirlikçi robotlarda anomalilerin erken tespiti operatör güvenliği ve "
    "üretim sürekliliği açısından kritiktir. Önceki çalışmada, bir UR10e kobot "
    "için fizik tabanlı kalıntı özkodlayıcısı ile veri güdümlü ham sinyal "
    "özkodlayıcısının skor düzeyinde birleşimi sunulmuş ve yalnızca çevrimdışı "
    "değerlendirilmişti. Bu makale o çerçeveyi başlangıç noktası alarak dört "
    "yönde genişletmekte ve gerçek hücrede çalışan bir dedektör olarak teslim "
    "etmektedir. Kalıntı tanımı, doğrulanmış ters dinamik çözücünün içermediği "
    "sürtünme terimiyle tamamlanmış, bilek eklemlerinde kalıntı yayılımı %87 ve "
    "%92 azalmıştır. Değerlendirme koşu-ayrık bölme ve beş eğitim tohumu "
    "üzerine oturtulmuş, bu koşulda birleşim iki tekil modeli de geçmiştir (F1 "
    "0,791'e karşılık 0,627 ve 0,614). Arızaların yalnız ölçülen kanallara "
    "uygulandığı ve kalıntının yeniden hesaplandığı fiziksel olarak tutarlı bir "
    "enjeksiyon protokolü ise birleşim marjını her tohumda +0,189'dan -0,003 "
    "PR-AUC'ye indirmektedir; tamamlayıcılığın kaynağı, fiziksel yayılımından "
    "elli bir kat küçük seçilmiş tek bir enjeksiyon genliğidir. Dedektör, "
    "öznitelik motoru çevrimdışı hatla kayan nokta düzeyinde örtüşen 500 Hz'lik "
    "bir ROS 2 düğümü olarak gerçeklenmiştir. Donanımda çevrimdışı türetilen "
    "eşik taşınmamış; sebebi her iki modelin de eğitim dağılımının çok dışında "
    "çalışması ve %5 ağırlıklı modelin birleşik skorun %29'unu sürüklemesidir. "
    "Hücrede yeniden kalibrasyondan sonra dedektör, operatörce doğrulanmış üç "
    "çarpışmayı yanlış alarm üretmeden yakalamış, düşük genlikli bir temas ise "
    "normal hareketin altında kalmıştır."
)


def front_matter(b):
    doc = b.doc

    def title(text):
        p = b.para("", style="Makale Başlığı",
                   align=WD_ALIGN_PARAGRAPH.CENTER, front=True,
                   space_before=0, space_after=0)
        style_run(p.add_run(text), 12, bold=True)
        return p

    title(TITLE_EN)
    b.para("", style="Normal", front=True)
    p = b.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, front=True)
    style_run(p.add_run("Adı SOYADI"), BODY_PT)
    style_run(p.add_run("1*"), BODY_PT - 3.5).font.superscript = True
    style_run(p.add_run(",  Adı SOYADI"), BODY_PT)
    style_run(p.add_run("2"), BODY_PT - 3.5).font.superscript = True
    style_run(p.add_run("  (Dergi editörlüğü tarafından basım aşamasında "
                        "yazılacaktır.)"), BODY_PT)
    for i in (1, 2):
        p = b.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER,
                   front=True)
        style_run(p.add_run(f"{i} "), SMALL_PT - 1)
        style_run(p.add_run(f"Yazar {i} Adresi, ORCID No : "
                            "https://orcid.org/"), SMALL_PT - 1)
    b.para("", style="Normal", front=True)

    abstract_block(b, "Keywords", "Abstract", KEYWORDS_EN, ABSTRACT_EN,
                   "Research Article")
    b.para("", style="Normal", front=True)
    title(TITLE_TR)
    b.para("", style="Normal", front=True)
    abstract_block(b, "Anahtar Kelimeler", "Öz", KEYWORDS_TR, ABSTRACT_TR,
                   "Araştırma Makalesi")
    b.para("", style="Normal", front=True)


def abstract_block(b, kw_head, ab_head, keywords, abstract, kind):
    tbl = b.doc.add_table(rows=3, cols=2)
    b.body.remove(tbl._tbl)
    b.anchor.addprevious(tbl._tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [4.6, 13.3]
    heads = [kw_head, ab_head]
    for c in range(2):
        cell = tbl.cell(0, c)
        cell.width = Cm(widths[c])
        par = cell.paragraphs[0]
        par.paragraph_format.space_before = Pt(2)
        par.paragraph_format.space_after = Pt(2)
        style_run(par.add_run(heads[c]), BODY_PT, bold=True)

    cell = tbl.cell(1, 0)
    cell.width = Cm(widths[0])
    for i, kw in enumerate(keywords):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_run(par.add_run(kw), SMALL_PT, italic=True)

    cell = tbl.cell(1, 1)
    cell.width = Cm(widths[1])
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(2)
    style_run(par.add_run(abstract), BODY_PT, italic=True)

    cell = tbl.cell(2, 0)
    par = cell.paragraphs[0]
    style_run(par.add_run(kind), SMALL_PT, bold=True)
    par = tbl.cell(2, 1).paragraphs[0]
    style_run(par.add_run("Received / Geliş :               "
                          "Accepted / Kabul :"), SMALL_PT, colour=GREY)
    Builder._fix_layout(tbl, widths)
    Builder._rule_table(tbl)


# ═════════════════════════════════════════════════════════════════════
# body
# ═════════════════════════════════════════════════════════════════════
def body(b):
    # ─────────────────────────────────────────── 1. Introduction ─────
    b.h1("1. Introduction")
    b.p(
        "Collaborative robots share their working volume with human operators, so "
        "anomalies such as a collision, a degrading motor or a failing sensor have "
        "to be detected before they become incidents. Two families of methods "
        "dominate. Physics-based residual methods compare a dynamic model of the "
        "robot with the measured forces and treat the difference as a fault "
        "indicator; they reduce dimensionality and suppress noise, but they also "
        "suppress low-amplitude sensor anomalies. Data-driven methods learn complex "
        "patterns directly from high-dimensional raw signals, but they lack the "
        "structure a physical model provides. The two families therefore fail on "
        "different fault types, and that asymmetry can be exploited by fusing them."
    )
    b.p(
        "In earlier work we proposed exactly such a fusion for a UR10e cobot "
        "(Yılmaz, Kahraman, Yılmaz, Yavuz and Yayan, 2026). A residual Long "
        "Short-Term Memory (LSTM) autoencoder operating on twelve intrinsic and "
        "extrinsic residual channels derived from a Functional Mock-up Unit (FMU) "
        "inverse dynamics model was combined, at score level, with a raw LSTM "
        "autoencoder operating on twenty-four direct sensor channels. That study was "
        "evaluated entirely offline on a recorded dataset, and its own future-work "
        "list named a single most important open item: running the system online and "
        "simultaneously with the robot."
    )
    b.p(
        "This paper closes that item. The framework of that study is taken as the "
        "starting point and extended in four directions, and the result is delivered "
        "as a detector running on the physical cell. The residual definition is "
        "completed with the friction term the validated inverse-dynamics solver does "
        "not contain; the evaluation is placed on a run-disjoint footing with five "
        "training seeds so that every number carries an uncertainty; the fault "
        "injection is made physically consistent, which turns out to change what the "
        "fusion is worth; and the normalisation and threshold are given a form that "
        "can be estimated without seeing the future, then re-estimated on the cell "
        "when the offline values prove not to transfer. The detector was then "
        "commissioned and validated against provoked collisions."
    )
    b.p(
        "The gap this study fills is therefore not a new architecture. Fusion "
        "frameworks for robot anomaly detection are numerous and are almost always "
        "evaluated on recorded data with injected faults. What is rarely reported is "
        "which parts of such a framework survive contact with hardware: whether the "
        "operating threshold transfers, whether the normalisation is even causal, "
        "whether the measured fusion advantage depends on how the faults were "
        "injected, and what the detector can and cannot see once it is running. "
        "This paper reports those measurements for one complete system, and states "
        "the resulting working envelope from both sides."
    )
    b.p("The contributions are:")
    for item in (
        "a residual definition completed with a Coulomb-plus-viscous friction term "
        "fitted outside the validated solver, which removes 87 % and 92 % of the "
        "residual spread at the two wrist joints where the inverse dynamics "
        "previously contributed almost nothing;",
        "a run-disjoint re-evaluation with five training seeds, so that the fusion "
        "advantage is reported with an uncertainty rather than as a single number;",
        "a physically consistent fault-injection protocol, and the measurement that "
        "the fusion advantage reported under the inherited protocol disappears under "
        "it — traced to a single injected amplitude fifty-one times smaller than its "
        "physical propagation;",
        "a causal formulation of the fusion normalisation and an operating threshold "
        "for the fused score, neither of which the offline method defines, together "
        "with the procedure that re-estimates both on the cell;",
        "a real-time ROS 2 implementation whose feature engine is numerically "
        "identical to the offline pipeline, friction term included, with a measured "
        "latency budget;",
        "a commissioning report in which the offline-derived threshold is shown not "
        "to transfer, the cause is traced to a normaliser rather than to the models, "
        "and the deployed detector is validated against operator-confirmed "
        "collisions with the sensitivity floor stated explicitly;",
        "an audit of the data path carried out before the extension, whose findings "
        "are corrected here and whose effect on the previously published figures is "
        "documented in Section 4.5.",
    ):
        p = b.para("", style="Paragraf", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_before=3, space_after=0)
        p.paragraph_format.left_indent = Cm(0.4)
        b.rich(p, "•  " + item)
    b.p(
        "The remainder of the paper is organised as follows. Section 2 reviews the "
        "related literature. Section 3 describes the platform and the data, the "
        "extended pipeline and the online implementation. Section 4 presents the "
        "offline and the real-cell findings. Section 5 discusses them and states "
        "the working envelope and the limitations, and Section 6 concludes."
    )

    # ──────────────────────────────────── 2. Literature review ───────
    b.h1("2. Literature Review")
    b.p(
        "Physics-based detection for manipulators is a mature field. Haddadin, De "
        "Luca and Albu-Schäffer (2017) survey collision detection, isolation and "
        "identification and establish the residual observer as the canonical tool. "
        "Li, Han and Xiong (2020) place a force/torque sensor at the bedplate and "
        "detect collisions from the resulting residual, while Zhang, Chen and Zou "
        "(2024) use an external torque observer for the same purpose. "
        "Katsampiris-Salgado et al. (2024) address high-payload collaborative "
        "assembly, where the model error itself becomes the limiting factor. The "
        "common weakness of this family is that whatever the model does not "
        "represent — friction, payload, joint elasticity — is indistinguishable from "
        "a fault, and whatever the model filters well is also filtered away when it "
        "is the fault."
    )
    b.p(
        "Data-driven detection for multivariate time series is surveyed by Darban, "
        "Webb, Pan, Aggarwal and Salehi (2024). Reconstruction-based autoencoders "
        "are the dominant design: Malhotra, Vig, Shroff and Agarwal (2015) "
        "introduced LSTM networks for time-series anomaly detection, and Malhotra, "
        "Ramakrishnan, Anand, Vig, Agarwal and Shroff (2016) extended the idea to a "
        "multi-sensor encoder–decoder. Park, Hoshi and Kemp (2018) applied an "
        "LSTM-based variational autoencoder to robot-assisted feeding, one of the "
        "few studies to report on a physical robot. These methods need no fault "
        "labels, but they inherit whatever bias the training distribution carries."
    )
    b.p(
        "Hybrid approaches attempt to combine the two. Liu et al. (2025) constrain "
        "an LSTM autoencoder with mechanism knowledge for pump operations, but do "
        "not study the fusion of two separate models' scores. Huang, Chen, Deng and "
        "Huang (2024) apply graph attention and a Transformer to multivariate "
        "anomaly detection without any physics-based preprocessing stage. Correia, "
        "Goos, Klein, Bäck and Kononova (2024) survey online model-based anomaly "
        "detection specifically, and identify threshold selection and non-stationary "
        "operating conditions as open research challenges — precisely the two "
        "problems that this paper measures on hardware."
    )
    b.p(
        "The transfer of a model calibrated on simulated or recorded data to a "
        "physical robot is studied mostly in the control and reinforcement learning "
        "literature, where it is known as the reality gap (Zhao, Queralta and "
        "Westerlund, 2020). For anomaly detection the equivalent question — does the "
        "decision threshold transfer? — is seldom asked, because most studies never "
        "define an operating threshold at all, reporting instead the best achievable "
        "F1 over all thresholds. That statistic requires the labels it is supposed "
        "to predict and therefore does not exist online. The present study makes "
        "this gap explicit and closes it with a causal threshold rule, then measures "
        "what happens when the rule is applied to real hardware."
    )

    # ───────────────────────────────────────────── 3. Method ─────────
    b.h1("3. Method")

    b.h2("3.1. Platform And Data Collection")
    b.p(
        "Experiments were carried out on a robotic quality inspection platform "
        "consisting of a Universal Robots UR10e collaborative robot (six joints, "
        "12.5 kg payload) mounted on a Festo linear axis, together with an "
        "inspection chassis frame, shown in Figure 1. The target application is "
        "automated visual quality inspection of automotive body panels. A "
        "force/torque sensor (FTS) is mounted at the end effector. The cell runs "
        "ROS 2 Humble with MoveIt 2, and the recorded motions were produced by "
        "eleven different motion planning algorithms so that the normal operating "
        "distribution is not dominated by a single planner."
    )
    b.figure(
        "fig1_platform.png",
        "Robotic Quality Inspection Platform | UR10e robot on a Festo linear "
        "axis, with the inspection chassis frame on the right.",
        wide=False, width_cm=6.2)
    b.p(
        "The source dataset is a 79-day export of joint states and wrench "
        "measurements, 1,124,432 samples nominally at 500 Hz. Model training was "
        "performed on a laptop with an Intel Core i7-13650HX processor and an "
        "NVIDIA RTX 4060 GPU; the deployment workstation is an Intel Core i9-14900 "
        "with 64 GB of memory and an NVIDIA RTX 4000 Ada GPU."
    )
    b.p(
        "The study uses no human subjects and required no ethics committee "
        "approval. All measurements were collected on laboratory equipment of the "
        "ESOGÜ Intelligent Systems Application and Research Centre, and the "
        "principles of research and publication ethics were observed throughout. "
        "The anomalies of the commissioning campaign were provoked deliberately by "
        "trained operators at reduced speed, with the robot's own protective-stop "
        "function active at all times; one of them, reported in Section 4.7, "
        "involved an operator's hand being caught between two links, and the "
        "protective stop halted the robot without injury."
    )

    b.h2("3.2. Data Investigation And Preparation")
    b.p(
        "The first stage of the audit was the data itself, and it produced the "
        "largest single correction. The export is not globally ordered by "
        "timestamp. Measured over the 1,124,432 samples, 321,841 consecutive pairs "
        "(28.6 %) are discontinuous, 96,107 of them with a negative or zero time "
        "step; the median step is the expected 2.00 ms but the range spans from "
        "−5.5·10⁹ ms to +3.6·10⁹ ms. Three quarters of the sliding windows "
        "(33,700 of 44,974, 74.9 %) therefore splice together two recording "
        "sessions that may be hours or days apart."
    )
    b.p(
        "This matters twice. It puts an arbitrary jump in the middle of the dynamics "
        "an LSTM autoencoder is asked to learn, and it corrupts the joint "
        "acceleration. Acceleration is obtained from the velocity channel by a "
        "Savitzky–Golay derivative (window 51, order 3) (Savitzky and Golay, 1964), "
        "which is a convolution: across a splice it produces a non-physical spike. "
        "Since acceleration enters the inverse dynamics directly through the M(q)q̈ "
        "term, the error propagates into the residual. Measured on the same 200,000 "
        "samples, the standard deviation of q̈ derived across splices is two to six "
        "times larger than when the derivative is taken inside continuous runs."
    )
    b.p(
        "The correction is to segment the export into physically continuous runs. "
        "Two consecutive samples belong to the same run only if the time step lies "
        "in (0, 8 ms] and the joint displacement stays below what the joint "
        "velocity limit allows over that step; the second condition catches gaps "
        "that a plausible-looking timestamp hides. "
        "Runs shorter than 150 samples are discarded, and each run is resampled onto "
        "a uniform 2 ms grid. Table 1 gives the outcome. The reduction in sample "
        "count is not a loss: the discarded samples are exactly those between which "
        "no physical continuity exists."
    )
    b.table(
        "Outcome Of Data Preparation.",
        ["Quantity", "Value"],
        [["Raw exported samples", "1,124,432"],
         ["Physically continuous runs", "600"],
         ["Prepared samples", "342,480  (685 s of robot time)"],
         ["Interpolated samples", "5.02 %"],
         ["Run length (median / min / max)", "227 / 155 / 35,860"]],
        widths=[4.8, 3.55], align_right=[1])

    b.h2("3.3. Current-To-Torque Calibration")
    b.p(
        "The UR ROS 2 driver writes the motor current, in amperes, into the effort "
        "field of the joint state message; the field is filled directly from the "
        "actual current reported by the controller. The inverse dynamics model, in "
        "contrast, produces newton metres. A conversion coefficient is therefore "
        "required before the residual of Equation (2) can be formed at all."
    )
    b.p(
        "In a first rebuild this coefficient was obtained by regressing measured "
        "current on model torque, and verified by repeating the same regression. "
        "That verification is circular: the coefficient is chosen so as to satisfy "
        "the equality that is then tested. The coefficient was therefore re-derived "
        "from a source independent of the residual under test. When the robot is "
        "nearly at rest (|q̇| < 0.005 rad/s) the joint torque is dominated by "
        "gravity, so g(q) ≈ k·(i − i₀) with g(q) taken from the closed-form gravity "
        "vector rather than from the model being audited. Of 197,090 quasi-static "
        "samples, 65,697 entered the regression. Table 2 lists the result."
    )
    b.table(
        "Measured Current-To-Torque Coefficients And Their Sensitivity.",
        ["Joint", "Nm/A", "R²", "Meas.", "ρ −50 %", "ρ +50 %"],
        [["shoulder_pan", "10.522", "—", "no", "0.968", "0.995"],
         ["shoulder_lift", "10.522", "0.984", "yes", "0.162", "0.637"],
         ["elbow", "9.130", "0.993", "yes", "0.023", "0.758"],
         ["wrist_1", "3.409", "0.000", "no", "0.890", "0.968"],
         ["wrist_2", "3.409", "0.012", "no", "0.957", "0.993"],
         ["wrist_3", "3.409", "0.010", "no", "0.994", "0.999"]],
        widths=[3.3, 2.5, 2.3, 2.2, 3.7, 3.9], wide=True,
        align_right=[1, 2, 4, 5],
        note="R² is that of the quasi-static gravity regression and “Meas.” marks "
             "the joints on which the coefficient could be measured directly. The "
             "last two columns give the correlation between the intrinsic residual "
             "recomputed with the coefficient perturbed by ∓50 % and the nominal "
             "one, over all 342,480 prepared samples. Coefficients that could not "
             "be measured are copied within the same module family or scaled by "
             "the torque-limit ratio, and the node warns about them at start-up.")

    b.p(
        "The consequence of the four assumed coefficients was measured rather than "
        "argued. Each coefficient was perturbed by ±50 % and the intrinsic residual "
        "recomputed, which requires no further evaluation of the inverse dynamics "
        "because the model torque is stored separately from the measurement. On the "
        "four joints whose coefficient could not be measured, the perturbation "
        "leaves the residual almost unchanged in shape: the correlation with the "
        "nominal residual is between 0.890 and 0.999 and the effect is close to a "
        "pure rescaling, with the standard deviation changing by factors of 0.55 to "
        "1.45. Since the autoencoders standardise their inputs, a transformation of "
        "that form cannot alter the learned dynamics. On the two joints whose "
        "coefficient was measured, the same perturbation is destructive: the "
        "correlation falls to 0.162 for shoulder_lift and to 0.023 for the elbow at "
        "−50 %. The asymmetry is itself diagnostic. Where the current term carries "
        "the gravity torque that the model is meant to cancel, mis-scaling it "
        "destroys the cancellation; where the model contributes little, as at the "
        "wrists, the coefficient acts as a scale factor only. The assumed "
        "coefficients therefore affect the physical interpretation of a fault "
        "amplitude in newton metres rather than the detection performance."
    )

    b.h2("3.4. Hybrid Residual Decomposition")
    b.p(
        "The inverse dynamics model packages the Newton–Euler formulation of the "
        "UR10e in the Functional Mock-up Interface format (Blochwitz et al., 2011):"
    )
    b.equation("τ̂_model  =  M(q)·q̈ + C(q, q̇)·q̇ + g(q)")
    b.p("The total residual is the difference between measurement and model,")
    b.equation("r_total  =  τ_meas − τ̂_model")
    b.p(
        "the extrinsic residual is the projection of the end-effector wrench into "
        "joint space,"
    )
    b.equation("r_ext  =  J(q)ᵀ · F_FTS")
    b.p("and the intrinsic residual is the difference of the two,")
    b.equation("r_int  =  r_total − r_ext")
    b.p(
        "where M(q) is the inertia matrix, C(q, q̇) the Coriolis and centrifugal "
        "matrix, g(q) the gravity vector, J(q) the geometric Jacobian, τ_meas the "
        "measured joint torque and F_FTS the measured wrench. This decomposition "
        "makes r_int sensitive to internal faults such as motor degradation or an "
        "encoder glitch, and r_ext sensitive to external interaction such as a "
        "collision or a payload change. Because Equation (4) is a definition, the "
        "numerical consistency of the decomposition can be tested directly: over "
        "200,000 samples the identity holds to a residue of 7.1·10⁻¹⁵, and the "
        "analytical Jacobian agrees with a numerical one to machine precision."
    )
    b.p(
        "Equation (3) requires an agreement of reference frames whose violation "
        "produces no diagnostic. "
        "The UR e-Series driver rotates the wrench into the tool frame and publishes "
        "it with the tool frame identifier, whereas J(q) is defined in the base "
        "frame. Multiplying the two directly produces a wrong residual with no error "
        "message. Measured over 40,000 samples, the correlation between the correct "
        "and the unrotated extrinsic residual is −0.995 on the first joint — the "
        "sign inverts — and −0.143 across all channels; on the second joint the "
        "amplitude triples. Both force and moment halves are therefore rotated into "
        "the base frame before the Jacobian transfer, in the offline pipeline and in "
        "the online engine through the same code path."
    )
    b.p(
        "Two properties of the solver were measured and are not corrected inside it. "
        "Over 500 random joint poses the inertia matrix is asymmetric in every pose "
        "and its symmetric part is not positive definite, the friction vector is "
        "identically zero and no payload model is present. The solver had been "
        "validated against the physical robot and is treated as fixed. Its cost is "
        "therefore measured rather than hidden: on wrist_2 and wrist_3 the standard "
        "deviation of the model torque is only 10 % and 2 % of the measured torque "
        "(0.095 and 0.030 Nm against 0.959 and 1.246 Nm), so in those two channels "
        "Equation (2) reduces to the measurement itself."
    )
    b.p(
        "The absent friction term is added outside the solver rather than inside it, "
        "which leaves the validated model untouched while restoring a component that "
        "is large in a geared manipulator. Equation (2) becomes"
    )
    b.equation("r_total  =  τ_meas − τ̂_model − τ̂_f(q̇)", "2a")
    b.p("with a two-parameter Coulomb-plus-viscous term per joint,")
    b.equation("τ̂_f  =  F_c · tanh(q̇ / ε) + F_v · q̇", "2b")
    b.p(
        "where tanh replaces the sign function because sign(q̇) is discontinuous at "
        "zero crossings and this robot spends most of its time at low speed; a sharp "
        "sign injects a step into the residual at every reversal, which is precisely "
        "the shape a detector is meant to flag. The smoothing width is ε = 0.02 rad/s. "
        "Coefficients are fitted by robust least squares (three iterations of 3σ "
        "clipping) on moving samples only and, decisively for the evaluation, "
        "exclusively on the training runs of the split defined in Section 3.7. "
        "Fitting them on all data would make the correction itself a leakage channel. "
        "Table 3 lists the result."
    )
    b.table(
        "Friction Coefficients, Fitted On Training Runs Only.",
        ["Joint", "F_c [Nm]", "F_v [Nm·s/rad]", "R²", "σ(r) before", "σ(r) after", "Δ"],
        [["shoulder_pan", "7.65", "35.76", "0.764", "13.20", "6.09", "−54 %"],
         ["shoulder_lift", "8.78", "26.19", "0.799", "11.91", "4.98", "−58 %"],
         ["elbow", "5.53", "19.97", "0.896", "7.66", "2.46", "−68 %"],
         ["wrist_1", "0.95", "1.29", "0.379", "1.41", "0.75", "−47 %"],
         ["wrist_2", "1.51", "1.58", "0.947", "1.53", "0.20", "−87 %"],
         ["wrist_3", "1.38", "2.05", "0.974", "1.90", "0.15", "−92 %"]],
        widths=[3.0, 2.1, 2.9, 1.9, 2.5, 2.4, 1.9], wide=True,
        align_right=[1, 2, 3, 4, 5, 6],
        note="σ(r) is the standard deviation of the total residual on moving samples "
             "(|q̇| > 0.02 rad/s) before and after the term is subtracted. The "
             "regression absorbs any velocity-correlated model error, not friction "
             "alone; these are not claimed to be the true tribological coefficients. "
             "The effect is largest exactly where the inverse dynamics contributes "
             "least: on wrist_2 and wrist_3 the residual loses 87 % and 92 % of its "
             "spread, so the admission above — that those channels carry no model "
             "information — no longer holds once the term is included. Whether that "
             "improves detection is a separate question, answered in Section 4.4."
    )
    b.p(
        "The online feature engine applies the identical expression from the same "
        "coefficient file, and the equivalence test of Section 3.9 is run with the "
        "term active."
    )

    b.h2("3.5. Dual LSTM Autoencoders")
    b.p(
        "Both models share a symmetric encoder–decoder architecture that compresses "
        "a sliding window of T = 100 samples (0.2 s) into a latent vector. A "
        "two-layer LSTM encoder reduces the input sequence to a fixed-size latent "
        "vector; the decoder repeats it over T steps and reconstructs it with "
        "symmetric LSTM layers. Dropout of 15 % is applied between layers. Both "
        "models are trained on normal operating data only, with a mean squared error "
        "loss, and the anomaly score of a window is its reconstruction error. Each "
        "model's own threshold is the 97th percentile of the reconstruction errors "
        "on the validation set."
    )
    b.p(
        "Windowing uses a stride of 25 samples, that is 75 % overlap. Two rules are "
        "enforced jointly: no window may span two runs, and windows containing "
        "samples invalidated by the derivative edge margin are dropped. Training and "
        "validation windows come from disjoint sets of physical runs (Section 3.7), "
        "so no sample is shared between them. Each model is trained five times with "
        "different seeds; every offline figure in this paper is reported as a mean "
        "over those five runs with its standard deviation, because a single training "
        "run cannot separate a difference between methods from the noise of "
        "initialisation. Table 4 reports the architectures and the training outcome."
    )
    b.table(
        "Model Architectures And Training Outcome (five seeds, mean ± s.d.).",
        ["", "Residual AE", "Raw AE"],
        [["Input channels", "12", "24"],
         ["Hidden / latent", "128 / 32", "256 / 64"],
         ["Parameters", "478,892", "1,907,032"],
         ["Train / val. windows", "6,756 / 2,162", "6,756 / 2,162"],
         ["Best epoch", "86 ± 41", "83 ± 42"],
         ["Best validation loss", "0.092 ± 0.016", "0.036 ± 0.005"],
         ["Threshold θ (P97)", "0.887 ± 0.412", "0.412 ± 0.038"]],
        widths=[2.95, 2.7, 2.7], align_right=[1, 2],
        note="Optimiser Adam (lr = 10⁻³, β = 0.9/0.999), batch size 256, gradient "
             "clipping at 1.0, ReduceLROnPlateau (factor 0.5, patience 8), early "
             "stopping with patience 25 over at most 300 epochs; each model trains "
             "in under a minute on the RTX 4060 laptop. Both validation losses are "
             "far below those of the reference study (0.181 and 0.320), and the "
             "reason is the composition of the validation set rather than a better "
             "model: the sixteen validation runs were chosen to lie inside the "
             "training joint range (Section 3.7), so validation resembles training "
             "more closely than an index cut does. The friction term does not "
             "explain it — the raw model never sees the residual channels and its "
             "loss is 0.036 with and without the term, while the residual model's "
             "loss is 0.092 with it against 0.062 without. A validation set that "
             "resembles training gives a narrow distribution, and θ is a percentile "
             "of that distribution; Section 4.7 measures what this costs on "
             "hardware. Across seeds the residual threshold varies by 46 % of its "
             "mean, the raw one by 9 %. Values here are five-seed means; Section 4.7 "
             "quotes the deployed seed, whose thresholds are 1.3723 and 0.3934.")

    b.h2("3.6. Score-Level Fusion And Causal Normalisation")
    b.p(
        "The normalised scores of the two models are combined by a weighted average,"
    )
    b.equation("S_fused  =  w_res · z_res  +  w_raw · z_raw")
    b.p(
        "with w_res + w_raw = 1. The weight is not tuned here: w_res = 0.95 is "
        "adopted a priori from the reference study and no evaluation set is "
        "consulted to choose it. This matters because the reference study selected "
        "that value as the argmax of Best F1 over its own test set; fixing it "
        "removes the selection entirely rather than moving it to another set. The "
        "sweep over [0, 1] in steps of 0.05 is still run, but Section 4.3 reports it "
        "as a sensitivity analysis, not as a search. The reference "
        "study normalises each score by a min–max transform whose bounds come from "
        "the fault-injected test set. That is not usable online for two reasons: the "
        "bounds contain future maxima, and those maxima are never seen in normal "
        "operation. Measured on that study's own test set, where the maxima are "
        "4341 and 7752, the entire normal operating range compresses to the order "
        "of 10⁻⁴, "
        "the fused threshold becomes stricter than either single threshold, and "
        "motor drift and encoder glitches are never detected at all."
    )
    b.p(
        "Both normalisations are affine, so the online engine implements a single "
        "code path,"
    )
    b.equation("z  =  (S − lo) / span")
    b.p(
        "and the bounds are taken from clean validation windows rather than from the "
        "test set. This keeps the published formula and reduces the deviation from "
        "the reference method to a change of bounding set only. The fused operating "
        "threshold extends the study's own rule to the fused score: the 97th "
        "percentile on clean validation windows. Offline this gives the expected "
        "3 % false-alarm rate. Section 4.7 shows that neither the bounds nor the "
        "threshold survive the move to hardware, and Section 3.11 gives the "
        "procedure that replaces them there."
    )

    b.h2("3.7. Data Split, Fault Injection And Evaluation Protocol")
    b.p(
        "The 600 physical runs are partitioned once, and every consumer reads the "
        "same partition: the friction fit of Section 3.4, the autoencoder training "
        "of Section 3.5, and the evaluation below. The partition is by run, not by "
        "row index. The reference study cut the recording at a fixed fraction and "
        "then built its test set from every window of the recording; because "
        "training used only the first part, four fifths of its evaluation windows "
        "were windows the models had already reconstructed. Measured on that "
        "arrangement, the reconstruction error of clean windows inside the training "
        "range was 0.007 against 0.486 outside it, a factor of seventy for both "
        "models alike."
    )
    b.p(
        "The run lengths make a naive split unusable: 521 of the 600 runs are "
        "shorter than one second, while the 32 longest carry 54 % of the samples and "
        "the single longest carries 10.5 %. Runs are therefore allocated in two "
        "passes. Runs long enough to host a fault are distributed by count so that "
        "each split receives its own independent fault events; the remaining short "
        "runs then fill the sample quota. Table 5 gives the outcome. Joint-range "
        "coverage was checked afterwards: validation and test lie inside the "
        "training range on all six joints."
    )
    b.table(
        "Run-Disjoint Split.",
        ["Split", "Runs", "Samples", "Share", "Windows", "Fault hosts"],
        [["Train", "489", "235,657", "68.8 %", "6,756", "74"],
         ["Validation", "16", "56,244", "16.4 %", "2,162", "16"],
         ["Test", "95", "50,579", "14.8 %", "1,504", "16"]],
        widths=[3.2, 2.4, 3.2, 2.6, 2.9, 3.6], wide=True,
        align_right=[1, 2, 3, 4, 5],
        note="Verified after allocation: zero samples are shared between any two "
             "splits once window overlap is accounted for, and no window reaches "
             "into a run belonging to another split. The clean-window score ratio "
             "between test and training falls from the seventy of the reference "
             "arrangement to 0.48.")
    b.p(
        "Ground-truth labels do not exist, so four synthetic fault scenarios are "
        "injected, listed in Table 6. Joint 3 (elbow) and joint 5 (wrist 2) are "
        "chosen because they carry the largest gravity load and the widest motion "
        "range. Faults are injected into every fault-hosting run of the target "
        "split separately, so each scenario contributes 16 independent events to the "
        "test set. Under the reference protocol the mask was a single contiguous "
        "block over the concatenated recording, which gave one event per scenario "
        "and, worse, tied each scenario to a fixed position in the recording. A "
        "window is labelled anomalous when at least 10 % of its samples overlap the "
        "fault mask."
    )
    b.table(
        "Synthetic Fault Scenarios.",
        ["Scenario", "Affected channel", "Amplitude", "Extent"],
        [["Motor drift", "Joint 3 torque, linear ramp", "15 Nm", "30–38 %"],
         ["Collision", "All FTS channels, Gaussian pulse", "30 N",
          "centre 65 %"],
         ["Encoder glitch", "Joint 5 position, step", "1.5 rad", "from 92 %"],
         ["Sensor noise", "All FTS channels, Gaussian noise", "3.5 N",
          "15–23 %"]],
        widths=[2.0, 3.45, 1.5, 1.4],
        note="Extent is expressed as a fraction of the run the fault is injected "
             "into. The amplitudes above are those of the reference study and apply "
             "to the measured channels.")
    b.p(
        "Two injection protocols are evaluated. The inherited protocol perturbs each "
        "representation space separately with its own amplitude, so the residual "
        "model receives a hand-chosen perturbation rather than one derived from the "
        "measurement. The physical protocol perturbs only the measured channels and "
        "recomputes the residual through the pipeline itself, so the disturbance in "
        "the residual is whatever the inverse dynamics and the Jacobian transfer "
        "produce. The calibration offset is frozen from the clean pipeline when the "
        "residual is recomputed; refitting it on injected data shifts the residual "
        "of untouched runs as well. Section 4.4 reports both protocols, because they "
        "do not agree."
    )

    b.h2("3.8. Verifying Where The Earlier Figures Came From")
    b.p(
        "Before extending a result it is worth establishing what the earlier "
        "numbers measured. The measurement-level checks of Sections 3.2 to 3.4 "
        "point at a consistent arrangement, and that arrangement is testable. The "
        "hypothesis is that the previously published figures come from a pipeline "
        "with mixed units: "
        "τ_meas in amperes, τ̂_model in newton metres, r_ext in newton metres and "
        "computed with an unrotated wrench, the three subtracted directly, and fault "
        "amplitudes specified in newton metres added as bare numbers to ampere "
        "channels. A further finding compounds it: the residual generator zeroes the "
        "model torque on joints whose fit quality falls below a threshold, and under "
        "mixed units that guard fires on all six joints, silently reducing the "
        "residual to the measurement itself without any message."
    )
    b.p(
        "The hypothesis makes a falsifiable prediction: rebuilt in that "
        "configuration, the signature findings of the reference study should return "
        "— specifically the blindness of the residual model to sensor noise and the "
        "roughly one-quarter share of anomalies detected by the residual model "
        "alone. The reproduction is scripted and was rerun for this paper; its "
        "result is reported in Section 4.5."
    )

    b.h2("3.9. Online Implementation")
    b.p(
        "The online system differs structurally from the offline pipeline in three "
        "ways, and its design is entirely the management of those three differences. "
        "Causality: nothing computed over a whole dataset may be used live. Time "
        "budget: at 500 Hz there are 2 ms per sample and, with a stride of 25, 50 ms "
        "per decision. Continuity: the offline pipeline may segment the data into "
        "runs, whereas the live system notices an interruption only after it has "
        "happened and must reset its state. Figure 2 shows the resulting data path."
    )
    b.figure(
        "fig2_architecture.png",
        "Data Path Of The Online Detector | every stage between the robot and "
        "the alarm runs once per decision, twenty times a second.")
    b.p(
        "The feature engine is the sample-by-sample equivalent of the offline "
        "computation. It keeps a 51-sample ring buffer and, for every new sample, "
        "emits the features of the sample at the centre of the buffer: the "
        "Savitzky–Golay derivative applied as an inner product with precomputed "
        "coefficients, the current-to-torque conversion, the inverse dynamics "
        "evaluation, the base-frame rotation and the Jacobian transfer. The centred "
        "derivative costs a structural delay of 25 samples (50 ms), which is the "
        "price of a noise-free acceleration and is identical to the offline "
        "definition."
    )
    b.p(
        "Equivalence was verified rather than assumed. The prepared dataset was "
        "pushed through the online engine sample by sample and compared with the "
        "offline feature file: over 8 runs and 61,524 samples the largest absolute "
        "differences are 1.4·10⁻¹⁴ for torque, exactly zero for model torque, "
        "8.0·10⁻¹⁵ for the total residual, 2.7·10⁻¹⁵ for the extrinsic and "
        "8.9·10⁻¹⁵ for the intrinsic residual. These are floating-point rounding "
        "levels, so the deployed node sees the features the models were trained on, "
        "not merely similar ones."
    )
    b.p(
        "Two optional rules were added on top of the reference method. An adaptive "
        "rule fires when the fused score exceeds median + k·1.4826·MAD of the last "
        "600 decisions, where 1.4826·MAD is a robust estimate of the standard "
        "deviation (Leys, Ley, Klein, Bernard and Licata, 2013); it exists because "
        "an absolute threshold cannot see a fault that raises the score a "
        "hundredfold inside a run whose baseline sits far below that threshold. The "
        "scale is floored so that a nearly constant baseline does not make the rule "
        "hypersensitive, the baseline is frozen while an alarm lasts, and the freeze "
        "expires after 3 s so that a long alarm is treated as a regime change rather "
        "than as an attack. The second rule requires two consecutive decisions "
        "(100 ms) before an alarm is raised. Both can be switched off with a single "
        "parameter, which restores the reference behaviour exactly."
    )
    b.p(
        "The detector is packaged as a ROS 2 node (Macenski, Foote, Gerkey, "
        "Lalancette and Woodall, 2022) that keeps the whole computation in a "
        "ROS-independent core; the node is a thin wrapper around it, so replay tests "
        "exercise the class the node actually runs. The models are exported to ONNX "
        "and executed with ONNX Runtime, with two threads per session so that a "
        "greedy inference cannot starve the 500 Hz callback loop. Every decision is "
        "written to a comma-separated log and every alarm to a line-buffered "
        "JSON-lines event log; these two files are the raw material from which the "
        "real-cell thresholds of Section 4.7 were re-measured."
    )
    b.p(
        "Deployment surfaced three connection-layer measurements, all of them "
        "silent. The wrench topic advertised in the controller configuration is not "
        "the one the running broadcaster uses, and subscribing to the wrong name "
        "yields no wrench and therefore no score at all. The joint state topic "
        "carries three different name sets from three publishers, of which the UR "
        "data arrives in a seven-element set in scrambled order at about 112 non-UR "
        "messages per second; the joint mapping must therefore be resolved per "
        "message and cached per name set, since resolving it once and applying it "
        "blindly reads five of six channels from the wrong joints. Finally the "
        "measured sample rate is 495 Hz rather than 500 Hz, which the fixed-step "
        "derivative tolerates at the one per cent level."
    )

    b.h2("3.10. Operator Interface")
    b.p(
        "The detector was integrated into the laboratory's existing web dashboard as "
        "a third tab. A collector subscribes to the node's decision, alarm and score "
        "topics, keeps a 60-second ring buffer and pushes it to the browser at 5 Hz. "
        "Each collector runs its own single-threaded executor: in a first version "
        "both collectors used the shared global executor, two threads then raced on "
        "the same wait set, the node disappeared from the graph, the subscriptions "
        "received nothing and no error was printed."
    )
    b.p(
        "Three interface decisions follow directly from measurements rather than "
        "from taste. Alarms are latched on the rising edge and held for five "
        "seconds, because the clearest verified real event of the campaign lasted "
        "0.25 s — five decisions — and level sampling at 5 Hz misses it entirely; in "
        "a controlled test with synthetic 0.25 s alarms, edge capture saw four "
        "events out of four and level sampling only two. The vertical axis is "
        "linear and framed on the live threshold rather than on a fixed range, "
        "because the score scale changes when the detector is recalibrated and a "
        "fixed frame leaves the whole trace flat against one edge; a logarithmic "
        "axis was rejected because it visually flattens the region above the "
        "threshold, which is the only region that matters. In the event table the "
        "peak score is the primary column and the entry score is secondary, "
        "because an entry value read without its peak led to a misjudged event "
        "during commissioning — event 3 of Table 14 enters at 13.75 and peaks at "
        "25.04."
        "file that the detector never touches, and they accumulate the labelled set "
        "on which Section 4.7 rests. Both the axis and the displayed per-model "
        "thresholds are read from the live decision stream rather than compiled "
        "in, so a recalibration cannot leave the display describing a "
        "configuration that is no longer running."
    )

    b.h2("3.11. Threshold Calibration On The Cell")
    b.p(
        "The offline procedure of Section 3.6 fixes (lo, span) and θ from clean "
        "validation windows. Section 4.7 shows that these do not transfer: the "
        "models are far outside their training distribution during ordinary "
        "operation of the real cell, so a threshold placed at the 97th percentile of "
        "an offline distribution sits below the cell's normal range. The deployed "
        "system therefore re-measures the same three quantities on the cell, using "
        "the same formula and changing only the set they are estimated from — clean "
        "decisions logged on the robot instead of clean validation windows."
    )
    b.p(
        "Two details are forced by measurement. First, the scale is the 97th "
        "percentile of each model's own clean scores, not the min–max range of the "
        "reference formula: on the cell the min–max range is set by the payload "
        "plateau of Section 4.7, which inflates the residual span to 30.8 and makes "
        "an event of ordinary magnitude invisible. Second, the calibration run must "
        "cover every task the cell performs. A calibration taken on the inspection "
        "cycle alone left the pick-and-place task outside its range and the detector "
        "alarmed on normal transport. The deployed configuration is calibrated on "
        "15,627 clean decisions spanning both tasks."
    )

    # ──────────────────────────────────────────── 4. Findings ────────
    b.h1("4. Findings")

    b.h2("4.1. Offline Performance Of The Extended Pipeline")
    b.p(
        "The extended pipeline is evaluated on 6,016 test windows of which 553 "
        "(9.2 %) are anomalous, drawn from 95 runs that contributed nothing to "
        "training, to the friction fit, or to the choice of any threshold. Every "
        "figure below is a mean over five training seeds with its standard "
        "deviation. Table 7 gives the result and Figure 3 shows the "
        "corresponding curves."
    )
    b.table(
        "Overall Performance On The Run-Disjoint Test Set (five seeds, mean ± s.d.).",
        ["Model", "AUC", "PR-AUC", "Best F1"],
        [["Residual LSTM AE", "0.878 ± 0.005", "0.653 ± 0.009", "0.628 ± 0.007"],
         ["Raw LSTM AE", "0.870 ± 0.008", "0.668 ± 0.009", "0.697 ± 0.001"],
         ["**Fusion (0.95/0.05)**", "**0.969 ± 0.004**", "**0.857 ± 0.009**",
          "**0.821 ± 0.006**"],
         ["Fusion MAX", "0.966 ± 0.004", "0.853 ± 0.010", "0.784 ± 0.012"],
         ["Residual norm threshold", "0.758 ± 0.000", "0.485 ± 0.000",
          "0.495 ± 0.000"],
         ["Isolation Forest", "0.753 ± 0.008", "0.418 ± 0.008", "0.415 ± 0.008"],
         ["One-Class SVM", "0.844 ± 0.001", "0.618 ± 0.001", "0.676 ± 0.000"]],
        widths=[2.9, 1.9, 1.9, 1.65],
        align_right=[1, 2, 3],
        note="Figures from the earlier study are not tabulated beside these, "
             "because that evaluation used a different and much larger window set "
             "under a protocol whose evaluation windows overlapped its training "
             "windows; Section 4.5 covers the comparison in one place. That study "
             "also lists an OR row with values identical to its MAX row. OR "
             "is a decision-level combination: it produces a binary outcome, not a "
             "ranking, so AUC, PR-AUC and Best F1 are undefined for it and the row "
             "in fact reports the MAX score. Its operating-point behaviour appears "
             "in Section 4.2 instead. Best F1 is the maximum over all thresholds "
             "and therefore requires the labels it is meant to predict; it is "
             "listed only for comparability, which is why Section 3.6 defines a "
             "causal operating threshold. Baselines are fitted on the clean "
             "training runs.")
    b.figure(
        "fig3_pr_roc.png",
        "Precision–Recall And ROC Curves On The Run-Disjoint Test Set | 6,016 test "
        "windows, 9.2 % anomalous, seed 1 of 5. Fusion dominates both single models "
        "on both curves; the margin is larger on the precision–recall panel because "
        "ROC is optimistic under class imbalance.")
    b.p(
        "At the causal operating point the fused detector reaches precision "
        "0.764 ± 0.086, recall 0.831 ± 0.055 and F1 0.791 ± 0.023, against "
        "F1 0.627 ± 0.006 for the raw model alone and 0.614 ± 0.016 for the residual "
        "model alone. The gain over the better single model is ΔF1 = +0.164. The "
        "clean-window score ratio between test and training runs is 0.48 for both "
        "models, against the factor of seventy that the reference split produced; "
        "the optimism that the earlier arrangement introduced is gone."
    )
    b.p(
        "The precision spread of ±0.086 is not noise in the models but in the "
        "threshold. Its value across the five seeds varies by 46 % of its mean, "
        "because it is the 97th percentile of a validation distribution that the "
        "validation distribution narrowed by the split design of Section 3.7. The "
        "ranking metrics, which do not depend on a threshold, vary by less than "
        "0.01."
    )

    b.h2("4.2. Fault-Type Asymmetry And Complementarity")
    b.p(
        "Table 8 breaks the result down by fault type, and reports AUC and Best F1 "
        "separately because they are not the same quantity and the distinction "
        "changes the reading: on motor drift the residual model has AUC 0.912 but "
        "Best F1 0.669, and quoting the second as if it were the first understates "
        "the ranking quality by a wide margin."
    )
    b.table(
        "Detection By Fault Type Under Both Injection Protocols (AUC, five seeds).",
        ["Scenario", "Inherited: residual", "Inherited: raw",
         "Physical: residual", "Physical: raw"],
        [["Motor drift", "0.912 ± 0.013", "0.652 ± 0.014", "0.846 ± 0.021",
          "0.652 ± 0.014"],
         ["Collision", "1.000 ± 0.000", "1.000 ± 0.000", "1.000 ± 0.000",
          "1.000 ± 0.000"],
         ["Encoder glitch", "1.000 ± 0.000", "0.832 ± 0.041", "0.783 ± 0.052",
          "0.832 ± 0.041"],
         ["Sensor noise", "0.682 ± 0.006", "1.000 ± 0.000", "1.000 ± 0.000",
          "1.000 ± 0.000"]],
        widths=[2.7, 3.1, 2.9, 3.1, 2.9], wide=True,
        align_right=[1, 2, 3, 4],
        note="The raw columns are identical between the two protocols because the "
             "raw model reads the measured channels, which both protocols perturb "
             "the same way. Only the residual columns move, and they move a great "
             "deal — the interpretation is in Section 4.4.")
    b.figure(
        "fig4_fusion_value.png",
        "Value Of The Fusion | (a) ranking quality against the residual weight, which "
        "is fixed a priori at 0.95 rather than selected. (b) detection "
        "by fault type under the inherited injection protocol.")
    b.p(
        "Under the inherited protocol the asymmetry that motivates the fusion is "
        "present and clean. On the encoder glitch the residual autoencoder is "
        "perfect and the raw model reaches 0.832; on sensor noise the roles invert, "
        "the residual model falling to 0.682 while the raw model is perfect. At "
        "window level, with each model using its own 97th-percentile threshold, "
        "27.6 % of anomalous windows are detected by the residual model only and "
        "28.0 % by the raw model only, against 24.1 % and 25.9 % in the reference "
        "study; 15.9 % are caught by neither. The recall of the OR combination is "
        "0.841. Table 9 gives the decision-level ablation."
    )
    b.table(
        "Decision-Level Ablation At The Operating Point (five seeds, mean ± s.d.).",
        ["Configuration", "Precision", "Recall", "F1"],
        [["Raw only  (w_res = 0.00)", "0.704 ± 0.015", "0.565 ± 0.002",
          "0.627 ± 0.006"],
         ["Residual only  (w_res = 1.00)", "0.697 ± 0.111", "0.562 ± 0.058",
          "0.614 ± 0.016"],
         ["OR  (decision level)", "0.701 ± 0.047", "0.841 ± 0.058",
          "0.762 ± 0.013"],
         ["**Fusion  (w_res = 0.95)**", "**0.764 ± 0.086**", "**0.831 ± 0.055**",
          "**0.791 ± 0.023**"]],
        widths=[3.05, 1.95, 1.75, 1.6], align_right=[1, 2, 3],
        note="Inherited injection protocol. The gain over the better single model "
             "is ΔF1 = +0.164; the reference study reports +0.161 at a different "
             "absolute level and under a protocol with overlapping windows.")

    b.h2("4.3. Sensitivity To The Fusion Weight")
    b.p(
        "The weight is fixed a priori at w_res = 0.95 (Section 3.6), so the sweep "
        "below measures sensitivity rather than selecting anything. Figure 4(a) "
        "shows a broad plateau over roughly [0.25, 0.95], reproducing the stable "
        "region reported by the reference study, with a sharp fall as the raw model "
        "is eliminated entirely. Under the inherited protocol the fused PR-AUC at "
        "w_res = 0.95 is 0.857 ± 0.009 against 0.653 ± 0.009 for the pure residual "
        "detector, so a five per cent weight on the raw model is associated with "
        "+0.204 PR-AUC. The margin over the better single model, "
        "+0.189 ± 0.009, is positive in every one of the five seeds."
    )
    b.p(
        "Two independent checks confirm that both models really run. Structurally, "
        "every rule of the reference method maps onto the deployed decision path. "
        "Operationally, the scoring method of each ONNX session was wrapped and "
        "counted over a replay of 6 runs and 5 scenarios: 13,160 decisions produced "
        "13,160 residual-model calls and 13,160 raw-model calls, so neither model is "
        "cached or skipped."
    )

    b.h2("4.4. What The Fusion Margin Depends On")
    b.p(
        "The margin reported above is a property of the inherited injection "
        "protocol, not of the two representation spaces. The protocol perturbs the "
        "residual channels directly with an amplitude chosen by hand, and one of "
        "those choices carries the entire complementarity result. For sensor noise "
        "the reference study adds 3.5 N to the wrench and, separately, 0.08 Nm to "
        "the extrinsic residual, describing the second as deliberately small so that "
        "the scenario 'affects the raw signal only'. Physically, a 3.5 N wrench "
        "disturbance propagates into the residual through r_ext = J(q)ᵀF; measured "
        "on this robot it produces a standard deviation of 4.1 Nm in the residual "
        "channels, fifty-one times the amplitude that was injected by hand. The "
        "inverse dynamics does not filter the wrench, it transforms it."
    )
    b.p(
        "The comparison is not one-sided. For the collision scenario the "
        "hand-chosen residual amplitude of 40 Nm agrees with the physical "
        "propagation, whose peak is 80.5 Nm — the same order. Sensor noise is the "
        "one scenario where the two disagree, and it is the scenario the reference "
        "study names as the test of complementarity."
    )
    b.p(
        "Table 10 crosses the two injection protocols with the presence of the "
        "friction term, five seeds each. The result is unambiguous: the fusion "
        "margin exists under the inherited protocol in both friction conditions and "
        "vanishes under the physical protocol in both."
    )
    b.table(
        "Fusion Margin Under Two Injection Protocols And Two Residual Definitions "
        "(PR-AUC, five seeds, mean ± s.d.).",
        ["Configuration", "Residual", "Raw", "Fusion", "Margin"],
        [["Inherited / no friction", "0.620 ± 0.007", "0.668 ± 0.009",
          "0.829 ± 0.007", "**+0.161 ± 0.005**"],
         ["Inherited / friction", "0.653 ± 0.009", "0.668 ± 0.009",
          "0.857 ± 0.009", "**+0.189 ± 0.009**"],
         ["Physical / no friction", "**0.750 ± 0.017**", "0.668 ± 0.009",
          "0.747 ± 0.016", "−0.003 ± 0.001"],
         ["Physical / friction", "0.730 ± 0.014", "0.668 ± 0.009",
          "0.727 ± 0.014", "−0.003 ± 0.001"]],
        widths=[3.6, 2.6, 2.6, 2.6, 3.0], wide=True,
        align_right=[1, 2, 3, 4],
        note="Margin is the fused PR-AUC minus the better of the two single models. "
             "The raw column is constant by construction, which is also a check "
             "that the ablation is controlled. Under the inherited protocol the "
             "margin is positive in all five seeds of both rows; under the physical "
             "protocol it is −0.003 in all five seeds of both rows, that is, the "
             "five per cent raw weight costs a little and returns nothing. "
             "Complementarity at window level collapses in the same direction: "
             "residual-only detections fall from 27.6 % to 4.8 % and raw-only from "
             "28.0 % to 0.1 %.")
    b.p(
        "The same table answers a second question. Adding the friction term improves "
        "the residual model under the inherited protocol by +0.034 ± 0.009 PR-AUC, "
        "in the same direction in all five seeds, but under the physical protocol "
        "the change is −0.020 ± 0.027, which is not significant and points the other "
        "way. The friction term's demonstrable benefit is therefore physical "
        "fidelity — the 87 % and 92 % reduction of wrist residual spread in Table 3, "
        "and fault amplitudes that remain interpretable in newton metres — not "
        "detection of these four synthetic faults. It also carries a measured cost: "
        "the seed-to-seed spread of the operating threshold rises from 17 % to 46 % "
        "of its mean in both protocols."
    )
    b.p(
        "None of this says the raw model is useless. It says that these four "
        "synthetic scenarios, once injected consistently, do not discriminate "
        "between the two representation spaces. Section 4.7 reports what the two "
        "models did during real events on hardware, which is the only evidence here "
        "that bears on real faults."
    )

    b.h2("4.5. Where The Earlier Figures Came From")
    b.p(
        "This subsection collects, in one place, the comparison with the study "
        "this work extends. Rebuilt in the mixed-unit configuration of Section 3.8, "
        "the pipeline reproduces the earlier figures point by point (Table 11), "
        "which identifies the arrangement they were produced under and therefore "
        "the physical scale on which they should be read."
    )
    b.table(
        "Reproduction Of The Earlier Pipeline.",
        ["Quantity", "Rebuild", "Ref.", "Δ"],
        [["Test windows", "179,888", "179,896", "8"],
         ["Fault windows", "14,401 (8.01 %)", "14,402 (8.0 %)", "1"],
         ["Residual-only share", "23.9 %", "24.1 %", "0.2 pt"],
         ["Sensor noise — residual", "0.267", "0.272", "0.005"],
         ["OR recall", "0.841", "0.854", "0.013"],
         ["Recall — raw", "0.601", "0.613", "0.012"],
         ["Encoder glitch — residual", "0.975", "0.986", "0.011"],
         ["Raw model AUC", "0.943", "0.952", "0.009"],
         ["Fusion MAX — AUC", "0.974", "0.976", "0.002"]],
        widths=[2.95, 1.95, 1.85, 1.60], align_right=[1, 2, 3],
        note="The decisive row is sensor noise. A residual model scoring 0.267 on "
             "that scenario can only arise on a pipeline where the residual has lost "
             "its physical meaning, that is where amperes and newton metres are "
             "subtracted directly. Agreement of the window counts to 8 and 1 windows "
             "in 179,896 confirms that the same raw dataset is being split the same "
             "way.")
    b.p(
        "The structural claim of the earlier study is unaffected. Complementarity "
        "under its own injection protocol is measured on both pipelines — 23.9 % "
        "residual-only on the reproduction and 27.6 % on the extended one — and the "
        "fusion margin survives on both. What changes is the physical scale on which "
        "the numbers should be read, and, separately, what happens to that margin "
        "once the faults are injected physically (Section 4.4). Establishing the "
        "scale is what made the rest of this work possible: fault amplitudes in "
        "newton metres, a friction term with meaningful coefficients, and a residual "
        "whose magnitude can be compared with a real contact force all require it."
    )

    b.h2("4.6. Latency And Compute Budget")
    b.p(
        "Per sample, the inverse dynamics evaluation costs 22 µs and the Jacobian "
        "about 30 µs, together 1.5 % of the 500 Hz budget. Per decision, the two "
        "ONNX forward passes cost 5.9 ms, that is 12 % of the 50 ms budget, measured "
        "on CPU in offline replay, and 8.5 ms (17 % of the budget) on the cell, "
        "where the ROS 2 executor and the logging layer share the same cores; on "
        "the deployment workstation with the CUDA provider both fall further. The "
        "detection latency is the sum of the 50 ms filter delay and the "
        "50 ms decision period, plus one more decision period when the two-consecutive "
        "rule is active, giving 120–150 ms end to end. At a tool speed of 0.25 m/s "
        "that corresponds to about 3.75 cm of additional travel."
    )
    b.p(
        "An end-to-end replay of the prepared data through the deployed class, over "
        "the five test runs long enough to carry a fault (27,135 samples, 0.9 "
        "minutes of robot time), gives a false-alarm rate of 0.1 % on clean data. "
        "Collision and sensor noise are detected in 5 of 5 runs with median "
        "latencies of 60 ms and 70 ms; motor drift in 2 of 5 at 2,905 ms; the "
        "encoder glitch in none. The latency split is structural — within a 200 ms "
        "window a slow ramp is almost a constant offset, and an autoencoder "
        "reconstructs a constant offset without difficulty — but the encoder-glitch "
        "result is not a latency effect. It is the physical-injection outcome of "
        "Table 8 seen at the decision level: a 1.5 rad step on joint 5 propagates "
        "into the residual through the Jacobian rather than as the hand-chosen "
        "8 Nm of the inherited protocol, and the propagated disturbance does not "
        "cross the threshold. The sample is small and is reported as such: five "
        "runs is what the run-disjoint test split contains at that length."
    )

    b.h2("4.7. Commissioning On The Real Robot")
    b.p(
        "All results up to this point were obtained on the recorded dataset. The "
        "detector was then commissioned on the physical UR10e cell, covering the two "
        "production use cases of the platform — the autonomous visual inspection "
        "cycle and a pick-and-place scenario — together with runs in which "
        "anomalies were provoked deliberately. Every decision was written to disk, every alarm "
        "block to an event log, and every session to a provenance record naming the "
        "model checksums, the thresholds and the residual definition in force. That "
        "record is not bookkeeping: a first attempt ran with superseded models "
        "because the software package shipped them as defaults, and nothing else "
        "would have revealed it. Table 12 summarises how the online system "
        "deviates from the reference method."
    )
    b.table(
        "Deviations Of The Online System From The Reference Method.",
        ["Deviation", "Status", "Reason"],
        [["D1  Normalisation bounds and threshold measured on the cell",
          "forced", "offline values do not transfer; measured below"],
         ["D2  Operating threshold for the fused score",
          "forced", "the reference method reports Best F1 only, which needs the "
          "labels it predicts"],
         ["D3  Adaptive median + k·MAD rule", "optional, **off**",
          "helps offline, catches no verified event on hardware"],
         ["D4  Two consecutive decisions before alarm", "optional, on",
          "suppresses single-decision noise triggers"]],
        widths=[6.3, 2.4, 5.6], wide=True,
        note="D3 and D4 are switched off by a single parameter each, which restores "
             "the reference behaviour exactly. D1 keeps the published min–max "
             "structure and changes only the set the quantities are estimated from.")

    b.p(
        "**The offline threshold does not transfer.** The fused threshold derived "
        "from clean validation windows by the rule of Section 3.6 is θ = 0.855, "
        "while the median fused score of a clean run on the cell is 0.97: 58.2 % of "
        "all decisions crossed it. The scores are not anomalous in scale — under "
        "the cell-calibrated configuration of Table 13 the same clean operation "
        "sits between 1.7 and 35.3 times below the threshold — so the failure is "
        "located unambiguously in where the threshold comes from, not in the "
        "models."
    )
    b.p(
        "The cause is that the models are far outside their training distribution "
        "during ordinary operation. Measured against the mean reconstruction error "
        "of the offline validation set, the residual model runs at eleven times that "
        "value on the cell and the raw model at eighty-six times. Expressed against "
        "each model's own P97 threshold the asymmetry is sharper: the residual model "
        "sits at 0.93 of its threshold, which is well calibrated, while the raw "
        "model sits at 7.28. The consequence is a distortion of the fusion itself. "
        "With w_raw = 0.05 the raw model was contributing 29 % of the fused score at "
        "the median and dominating it at peaks, because its normalising constant was "
        "twenty-five times too small for the cell. Re-estimating both scales on the "
        "cell (Section 3.11) returns its share to 8–15 %, in line with its weight."
    )
    b.p(
        "One further correction was forced by measurement. The min–max form of the "
        "reference formula, applied to real-cell data, locks the residual span to "
        "30.8 because the calibration run contains a payload-carrying phase — an "
        "85-second plateau of fused score whose interior barely moves (median 14.20, "
        "p90 14.39, maximum 14.5). The inverse dynamics has no payload model, so "
        "carrying anything produces a sustained bias, and a min–max scale fitted "
        "over it makes ordinary events invisible. The 97th percentile of each "
        "model's clean scores is used instead."
    )
    b.p(
        "With the cell-calibrated configuration — residual span 25.92, raw span "
        "22.37, θ = 5.0, estimated from 15,627 clean decisions spanning both tasks — "
        "the detector produced no alarm across 14.3 minutes of ordinary operation "
        "(Table 13)."
    )
    b.table(
        "Clean Operation Under The Cell-Calibrated Configuration.",
        ["Run", "Duration", "Median", "Maximum", "Margin to θ", "Alarms"],
        [["Inspection cycle", "6.2 min", "0.055", "1.28", "3.9×", "**0**"],
         ["Provoked hand contact", "1.3 min", "0.023", "0.14", "35.3×", "0"],
         ["Pick and place", "6.9 min", "0.075", "2.92", "1.7×", "**0**"]],
        widths=[4.0, 2.6, 2.4, 2.6, 3.0, 2.4], wide=True,
        align_right=[1, 2, 3, 4, 5],
        note="Margin is θ divided by the run's peak. The second row is a provoked "
             "anomaly that the detector did not register; it is discussed below and "
             "is the working limit of the system.")

    b.p(
        "**Collisions are detected.** In a session of three runs totalling 10.0 "
        "minutes, two collisions were provoked against the vacuum gripper. Three "
        "alarm blocks were raised; the operator labelled all three as genuine "
        "through the interface, and no unlabelled block occurred. Table 14 lists "
        "them and Figure 5 shows the operator interface immediately after the "
        "strongest one."
    )
    b.table(
        "Provoked Collisions On The Physical Cell, All Operator-Confirmed.",
        ["Event", "Duration", "Peak", "Entry", "Trigger", "Residual", "Raw"],
        [["1", "3.65 s", "**11.30**", "6.63", "both", "170.1  (124×θ)",
          "175.9  (447×θ)"],
         ["2", "0.25 s", "**16.13**", "14.60", "both", "369.8  (270×θ)",
          "466.5  (1186×θ)"],
         ["3", "44.95 s", "**25.04**", "13.75", "both", "348.9  (254×θ)",
          "427.7  (1087×θ)"]],
        widths=[1.4, 2.0, 2.0, 1.8, 1.7, 4.4, 4.6], wide=True,
        align_right=[1, 2, 3, 5, 6],
        note="Residual and raw columns give the model's own reconstruction error "
             "and its ratio to that model's P97 threshold. Both models crossed "
             "their own thresholds on all three events, and all three were raised "
             "by the absolute rule with the adaptive rule disabled. Events 2 and 3 "
             "occur seven seconds apart at almost the same pose; the 44.95-second "
             "duration of the third is the robot remaining in contact, not a "
             "detection that failed to clear.")
    b.figure(
        "fig5_collision.png",
        "Operator Interface After A Provoked Collision | the 44.95-second block of "
        "event 3 is the wide plateau at 25; the narrow spike to its left is event 2, "
        "seven seconds earlier. The threshold line sits at 5. The breakdown panel "
        "gives each model's reconstruction error against its own threshold, and the "
        "event table shows all three blocks confirmed by the operator with no "
        "unlabelled block. Replayed from the logged decisions of the session.")
    b.p(
        "The separation is genuine rather than a threshold artefact. The weakest "
        "confirmed event peaks at 11.30 while the highest excursion in any clean run "
        "of that day reaches 4.14, so any threshold in the interval (4.14, 11.30) "
        "gives three detections and no false alarm. The deployed value of 5.0 sits "
        "at the lower end of that interval, deliberately, to leave margin for events "
        "weaker than those observed. Table 15 shows how the interval closes."
    )
    b.table(
        "Threshold Sweep Recounted From All Logged Decisions.",
        ["Threshold", "Clean blocks / min", "Confirmed events caught"],
        [["3.0", "0.00", "3 / 3  (fragmented)"],
         ["**5.0**  (deployed)", "**0.00**", "**3 / 3**"],
         ["8.0", "0.00", "3 / 3"],
         ["11.3", "0.00", "2 / 3"],
         ["16.0", "0.00", "1 / 3"],
         ["26.0", "0.00", "0 / 3"]],
        widths=[2.55, 2.9, 2.9], align_right=[1, 2],
        note="Recounted over 14.3 minutes of clean operation and 10.0 minutes "
             "containing the confirmed events, under the same two-consecutive-"
             "decision rule the node applies.")

    b.p(
        "**A low-amplitude contact is not detected, and no threshold recovers it.** "
        "The hand contact provoked in the second row of Table 13 raised the fused "
        "score to 6.8 times that run's own median, while the clean inspection cycle "
        "routinely reaches 26.7 times its own median. Expressed on the common scale, "
        "the peak of the provoked contact falls at the 75.6th percentile of ordinary "
        "inspection motion: it is smaller than three quarters of what the robot does "
        "normally. Four decision rules were evaluated against this event on the "
        "logged decisions — an absolute threshold on two different normalisations, "
        "the adaptive median + k·MAD rule swept over k ∈ [4, 32], and subtraction of "
        "a rolling median over windows of 10, 30 and 60 seconds. In every case the "
        "clean inspection run produced at least as many blocks as the run containing "
        "the event. The signal required to separate them is not present in these two "
        "scores."
    )
    b.p(
        "The working envelope of the deployed system is therefore explicit: events "
        "of collision magnitude are detected reliably, with peaks of 11 to 25 against "
        "a threshold of 5 and a separation factor of 3.9 over ordinary operation; "
        "slow low-amplitude contact is not."
    )

    b.h2("4.8. Operator Labelling")
    b.p(
        "The three events of Table 14 were classified through the interface of "
        "Figure 5 while the operator stood at the cell, and the classification is "
        "what makes the false-alarm count of Section 4.7 a measurement rather than "
        "an impression: the interface reports three events, three confirmed, zero "
        "false alarms and zero unlabelled. Labels are written atomically to a file "
        "the detector never reads, so a label can never influence a decision. "
        "Together with the per-session provenance record they make a logged session "
        "re-analysable months later, which is what allowed the threshold failure of "
        "Section 4.7 to be attributed to configuration rather than to the models."
    )
    b.p(
        "The breakdown panel in Figure 5 also records which model drove each "
        "decision. On all three confirmed collisions both models exceeded their own "
        "thresholds, the residual model by factors of 124 to 270 and the raw model "
        "by 447 to 1186. Together with the undetected hand contact of Section 4.7 "
        "this is the only evidence here on how the two spaces behave during "
        "genuine faults, and it does not resolve the question raised in Section "
        "4.4: events this large saturate both models, so they cannot discriminate "
        "between them, and the one event small enough to discriminate was seen by "
        "neither."
    )

    # ────────────────────────────────────────── 5. Discussion ────────
    b.h1("5. Discussion")

    b.h2("5.1. What The Extension Changed Offline")
    b.p(
        "The extended pipeline scores below the earlier study on the absolute "
        "metrics, and the reasons are measurable rather than mysterious. Fault "
        "amplitudes are now on the true physical scale, so a ramp specified as "
        "15 Nm is injected as 15 Nm rather than added as a bare number to a channel "
        "expressed in amperes. The dataset was cleared of discontinuities that a "
        "reconstruction model reads as anomalous. And the split is run-disjoint, so "
        "the evaluation no longer measures partly on windows the models had already "
        "reconstructed: the clean-window score ratio between test and training runs "
        "is 0.48, against the factor of seventy that the earlier arrangement "
        "produced."
    )
    b.p(
        "What survives the extension is the structure. The fusion still exceeds "
        "both single models and every baseline; the weight sweep is still broad "
        "and flat over the same interval; the complementarity ratios remain close "
        "to those of the earlier study, 27.6 % against 24.1 % residual-only and "
        "28.0 % against 25.9 % raw-only. Reporting five seeds rather than one also "
        "shows where the remaining uncertainty lives: the ranking metrics vary by "
        "less than 0.01 across seeds, while the operating-point precision varies "
        "by ±0.086, because the threshold is a percentile of a validation "
        "distribution narrowed by the split design rather than by the models."
    )

    b.h2("5.2. The Fusion Margin Is A Property Of The Injection Protocol")
    b.p(
        "The central offline finding of this paper is negative, and it concerns the "
        "framework's own evaluation rather than its architecture. Under the "
        "inherited injection protocol the fusion gains +0.189 ± 0.009 PR-AUC over "
        "the better single model, positive in every seed. Under a protocol in which "
        "the fault is applied only to the measured channels and the residual is "
        "recomputed by the pipeline, the same gain is −0.003 ± 0.001, also in every "
        "seed. Adding or removing the friction term does not change this in either "
        "direction (Table 10)."
    )
    b.p(
        "The mechanism is a single amplitude. The inherited protocol perturbs the "
        "residual space with a hand-chosen number, and for sensor noise that number "
        "is 0.08 Nm, described in the reference study as deliberately small so that "
        "the scenario affects the raw signal only. But the extrinsic residual is "
        "r_ext = J(q)ᵀF, a linear map of the wrench; it does not attenuate wrench "
        "noise, it transforms it. Propagated physically, 3.5 N of wrench noise "
        "produces 4.1 Nm of residual disturbance — fifty-one times the injected "
        "amplitude — and the residual model detects it perfectly. The complementarity "
        "that motivates the fusion was not tested by that scenario; it was defined "
        "by it."
    )
    b.p(
        "This should be read narrowly. The collision scenario's hand-chosen residual "
        "amplitude of 40 Nm agrees with the physical propagation, whose peak is "
        "80.5 Nm, so the reference study's choices are not uniformly arbitrary. And "
        "the finding does not show that a raw-signal model is useless — it shows "
        "that four analytic perturbations, once applied consistently, do not "
        "discriminate between these two representation spaces. The evidence that "
        "bears on real faults is in Section 4.7, where both models crossed their own "
        "thresholds on all three confirmed collisions."
    )

    b.h2("5.3. Why The Offline Threshold Does Not Transfer")
    b.p(
        "A threshold placed at the 97th percentile of an offline validation "
        "distribution declared most decisions on hardware to be alarms. The scores "
        "themselves are on an ordinary scale — under the cell-calibrated "
        "configuration the same clean operation sits between 1.7 and 35.3 times "
        "below the threshold (Table 13) — so the fault is in the estimator, not in "
        "the models."
    )
    b.p(
        "The measurement that explains it is the distance of the models from their "
        "training distribution during ordinary operation: eleven times the "
        "validation mean error for the residual model and eighty-six times for the "
        "raw model. A percentile of a distribution the deployment never occupies "
        "carries no information about the deployment. The asymmetry between the two "
        "models is the practically important part. The residual model, whose input "
        "has had the trajectory-dependent dynamics removed by the physical model, "
        "sits at 0.93 of its own threshold on the cell — well calibrated. The raw "
        "model, reading the signals directly, sits at 7.28. With a five per cent "
        "weight it was nevertheless contributing 29 % of the fused score, because a "
        "normalising constant twenty-five times too small turns a small weight into "
        "a large influence. This is a general trap in score-level fusion: the weights "
        "mean what they say only while the normalisers are valid."
    )
    b.p(
        "The split design itself contributes, and this is the trade-off at the "
        "centre of the extension. To keep validation and test inside the training "
        "joint range — necessary, or a high score on the cell would mean 'unseen "
        "pose' rather than 'anomaly' — the sixteen validation runs were selected for "
        "similarity to training. Validation loss falls accordingly, to 0.036 for the "
        "raw model against 0.320 in the reference study, and with it the width of "
        "the distribution θ is drawn from. The same choice that removed the leakage "
        "therefore produced a threshold too narrow for the cell. This is not a "
        "defect of the models: their loss is identical with and without the friction "
        "term, so it cannot be attributed to the residual definition either. It is a "
        "property of estimating an operating point on data selected to resemble "
        "training."
    )
    b.p(
        "The remedy used here keeps the published formula and changes only the set "
        "the quantities are estimated from: clean decisions logged on the cell, "
        "spanning every task the cell performs, with a robust percentile in place of "
        "the min–max range. The last detail is not cosmetic. Applied to real-cell "
        "data the min–max range is set by an unmodelled payload — the inverse "
        "dynamics has no payload term, so carrying a workpiece produces a sustained "
        "residual bias, an 85-second plateau whose interior varies by 0.3 — and a "
        "span fitted over it makes ordinary events invisible."
    )

    b.h2("5.4. What The Friction Term Bought")
    b.p(
        "Adding a Coulomb-plus-viscous term outside the solver removes 87 % and 92 % "
        "of the residual spread on wrist_2 and wrist_3, the two channels where the "
        "inverse dynamics previously contributed almost nothing and the residual was "
        "in practice the raw measurement. That is a real gain in physical fidelity "
        "and it makes fault amplitudes interpretable in newton metres on all six "
        "joints."
    )
    b.p(
        "It is not a detection gain. Under the inherited protocol the residual model "
        "improves by +0.034 ± 0.009 PR-AUC, consistent across seeds; under the "
        "physical protocol the change is −0.020 ± 0.027, not significant. It also "
        "costs threshold stability: the seed-to-seed spread of the operating "
        "threshold rises from 17 % to 46 % of its mean. The honest summary is that "
        "the term corrects the model, and that correcting the model does not by "
        "itself improve the detection of these particular synthetic faults. It is "
        "retained in the deployed configuration because real faults are not these "
        "four, and because a residual that means what it says is easier to reason "
        "about when something unexpected happens."
    )

    b.h2("5.5. The Working Envelope")
    b.p(
        "The deployed detector catches collisions and does not catch slow "
        "low-amplitude contact, and both halves of that sentence are measured. "
        "Three operator-confirmed collisions peaked between 11.30 and 25.04 against "
        "a threshold of 5.0, with no unlabelled alarm block in ten minutes and no "
        "alarm at all in fourteen minutes of ordinary operation; any threshold "
        "between 4.14 and 11.30 gives the same result. A deliberately provoked hand "
        "contact, by contrast, raised the fused score to 6.8 times that run's median "
        "while ordinary inspection motion routinely reaches 26.7 times its own — the "
        "provoked event sits at the 75.6th percentile of normal operation. Four "
        "decision rules were evaluated against it and none separated it."
    )
    b.p(
        "The limit is therefore not the threshold but the representation. An event "
        "smaller than the variation that normal motion already produces cannot be "
        "isolated by any monotone rule on these two scores. Closing that gap needs "
        "either a model trained on the cell's own trajectories, so that normal "
        "operation stops being out-of-distribution, or an additional channel that "
        "responds to contact rather than to dynamics. Stating the envelope is more "
        "useful than a single headline number: a safety function whose sensitivity "
        "floor is unknown is harder to deploy than one whose floor is published."
    )

    b.h2("5.6. Limitations")
    for t in [
        "A single robot (UR10e) and a single task profile; no validation on other "
        "robot types has been performed.",
        "Synthetic faults, even injected physically, remain analytic perturbations. "
        "The collision pulse is distributed over the wrench channels rather than "
        "propagated from a real contact point, and the four scenarios were shown in "
        "Section 4.4 not to discriminate between the two representation spaces once "
        "injected consistently.",
        "The current-to-torque coefficient could be measured directly on only two of "
        "six joints; the other four are derived by assumption, which affects the "
        "physical interpretation rather than the detection.",
        "The friction coefficients are fitted by a regression that absorbs any "
        "velocity-correlated model error, not friction alone. No independent "
        "tribological validation was performed.",
        "The inverse dynamics has no payload model. Carrying a workpiece produces a "
        "sustained residual bias that the deployed threshold sits above but does not "
        "explain away.",
        "The operating threshold is measured on the cell, which is necessary, but it "
        "is measured on a calibration run of 15,627 decisions and validated on "
        "twenty-four minutes of operation. The envelope it defines is a "
        "commissioning-scale result with respect to thermal drift, wear, "
        "part-to-part variation and operator differences.",
        "Only three confirmed anomalous events were available for the deployed "
        "configuration. The threshold interval (4.14, 11.30) is bounded below by "
        "observed clean operation and above by the weakest of those three events; a "
        "weaker real event would narrow it.",
        "Online measurements were taken on CPU; they should be repeated with the GPU "
        "provider on the deployment workstation.",
    ]:
        pp = b.para("", style="Paragraf", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=3, space_after=0)
        pp.paragraph_format.left_indent = Cm(0.4)
        b.rich(pp, "•  " + t)
    b.p("")

    b.h1("6. Conclusions")
    b.p(
        "This paper took a published score-level fusion framework for cobot anomaly "
        "detection, extended it, and delivered it as a detector running on a "
        "physical UR10e cell. Every measurement taken on the way is reported, "
        "including the ones that are unfavourable to the framework."
    )
    b.p(
        "The starting framework was rebuilt independently and its architectural and "
        "procedural components match exactly. Establishing the physical scale of the "
        "signals was a precondition for everything that follows: the driver writes "
        "motor current, not torque, into the effort field, so the earlier figures "
        "were produced on a mixed-unit pipeline. That pipeline was reproduced to "
        "within 0.013 on every quantity, which fixes the scale on which the earlier "
        "numbers should be read and makes fault amplitudes, friction coefficients "
        "and residual magnitudes physically meaningful for the extension."
    )
    b.p(
        "On that basis the pipeline was extended. A friction term absent from the "
        "validated solver was added outside it, removing 87 % and 92 % of the "
        "residual spread on the two wrist joints where the inverse dynamics had "
        "previously contributed almost nothing, and the evaluation was placed on a "
        "run-disjoint split with five training seeds. Under the inherited injection "
        "protocol the fusion retains its advantage: F1 0.791 ± 0.023 against 0.627 ± 0.006 and "
        "0.614 ± 0.016 for the single models, a gain of +0.164, with complementarity "
        "ratios close to those of the earlier study (27.6 % against 24.1 %)."
    )
    b.p(
        "The injection protocol itself was then examined, and this is the paper's "
        "principal negative result. When the same four faults are injected only into "
        "the measured "
        "channels and the residual is recomputed by the pipeline, the fusion margin "
        "falls from +0.189 ± 0.009 to −0.003 ± 0.001 PR-AUC, in every seed and "
        "irrespective of the friction term. The mechanism is a single amplitude: the "
        "reference protocol perturbs the residual space for the sensor-noise "
        "scenario with 0.08 Nm, chosen so that the scenario affects the raw signal "
        "only, whereas the physical propagation of the same 3.5 N wrench disturbance "
        "through r_ext = J(q)ᵀF produces 4.1 Nm. The scenario the reference study "
        "names as the test of complementarity does not test it. The complementarity "
        "of the two representation spaces is therefore reported here as a property "
        "of the evaluation protocol, not of the models."
    )
    b.p(
        "The system was implemented as a ROS 2 node running at 500 Hz whose feature "
        "engine is numerically identical to the offline pipeline to floating-point "
        "rounding, friction term included, in which both models run on every "
        "decision and whose inference consumes 12 % of the decision budget."
    )
    b.p(
        "On hardware the offline-derived threshold did not transfer, declaring 58 % "
        "of decisions to be alarms. The cause was measured rather than assumed: during ordinary "
        "operation the residual model runs at eleven times its validation mean error "
        "and the raw model at eighty-six times, so a percentile of the validation "
        "distribution carries no information about the deployment. The asymmetry is "
        "the practically useful part — the residual model sits at 0.93 of its own "
        "threshold on the cell while the raw model sits at 7.28, which let a "
        "five-per-cent-weighted model supply 29 % of the fused score. Re-estimating "
        "both scales and the threshold on clean cell decisions, with a robust "
        "percentile in place of the min–max range, restores the intended weighting."
    )
    b.p(
        "With that configuration the detector was validated against provoked "
        "collisions. Three operator-confirmed events were caught with peaks of "
        "11.30, 16.13 and 25.04 against a threshold of 5.0, with no unlabelled alarm "
        "block in ten minutes of event-bearing operation and no alarm at all in "
        "fourteen minutes of ordinary operation; any threshold in (4.14, 11.30) "
        "gives the same result. The working envelope is bounded on the other side "
        "too: a deliberately provoked low-amplitude hand contact reached only 6.8 "
        "times its own run median while ordinary inspection motion reaches 26.7 "
        "times, placing it at the 75.6th percentile of normal operation. Four "
        "decision rules were evaluated against that event and none separated it."
    )
    b.p(
        "Two directions follow. The sensitivity floor is a property of the "
        "representation, not of the threshold, so closing it requires models trained "
        "on the cell's own trajectories — which would also remove the "
        "out-of-distribution condition behind the threshold-transfer failure — or an "
        "additional channel that responds to contact rather than to dynamics. The "
        "second is a payload term outside the solver, following the same pattern as "
        "the friction term, since the unmodelled payload is what forces the "
        "normalising scale up during transport. Both are supported by the "
        "instrumentation already in place: every decision, every alarm block, the "
        "operator's label and the provenance of the configuration that produced them "
        "are written to disk."
    )

    b.h1("Acknowledgement")
    b.p(
        "The project is supported by the KDT Joint Undertaking (101140216) and its "
        "members, including additional funding from Vinnova (Sweden), "
        "Österreichische Forschungsförderungsgesellschaft mbH – FFG (Austria), "
        "Business Finland (Finland), Ministry of Universities and Research (Italy), "
        "FCT (Portugal) and TÜBİTAK (124N448) (Türkiye). The measurements were "
        "carried out at the Autonomous Systems and Reliability Laboratory of the "
        "ESOGÜ Intelligent Systems Application and Research Centre."
    )

    b.h1("Contribution Of Researchers")
    b.p(
        "Author 1: design and implementation of the offline pipeline, the audit and "
        "the ROS 2 detector node, commissioning measurements, preparation of the "
        "manuscript. Author 2: … . Author 3: … . (To be completed by the authors.)"
    )

    b.h1("Conflict Of Interest")
    b.p("No conflict of interest has been declared by the authors.")

    b.h1("References")
    for ref in REFERENCES:
        p = b.para("", style="Kaynaklar", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_before=6, space_after=0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        b.rich(p, ref, size=BODY_PT)


REFERENCES = [
    "Blochwitz, T., Otter, M., Arnold, M., Bausch, C., Clauß, C., Elmqvist, H., "
    "… Wolf, S. (2011). The Functional Mockup Interface for tool independent "
    "exchange of simulation models. *Proceedings of the 8th International Modelica "
    "Conference*, 105–114, Dresden, Germany.",

    "Correia, L., Goos, J. C., Klein, P., Bäck, T. & Kononova, A. V. (2024). "
    "Online model-based anomaly detection in multivariate time series: Taxonomy, "
    "survey, research challenges and future directions. *Engineering Applications "
    "of Artificial Intelligence, 138*, 109323.",

    "Darban, Z. Z., Webb, G. I., Pan, S., Aggarwal, C. C. & Salehi, M. (2024). "
    "Deep learning for time series anomaly detection: A survey. *ACM Computing "
    "Surveys, 57*(1), 1–42.",

    "Haddadin, S., De Luca, A. & Albu-Schäffer, A. (2017). Robot collisions: A "
    "survey on detection, isolation, and identification. *IEEE Transactions on "
    "Robotics, 33*(6), 1292–1312.",

    "Huang, X., Chen, N., Deng, Z. & Huang, S. (2024). Multivariate time series "
    "anomaly detection via dynamic graph attention network and Informer. *Applied "
    "Intelligence, 54*, 7636–7658.",

    "Katsampiris-Salgado, K., Dimitropoulos, N., Gkrizis, C., Michalos, G. & "
    "Makris, S. (2024). Collision detection for collaborative assembly operations "
    "on high-payload robots. *Robotics and Computer-Integrated Manufacturing, 87*, "
    "102708.",

    "Leys, C., Ley, C., Klein, O., Bernard, P. & Licata, L. (2013). Detecting "
    "outliers: Do not use standard deviation around the mean, use absolute "
    "deviation around the median. *Journal of Experimental Social Psychology, "
    "49*(4), 764–766.",

    "Li, W., Han, Y. & Xiong, Z. (2020). Collision detection of robots based on a "
    "force/torque sensor at the bedplate. *IEEE Transactions on Industrial "
    "Electronics, 67*(12), 12440–12449.",

    "Liu, K., Wang, L., Zhang, X., Sun, Y. & Li, J. (2025). Anomaly detection in "
    "multidimensional time series for water injection pump operations based on "
    "LSTMA-AE and mechanism constraints. *Scientific Reports, 15*.",

    "Macenski, S., Foote, T., Gerkey, B., Lalancette, C. & Woodall, W. (2022). "
    "Robot Operating System 2: Design, architecture, and uses in the wild. "
    "*Science Robotics, 7*(66), eabm6074.",

    "Malhotra, P., Ramakrishnan, A., Anand, G., Vig, L., Agarwal, P. & Shroff, G. "
    "(2016). LSTM-based encoder-decoder for multi-sensor anomaly detection. "
    "*arXiv preprint arXiv:1607.00148*.",

    "Malhotra, P., Vig, L., Shroff, G. & Agarwal, P. (2015). Long short term "
    "memory networks for anomaly detection in time series. *Proceedings of the "
    "European Symposium on Artificial Neural Networks (ESANN)*, 89–94, Bruges, "
    "Belgium.",

    "Park, D., Hoshi, Y. & Kemp, C. C. (2018). A multimodal anomaly detector for "
    "robot-assisted feeding using an LSTM-based variational autoencoder. *IEEE "
    "Robotics and Automation Letters, 3*(3), 1544–1551.",

    "Savitzky, A. & Golay, M. J. E. (1964). Smoothing and differentiation of data "
    "by simplified least squares procedures. *Analytical Chemistry, 36*(8), "
    "1627–1639.",

    "Yılmaz, C. S., Kahraman, S., Yılmaz, M., Yavuz, H. S. & Yayan, U. (2026). "
    "FMU tabanlı kalıntı ayrıştırma ve ikili LSTM özkodlayıcı birleşimi ile "
    "işbirlikçi robotlarda anomali tespiti [FMU-based residual decomposition and "
    "dual LSTM autoencoder fusion for anomaly detection in collaborative robots]. "
    "⟨KONFERANS ADI VE YERİ — yazarlar tarafından tamamlanacaktır⟩ kurultayında "
    "sunulmuş bildiri, Türkiye.",

    "Zhang, T., Chen, Y. & Zou, Y. (2024). Robot collision detection based on "
    "external torque observer. *Journal of South China University of Technology, "
    "52*(3), 84–92.",

    "Zhao, W., Queralta, J. P. & Westerlund, T. (2020). Sim-to-real transfer in "
    "deep reinforcement learning for robotics: A survey. *2020 IEEE Symposium "
    "Series on Computational Intelligence (SSCI)*, 737–744, Canberra, Australia.",
]


def main():
    doc, anchor = open_template()
    b = Builder(doc, anchor)
    front_matter(b)
    body(b)
    b._leave_wide()
    doc.save(str(OUT))
    print(f"written: {OUT}")
    print(f"  figures: {b.fig_no}   tables: {b.tab_no}   equations: {b.eq_no}")


if __name__ == "__main__":
    main()
