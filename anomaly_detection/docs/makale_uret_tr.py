#!/usr/bin/env python3
"""
makale_uret_tr.py
=================
Makalenin TÜRKÇE sürümünü üretir.

`makale_uret.py`'nin bütün yerleşim makinesini (Builder, şablon açma, tablo/şekil
düzeni, iki sütun yönetimi, geniş blok mantığı) olduğu gibi kullanır; değişen tek
şey içeriktir. Böylece iki sürüm arasında biçim farkı oluşmaz ve şablonla ilgili
bir düzeltme tek yerde yapılır.

Türkçe sürümde ön sayfa sırası ters çevrilir: dergi Türkçe olduğu için önce
Türkçe başlık ve Öz, sonra İngilizce başlık ve Abstract gelir.

    python3 makale_uret_tr.py
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

import makale_uret as en
from makale_uret import (
    ABSTRACT_EN, ABSTRACT_TR, BODY_PT, Builder, KEYWORDS_EN, KEYWORDS_TR,
    SMALL_PT, TITLE_EN, TITLE_TR, abstract_block, open_template, style_run,
)

OUT = en.HERE / "ESOGU_MMF_Makale_TR.docx"


class BuilderTR(Builder):
    """Yalnız başlık etiketleri Türkçe; yerleşim davranışı değişmez."""

    FIG_LABEL = "Şekil"
    TAB_LABEL = "Tablo"


# ─────────────────────────────── ön sayfa ───────────────────────────────

def front_matter(b: Builder) -> None:
    """İngilizce sürümle aynı, yalnız dil sırası ters."""

    def title(text):
        p = b.para("", style="Makale Başlığı",
                   align=WD_ALIGN_PARAGRAPH.CENTER, front=True)
        style_run(p.add_run(text), BODY_PT + 2, bold=True)

    title(TITLE_TR)
    p = b.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, front=True)
    style_run(p.add_run("Adı SOYADI"), BODY_PT)
    style_run(p.add_run("1*"), BODY_PT - 3.5).font.superscript = True
    style_run(p.add_run(",  Adı SOYADI"), BODY_PT)
    style_run(p.add_run("2"), BODY_PT - 3.5).font.superscript = True
    style_run(p.add_run("  (Dergi editörlüğü tarafından basım aşamasında "
                        "yazılacaktır.)"), BODY_PT)
    for i in (1, 2):
        p = b.para("", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER,
                   front=True)
        style_run(p.add_run(f"{i} "), SMALL_PT - 1)
        style_run(p.add_run(f"Yazar {i} Adresi, ORCID No : "
                            "https://orcid.org/"), SMALL_PT - 1)
    b.para("", style="Normal", front=True)

    abstract_block(b, "Anahtar Kelimeler", "Öz", KEYWORDS_TR, ABSTRACT_TR,
                   "Araştırma Makalesi")
    b.para("", style="Normal", front=True)
    title(TITLE_EN)
    b.para("", style="Normal", front=True)
    abstract_block(b, "Keywords", "Abstract", KEYWORDS_EN, ABSTRACT_EN,
                   "Research Article")
    b.para("", style="Normal", front=True)


# ─────────────────────────────── gövde ───────────────────────────────

def body(b: Builder) -> None:
    _giris(b)
    _literatur(b)
    _yontem(b)
    _bulgular(b)
    _tartisma(b)
    _sonuclar(b)
    _arka(b)


def _madde(b, metin):
    p = b.para("", style="Paragraf", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
               space_before=3, space_after=0)
    p.paragraph_format.left_indent = Cm(0.4)
    b.rich(p, "•  " + metin)


# ── 1. Giriş ──
def _giris(b: Builder) -> None:
    b.h1("1. Giriş")
    b.p(
        "İşbirlikçi robotlar çalışma hacmini insan operatörlerle paylaşır; bu "
        "yüzden çarpışma, bozulan bir motor ya da arızalanan bir algılayıcı gibi "
        "anomalilerin olaya dönüşmeden tespit edilmesi gerekir. İki yöntem ailesi "
        "öne çıkar. Fizik tabanlı kalıntı yöntemleri robotun dinamik modelini "
        "ölçülen kuvvetlerle karşılaştırır ve farkı arıza göstergesi sayar; "
        "boyutu düşürür ve gürültüyü bastırırlar, ama düşük genlikli algılayıcı "
        "anomalilerini de bastırırlar. Veri güdümlü yöntemler karmaşık örüntüleri "
        "doğrudan yüksek boyutlu ham sinyallerden öğrenir, ancak fiziksel bir "
        "modelin sağladığı yapıdan yoksundurlar. İki aile bu nedenle farklı arıza "
        "tiplerinde başarısız olur ve bu asimetri, ikisini birleştirerek "
        "değerlendirilebilir."
    )
    b.p(
        "Önceki çalışmamızda bir UR10e kobot için tam olarak böyle bir birleşim "
        "önerilmişti (Yılmaz, Kahraman, Yılmaz, Yavuz ve Yayan, 2026). Fonksiyonel "
        "Model Birimi (Functional Mock-up Unit, FMU) ters dinamik modelinden "
        "türetilen on iki içsel ve dışsal kalıntı kanalı üzerinde çalışan bir "
        "kalıntı Uzun Kısa Süreli Bellek (Long Short-Term Memory, LSTM) "
        "özkodlayıcısı, yirmi dört doğrudan algılayıcı kanalı üzerinde çalışan bir "
        "ham LSTM özkodlayıcısıyla skor düzeyinde birleştirilmişti. O çalışma "
        "tümüyle çevrimdışı, kayıtlı bir veri kümesi üzerinde değerlendirilmiş ve "
        "kendi gelecek çalışma listesinde tek bir maddeyi en önemli açık iş olarak "
        "adlandırmıştı: sistemi robotla eş zamanlı ve çevrimiçi çalıştırmak."
    )
    b.p(
        "Bu makale o maddeyi kapatmaktadır. Anılan çalışmanın çerçevesi başlangıç "
        "noktası alınmakta, dört yönde genişletilmekte ve sonuç, gerçek hücrede "
        "çalışan bir dedektör olarak teslim edilmektedir. Kalıntı tanımı, "
        "doğrulanmış ters dinamik çözücünün içermediği sürtünme terimiyle "
        "tamamlanmaktadır; değerlendirme, her sayının bir belirsizlik taşıması için "
        "koşu-ayrık bir bölme ve beş eğitim tohumu üzerine oturtulmaktadır; arıza "
        "enjeksiyonu fiziksel olarak tutarlı hale getirilmekte, ki bunun birleşimin "
        "değerini değiştirdiği görülmektedir; ve normalleştirme ile eşiğe, geleceği "
        "görmeden kestirilebilecek bir biçim verilmekte, çevrimdışı değerler "
        "taşınmadığında ise bunlar hücrede yeniden kestirilmektedir. Dedektör "
        "ardından devreye alınmış ve provoke edilen çarpışmalara karşı "
        "doğrulanmıştır."
    )
    b.p(
        "Bu çalışmanın doldurduğu boşluk dolayısıyla yeni bir mimari değildir. "
        "Robot anomali tespiti için birleşim çerçeveleri çoktur ve neredeyse her "
        "zaman, arıza enjekte edilmiş kayıtlı veri üzerinde değerlendirilirler. "
        "Nadiren raporlanan şey, böyle bir çerçevenin hangi parçalarının donanımla "
        "temasa dayandığıdır: çalışma eşiği taşınıyor mu, normalleştirme nedensel "
        "mi, ölçülen birleşim üstünlüğü arızaların nasıl enjekte edildiğine bağlı "
        "mı, ve dedektör çalışır hale geldiğinde neyi görüp neyi göremiyor. Bu "
        "makale bu ölçümleri tek bir eksiksiz sistem için raporlamakta ve ortaya "
        "çıkan çalışma zarfını iki yandan da belirtmektedir."
    )
    b.p("Katkılar şunlardır:")
    for it in (
        "doğrulanmış çözücünün dışında uydurulan Coulomb+viskoz bir sürtünme "
        "terimiyle tamamlanan bir kalıntı tanımı; bu terim, ters dinamiğin daha "
        "önce neredeyse hiçbir katkı sağlamadığı iki bilek ekleminde kalıntı "
        "yayılımının %87 ve %92'sini gidermektedir;",
        "beş eğitim tohumuyla koşu-ayrık bir yeniden değerlendirme; böylece "
        "birleşimin üstünlüğü tek bir sayı olarak değil, bir belirsizlikle "
        "raporlanmaktadır;",
        "fiziksel olarak tutarlı bir arıza enjeksiyon protokolü ve devralınan "
        "protokol altında raporlanan birleşim üstünlüğünün bu protokolde ortadan "
        "kalktığı ölçümü — kaynağı, fiziksel yayılımından elli bir kat küçük "
        "seçilmiş tek bir enjeksiyon genliğidir;",
        "çevrimdışı yöntemin tanımlamadığı iki büyüklüğün, birleşim "
        "normalleştirmesinin nedensel bir formülasyonu ile birleşik skor için bir "
        "çalışma eşiğinin verilmesi ve bunları hücrede yeniden kestiren yordam;",
        "öznitelik motoru çevrimdışı hatla — sürtünme terimi dâhil — kayan nokta "
        "düzeyinde örtüşen, ölçülmüş gecikme bütçesine sahip gerçek zamanlı bir "
        "ROS 2 gerçeklemesi;",
        "çevrimdışı türetilen eşiğin taşınmadığının gösterildiği, sebebin "
        "modellere değil bir normalleştiriciye dayandırıldığı ve dağıtılan "
        "dedektörün operatörce doğrulanmış çarpışmalarla, duyarlılık tabanı açıkça "
        "belirtilerek doğrulandığı bir devreye alma raporu;",
        "genişletmeden önce yürütülen ve bulguları burada düzeltilen, önceki "
        "yayımlanmış sayılara etkisi Bölüm 4.5'te belgelenen bir veri yolu "
        "denetimi.",
    ):
        _madde(b, it)
    b.p(
        "Makalenin kalanı şöyle düzenlenmiştir. Bölüm 2 ilgili literatürü "
        "gözden geçirir. Bölüm 3 platformu ve veriyi, genişletilmiş hattı ve "
        "çevrimiçi gerçeklemeyi anlatır. Bölüm 4 çevrimdışı ve gerçek hücre "
        "bulgularını sunar. Bölüm 5 bunları tartışır, çalışma zarfını ve "
        "kısıtları belirtir; Bölüm 6 sonuçlandırır."
    )


# ── 2. Literatür ──
def _literatur(b: Builder) -> None:
    b.h1("2. Literatür Taraması")
    b.p(
        "Manipülatörler için fizik tabanlı tespit olgun bir alandır. Haddadin, De "
        "Luca ve Albu-Schäffer (2017) çarpışma tespiti, yalıtımı ve tanımlamasını "
        "derlemekte ve kalıntı gözlemleyicisini kanonik araç olarak "
        "yerleştirmektedir. Li, Han ve Xiong (2020) tabana bir kuvvet/tork "
        "algılayıcısı yerleştirip çarpışmayı ortaya çıkan kalıntıdan tespit "
        "ederken, Zhang, Chen ve Zou (2024) aynı amaçla dışsal bir tork "
        "gözlemleyicisi kullanmaktadır. Katsampiris-Salgado vd. (2024) model "
        "hatasının kendisinin sınırlayıcı etken hâline geldiği yüksek yükli "
        "işbirlikçi montajı ele almaktadır. Bu ailenin ortak zayıflığı şudur: "
        "modelin temsil etmediği her şey — sürtünme, yük, eklem esnekliği — "
        "arızadan ayırt edilemez; modelin iyi süzdüğü her şey ise, arıza o "
        "olduğunda da süzülüp gider."
    )
    b.p(
        "Çok değişkenli zaman serileri için veri güdümlü tespit Darban, Webb, Pan, "
        "Aggarwal ve Salehi (2024) tarafından derlenmiştir. Yeniden yapılanma "
        "temelli özkodlayıcılar baskın tasarımdır: Malhotra, Vig, Shroff ve "
        "Agarwal (2015) zaman serisi anomali tespiti için LSTM ağlarını tanıtmış, "
        "Malhotra, Ramakrishnan, Anand, Vig, Agarwal ve Shroff (2016) fikri çok "
        "algılayıcılı bir kodlayıcı–çözücüye genişletmiştir. Park, Hoshi ve Kemp "
        "(2018) robot destekli beslemeye LSTM tabanlı değişimsel bir özkodlayıcı "
        "uygulamıştır; bu, fiziksel bir robot üzerinde raporlanan az sayıdaki "
        "çalışmadan biridir. Bu yöntemler arıza etiketi gerektirmez, ancak eğitim "
        "dağılımının taşıdığı her yanlılığı devralırlar."
    )
    b.p(
        "Melez yaklaşımlar ikisini birleştirmeye çalışır. Liu vd. (2025) pompa "
        "işletimi için bir LSTM özkodlayıcısını mekanizma bilgisiyle "
        "kısıtlamakta, ancak iki ayrı modelin skorlarının birleşimini "
        "incelememektedir. Huang, Chen, Deng ve Huang (2024) çok değişkenli "
        "anomali tespitine, fizik tabanlı hiçbir ön işleme aşaması olmaksızın "
        "çizge dikkati ve Transformer uygulamaktadır. Correia, Goos, Klein, Bäck "
        "ve Kononova (2024) özellikle çevrimiçi model tabanlı anomali tespitini "
        "derlemekte ve eşik seçimi ile durağan olmayan çalışma koşullarını açık "
        "araştırma sorunları olarak belirlemektedir — bu makalenin donanım "
        "üzerinde ölçtüğü tam da bu iki sorundur."
    )
    b.p(
        "Benzetimli ya da kayıtlı veri üzerinde kalibre edilmiş bir modelin "
        "fiziksel bir robota aktarılması çoğunlukla denetim ve pekiştirmeli öğrenme "
        "literatüründe, gerçeklik boşluğu (reality gap) adıyla incelenmiştir "
        "(Zhao, Queralta ve Westerlund, 2020). Anomali tespitinde eşdeğer soru — "
        "karar eşiği taşınıyor mu? — nadiren sorulur, çünkü çalışmaların çoğu bir "
        "çalışma eşiği hiç tanımlamaz; onun yerine tüm eşikler üzerinden ulaşılan "
        "en yüksek F1'i raporlar. Bu istatistik, tahmin etmesi gereken etiketleri "
        "gerektirir ve dolayısıyla çevrimiçi karşılığı yoktur. Bu çalışma söz "
        "konusu boşluğu açık hâle getirmekte, nedensel bir eşik kuralıyla "
        "kapatmakta ve kuralın gerçek donanıma uygulandığında ne olduğunu "
        "ölçmektedir."
    )


# ── 3. Yöntem ──
def _yontem(b: Builder) -> None:
    b.h1("3. Yöntem")

    b.h2("3.1. Platform ve Veri Toplama")
    b.p(
        "Deneyler, Festo doğrusal eksen üzerine monte edilmiş bir Universal Robots "
        "UR10e işbirlikçi robotundan (altı eklem, 12,5 kg yük kapasitesi) ve bir "
        "muayene şasi çerçevesinden oluşan, Şekil 1'de görülen robotik kalite "
        "muayene platformunda yürütülmüştür. Hedef uygulama otomotiv gövde "
        "panellerinin otomatik görsel kalite muayenesidir. Uç işlevciye bir "
        "kuvvet/tork algılayıcısı (KTS) monte edilmiştir. Hücre ROS 2 Humble ve "
        "MoveIt 2 ile çalışmakta, kaydedilen hareketler on bir farklı hareket "
        "planlama algoritmasıyla üretilmiştir; böylece normal çalışma dağılımı tek "
        "bir planlayıcının hâkimiyetinde değildir."
    )
    b.figure(
        "fig1_platform.png",
        "Robotik Kalite Muayene Platformu | Festo doğrusal eksen üzerindeki UR10e "
        "robot, sağda muayene şasi çerçevesi.",
        wide=False, width_cm=8.35)
    b.p(
        "Kaynak veri kümesi, eklem durumlarının ve wrench ölçümlerinin 79 günlük "
        "bir dışa aktarımıdır; nominal 500 Hz'de 1.124.432 örnek içerir. Model "
        "eğitimi Intel Core i7-13650HX işlemcili ve NVIDIA RTX 4060 GPU'lu bir "
        "dizüstü bilgisayarda yapılmış, dağıtım iş istasyonu ise 64 GB bellekli "
        "Intel Core i9-14900 ve NVIDIA RTX 4000 Ada'dır."
    )
    b.p(
        "Çalışmada insan denek kullanılmamış ve etik kurul onayı gerekmemiştir. "
        "Tüm ölçümler ESOGÜ Akıllı Sistemler Uygulama ve Araştırma Merkezi'nin "
        "laboratuvar donanımında toplanmış, araştırma ve yayın etiği ilkelerine "
        "uyulmuştur. Devreye alma kampanyasındaki anomaliler, robotun kendi "
        "koruyucu durdurma işlevi her an etkin hâldeyken, düşürülmüş hızda ve "
        "eğitimli operatörlerce kasten provoke edilmiştir."
    )

    b.h2("3.2. Veri İncelemesi ve Hazırlama")
    b.p(
        "Yeniden kurmanın ilk aşaması verinin kendisiydi ve en büyük tekil "
        "düzeltmeyi orada üretti. Dışa aktarım zaman damgasına göre genel olarak "
        "sıralı değildir. 1.124.432 örnek üzerinde ölçüldüğünde, ardışık çiftlerin "
        "321.841'i (%28,6) süreksizdir; bunların 96.107'sinde zaman adımı negatif "
        "ya da sıfırdır. Medyan adım beklenen 2,00 ms olmakla birlikte aralık "
        "−5,5·10⁹ ms ile +3,6·10⁹ ms arasında değişmektedir. Kayan pencerelerin "
        "dörtte üçü (44.974'ün 33.700'ü, %74,9) bu nedenle saatler ya da günler "
        "arayla alınmış iki kayıt oturumunu birbirine eklemektedir."
    )
    b.p(
        "Bunun iki ayrı sonucu vardır. Bir LSTM özkodlayıcısının öğrenmesi "
        "istenen dinamiğin ortasına keyfî bir sıçrama koyar ve eklem ivmesini "
        "bozar. İvme, hız kanalından Savitzky–Golay türeviyle (pencere 51, derece "
        "3) elde edilir (Savitzky ve Golay, 1964); bu bir evrişimdir ve bir ek "
        "yerinin üzerinden geçtiğinde fiziksel olmayan bir sivri uç üretir. İvme "
        "M(q)q̈ terimi üzerinden doğrudan ters dinamiğe girdiği için hata kalıntıya "
        "yayılır. Aynı 200.000 örnek üzerinde ölçüldüğünde, ek yerlerinin üzerinden "
        "türetilen q̈'nin standart sapması, türev kesintisiz koşuların içinde "
        "alındığındakinin iki ilâ altı katıdır."
    )
    b.p(
        "Düzeltme, dışa aktarımı fiziksel olarak kesintisiz koşulara bölmektir. "
        "İki ardışık örnek yalnızca zaman adımı (0, 8 ms] aralığındaysa ve eklem "
        "yer değiştirmesi eklem hız sınırının o adımda izin verdiğinin altında "
        "kalıyorsa aynı koşuya aittir; ikinci koşul, makul görünen bir zaman "
        "damgasının gizlediği boşlukları yakalar. 150 örnekten kısa koşular atılır "
        "ve her koşu düzgün bir 2 ms ızgarasına yeniden örneklenir. Tablo 1 sonucu "
        "vermektedir. Örnek sayısındaki azalma bir kayıp değildir: atılan örnekler "
        "tam olarak aralarında fiziksel süreklilik bulunmayanlardır."
    )
    b.table(
        "Veri Hazırlamanın Sonucu.",
        ["Büyüklük", "Değer"],
        [["Ham dışa aktarılan örnek", "1.124.432"],
         ["Fiziksel olarak kesintisiz koşu", "600"],
         ["Hazırlanan örnek", "342.480  (685 s robot zamanı)"],
         ["Aradeğerlenen örnek", "%5,02"],
         ["Koşu uzunluğu (medyan / min / maks)", "227 / 155 / 35.860"]],
        widths=[4.8, 3.55], align_right=[1])

    b.h2("3.3. Akım–Tork Kalibrasyonu")
    b.p(
        "UR ROS 2 sürücüsü, eklem durumu mesajının effort alanına motor akımını "
        "amper cinsinden yazar; alan doğrudan denetleyicinin bildirdiği gerçek "
        "akımla doldurulur. Ters dinamik modeli ise newton metre üretir. Denklem "
        "(2)'nin kalıntısı oluşturulabilmesi için bu nedenle önce bir dönüşüm "
        "katsayısı gereklidir."
    )
    b.p(
        "İlk yeniden kurulumda bu katsayı, ölçülen akımın model torkuna "
        "regresyonuyla elde edilmiş ve aynı regresyon tekrarlanarak "
        "doğrulanmıştı. Bu doğrulama döngüseldir: katsayı, sonradan sınanan "
        "eşitliği sağlayacak biçimde seçilmektedir. Katsayı bu yüzden, sınanan "
        "kalıntıdan bağımsız bir kaynaktan yeniden türetilmiştir. Robot neredeyse "
        "duruyorken (|q̇| < 0,005 rad/s) eklem torku yerçekimince belirlenir, "
        "dolayısıyla g(q) ≈ k·(i − i₀) yazılabilir; burada g(q), denetlenen "
        "modelden değil kapalı biçimli yerçekimi vektöründen alınır. 197.090 "
        "yarı-statik örneğin 65.697'si regresyona girmiştir. Tablo 2 sonucu "
        "listelemektedir."
    )
    b.table(
        "Ölçülen Akım–Tork Katsayıları ve Duyarlılıkları.",
        ["Eklem", "Nm/A", "R²", "Ölç.", "ρ −%50", "ρ +%50"],
        [["shoulder_pan", "10,522", "—", "hayır", "0,968", "0,995"],
         ["shoulder_lift", "10,522", "0,984", "evet", "0,162", "0,637"],
         ["elbow", "9,130", "0,993", "evet", "0,023", "0,758"],
         ["wrist_1", "3,409", "0,000", "hayır", "0,890", "0,968"],
         ["wrist_2", "3,409", "0,012", "hayır", "0,957", "0,993"],
         ["wrist_3", "3,409", "0,010", "hayır", "0,994", "0,999"]],
        widths=[3.3, 2.5, 2.3, 2.2, 3.7, 3.9], wide=True,
        align_right=[1, 2, 4, 5],
        note="R², yarı-statik yerçekimi regresyonunun R²'sidir; “Ölç.” katsayının "
             "doğrudan ölçülebildiği eklemleri işaretler. Son iki sütun, katsayı "
             "∓%50 bozulduğunda yeniden hesaplanan içsel kalıntının nominal olanla "
             "korelasyonunu, hazırlanan 342.480 örneğin tamamı üzerinden verir. "
             "Ölçülemeyen katsayılar aynı modül ailesinden kopyalanır ya da tork "
             "sınırı oranıyla ölçeklenir; düğüm başlangıçta bunlar için uyarı "
             "basar.")
    b.p(
        "Varsayılan dört katsayının sonucu tartışılmak yerine ölçülmüştür. Her "
        "katsayı ±%50 bozulup içsel kalıntı yeniden hesaplanmıştır; model torku "
        "ölçümden ayrı saklandığı için bu, ters dinamiğin yeniden "
        "değerlendirilmesini gerektirmez. Katsayısı ölçülemeyen dört eklemde "
        "bozma, kalıntıyı biçim olarak neredeyse değiştirmez: nominal kalıntıyla "
        "korelasyon 0,890 ile 0,999 arasındadır ve etki, standart sapmanın 0,55 "
        "ilâ 1,45 katına değişmesiyle sınırlı, saf bir yeniden ölçeklemeye "
        "yakındır. Özkodlayıcılar girdilerini standartlaştırdığından bu biçimde bir "
        "dönüşüm öğrenilen dinamiği değiştiremez. Katsayısı ölçülen iki eklemde ise "
        "aynı bozma yıkıcıdır: korelasyon −%50'de shoulder_lift için 0,162'ye, "
        "elbow için 0,023'e düşer. Asimetrinin kendisi tanısaldır. Akım terimi, "
        "modelin iptal etmesi beklenen yerçekimi torkunu taşıdığında yanlış "
        "ölçekleme iptali bozar; modelin az katkı verdiği bileklerde ise katsayı "
        "yalnızca bir ölçek çarpanı gibi davranır. Varsayılan katsayılar bu nedenle "
        "tespit başarımını değil, bir arıza genliğinin newton metre cinsinden "
        "fiziksel yorumunu etkiler."
    )

    b.h2("3.4. Melez Kalıntı Ayrıştırma ve Sürtünme")
    b.p(
        "Ters dinamik modeli, UR10e'nin Newton–Euler formülasyonunu Fonksiyonel "
        "Model Arayüzü biçiminde paketler (Blochwitz vd., 2011):"
    )
    b.equation("τ̂_model  =  M(q)·q̈ + C(q, q̇)·q̇ + g(q)")
    b.p("Toplam kalıntı ölçüm ile model arasındaki farktır,")
    b.equation("r_total  =  τ_ölç − τ̂_model")
    b.p("dışsal kalıntı uç işlevci wrench'inin eklem uzayına izdüşümüdür,")
    b.equation("r_ext  =  J(q)ᵀ · F_KTS")
    b.p("ve içsel kalıntı ikisinin farkıdır,")
    b.equation("r_int  =  r_total − r_ext")
    b.p(
        "burada M(q) atalet matrisi, C(q, q̇) Coriolis ve merkezkaç matrisi, g(q) "
        "yerçekimi vektörü, J(q) geometrik Jacobian, τ_ölç ölçülen eklem torku ve "
        "F_KTS ölçülen wrench'tir. Bu ayrıştırma r_int'i motor bozulması ya da "
        "gizyazar hatası gibi içsel arızalara, r_ext'i ise çarpışma ya da yük "
        "değişimi gibi dış etkileşimlere duyarlı kılar. Denklem (4) bir tanım "
        "olduğundan ayrıştırmanın sayısal tutarlılığı doğrudan sınanabilir: "
        "200.000 örnek üzerinde özdeşlik 7,1·10⁻¹⁵ artıkla sağlanmakta ve analitik "
        "Jacobian sayısal olanla makine hassasiyetinde örtüşmektedir."
    )
    b.p(
        "Denklem (3), ihlali hiçbir tanı üretmeyen bir referans çerçevesi "
        "uyuşmasını gerektirir. UR e-Serisi sürücüsü wrench'i takım çerçevesine "
        "döndürüp takım çerçevesi kimliğiyle yayınlarken, J(q) taban çerçevesinde "
        "tanımlıdır. İkisini doğrudan çarpmak, hiçbir hata mesajı olmadan yanlış "
        "bir kalıntı üretir. 40.000 örnek üzerinde ölçüldüğünde doğru ve "
        "döndürülmemiş dışsal kalıntı arasındaki korelasyon birinci eklemde "
        "−0,995'tir — işaret tersine dönmektedir — ve tüm kanallar genelinde "
        "−0,143'tür; ikinci eklemde genlik üçe katlanmaktadır. Bu nedenle wrench'in "
        "hem kuvvet hem moment yarısı, Jacobian aktarımından önce ve hem çevrimdışı "
        "hatta hem çevrimiçi motorda aynı kod yolundan geçerek taban çerçevesine "
        "döndürülmektedir."
    )
    b.p(
        "Çözücünün iki özelliği ölçülmüş ve içinde düzeltilmemiştir. 500 rastgele "
        "eklem pozunda atalet matrisi her pozda simetrik değildir ve simetrik "
        "kısmı pozitif tanımlı değildir; sürtünme vektörü birebir sıfırdır ve bir "
        "yük modeli bulunmamaktadır. Çözücü fiziksel robota karşı doğrulanmış olup "
        "sabit kabul edilmektedir. Bedeli bu yüzden gizlenmek yerine ölçülmüştür: "
        "wrist_2 ve wrist_3'te model torkunun standart sapması ölçülen torkun "
        "yalnızca %10 ve %2'sidir (0,959 ve 1,246 Nm'ye karşılık 0,095 ve 0,030 "
        "Nm); dolayısıyla bu iki kanalda Denklem (2) ölçümün kendisine indirgenir."
    )
    b.p(
        "Eksik sürtünme terimi, doğrulanmış modele dokunmamak için çözücünün içine "
        "değil dışına eklenmektedir; bu, redüktörlü bir manipülatörde büyük olan "
        "bir bileşeni geri kazandırır. Denklem (2) şu hâle gelir:"
    )
    b.equation("r_total  =  τ_ölç − τ̂_model − τ̂_f(q̇)", "2a")
    b.p("eklem başına iki parametreli Coulomb+viskoz bir terimle,")
    b.equation("τ̂_f  =  F_c · tanh(q̇ / ε) + F_v · q̇", "2b")
    b.p(
        "burada tanh, işaret fonksiyonunun yerini almaktadır; çünkü sign(q̇) sıfır "
        "geçişlerinde süreksizdir ve bu robot zamanının çoğunu düşük hızda "
        "geçirir. Keskin bir işaret, her yön değişiminde kalıntıya bir basamak "
        "enjekte eder — ki bu tam olarak bir dedektörün işaretlemesi beklenen "
        "biçimdir. Yumuşatma genişliği ε = 0,02 rad/s'dir. Katsayılar sağlam en "
        "küçük kareler yöntemiyle (üç yinelemeli 3σ kırpma), yalnızca hareket eden "
        "örnekler üzerinde ve — değerlendirme açısından belirleyici olarak — "
        "yalnızca Bölüm 3.7'de tanımlanan bölmenin eğitim koşularında "
        "uydurulmuştur. Tüm veri üzerinde uydurmak, düzeltmenin kendisini bir "
        "sızıntı kanalı hâline getirirdi. Tablo 3 sonucu listelemektedir."
    )
    b.table(
        "Sürtünme Katsayıları, Yalnız Eğitim Koşularında Uydurulmuş.",
        ["Eklem", "F_c [Nm]", "F_v [Nm·s/rad]", "R²", "σ(r) önce", "σ(r) sonra", "Δ"],
        [["shoulder_pan", "7,65", "35,76", "0,764", "13,20", "6,09", "−%54"],
         ["shoulder_lift", "8,78", "26,19", "0,799", "11,91", "4,98", "−%58"],
         ["elbow", "5,53", "19,97", "0,896", "7,66", "2,46", "−%68"],
         ["wrist_1", "0,95", "1,29", "0,379", "1,41", "0,75", "−%47"],
         ["wrist_2", "1,51", "1,58", "0,947", "1,53", "0,20", "−%87"],
         ["wrist_3", "1,38", "2,05", "0,974", "1,90", "0,15", "−%92"]],
        widths=[3.0, 2.1, 2.9, 1.9, 2.5, 2.4, 1.9], wide=True,
        align_right=[1, 2, 3, 4, 5, 6],
        note="σ(r), terim çıkarılmadan önce ve sonra hareket eden örneklerde "
             "(|q̇| > 0,02 rad/s) toplam kalıntının standart sapmasıdır. Regresyon "
             "yalnız sürtünmeyi değil, hıza bağlı her model hatasını soğurur; "
             "bunların gerçek tribolojik katsayılar olduğu iddia edilmemektedir. "
             "Etki tam da ters dinamiğin en az katkı verdiği yerde en büyüktür: "
             "wrist_2 ve wrist_3'te kalıntı yayılımının %87 ve %92'si "
             "kaybolmaktadır; dolayısıyla yukarıdaki itiraf — bu kanalların model "
             "bilgisi taşımadığı — terim dâhil edildiğinde artık geçerli değildir. "
             "Bunun tespiti iyileştirip iyileştirmediği ayrı bir sorudur ve Bölüm "
             "4.4'te yanıtlanmaktadır.")
    b.p(
        "Çevrimiçi öznitelik motoru aynı ifadeyi aynı katsayı dosyasından "
        "uygulamakta ve Bölüm 3.9'daki eşdeğerlik sınaması terim etkinken "
        "çalıştırılmaktadır."
    )

    b.h2("3.5. İkili LSTM Özkodlayıcı")
    b.p(
        "Her iki model de, T = 100 örneklik (0,2 s) kayan bir pencereyi saklı bir "
        "vektöre sıkıştıran simetrik bir kodlayıcı–çözücü mimarisini paylaşır. İki "
        "katmanlı bir LSTM kodlayıcı girdi dizisini sabit boyutlu saklı vektöre "
        "indirger; çözücü onu T adım boyunca tekrarlar ve simetrik LSTM "
        "katmanlarıyla yeniden inşa eder. Katmanlar arasında %15 sönümleme "
        "uygulanır. Her iki model de yalnız normal çalışma verisiyle, ortalama "
        "karesel hata kaybıyla eğitilir ve bir pencerenin anomali skoru onun "
        "yeniden yapılanma hatasıdır. Her modelin kendi eşiği, doğrulama kümesindeki "
        "yeniden yapılanma hatalarının 97. persentilidir."
    )
    b.p(
        "Pencereleme 25 örneklik adım, yani %75 örtüşme kullanır. İki kural birlikte "
        "uygulanır: hiçbir pencere iki koşuya yayılamaz ve türev kenar payınca "
        "geçersizleştirilmiş örnek içeren pencereler elenir. Eğitim ve doğrulama "
        "pencereleri ayrık fiziksel koşu kümelerinden gelir (Bölüm 3.7), dolayısıyla "
        "aralarında paylaşılan hiçbir örnek yoktur. Her model farklı tohumlarla beş "
        "kez eğitilmiştir; bu makaledeki her çevrimdışı sayı, bu beş koşunun "
        "ortalaması ve standart sapmasıyla verilmektedir, çünkü tek bir eğitim "
        "koşusu yöntemler arasındaki farkı ilklendirme gürültüsünden ayıramaz. "
        "Tablo 4 mimarileri ve eğitim sonucunu vermektedir."
    )
    b.table(
        "Model Mimarileri ve Eğitim Sonucu (beş tohum, ort ± s.s.).",
        ["", "Kalıntı ÖK", "Ham ÖK"],
        [["Girdi kanalı", "12", "24"],
         ["Gizli / saklı", "128 / 32", "256 / 64"],
         ["Parametre", "478.892", "1.907.032"],
         ["Eğitim / doğrulama penceresi", "6.756 / 2.162", "6.756 / 2.162"],
         ["En iyi dönem", "86 ± 41", "83 ± 42"],
         ["En iyi doğrulama kaybı", "0,092 ± 0,016", "0,036 ± 0,005"],
         ["Eşik θ (P97)", "0,887 ± 0,412", "0,412 ± 0,038"]],
        widths=[2.95, 2.7, 2.7], align_right=[1, 2],
        note="Adam (lr = 10⁻³, β = 0,9/0,999), yığın 256, gradyan kırpma 1,0, "
             "ReduceLROnPlateau (faktör 0,5, sabır 8), en fazla 300 dönem boyunca "
             "sabır 25 ile erken durdurma; her model RTX 4060 dizüstünde bir "
             "dakikanın altında eğitilmektedir. Her iki doğrulama kaybı da referans "
             "çalışmanınkilerin (0,181 ve 0,320) çok altındadır; sebebi daha iyi bir "
             "model değil, doğrulama kümesinin bileşimidir: on altı doğrulama koşusu "
             "eğitim eklem aralığının içinde kalacak biçimde seçilmiştir (Bölüm "
             "3.7), dolayısıyla doğrulama, indeks kesiğine göre eğitime daha çok "
             "benzemektedir. Sürtünme terimi bunu açıklamaz — ham model kalıntı "
             "kanallarını hiç görmez ve kaybı terimle de terimsiz de 0,036'dır; "
             "kalıntı modelinin kaybı ise terimle 0,092, terimsiz 0,062'dir. "
             "Eğitime benzeyen bir doğrulama kümesi dar bir dağılım verir ve θ o "
             "dağılımın bir persentilidir; Bölüm 4.7 bunun donanımda neye mal "
             "olduğunu ölçmektedir. Tohumlar arasında kalıntı eşiği ortalamasının "
             "%46'sı, ham eşiği %9'u kadar değişmektedir. Buradaki değerler beş "
             "tohum ortalamasıdır; Bölüm 4.7 dağıtılan tohumu kullanır, onun "
             "eşikleri 1,3723 ve 0,3934'tür.")

    b.h2("3.6. Skor Düzeyinde Birleşim ve Nedensel Normalleştirme")
    b.p("İki modelin normalleştirilmiş skorları ağırlıklı ortalamayla birleştirilir,")
    b.equation("S_bir  =  w_kal · z_kal  +  w_ham · z_ham")
    b.p(
        "burada w_kal + w_ham = 1'dir. Ağırlık burada ayarlanmamaktadır: w_kal = "
        "0,95 referans çalışmadan önsel olarak benimsenmiş olup seçimi için hiçbir "
        "değerlendirme kümesine başvurulmamıştır. Bu önemlidir, çünkü referans "
        "çalışma bu değeri kendi test kümesindeki en yüksek F1'in argmaksı olarak "
        "seçmişti; sabitlemek, seçimi başka bir kümeye taşımak yerine tamamen "
        "ortadan kaldırır. [0, 1] aralığında 0,05 adımlı tarama yine "
        "yapılmaktadır, ancak Bölüm 4.3 bunu bir arama olarak değil bir duyarlılık "
        "analizi olarak raporlar. Referans çalışma her skoru, sınırları arıza "
        "enjekte edilmiş test kümesinden gelen bir min–maks dönüşümüyle "
        "normalleştirir. Bu çevrimiçi kullanılamaz, iki sebeple: sınırlar gelecekteki "
        "maksimumları içerir ve o maksimumlar normal çalışmada hiç görülmez. O "
        "çalışmanın kendi test kümesinde ölçüldüğünde, maksimumların 4341 ve 7752 "
        "olması yüzünden tüm normal çalışma aralığı 10⁻⁴ mertebesine sıkışmakta, "
        "birleşik eşik her iki tekil eşikten de katı hâle gelmekte ve motor kayması "
        "ile gizyazar hataları hiç tespit edilememektedir."
    )
    b.p(
        "Her iki normalleştirme de afin olduğundan çevrimiçi motor tek bir kod yolu "
        "uygular,"
    )
    b.equation("z  =  (S − lo) / genişlik")
    b.p(
        "ve sınırlar test kümesi yerine temiz doğrulama pencerelerinden alınır. Bu, "
        "yayımlanmış formülü korur ve referans yöntemden sapmayı yalnızca sınırların "
        "alındığı kümenin değişmesine indirger. Birleşik çalışma eşiği, çalışmanın "
        "kendi kuralını birleşik skora genişletir: temiz doğrulama pencerelerinin "
        "97. persentili. Çevrimdışı bu, beklenen %3 yanlış alarm oranını verir. "
        "Bölüm 4.7, ne sınırların ne de eşiğin donanıma taşındığını göstermekte; "
        "Bölüm 3.11 orada bunların yerini alan yordamı vermektedir."
    )

    b.h2("3.7. Veri Bölmesi, Arıza Enjeksiyonu ve Değerlendirme Protokolü")
    b.p(
        "600 fiziksel koşu bir kez bölünür ve her tüketici aynı bölmeyi okur: Bölüm "
        "3.4'ün sürtünme uydurması, Bölüm 3.5'in özkodlayıcı eğitimi ve aşağıdaki "
        "değerlendirme. Bölme satır indeksine göre değil koşuya göredir. Referans "
        "çalışma kaydı sabit bir oranda kesmiş, sonra test kümesini kaydın her "
        "penceresinden kurmuştu; eğitim yalnız ilk kısmı kullandığı için "
        "değerlendirme pencerelerinin beşte dördü, modellerin daha önce yeniden "
        "inşa ettiği pencerelerdi. O düzen üzerinde ölçüldüğünde, eğitim aralığının "
        "içindeki temiz pencerelerin yeniden yapılanma hatası 0,007, dışındakilerin "
        "0,486 idi; her iki model için de yetmiş kat."
    )
    b.p(
        "Koşu uzunlukları saf bir bölmeyi kullanılamaz kılmaktadır: 600 koşunun "
        "521'i bir saniyeden kısayken, en uzun 32 koşu örneklerin %54'ünü ve tek "
        "başına en uzun olan %10,5'ini taşımaktadır. Koşular bu yüzden iki geçişte "
        "dağıtılır. Arıza taşıyacak kadar uzun koşular, her bölmenin kendi bağımsız "
        "arıza olaylarını alması için sayıca dağıtılır; kalan kısa koşular ardından "
        "örnek kotasını doldurur. Tablo 5 sonucu vermektedir. Eklem aralığı "
        "kapsaması sonradan denetlenmiştir: doğrulama ve test, altı eklemin "
        "hepsinde eğitim aralığının içinde kalmaktadır."
    )
    b.table(
        "Koşu-Ayrık Bölme.",
        ["Bölme", "Koşu", "Örnek", "Pay", "Pencere", "Arıza taşıyıcı"],
        [["Eğitim", "489", "235.657", "%68,8", "6.756", "74"],
         ["Doğrulama", "16", "56.244", "%16,4", "2.162", "16"],
         ["Test", "95", "50.579", "%14,8", "1.504", "16"]],
        widths=[3.2, 2.4, 3.2, 2.6, 2.9, 3.6], wide=True,
        align_right=[1, 2, 3, 4, 5],
        note="Dağıtımdan sonra doğrulanmıştır: pencere örtüşmesi de hesaba "
             "katıldığında herhangi iki bölme arasında paylaşılan hiçbir örnek "
             "yoktur ve hiçbir pencere başka bir bölmeye ait koşuya uzanmaz. Temiz "
             "pencere skorunun test/eğitim oranı, referans düzenindeki yetmişten "
             "0,48'e düşmektedir.")
    b.p(
        "Gerçek etiketler mevcut olmadığından dört sentetik arıza senaryosu "
        "enjekte edilmiştir; Tablo 6'da listelenmişlerdir. Eklem 3 (elbow) ve eklem "
        "5 (wrist 2), en büyük yerçekimi yükünü ve en geniş hareket aralığını "
        "taşıdıkları için seçilmiştir. Arızalar hedef bölmenin arıza taşıyan her "
        "koşusuna ayrı ayrı enjekte edilir; böylece her senaryo test kümesine 16 "
        "bağımsız olay katar. Referans protokolde maske, birleştirilmiş kayıt "
        "üzerinde tek bir bitişik bloktu; bu, senaryo başına bir olay veriyor ve "
        "daha kötüsü, her senaryoyu kayıttaki sabit bir konuma bağlıyordu. Bir "
        "pencere, örneklerinin en az %10'u arıza maskesine denk geldiğinde anomali "
        "olarak etiketlenir."
    )
    b.table(
        "Sentetik Arıza Senaryoları.",
        ["Senaryo", "Etkilenen kanal", "Genlik", "Aralık"],
        [["Motor kayması", "Eklem 3 torku, doğrusal rampa", "15 Nm", "%30–38"],
         ["Çarpışma", "Tüm KTS kanalları, Gauss darbesi", "30 N", "merkez %65"],
         ["Gizyazar hatası", "Eklem 5 konumu, basamak", "1,5 rad", "%92'den"],
         ["Sensör gürültüsü", "Tüm KTS kanalları, Gauss gürültüsü", "3,5 N",
          "%15–23"]],
        widths=[2.0, 3.45, 1.5, 1.4],
        note="Aralık, arızanın enjekte edildiği koşunun bir oranı olarak ifade "
             "edilmiştir. Yukarıdaki genlikler referans çalışmanınkilerdir ve "
             "ölçülen kanallara uygulanır.")
    b.p(
        "İki enjeksiyon protokolü değerlendirilmektedir. Devralınan protokol her "
        "temsil uzayını kendi genliğiyle ayrı ayrı bozar; dolayısıyla kalıntı "
        "modeli, ölçümden türetilmiş değil elle seçilmiş bir bozulma alır. Fiziksel "
        "protokol yalnız ölçülen kanalları bozar ve kalıntıyı hattın kendisiyle "
        "yeniden hesaplar; böylece kalıntıdaki bozulma, ters dinamiğin ve Jacobian "
        "aktarımının ürettiği her neyse odur. Kalıntı yeniden hesaplanırken "
        "kalibrasyon ofseti temiz hattan dondurulur; enjekte edilmiş veride yeniden "
        "uydurmak, dokunulmamış koşuların kalıntısını da kaydırır. Bölüm 4.4 her iki "
        "protokolü de raporlar, çünkü birbirleriyle uyuşmamaktadırlar."
    )

    b.h2("3.8. Önceki Sayıların Nereden Geldiğinin Doğrulanması")
    b.p(
        "Bir sonucu genişletmeden önce, önceki sayıların neyi ölçtüğünü saptamak "
        "gerekir. Bölüm 3.2–3.4 arasındaki ölçüm düzeyi denetimleri tutarlı bir "
        "düzene işaret etmekte ve o düzen sınanabilir olmaktadır. Hipotez, "
        "önceden yayımlanmış sayıların birimleri karışık bir hattan geldiğidir: "
        "τ_ölç amper, τ̂_model newton metre, r_ext newton metre ve döndürülmemiş bir "
        "wrench ile hesaplanmış, üçü doğrudan çıkarılmış ve newton metre cinsinden "
        "belirtilen arıza genlikleri amper kanallarına çıplak sayılar olarak "
        "eklenmiştir. Bir bulgu daha bunu ağırlaştırır: kalıntı üreteci, uyum "
        "kalitesi bir eşiğin altında kalan eklemlerde model torkunu sıfırlar ve "
        "karışık birimler altında bu koruma altı eklemin hepsinde devreye girerek, "
        "hiçbir mesaj vermeden kalıntıyı ölçümün kendisine indirger."
    )
    b.p(
        "Hipotez yanlışlanabilir bir öngörüde bulunur: o yapılandırmada yeniden "
        "kurulduğunda referans çalışmanın imza bulguları geri dönmelidir — "
        "özellikle kalıntı modelinin sensör gürültüsüne körlüğü ve anomalilerin "
        "yaklaşık dörtte birinin yalnız kalıntı modeliyle tespit edilmesi. Yeniden "
        "üretim betiklenmiş olup bu makale için yeniden koşturulmuştur; sonucu "
        "Bölüm 4.5'te raporlanmaktadır."
    )

    b.h2("3.9. Çevrimiçi Gerçekleme")
    b.p(
        "Çevrimiçi sistem çevrimdışı hattan üç yapısal noktada ayrılır ve tasarımı "
        "bütünüyle bu üç farkın yönetiminden ibarettir. Nedensellik: bütün bir veri "
        "kümesi üzerinden hesaplanan hiçbir şey canlıda kullanılamaz. Zaman "
        "bütçesi: 500 Hz'de örnek başına 2 ms, 25 adımla karar başına 50 ms vardır. "
        "Süreklilik: çevrimdışı hat veriyi koşulara bölebilirken canlı sistem bir "
        "kesintiyi ancak olduktan sonra fark eder ve durumunu sıfırlamak zorundadır. "
        "Şekil 2 ortaya çıkan veri yolunu göstermektedir."
    )
    b.figure(
        "fig2_architecture.png",
        "Çevrimiçi Dedektörün Veri Yolu | robotla alarm arasındaki her aşama karar "
        "başına bir kez, saniyede yirmi kez çalışır.")
    b.p(
        "Öznitelik motoru çevrimdışı hesabın örnek-örnek eşdeğeridir. 51 örneklik "
        "bir halka tampon tutar ve her yeni örnek için tamponun ortasındaki örneğin "
        "özniteliklerini üretir: önceden hesaplanmış katsayılarla iç çarpım olarak "
        "uygulanan Savitzky–Golay türevi, akım–tork dönüşümü, ters dinamik "
        "değerlendirmesi, taban çerçevesi döndürmesi ve Jacobian aktarımı. Merkezli "
        "türev 25 örneklik (50 ms) yapısal bir gecikmeye mal olur; bu, gürültüsüz "
        "bir ivmenin bedelidir ve çevrimdışı tanımla birebir aynıdır."
    )
    b.p(
        "Eşdeğerlik varsayılmak yerine doğrulanmıştır. Hazırlanan veri kümesi "
        "çevrimiçi motordan örnek örnek geçirilip çevrimdışı öznitelik dosyasıyla "
        "karşılaştırılmıştır: 8 koşu ve 61.524 örnek üzerinde en büyük mutlak "
        "farklar tork için 1,4·10⁻¹⁴, model torku için tam sıfır, toplam kalıntı "
        "için 8,0·10⁻¹⁵, dışsal için 2,7·10⁻¹⁵ ve içsel kalıntı için 8,9·10⁻¹⁵'tir. "
        "Bunlar kayan nokta yuvarlama düzeyleridir; dolayısıyla dağıtılan düğüm, "
        "modellerin üzerinde eğitildiği özniteliklerin benzerini değil aynısını "
        "görmektedir. Sınama sürtünme terimi etkinken yapılmıştır."
    )
    b.p(
        "Referans yöntemin üzerine iki isteğe bağlı kural eklenmiştir. "
        "Uyarlanabilir kural, birleşik skor son 600 kararın medyan + k·1,4826·MAD "
        "değerini aştığında tetiklenir; burada 1,4826·MAD standart sapmanın sağlam "
        "bir kestirimidir (Leys, Ley, Klein, Bernard ve Licata, 2013). Bu kural, "
        "mutlak bir eşiğin, taban çizgisi o eşiğin çok altında oturan bir koşu "
        "içinde skoru yüz kat yükselten bir arızayı göremeyeceği için vardır. Ölçek, "
        "neredeyse sabit bir taban çizgisinin kuralı aşırı hassas kılmaması için "
        "tabanlanır; taban çizgisi alarm sürerken dondurulur ve dondurma 3 s sonra "
        "sona erer, böylece uzun süren bir alarm saldırı değil rejim değişimi "
        "sayılır. İkinci kural, alarm verilmeden önce iki ardışık karar (100 ms) "
        "gerektirir. Her ikisi de tek bir parametreyle kapatılabilir; bu, referans "
        "davranışı birebir geri getirir."
    )
    b.p(
        "Dedektör, tüm hesabı ROS'tan bağımsız bir çekirdekte tutan bir ROS 2 "
        "düğümü olarak paketlenmiştir (Macenski, Foote, Gerkey, Lalancette ve "
        "Woodall, 2022); düğüm bunun ince bir sarmalayıcısıdır, dolayısıyla tekrar "
        "oynatma sınamaları düğümün gerçekten çalıştırdığı sınıfı sınar. Modeller "
        "ONNX'e aktarılır ve ONNX Runtime ile çalıştırılır; oturum başına iki iş "
        "parçacığı verilerek aç gözlü bir çıkarımın 500 Hz geri çağırma döngüsünü "
        "aç bırakması engellenir. Her karar virgülle ayrılmış bir kayda, her alarm "
        "satır tamponlu bir JSON-lines olay kaydına yazılır; Bölüm 4.7'nin gerçek "
        "hücre eşikleri bu iki dosyadan yeniden ölçülmüştür."
    )
    b.p(
        "Dağıtım, üçü de sessiz olan üç bağlantı katmanı ölçümünü açığa "
        "çıkarmıştır. Denetleyici yapılandırmasında ilan edilen wrench konusu, "
        "çalışan yayıncının kullandığı konu değildir ve yanlış ada abone olmak "
        "hiçbir wrench, dolayısıyla hiçbir skor vermez. Eklem durumu konusu üç "
        "yayıncıdan üç farklı ad kümesi taşır; bunların içinde UR verisi saniyede "
        "yaklaşık 112 UR dışı mesajla birlikte, yedi elemanlı ve karışık sıralı bir "
        "kümede gelir. Eklem eşlemesi bu yüzden mesaj başına çözülüp ad kümesi "
        "başına önbelleklenmelidir; bir kez çözüp körü körüne uygulamak altı "
        "kanalın beşini yanlış eklemden okur. Son olarak ölçülen örnekleme hızı 500 "
        "Hz değil 495 Hz'dir; sabit adımlı türev bunu yüzde bir düzeyinde tolere "
        "eder."
    )

    b.h2("3.10. Operatör Arayüzü")
    b.p(
        "Dedektör, laboratuvarın mevcut web panosuna üçüncü bir sekme olarak "
        "eklenmiştir. Bir toplayıcı düğümün karar, alarm ve skor konularına abone "
        "olur, 60 saniyelik bir halka tampon tutar ve bunu tarayıcıya 5 Hz'de "
        "gönderir. Her toplayıcı kendi tek iş parçacıklı yürütücüsünü çalıştırır: "
        "ilk sürümde her iki toplayıcı da paylaşılan genel yürütücüyü kullanıyordu; "
        "iki iş parçacığı aynı bekleme kümesinde yarıştı, düğüm çizgeden kayboldu, "
        "abonelikler hiçbir şey almadı ve hiçbir hata basılmadı."
    )
    b.p(
        "Üç arayüz kararı zevkten değil doğrudan ölçümlerden gelmektedir. Alarmlar "
        "yükselen kenarda mandallanır ve beş saniye tutulur; çünkü kampanyanın en "
        "net doğrulanmış olayı 0,25 saniye — beş karar — sürmüştür ve 5 Hz'de "
        "seviye örneklemesi bunu tamamen kaçırır; sentetik 0,25 s alarmlarla yapılan "
        "kontrollü bir sınamada kenar yakalama dört olayın dördünü, seviye "
        "örneklemesi yalnız ikisini görmüştür. Dikey eksen doğrusaldır ve sabit bir "
        "aralık yerine canlı eşiğe göre çerçevelenir; çünkü dedektör yeniden "
        "kalibre edildiğinde skor ölçeği değişir ve sabit bir çerçeve tüm izi bir "
        "kenara yapıştırır. Logaritmik eksen reddedilmiştir, çünkü eşiğin "
        "üstündeki — tek önemli olan — bölgeyi görsel olarak düzleştirir. Olay "
        "tablosunda tepe skor birincil, giriş skoru ikincil sütundur; çünkü tepe "
        "değeri görülmeden okunan bir giriş değeri, devreye alma sırasında yanlış "
        "değerlendirilen bir olaya yol açmıştır — Tablo 14'ün 3. olayı 13,75'te "
        "girip 25,04'te tepe yapmaktadır. Operatör etiketleri, dedektörün hiç "
        "dokunmadığı ayrı bir dosyaya atomik olarak yazılır ve Bölüm 4.7'nin "
        "dayandığı etiketli kümeyi oluşturur. Hem eksen hem de gösterilen model "
        "eşikleri, derlenmiş değerlerden değil canlı karar akışından okunur; böylece "
        "bir yeniden kalibrasyon, ekranı artık çalışmayan bir yapılandırmayı "
        "anlatır hâlde bırakamaz."
    )

    b.h2("3.11. Hücrede Eşik Kalibrasyonu")
    b.p(
        "Bölüm 3.6'nın çevrimdışı yordamı (lo, genişlik) ve θ'yı temiz doğrulama "
        "pencerelerinden sabitler. Bölüm 4.7 bunların taşınmadığını göstermektedir: "
        "modeller, gerçek hücrenin olağan çalışması sırasında eğitim dağılımlarının "
        "çok dışındadır; dolayısıyla çevrimdışı bir dağılımın 97. persentiline "
        "yerleştirilen bir eşik, hücrenin normal aralığının altında kalır. "
        "Dağıtılan sistem bu yüzden aynı üç büyüklüğü hücrede yeniden ölçer; aynı "
        "formülü kullanır ve yalnız kestirildikleri kümeyi değiştirir — temiz "
        "doğrulama pencereleri yerine robotta kaydedilmiş temiz kararlar."
    )
    b.p(
        "İki ayrıntı ölçümle zorunlu kılınmıştır. Birincisi, ölçek referans "
        "formülün min–maks aralığı değil, her modelin kendi temiz skorlarının 97. "
        "persentilidir: hücrede min–maks aralığı Bölüm 4.7'deki yük platosunca "
        "belirlenir, bu da kalıntı genişliğini 30,8'e şişirir ve olağan büyüklükteki "
        "bir olayı görünmez kılar. İkincisi, kalibrasyon koşusu hücrenin yaptığı her "
        "görevi kapsamalıdır. Yalnız muayene çevriminden alınan bir kalibrasyon, "
        "pick-and-place görevini aralığının dışında bırakmış ve dedektör normal "
        "taşıma sırasında alarm vermiştir. Dağıtılan yapılandırma, her iki görevi "
        "kapsayan 15.627 temiz karar üzerinde kalibre edilmiştir."
    )


# ── 4. Bulgular ──
def _bulgular(b: Builder) -> None:
    b.h1("4. Bulgular")

    b.h2("4.1. Genişletilmiş Hattın Çevrimdışı Başarımı")
    b.p(
        "Genişletilmiş hat, 553'ü (%9,2) anomali olan 6.016 test penceresi "
        "üzerinde değerlendirilmiştir; bu pencereler, ne eğitime, ne sürtünme "
        "uydurmasına, ne de herhangi bir eşiğin seçimine katkı vermemiş 95 koşudan "
        "gelmektedir. Aşağıdaki her sayı beş eğitim tohumunun ortalaması ve standart "
        "sapmasıdır. Tablo 7 sonucu, Şekil 3 ise ilgili eğrileri vermektedir."
    )
    b.table(
        "Koşu-Ayrık Test Kümesinde Genel Başarım (beş tohum, ort ± s.s.).",
        ["Model", "AUC", "PR-AUC", "En iyi F1"],
        [["Kalıntı LSTM ÖK", "0,878 ± 0,005", "0,653 ± 0,009", "0,628 ± 0,007"],
         ["Ham LSTM ÖK", "0,870 ± 0,008", "0,668 ± 0,009", "0,697 ± 0,001"],
         ["**Birleşim (0,95/0,05)**", "**0,969 ± 0,004**", "**0,857 ± 0,009**",
          "**0,821 ± 0,006**"],
         ["Birleşim MAKS", "0,966 ± 0,004", "0,853 ± 0,010", "0,784 ± 0,012"],
         ["Kalıntı norm eşiği", "0,758 ± 0,000", "0,485 ± 0,000", "0,495 ± 0,000"],
         ["Isolation Forest", "0,753 ± 0,008", "0,418 ± 0,008", "0,415 ± 0,008"],
         ["One-Class SVM", "0,844 ± 0,001", "0,618 ± 0,001", "0,676 ± 0,000"]],
        widths=[2.9, 1.9, 1.9, 1.65], align_right=[1, 2, 3],
        note="Önceki çalışmanın sayıları bunların yanına konmamıştır; çünkü o "
             "değerlendirme, eğitim pencereleriyle örtüşen değerlendirme "
             "pencerelerine sahip bir protokol altında ve çok daha büyük, farklı bir "
             "pencere kümesi üzerinde yapılmıştır. Karşılaştırma tek bir yerde, "
             "Bölüm 4.5'te ele alınmaktadır. O çalışma ayrıca MAKS satırıyla birebir "
             "aynı değerleri taşıyan bir OR satırı listeler. OR bir karar düzeyi "
             "birleşimidir: sıralama değil ikili çıktı üretir, dolayısıyla AUC, "
             "PR-AUC ve En iyi F1 onun için tanımsızdır ve satır aslında MAKS "
             "skorunu raporlamaktadır. Çalışma noktası davranışı bunun yerine Bölüm "
             "4.2'de verilmektedir. En iyi F1 tüm eşikler üzerinden maksimumdur ve "
             "bu yüzden tahmin etmesi gereken etiketleri gerektirir; yalnız "
             "karşılaştırılabilirlik için listelenmiştir, ki Bölüm 3.6 bu nedenle "
             "nedensel bir çalışma eşiği tanımlar. Referans yöntemler temiz eğitim "
             "koşularında uydurulmuştur.")
    b.figure(
        "fig3_pr_roc.png",
        "Koşu-Ayrık Test Kümesinde Kesinlik–Geri Çağırma ve ROC Eğrileri | 6.016 "
        "test penceresi, %9,2 anomali, beş tohumdan 1'i. Birleşim her iki eğride de "
        "iki tekil modeli baskılar; kesinlik–geri çağırma panelindeki fark daha "
        "büyüktür, çünkü ROC sınıf dengesizliğinde iyimserdir.")
    b.p(
        "Nedensel çalışma noktasında birleşik dedektör 0,764 ± 0,086 kesinlik, "
        "0,831 ± 0,055 geri çağırma ve 0,791 ± 0,023 F1 değerine ulaşmaktadır; buna "
        "karşılık yalnız ham model için F1 0,627 ± 0,006 ve yalnız kalıntı modeli "
        "için 0,614 ± 0,016'dır. Daha iyi tekil modele göre kazanç ΔF1 = +0,164'tür. "
        "Temiz pencere skorunun test/eğitim oranı her iki model için de 0,48'dir; "
        "referans bölmesinin ürettiği yetmiş katın yerine. Önceki düzenin getirdiği "
        "iyimserlik ortadan kalkmıştır."
    )
    b.p(
        "±0,086'lık kesinlik yayılımı modellerdeki değil eşikteki gürültüdür. "
        "Eşiğin değeri beş tohum arasında ortalamasının %46'sı kadar değişmektedir; "
        "çünkü o, Bölüm 3.7'nin bölme tasarımınca daraltılmış bir doğrulama "
        "dağılımının 97. persentilidir. Eşikten bağımsız olan sıralama ölçütleri "
        "0,01'den az değişmektedir."
    )

    b.h2("4.2. Arıza Tipi Asimetrisi ve Tamamlayıcılık")
    b.p(
        "Tablo 8 sonucu arıza tipine göre ayırmakta ve AUC ile En iyi F1'i ayrı "
        "raporlamaktadır; çünkü bunlar aynı büyüklük değildir ve ayrım okumayı "
        "değiştirir: motor kaymasında kalıntı modelinin AUC'si 0,912 iken En iyi "
        "F1'i 0,669'dur ve ikincisini birincisiymiş gibi aktarmak sıralama "
        "kalitesini geniş bir farkla eksik gösterir."
    )
    b.table(
        "Her İki Enjeksiyon Protokolü Altında Arıza Tipine Göre Tespit "
        "(AUC, beş tohum).",
        ["Senaryo", "Devralınan: kalıntı", "Devralınan: ham",
         "Fiziksel: kalıntı", "Fiziksel: ham"],
        [["Motor kayması", "0,912 ± 0,013", "0,652 ± 0,014", "0,846 ± 0,021",
          "0,652 ± 0,014"],
         ["Çarpışma", "1,000 ± 0,000", "1,000 ± 0,000", "1,000 ± 0,000",
          "1,000 ± 0,000"],
         ["Gizyazar hatası", "1,000 ± 0,000", "0,832 ± 0,041", "0,783 ± 0,052",
          "0,832 ± 0,041"],
         ["Sensör gürültüsü", "0,682 ± 0,006", "1,000 ± 0,000", "1,000 ± 0,000",
          "1,000 ± 0,000"]],
        widths=[2.7, 3.1, 2.9, 3.1, 2.9], wide=True,
        align_right=[1, 2, 3, 4],
        note="Ham sütunlar iki protokolde birebir aynıdır; çünkü ham model ölçülen "
             "kanalları okur ve her iki protokol de o kanalları aynı biçimde bozar. "
             "Yalnız kalıntı sütunları değişmekte ve epeyce değişmektedir — yorumu "
             "Bölüm 4.4'tedir.")
    b.figure(
        "fig4_fusion_value.png",
        "Birleşimin Değeri | (a) kalıntı ağırlığına karşı sıralama kalitesi; "
        "ağırlık seçilmemiş, 0,95'te önsel olarak sabitlenmiştir. (b) devralınan "
        "enjeksiyon protokolü altında arıza tipine göre tespit.")
    b.p(
        "Devralınan protokol altında birleşimi güdüleyen asimetri mevcut ve "
        "nettir. Gizyazar hatasında kalıntı özkodlayıcısı kusursuzken ham model "
        "0,832'ye ulaşır; sensör gürültüsünde roller tersine döner, kalıntı modeli "
        "0,682'ye düşerken ham model kusursuzdur. Pencere düzeyinde, her model kendi "
        "97. persentil eşiğini kullanırken, anomali pencerelerinin %27,6'sı yalnız "
        "kalıntı modeliyle ve %28,0'ı yalnız ham modelle tespit edilmektedir; "
        "referans çalışmada bu oranlar %24,1 ve %25,9'dur. %15,9'u ise hiçbiri "
        "tarafından yakalanmamaktadır. OR birleşiminin geri çağırması 0,841'dir. "
        "Tablo 9 karar düzeyi ablasyonunu vermektedir."
    )
    b.table(
        "Çalışma Noktasında Karar Düzeyi Ablasyonu (beş tohum, ort ± s.s.).",
        ["Yapılandırma", "Kesinlik", "Geri çağırma", "F1"],
        [["Yalnız ham  (w_kal = 0,00)", "0,704 ± 0,015", "0,565 ± 0,002",
          "0,627 ± 0,006"],
         ["Yalnız kalıntı  (w_kal = 1,00)", "0,697 ± 0,111", "0,562 ± 0,058",
          "0,614 ± 0,016"],
         ["OR  (karar düzeyi)", "0,701 ± 0,047", "0,841 ± 0,058", "0,762 ± 0,013"],
         ["**Birleşim  (w_kal = 0,95)**", "**0,764 ± 0,086**", "**0,831 ± 0,055**",
          "**0,791 ± 0,023**"]],
        widths=[3.05, 1.95, 1.75, 1.6], align_right=[1, 2, 3],
        note="Devralınan enjeksiyon protokolü. Daha iyi tekil modele göre kazanç "
             "ΔF1 = +0,164'tür; referans çalışma farklı bir mutlak düzeyde ve "
             "örtüşen pencereli bir protokol altında +0,161 raporlamaktadır.")

    b.h2("4.3. Birleşim Ağırlığına Duyarlılık")
    b.p(
        "Ağırlık w_kal = 0,95'te önsel olarak sabitlenmiştir (Bölüm 3.6); "
        "dolayısıyla aşağıdaki tarama bir seçim yapmaz, duyarlılık ölçer. Şekil "
        "4(a) yaklaşık [0,25, 0,95] aralığında geniş bir plato göstermektedir; bu, "
        "referans çalışmanın raporladığı kararlı bölgeyi yeniden üretmekte ve ham "
        "model tamamen elendiğinde keskin bir düşüşle bitmektedir. Devralınan "
        "protokol altında w_kal = 0,95'teki birleşik PR-AUC 0,857 ± 0,009 iken saf "
        "kalıntı dedektörü için 0,653 ± 0,009'dur; yani ham modele verilen yüzde "
        "beşlik ağırlık +0,204 PR-AUC ile ilişkilidir. Daha iyi tekil modele göre "
        "marj olan +0,189 ± 0,009, beş tohumun her birinde pozitiftir."
    )
    b.p(
        "İki bağımsız denetim her iki modelin de gerçekten çalıştığını "
        "doğrulamaktadır. Yapısal olarak, referans yöntemin her kuralı dağıtılan "
        "karar yoluna eşlenmektedir. İşlevsel olarak, her ONNX oturumunun skorlama "
        "yöntemi sarmalanıp 6 koşu ve 5 senaryo boyunca sayılmıştır: 13.160 karar, "
        "13.160 kalıntı modeli çağrısı ve 13.160 ham model çağrısı üretmiştir; "
        "dolayısıyla hiçbir model önbelleklenmemekte ya da atlanmamaktadır."
    )

    b.h2("4.4. Birleşim Marjı Neye Bağlıdır")
    b.p(
        "Yukarıda raporlanan marj, iki temsil uzayının değil devralınan enjeksiyon "
        "protokolünün bir özelliğidir. Protokol kalıntı kanallarını elle seçilmiş "
        "bir genlikle doğrudan bozmakta ve bu seçimlerden biri tamamlayıcılık "
        "sonucunun tamamını taşımaktadır. Sensör gürültüsü için referans çalışma "
        "wrench'e 3,5 N ve ayrıca dışsal kalıntıya 0,08 Nm eklemekte, ikincisini "
        "senaryonun 'yalnız ham sinyali etkilemesi' için kasten küçük tutulmuş "
        "olarak tanımlamaktadır. Fiziksel olarak, 3,5 N'luk bir wrench bozulması "
        "kalıntıya r_ext = J(q)ᵀF üzerinden yayılır; bu robotta ölçüldüğünde kalıntı "
        "kanallarında 4,1 Nm'lik bir standart sapma üretmektedir — elle enjekte "
        "edilen genliğin elli bir katı. Ters dinamik wrench'i süzmez, dönüştürür."
    )
    b.p(
        "Karşılaştırma tek yanlı değildir. Çarpışma senaryosunda elle seçilen 40 "
        "Nm'lik kalıntı genliği, tepe değeri 80,5 Nm olan fiziksel yayılımla "
        "uyuşmaktadır — aynı mertebe. Sensör gürültüsü ikisinin uyuşmadığı tek "
        "senaryodur ve referans çalışmanın tamamlayıcılığın sınavı olarak "
        "adlandırdığı senaryo da odur."
    )
    b.p(
        "Tablo 10, iki enjeksiyon protokolünü sürtünme teriminin varlığıyla "
        "çaprazlamakta ve her hücrede beş tohum kullanmaktadır. Sonuç açıktır: "
        "birleşim marjı devralınan protokolde her iki sürtünme koşulunda da mevcut, "
        "fiziksel protokolde her ikisinde de yoktur."
    )
    b.table(
        "İki Enjeksiyon Protokolü ve İki Kalıntı Tanımı Altında Birleşim Marjı "
        "(PR-AUC, beş tohum, ort ± s.s.).",
        ["Yapılandırma", "Kalıntı", "Ham", "Birleşim", "Marj"],
        [["Devralınan / sürtünmesiz", "0,620 ± 0,007", "0,668 ± 0,009",
          "0,829 ± 0,007", "**+0,161 ± 0,005**"],
         ["Devralınan / sürtünmeli", "0,653 ± 0,009", "0,668 ± 0,009",
          "0,857 ± 0,009", "**+0,189 ± 0,009**"],
         ["Fiziksel / sürtünmesiz", "**0,750 ± 0,017**", "0,668 ± 0,009",
          "0,747 ± 0,016", "−0,003 ± 0,001"],
         ["Fiziksel / sürtünmeli", "0,730 ± 0,014", "0,668 ± 0,009",
          "0,727 ± 0,014", "−0,003 ± 0,001"]],
        widths=[3.6, 2.6, 2.6, 2.6, 3.0], wide=True,
        align_right=[1, 2, 3, 4],
        note="Marj, birleşik PR-AUC eksi iki tekil modelin iyisidir. Ham sütun "
             "yapısı gereği sabittir; bu aynı zamanda ablasyonun kontrollü "
             "olduğunun bir denetimidir. Devralınan protokolde marj her iki satırın "
             "beş tohumunda da pozitiftir; fiziksel protokolde her iki satırın beş "
             "tohumunda da −0,003'tür, yani yüzde beşlik ham ağırlık az da olsa "
             "maliyetlidir ve hiçbir şey getirmez. Pencere düzeyinde tamamlayıcılık "
             "aynı yönde çökmektedir: yalnız kalıntıyla tespitler %27,6'dan %4,8'e, "
             "yalnız hamla tespitler %28,0'dan %0,1'e düşmektedir.")
    b.p(
        "Aynı tablo ikinci bir soruyu da yanıtlar. Sürtünme terimini eklemek, "
        "devralınan protokol altında kalıntı modelini +0,034 ± 0,009 PR-AUC "
        "iyileştirmekte ve beş tohumun hepsinde aynı yönde davranmaktadır; buna "
        "karşılık fiziksel protokolde değişim −0,020 ± 0,027'dir, anlamlı değildir "
        "ve ters yönü göstermektedir. Sürtünme teriminin gösterilebilir faydası bu "
        "nedenle fiziksel sadakattir — Tablo 3'teki %87 ve %92'lik bilek kalıntısı "
        "azalması ve newton metre cinsinden yorumlanabilir kalan arıza genlikleri — "
        "bu dört sentetik arızanın tespiti değil. Ayrıca ölçülmüş bir bedeli "
        "vardır: çalışma eşiğinin tohumdan tohuma yayılımı, her iki protokolde de "
        "ortalamasının %17'sinden %46'sına çıkmaktadır."
    )
    b.p(
        "Bunların hiçbiri ham modelin işe yaramaz olduğunu söylemez. Söylediği "
        "şudur: bu dört sentetik senaryo, tutarlı biçimde enjekte edildiğinde iki "
        "temsil uzayı arasında ayrım yapmamaktadır. Bölüm 4.7, iki modelin donanım "
        "üzerinde gerçek olaylar sırasında ne yaptığını raporlamaktadır ve buradaki "
        "gerçek arızalara ilişkin tek kanıt odur."
    )

    b.h2("4.5. Önceki Sayıların Kaynağı")
    b.p(
        "Bu alt bölüm, bu çalışmanın genişlettiği çalışmayla karşılaştırmayı tek "
        "bir yerde toplamaktadır. Bölüm 3.8'in karışık birimli yapılandırmasında "
        "yeniden kurulduğunda hat, önceki sayıları nokta nokta yeniden "
        "üretmektedir (Tablo 11); bu, onların hangi düzende üretildiğini ve "
        "dolayısıyla hangi fiziksel ölçekte okunmaları gerektiğini belirlemektedir."
    )
    b.table(
        "Önceki Hattın Yeniden Üretimi.",
        ["Büyüklük", "Yeniden kurulum", "Önceki", "Δ"],
        [["Test penceresi", "179.888", "179.896", "8"],
         ["Arıza penceresi", "14.401 (%8,01)", "14.402 (%8,0)", "1"],
         ["Yalnız kalıntı payı", "%23,9", "%24,1", "0,2 puan"],
         ["Sensör gürültüsü — kalıntı", "0,267", "0,272", "0,005"],
         ["OR geri çağırma", "0,841", "0,854", "0,013"],
         ["Geri çağırma — ham", "0,601", "0,613", "0,012"],
         ["Gizyazar hatası — kalıntı", "0,975", "0,986", "0,011"],
         ["Ham model AUC", "0,943", "0,952", "0,009"],
         ["Birleşim MAKS — AUC", "0,974", "0,976", "0,002"]],
        widths=[3.3, 2.0, 1.6, 1.4], align_right=[1, 2, 3],
        note="Belirleyici satır sensör gürültüsüdür. O senaryoda 0,267 alan bir "
             "kalıntı modeli, ancak kalıntının fiziksel anlamını yitirdiği, yani "
             "amper ile newton metrenin doğrudan çıkarıldığı bir hatta ortaya "
             "çıkabilir. Pencere sayılarının 179.896'da 8 ve 1 pencereye kadar "
             "uyuşması, aynı ham veri kümesinin aynı biçimde bölündüğünü "
             "doğrulamaktadır.")
    b.p(
        "Önceki çalışmanın yapısal iddiası bundan etkilenmemektedir. Kendi "
        "enjeksiyon protokolü altında tamamlayıcılık her iki hatta da ölçülmekte — "
        "yeniden üretimde %23,9, genişletilmiş hatta %27,6 yalnız kalıntı — ve "
        "birleşim marjı ikisinde de ayakta kalmaktadır. Değişen şey, sayıların "
        "hangi fiziksel ölçekte okunması gerektiği ve ayrıca arızalar fiziksel "
        "olarak enjekte edildiğinde o marja ne olduğudur (Bölüm 4.4). Ölçeği "
        "saptamak, bu çalışmanın geri kalanını mümkün kılan şeydir: newton metre "
        "cinsinden arıza genlikleri, anlamlı katsayılara sahip bir sürtünme terimi "
        "ve büyüklüğü gerçek bir temas kuvvetiyle karşılaştırılabilir bir kalıntı, "
        "hepsi bunu gerektirir."
    )

    b.h2("4.6. Gecikme ve Hesap Bütçesi")
    b.p(
        "Örnek başına ters dinamik değerlendirmesi 22 µs, Jacobian ise yaklaşık 30 "
        "µs tutmakta; ikisi birlikte 500 Hz bütçesinin %1,5'idir. Karar başına iki "
        "ONNX ileri geçişi 5,9 ms, yani 50 ms bütçesinin %12'sidir; bu değer "
        "çevrimdışı tekrar oynatmada CPU üzerinde ölçülmüştür. Hücrede ise 8,5 "
        "ms'dir (bütçenin %17'si), çünkü ROS 2 yürütücüsü ve kayıt katmanı aynı "
        "çekirdekleri paylaşmaktadır; dağıtım iş istasyonunda CUDA sağlayıcısıyla "
        "her ikisi de daha da düşmektedir. Tespit gecikmesi 50 ms süzgeç "
        "gecikmesi ile 50 ms karar periyodunun toplamıdır; iki ardışık karar kuralı "
        "etkinken bir karar periyodu daha eklenir ve uçtan uca 120–150 ms eder. "
        "0,25 m/s'lik takım hızında bu, yaklaşık 3,75 cm ek yol demektir."
    )
    b.p(
        "Hazırlanan verinin dağıtılan sınıftan uçtan uca tekrar oynatılması — bir "
        "arıza taşıyacak kadar uzun beş test koşusu üzerinde, 27.135 örnek, 0,9 "
        "dakika robot zamanı — temiz veride %0,1 yanlış alarm oranı vermektedir. "
        "Çarpışma ve sensör gürültüsü 5 koşunun 5'inde, sırasıyla 60 ms ve 70 ms "
        "medyan gecikmeyle tespit edilmektedir; motor kayması 5 koşunun 2'sinde "
        "2.905 ms'de; gizyazar hatası ise hiçbirinde. Gecikme ayrımı yapısaldır — "
        "200 ms'lik bir pencere içinde yavaş bir rampa neredeyse sabit bir kaymadır "
        "ve bir özkodlayıcı sabit bir kaymayı zorlanmadan yeniden inşa eder — ancak "
        "gizyazar sonucu bir gecikme etkisi değildir. O, Tablo 8'in fiziksel "
        "enjeksiyon sonucunun karar düzeyinde görünen hâlidir: eklem 5'teki 1,5 "
        "rad'lık bir basamak kalıntıya, devralınan protokolün elle seçtiği 8 Nm "
        "olarak değil Jacobian üzerinden yayılmaktadır ve yayılan bozulma eşiği "
        "aşmamaktadır. Örneklem küçüktür ve öyle raporlanmaktadır: koşu-ayrık test "
        "bölmesinin o uzunlukta içerdiği koşu sayısı beştir."
    )

    b.h2("4.7. Gerçek Robotta Devreye Alma")
    b.p(
        "Buraya kadarki tüm sonuçlar kayıtlı veri kümesi üzerinde elde edilmiştir. "
        "Dedektör ardından fiziksel UR10e hücresinde devreye alınmış; platformun "
        "iki üretim senaryosu — otonom görsel muayene çevrimi ve bir pick-and-place "
        "senaryosu — ile anomalilerin kasten provoke edildiği koşular "
        "kapsanmıştır. Her karar diske, her alarm bloğu bir olay kaydına ve her "
        "oturum; model sağlama toplamlarını, eşikleri ve yürürlükteki kalıntı "
        "tanımını adlandıran bir köken kaydına yazılmıştır. Bu kayıt bir "
        "defter tutma işi değildir: bir ilk deneme, yazılım paketi onları "
        "varsayılan olarak dağıttığı için geçersiz kılınmış modellerle koşmuş ve "
        "bunu başka hiçbir şey açığa çıkarmazdı. Tablo 12, çevrimiçi sistemin "
        "referans yöntemden nasıl saptığını özetlemektedir."
    )
    b.table(
        "Çevrimiçi Sistemin Referans Yöntemden Sapmaları.",
        ["Sapma", "Durum", "Gerekçe"],
        [["D1  Normalleştirme sınırları ve eşik hücrede ölçüldü", "zorunlu",
          "çevrimdışı değerler taşınmıyor; aşağıda ölçülmüştür"],
         ["D2  Birleşik skor için çalışma eşiği", "zorunlu",
          "referans yöntem yalnız En iyi F1 raporlar, o da tahmin ettiği "
          "etiketleri gerektirir"],
         ["D3  Uyarlanabilir medyan + k·MAD kuralı", "isteğe bağlı, **kapalı**",
          "çevrimdışı yardımcı olur, donanımda doğrulanmış hiçbir olayı yakalamaz"],
         ["D4  Alarmdan önce iki ardışık karar", "isteğe bağlı, açık",
          "tek karardan doğan gürültü tetiklerini bastırır"]],
        widths=[6.3, 2.4, 5.6], wide=True,
        note="D3 ve D4 birer parametreyle kapatılabilir; bu, referans davranışı "
             "birebir geri getirir. D1 yayımlanmış min–maks yapısını korur ve "
             "yalnız büyüklüklerin kestirildiği kümeyi değiştirir.")
    b.p(
        "**Çevrimdışı eşik taşınmamaktadır.** Bölüm 3.6'nın kuralıyla temiz "
        "doğrulama pencerelerinden türetilen birleşik eşik θ = 0,855 iken, hücrede "
        "temiz bir koşunun medyan birleşik skoru 0,97'dir: tüm kararların %58,2'si "
        "eşiği aşmıştır. Skorlar ölçek bakımından olağandışı değildir — Tablo "
        "13'ün hücrede kalibre edilmiş yapılandırması altında aynı temiz çalışma, "
        "eşiğin 1,7 ilâ 35,3 katı altında kalmaktadır — dolayısıyla başarısızlık, "
        "modellerde değil eşiğin nereden geldiğinde açıkça yerini bulmaktadır."
    )
    b.p(
        "Sebep, modellerin olağan çalışma sırasında eğitim dağılımlarının çok "
        "dışında olmasıdır. Çevrimdışı doğrulama kümesinin ortalama yeniden "
        "yapılanma hatasına göre ölçüldüğünde, kalıntı modeli hücrede o değerin on "
        "bir katında, ham model ise seksen altı katında çalışmaktadır. Her modelin "
        "kendi P97 eşiğine göre ifade edildiğinde asimetri daha da keskindir: "
        "kalıntı modeli eşiğinin 0,93'ünde oturmakta, ki bu iyi kalibre demektir; "
        "ham model ise 7,28'inde. Sonuç, birleşimin kendisinin bozulmasıdır. "
        "w_ham = 0,05 iken ham model medyanda birleşik skorun %29'unu sağlamakta ve "
        "tepe noktalarında ona hâkim olmaktadır; çünkü normalleştirme sabiti hücre "
        "için yirmi beş kat küçüktür. Her iki ölçeğin hücrede yeniden kestirilmesi "
        "(Bölüm 3.11) bu payı ağırlığıyla uyumlu biçimde %8–15'e geri "
        "döndürmektedir."
    )
    b.p(
        "Bir düzeltme daha ölçümle zorunlu kılınmıştır. Referans formülün min–maks "
        "biçimi gerçek hücre verisine uygulandığında kalıntı genişliğini 30,8'e "
        "kilitlemektedir; çünkü kalibrasyon koşusu yük taşıyan bir evre içerir — "
        "birleşik skorun 85 saniyelik, içi neredeyse hiç kıpırdamayan bir platosu "
        "(medyan 14,20, p90 14,39, maksimum 14,5). Ters dinamikte yük modeli "
        "bulunmadığı için herhangi bir şey taşımak sürekli bir yanlılık "
        "üretmektedir ve bunun üzerine uydurulan bir min–maks ölçeği olağan "
        "olayları görünmez kılar. Onun yerine her modelin temiz skorlarının 97. "
        "persentili kullanılmaktadır."
    )
    b.p(
        "Hücrede kalibre edilmiş yapılandırmayla — kalıntı genişliği 25,92, ham "
        "genişlik 22,37, θ = 5,0, her iki görevi kapsayan 15.627 temiz karardan "
        "kestirilmiş — dedektör 14,3 dakikalık olağan çalışma boyunca hiç alarm "
        "üretmemiştir (Tablo 13)."
    )
    b.table(
        "Hücrede Kalibre Edilmiş Yapılandırmada Temiz Çalışma.",
        ["Koşu", "Süre", "Medyan", "Maksimum", "θ'ya pay", "Alarm"],
        [["Muayene çevrimi", "6,2 dk", "0,055", "1,28", "3,9×", "**0**"],
         ["Provoke edilen el teması", "1,3 dk", "0,023", "0,14", "35,3×", "0"],
         ["Pick and place", "6,9 dk", "0,075", "2,92", "1,7×", "**0**"]],
        widths=[4.0, 2.6, 2.4, 2.6, 3.0, 2.4], wide=True,
        align_right=[1, 2, 3, 4, 5],
        note="Pay, θ'nın koşunun tepe değerine bölümüdür. İkinci satır, dedektörün "
             "kaydetmediği provoke edilmiş bir anomalidir; aşağıda tartışılmakta ve "
             "sistemin çalışma sınırını oluşturmaktadır.")
    b.p(
        "**Çarpışmalar tespit edilmektedir.** Toplam 10,0 dakikalık üç koşuluk bir "
        "oturumda vakum tutucuya karşı iki çarpışma provoke edilmiştir. Üç alarm "
        "bloğu oluşmuş; operatör arayüz üzerinden üçünü de gerçek olarak "
        "etiketlemiş ve etiketsiz hiçbir blok meydana gelmemiştir. Tablo 14 "
        "bunları listelemekte, Şekil 5 ise en güçlüsünün hemen ardındaki operatör "
        "arayüzünü göstermektedir."
    )
    b.table(
        "Fiziksel Hücrede Provoke Edilen Çarpışmalar, Tümü Operatörce Doğrulanmış.",
        ["Olay", "Süre", "Tepe", "Giriş", "Tetikleyen", "Kalıntı", "Ham"],
        [["1", "3,65 s", "**11,30**", "6,63", "ikisi", "170,1  (124×θ)",
          "175,9  (447×θ)"],
         ["2", "0,25 s", "**16,13**", "14,60", "ikisi", "369,8  (270×θ)",
          "466,5  (1186×θ)"],
         ["3", "44,95 s", "**25,04**", "13,75", "ikisi", "348,9  (254×θ)",
          "427,7  (1087×θ)"]],
        widths=[1.4, 2.0, 2.0, 1.8, 1.7, 4.4, 4.6], wide=True,
        align_right=[1, 2, 3, 5, 6],
        note="Kalıntı ve ham sütunları, modelin kendi yeniden yapılanma hatasını ve "
             "bunun o modelin P97 eşiğine oranını vermektedir; oranlar dağıtılan "
             "tohumun eşikleriyle (1,3723 ve 0,3934) hesaplanmıştır. Her üç olayda "
             "da iki model kendi eşiklerini aşmış ve üçü de uyarlanabilir kural "
             "kapalıyken mutlak kuralla verilmiştir. 2. ve 3. olaylar neredeyse aynı "
             "pozda, yedi saniye arayla gerçekleşmektedir; üçüncüsünün 44,95 "
             "saniyelik süresi, robotun temasta kalmasıdır, temizlenemeyen bir "
             "tespit değil.")
    b.figure(
        "fig5_collision.png",
        "Provoke Edilen Bir Çarpışmadan Sonra Operatör Arayüzü | 3. olayın 44,95 "
        "saniyelik bloğu 25 seviyesindeki geniş platodur; solundaki dar sivri uç, "
        "yedi saniye önceki 2. olaydır. Eşik çizgisi 5'tedir. Ayrıştırma paneli her "
        "modelin yeniden yapılanma hatasını kendi eşiğine karşı vermekte, olay "
        "tablosu ise üç bloğun da operatörce doğrulandığını ve etiketsiz blok "
        "bulunmadığını göstermektedir. Oturumun kaydedilmiş kararlarından yeniden "
        "oynatılmıştır.")
    b.p(
        "Ayrım, bir eşik yapaylığı değil gerçektir. En zayıf doğrulanmış olay "
        "11,30'da tepe yaparken o günün herhangi bir temiz koşusundaki en yüksek "
        "sapma 4,14'e ulaşmaktadır; dolayısıyla (4,14 · 11,30) aralığındaki her "
        "eşik üç tespit ve sıfır yanlış alarm vermektedir. Dağıtılan 5,0 değeri, "
        "gözlenenlerden daha zayıf olaylara pay bırakmak üzere bilinçli olarak bu "
        "aralığın alt ucundadır. Tablo 15 aralığın nasıl kapandığını göstermektedir."
    )
    b.table(
        "Tüm Kaydedilmiş Kararlardan Yeniden Sayılan Eşik Taraması.",
        ["Eşik", "Temiz blok / dk", "Yakalanan doğrulanmış olay"],
        [["3,0", "0,00", "3 / 3  (parçalı)"],
         ["**5,0**  (dağıtılan)", "**0,00**", "**3 / 3**"],
         ["8,0", "0,00", "3 / 3"],
         ["11,3", "0,00", "2 / 3"],
         ["16,0", "0,00", "1 / 3"],
         ["26,0", "0,00", "0 / 3"]],
        widths=[2.55, 2.9, 2.9], align_right=[1, 2],
        note="14,3 dakikalık temiz çalışma ve doğrulanmış olayları içeren 10,0 "
             "dakika üzerinden, düğümün uyguladığı aynı iki-ardışık-karar kuralıyla "
             "yeniden sayılmıştır.")
    b.p(
        "**Düşük genlikli bir temas tespit edilmemektedir ve hiçbir eşik onu geri "
        "kazandırmaz.** Tablo 13'ün ikinci satırında provoke edilen el teması, "
        "birleşik skoru o koşunun kendi medyanının 6,8 katına çıkarmıştır; oysa "
        "temiz muayene çevrimi rutin olarak kendi medyanının 26,7 katına "
        "ulaşmaktadır. Ortak ölçekte ifade edildiğinde, provoke edilen temasın tepe "
        "değeri olağan muayene hareketinin 75,6. persentiline düşmektedir: robotun "
        "normalde yaptığının dörtte üçünden küçüktür. Kaydedilmiş kararlar "
        "üzerinde bu olaya karşı dört karar kuralı değerlendirilmiştir — iki farklı "
        "normalleştirmede mutlak eşik, k ∈ [4, 32] aralığında taranan uyarlanabilir "
        "medyan + k·MAD kuralı ve 10, 30 ve 60 saniyelik pencerelerde kayan medyan "
        "çıkarımı. Her durumda temiz muayene koşusu, olayı içeren koşu kadar ya da "
        "ondan çok blok üretmiştir. Onları ayırmak için gereken sinyal bu iki skorda "
        "mevcut değildir."
    )
    b.p(
        "Dağıtılan sistemin çalışma zarfı bu nedenle açıktır: çarpışma "
        "büyüklüğündeki olaylar, 5'lik bir eşiğe karşı 11 ilâ 25 tepe değeriyle ve "
        "olağan çalışma üzerinde 3,9 katlık bir ayrım payıyla güvenilir biçimde "
        "tespit edilmekte; yavaş ve düşük genlikli temas ise edilmemektedir."
    )

    b.h2("4.8. Operatör Etiketlemesi")
    b.p(
        "Tablo 14'ün üç olayı, operatör hücrenin başında dururken Şekil 5'teki "
        "arayüz üzerinden sınıflandırılmıştır ve Bölüm 4.7'nin yanlış alarm "
        "sayısını bir izlenim değil bir ölçüm yapan şey de bu sınıflandırmadır: "
        "arayüz üç olay, üç doğrulanmış, sıfır yanlış alarm ve sıfır etiketsiz "
        "bildirmektedir. Etiketler, dedektörün hiç okumadığı bir dosyaya atomik "
        "olarak yazılır; dolayısıyla bir etiket hiçbir zaman bir kararı "
        "etkileyemez. Oturum başına köken kaydıyla birlikte bunlar, kaydedilmiş bir "
        "oturumu aylar sonra yeniden çözümlenebilir kılmaktadır ve Bölüm 4.7'deki "
        "eşik başarısızlığının modellere değil yapılandırmaya atfedilebilmesini "
        "sağlayan da budur."
    )
    b.p(
        "Şekil 5'teki ayrıştırma paneli her kararı hangi modelin sürüklediğini de "
        "kaydeder. Doğrulanmış üç çarpışmanın hepsinde her iki model de kendi "
        "eşiğini aşmıştır; kalıntı modeli 124 ilâ 270 kat, ham model 447 ilâ 1186 "
        "kat. Bölüm 4.7'deki tespit edilemeyen el temasıyla birlikte bu, iki uzayın "
        "gerçek arızalar sırasında nasıl davrandığına dair buradaki tek kanıttır ve "
        "Bölüm 4.4'te ortaya konan soruyu çözmemektedir: bu büyüklükteki olaylar her "
        "iki modeli de doyurmakta, dolayısıyla aralarında ayrım yapamamakta; ayrım "
        "yapacak kadar küçük olan tek olayı ise ikisi de görememektedir."
    )


# ── 5. Tartışma ──
def _tartisma(b: Builder) -> None:
    b.h1("5. Tartışma")

    b.h2("5.1. Genişletmenin Çevrimdışı Değiştirdikleri")
    b.p(
        "Genişletilmiş hat mutlak ölçütlerde önceki çalışmanın altında kalmakta ve "
        "sebepleri gizemli değil ölçülebilirdir. Arıza genlikleri artık gerçek "
        "fiziksel ölçektedir; 15 Nm olarak belirtilen bir rampa, amper cinsinden "
        "bir kanala çıplak sayı olarak eklenmek yerine 15 Nm olarak enjekte "
        "edilmektedir. Veri kümesi, bir yeniden yapılanma modelinin anomali diye "
        "okuduğu süreksizliklerden arındırılmıştır. Ve bölme koşu-ayrıktır; "
        "dolayısıyla değerlendirme artık kısmen modellerin daha önce yeniden inşa "
        "ettiği pencereler üzerinde ölçüm yapmamaktadır: temiz pencere skorunun "
        "test/eğitim oranı, önceki düzenin ürettiği yetmiş kata karşılık 0,48'dir."
    )
    b.p(
        "Genişletmeden sağ çıkan şey yapıdır. Birleşim hâlâ iki tekil modeli ve "
        "bütün referans yöntemleri geçmektedir; ağırlık taraması hâlâ aynı aralıkta "
        "geniş ve düzdür; tamamlayıcılık oranları önceki çalışmanınkilere yakın "
        "kalmaktadır, yalnız kalıntıda %27,6'ya karşı %24,1 ve yalnız hamda "
        "%28,0'a karşı %25,9. Tek yerine beş tohum raporlamak, kalan belirsizliğin "
        "nerede yaşadığını da göstermektedir: sıralama ölçütleri tohumlar arasında "
        "0,01'den az değişirken çalışma noktası kesinliği ±0,086 değişmektedir, "
        "çünkü eşik, modellerce değil bölme tasarımınca daraltılmış bir doğrulama "
        "dağılımının persentilidir."
    )

    b.h2("5.2. Birleşim Marjı Enjeksiyon Protokolünün Bir Özelliğidir")
    b.p(
        "Bu makalenin merkezî çevrimdışı bulgusu olumsuzdur ve çerçevenin "
        "mimarisini değil kendi değerlendirmesini ilgilendirmektedir. Devralınan "
        "enjeksiyon protokolü altında birleşim, daha iyi tekil modele göre "
        "+0,189 ± 0,009 PR-AUC kazanmakta ve bu her tohumda pozitiftir. Arızanın "
        "yalnız ölçülen kanallara uygulandığı ve kalıntının hatça yeniden "
        "hesaplandığı bir protokolde ise aynı kazanç −0,003 ± 0,001'dir, yine her "
        "tohumda. Sürtünme terimini eklemek ya da çıkarmak bunu hiçbir yönde "
        "değiştirmemektedir (Tablo 10)."
    )
    b.p(
        "Mekanizma tek bir genliktir. Devralınan protokol kalıntı uzayını elle "
        "seçilmiş bir sayıyla bozmakta ve sensör gürültüsü için o sayı 0,08 "
        "Nm'dir; referans çalışmada bu, senaryonun yalnız ham sinyali etkilemesi "
        "için kasten küçük tutulmuş olarak tanımlanmaktadır. Ama dışsal kalıntı "
        "r_ext = J(q)ᵀF'dir, wrench'in doğrusal bir dönüşümüdür; wrench gürültüsünü "
        "zayıflatmaz, dönüştürür. Fiziksel olarak yayıldığında 3,5 N'luk wrench "
        "gürültüsü 4,1 Nm'lik bir kalıntı bozulması üretmekte — enjekte edilen "
        "genliğin elli bir katı — ve kalıntı modeli bunu kusursuz tespit "
        "etmektedir. Birleşimi güdüleyen tamamlayıcılık o senaryoyla sınanmamış; "
        "onun tarafından tanımlanmıştır."
    )
    b.p(
        "Bu dar okunmalıdır. Çarpışma senaryosunun elle seçilen 40 Nm'lik kalıntı "
        "genliği, tepe değeri 80,5 Nm olan fiziksel yayılımla uyuşmaktadır; "
        "dolayısıyla referans çalışmanın seçimleri baştan sona keyfî değildir. Ve "
        "bulgu, ham sinyal modelinin işe yaramaz olduğunu göstermez — dört analitik "
        "bozulmanın, tutarlı biçimde uygulandığında bu iki temsil uzayı arasında "
        "ayrım yapmadığını gösterir. Gerçek arızalara ilişkin kanıt Bölüm 4.7'de, "
        "doğrulanmış üç çarpışmanın hepsinde her iki modelin de kendi eşiğini "
        "aştığı yerdedir."
    )

    b.h2("5.3. Çevrimdışı Eşik Neden Taşınmıyor")
    b.p(
        "Çevrimdışı bir doğrulama dağılımının 97. persentiline yerleştirilen bir "
        "eşik, donanımdaki kararların çoğunu alarm ilan etmiştir. Skorların "
        "kendisi olağan bir ölçektedir — hücrede kalibre edilmiş yapılandırma "
        "altında aynı temiz çalışma eşiğin 1,7 ilâ 35,3 katı altında kalmaktadır "
        "(Tablo 13) — dolayısıyla hata modellerde değil kestiricidedir."
    )
    b.p(
        "Bunu açıklayan ölçüm, modellerin olağan çalışma sırasında eğitim "
        "dağılımlarına olan uzaklığıdır: kalıntı modeli için doğrulama ortalama "
        "hatasının on bir katı, ham model için seksen altı katı. Dağıtımın hiç "
        "işgal etmediği bir dağılımın persentili, dağıtım hakkında bilgi taşımaz. "
        "İki model arasındaki asimetri pratik açıdan önemli olan kısımdır. Girdisi "
        "fiziksel modelce yörüngeye bağlı dinamiklerden arındırılmış olan kalıntı "
        "modeli, hücrede kendi eşiğinin 0,93'ünde oturmaktadır — iyi kalibre. "
        "Sinyalleri doğrudan okuyan ham model ise 7,28'inde. Buna rağmen yüzde "
        "beşlik bir ağırlıkla birleşik skorun %29'unu sağlamaktadır; çünkü yirmi "
        "beş kat küçük bir normalleştirme sabiti, küçük bir ağırlığı büyük bir "
        "etkiye çevirir. Bu, skor düzeyinde birleşimde genel bir tuzaktır: "
        "ağırlıklar, ancak normalleştiriciler geçerli olduğu sürece söyledikleri "
        "şeyi ifade ederler."
    )
    b.p(
        "Bölme tasarımının kendisi de katkıda bulunmaktadır ve genişletmenin "
        "merkezindeki ödünleşim budur. Doğrulama ve testi eğitim eklem aralığının "
        "içinde tutmak için — ki bu zorunludur, aksi hâlde hücredeki yüksek bir "
        "skor 'anomali' değil 'görülmemiş poz' anlamına gelirdi — on altı doğrulama "
        "koşusu eğitime benzerlikleri gözetilerek seçilmiştir. Doğrulama kaybı "
        "buna bağlı olarak düşmekte, ham model için referans çalışmadaki 0,320'ye "
        "karşılık 0,036'ya inmekte ve onunla birlikte θ'nın çekildiği dağılımın "
        "genişliği de daralmaktadır. Sızıntıyı ortadan kaldıran seçim, bu nedenle "
        "hücre için fazla dar bir eşik üretmiştir. Bu, modellerin bir kusuru "
        "değildir: kayıpları sürtünme terimiyle de terimsiz de aynıdır, dolayısıyla "
        "kalıntı tanımına da atfedilemez. Bu, bir çalışma noktasını eğitime "
        "benzemesi için seçilmiş veri üzerinde kestirmenin bir özelliğidir."
    )
    b.p(
        "Burada kullanılan çözüm yayımlanmış formülü korumakta ve yalnız "
        "büyüklüklerin kestirildiği kümeyi değiştirmektedir: hücrenin yaptığı her "
        "görevi kapsayan, robotta kaydedilmiş temiz kararlar ve min–maks aralığın "
        "yerine sağlam bir persentil. Son ayrıntı süs değildir. Gerçek hücre "
        "verisine uygulandığında min–maks aralığı modellenmemiş bir yükçe "
        "belirlenir — ters dinamikte yük terimi yoktur, dolayısıyla bir iş parçası "
        "taşımak sürekli bir kalıntı yanlılığı, içi 0,3 kadar değişen 85 saniyelik "
        "bir plato üretir — ve bunun üzerine uydurulan bir genişlik olağan olayları "
        "görünmez kılar."
    )

    b.h2("5.4. Sürtünme Terimi Ne Kazandırdı")
    b.p(
        "Çözücünün dışına bir Coulomb+viskoz terim eklemek, ters dinamiğin daha "
        "önce neredeyse hiç katkı vermediği ve kalıntının pratikte ham ölçüm olduğu "
        "iki kanalda — wrist_2 ve wrist_3 — kalıntı yayılımının %87 ve %92'sini "
        "gidermektedir. Bu, fiziksel sadakatte gerçek bir kazançtır ve arıza "
        "genliklerini altı eklemin hepsinde newton metre cinsinden yorumlanabilir "
        "kılmaktadır."
    )
    b.p(
        "Bir tespit kazancı değildir. Devralınan protokol altında kalıntı modeli "
        "+0,034 ± 0,009 PR-AUC iyileşmekte ve tohumlar arasında tutarlı "
        "davranmaktadır; fiziksel protokolde ise değişim −0,020 ± 0,027'dir ve "
        "anlamlı değildir. Ayrıca eşik kararlılığına da mal olmaktadır: çalışma "
        "eşiğinin tohumdan tohuma yayılımı ortalamasının %17'sinden %46'sına "
        "çıkmaktadır. Dürüst özet şudur: terim modeli düzeltmektedir ve modeli "
        "düzeltmek, bu belirli sentetik arızaların tespitini kendiliğinden "
        "iyileştirmemektedir. Dağıtılan yapılandırmada korunmasının sebebi, gerçek "
        "arızaların bu dördü olmaması ve söylediği şeyi ifade eden bir kalıntının, "
        "beklenmedik bir şey olduğunda üzerinde akıl yürütmeyi kolaylaştırmasıdır."
    )

    b.h2("5.5. Çalışma Zarfı")
    b.p(
        "Dağıtılan dedektör çarpışmaları yakalamakta, yavaş ve düşük genlikli "
        "teması yakalamamaktadır; bu cümlenin her iki yarısı da ölçülmüştür. "
        "Operatörce doğrulanmış üç çarpışma 5,0'lık bir eşiğe karşı 11,30 ile 25,04 "
        "arasında tepe yapmış; on dakikada etiketsiz hiçbir alarm bloğu ve on dört "
        "dakikalık olağan çalışmada hiç alarm oluşmamıştır. 4,14 ile 11,30 "
        "arasındaki her eşik aynı sonucu vermektedir. Buna karşılık kasten provoke "
        "edilen bir el teması, birleşik skoru o koşunun medyanının 6,8 katına "
        "çıkarmıştır; oysa olağan muayene hareketi rutin olarak kendi medyanının "
        "26,7 katına ulaşmaktadır — provoke edilen olay normal çalışmanın 75,6. "
        "persentilinde oturmaktadır. Buna karşı dört karar kuralı değerlendirilmiş "
        "ve hiçbiri onu ayırmamıştır."
    )
    b.p(
        "Sınır bu nedenle eşik değil temsildir. Normal hareketin zaten ürettiği "
        "değişkenlikten küçük bir olay, bu iki skor üzerindeki hiçbir monoton "
        "kuralla yalıtılamaz. Bu boşluğu kapatmak için ya hücrenin kendi "
        "yörüngeleriyle eğitilmiş bir model gerekir — böylece olağan çalışma "
        "dağılım dışı olmaktan çıkar — ya da dinamiğe değil temasa tepki veren ek "
        "bir kanal. Zarfı belirtmek tek bir başlık sayısı vermekten daha "
        "yararlıdır: duyarlılık tabanı bilinmeyen bir güvenlik işlevini "
        "dağıtmak, tabanı yayımlanmış olandan daha zordur."
    )

    b.h2("5.6. Kısıtlar")
    for t in [
        "Tek bir robot (UR10e) ve tek bir görev profili; başka robot tipleri "
        "üzerinde doğrulama yapılmamıştır.",
        "Sentetik arızalar, fiziksel olarak enjekte edildiklerinde bile analitik "
        "bozulmalar olarak kalmaktadır. Çarpışma darbesi gerçek bir temas "
        "noktasından yayılmak yerine wrench kanallarına dağıtılmaktadır ve dört "
        "senaryonun, tutarlı biçimde enjekte edildiğinde iki temsil uzayı arasında "
        "ayrım yapmadığı Bölüm 4.4'te gösterilmiştir.",
        "Akım–tork katsayısı altı eklemin yalnız ikisinde doğrudan "
        "ölçülebilmiştir; diğer dördü varsayımla türetilmiştir ve bu, tespiti "
        "değil fiziksel yorumu etkilemektedir.",
        "Sürtünme katsayıları, yalnız sürtünmeyi değil hıza bağlı her model "
        "hatasını soğuran bir regresyonla uydurulmuştur. Bağımsız bir tribolojik "
        "doğrulama yapılmamıştır.",
        "Ters dinamikte yük modeli yoktur. Bir iş parçası taşımak, dağıtılan "
        "eşiğin üstünde kaldığı ama açıklamadığı sürekli bir kalıntı yanlılığı "
        "üretmektedir.",
        "Çalışma eşiği hücrede ölçülmektedir, ki bu zorunludur; ancak 15.627 "
        "kararlık bir kalibrasyon koşusunda ölçülmüş ve yirmi dört dakikalık "
        "çalışmada doğrulanmıştır. Tanımladığı zarf; sürücülerde termal "
        "sürüklenme, aşınma, parçadan parçaya değişkenlik ve operatör farkları "
        "bakımından devreye alma ölçeğinde bir sonuçtur.",
        "Dağıtılan yapılandırma için yalnız üç doğrulanmış anomali olayı mevcuttu. "
        "(4,14 · 11,30) eşik aralığı aşağıdan gözlenen temiz çalışmayla, yukarıdan "
        "bu üç olayın en zayıfıyla sınırlıdır; daha zayıf bir gerçek olay aralığı "
        "daraltacaktır.",
        "Çevrimiçi ölçümler CPU üzerinde alınmıştır; dağıtım iş istasyonunda GPU "
        "sağlayıcısıyla tekrarlanmalıdır.",
    ]:
        _madde(b, t)
    b.p("")


# ── 6. Sonuçlar ──
def _sonuclar(b: Builder) -> None:
    b.h1("6. Sonuçlar")
    b.p(
        "Bu makale, kobot anomali tespiti için yayımlanmış bir skor düzeyi birleşim "
        "çerçevesini almış, genişletmiş ve fiziksel bir UR10e hücresinde çalışan "
        "bir dedektör olarak teslim etmiştir. Yol boyunca alınan her ölçüm, "
        "çerçevenin aleyhine olanlar dâhil, raporlanmıştır."
    )
    b.p(
        "Başlangıç çerçevesi bağımsız olarak yeniden kurulmuş ve mimari ile "
        "yordamsal bileşenleri birebir örtüşmüştür. Sinyallerin fiziksel ölçeğini "
        "saptamak, sonrasındaki her şeyin ön koşuluydu: sürücü effort alanına tork "
        "değil motor akımı yazmakta, dolayısıyla önceki sayılar birimleri karışık "
        "bir hatta üretilmiştir. O hat, her büyüklükte 0,013'ün içinde yeniden "
        "üretilmiştir; bu, önceki sayıların hangi ölçekte okunması gerektiğini "
        "sabitlemekte ve arıza genliklerini, sürtünme katsayılarını ve kalıntı "
        "büyüklüklerini genişletme için fiziksel olarak anlamlı kılmaktadır."
    )
    b.p(
        "Bu temel üzerine hat genişletilmiştir. Doğrulanmış çözücüde bulunmayan bir "
        "sürtünme terimi onun dışına eklenmiş, ters dinamiğin daha önce neredeyse "
        "hiç katkı vermediği iki bilek ekleminde kalıntı yayılımının %87 ve %92'si "
        "giderilmiş; değerlendirme beş eğitim tohumuyla koşu-ayrık bir bölme "
        "üzerine oturtulmuştur. Devralınan enjeksiyon protokolü altında birleşim "
        "üstünlüğünü korumaktadır: tekil modeller için 0,627 ± 0,006 ve "
        "0,614 ± 0,016'ya karşılık F1 0,791 ± 0,023, yani +0,164'lük bir kazanç; "
        "tamamlayıcılık oranları önceki çalışmanınkilere yakındır."
    )
    b.p(
        "Enjeksiyon protokolünün kendisi ardından incelenmiştir ve bu, makalenin "
        "asıl olumsuz sonucudur. Aynı dört arıza yalnız ölçülen kanallara enjekte "
        "edilip kalıntı hatça yeniden hesaplandığında, birleşim marjı "
        "+0,189 ± 0,009'dan −0,003 ± 0,001 PR-AUC'ye düşmektedir; her tohumda ve "
        "sürtünme teriminden bağımsız olarak. Mekanizma tek bir genliktir: referans "
        "protokol, sensör gürültüsü senaryosu için kalıntı uzayını 0,08 Nm ile "
        "bozmakta ve bu, senaryonun yalnız ham sinyali etkilemesi için seçilmiştir; "
        "oysa aynı 3,5 N'luk wrench bozulmasının r_ext = J(q)ᵀF üzerinden fiziksel "
        "yayılımı 4,1 Nm üretmektedir. Referans çalışmanın tamamlayıcılığın sınavı "
        "olarak adlandırdığı senaryo onu sınamamaktadır. İki temsil uzayının "
        "tamamlayıcılığı bu nedenle burada modellerin değil, değerlendirme "
        "protokolünün bir özelliği olarak raporlanmaktadır."
    )
    b.p(
        "Sistem, öznitelik motoru — sürtünme terimi dâhil — çevrimdışı hatla kayan "
        "nokta düzeyinde örtüşen, her kararda iki modeli birden çalıştıran ve "
        "çıkarımı karar bütçesinin %12'sini tüketen 500 Hz'lik bir ROS 2 düğümü "
        "olarak gerçeklenmiştir."
    )
    b.p(
        "Donanımda çevrimdışı türetilen eşik taşınmamış, kararların %58'ini alarm "
        "ilan etmiştir. Sebep varsayılmak yerine ölçülmüştür: olağan çalışma "
        "sırasında kalıntı modeli doğrulama ortalama hatasının on bir katında, ham "
        "model seksen altı katında çalışmaktadır; dolayısıyla doğrulama dağılımının "
        "bir persentili dağıtım hakkında bilgi taşımaz. Pratik açıdan yararlı olan "
        "kısım asimetridir — kalıntı modeli hücrede kendi eşiğinin 0,93'ünde "
        "otururken ham model 7,28'inde oturmakta, bu da yüzde beş ağırlıklı bir "
        "modelin birleşik skorun %29'unu sağlamasına yol açmaktadır. Her iki ölçeğin "
        "ve eşiğin temiz hücre kararları üzerinde, min–maks aralığın yerine sağlam "
        "bir persentille yeniden kestirilmesi amaçlanan ağırlıklandırmayı geri "
        "getirmektedir."
    )
    b.p(
        "Bu yapılandırmayla dedektör provoke edilen çarpışmalara karşı "
        "doğrulanmıştır. Operatörce doğrulanmış üç olay, 5,0'lık bir eşiğe karşı "
        "11,30, 16,13 ve 25,04 tepe değerleriyle yakalanmış; olay içeren on "
        "dakikalık çalışmada etiketsiz hiçbir alarm bloğu ve on dört dakikalık "
        "olağan çalışmada hiç alarm oluşmamıştır. (4,14 · 11,30) aralığındaki her "
        "eşik aynı sonucu vermektedir. Çalışma zarfı diğer yandan da sınırlıdır: "
        "kasten provoke edilen düşük genlikli bir el teması, olağan muayene "
        "hareketinin 26,7 katına karşılık yalnız kendi koşu medyanının 6,8 katına "
        "ulaşmış ve normal çalışmanın 75,6. persentiline yerleşmiştir. O olaya karşı "
        "dört karar kuralı değerlendirilmiş ve hiçbiri onu ayırmamıştır."
    )
    b.p(
        "Buradan iki yön çıkmaktadır. Duyarlılık tabanı eşiğin değil temsilin bir "
        "özelliğidir; dolayısıyla onu kapatmak, hücrenin kendi yörüngeleriyle "
        "eğitilmiş modeller — ki bu aynı zamanda eşik transferi başarısızlığının "
        "ardındaki dağılım dışı koşulu da ortadan kaldırırdı — ya da dinamiğe değil "
        "temasa tepki veren ek bir kanal gerektirir. İkincisi, sürtünme teriminin "
        "izlediği aynı desenle çözücünün dışına bir yük terimidir; çünkü taşıma "
        "sırasında normalleştirme ölçeğini yukarı zorlayan şey modellenmemiş "
        "yüktür. Her ikisi de hâlihazırda kurulu araçlarla desteklenmektedir: her "
        "karar, her alarm bloğu, operatörün etiketi ve bunları üreten "
        "yapılandırmanın kökeni diske yazılmaktadır."
    )


# ── arka bölümler ──
def _arka(b: Builder) -> None:
    b.h1("Teşekkür")
    b.p(
        "Proje, KDT Ortak Girişimi (101140216) ve üyeleri tarafından "
        "desteklenmektedir; ayrıca Vinnova (İsveç), Österreichische "
        "Forschungsförderungsgesellschaft mbH – FFG (Avusturya), Business Finland "
        "(Finlandiya), Üniversiteler ve Araştırma Bakanlığı (İtalya), FCT "
        "(Portekiz) ve TÜBİTAK (124N448) (Türkiye) tarafından ek finansman "
        "sağlanmıştır. Ölçümler ESOGÜ Akıllı Sistemler Uygulama ve Araştırma "
        "Merkezi'nin Otonom Sistemler ve Güvenilirlik Laboratuvarı'nda "
        "gerçekleştirilmiştir."
    )

    b.h1("Araştırmacıların Katkı Oranı")
    b.p(
        "Yazar 1: çevrimdışı hattın ve ROS 2 dedektör düğümünün tasarımı ve "
        "gerçeklenmesi, devreye alma ölçümleri, makalenin hazırlanması. Yazar 2: … "
        ". Yazar 3: … . (Yazarlar tarafından tamamlanacaktır.)"
    )

    b.h1("Çıkar Çatışması")
    b.p("Yazarlar tarafından herhangi bir çıkar çatışması bildirilmemiştir.")

    b.h1("Kaynaklar")
    for ref in en.REFERENCES:
        p = b.para("", style="Kaynaklar", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_before=6, space_after=0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        b.rich(p, ref, size=BODY_PT)


def main() -> None:
    doc, anchor = open_template()
    b = BuilderTR(doc, anchor)
    front_matter(b)
    body(b)
    b._leave_wide()
    doc.save(str(OUT))
    print(f"yazıldı: {OUT}")
    print(f"  şekil: {b.fig_no}   tablo: {b.tab_no}   denklem: {b.eq_no}")


if __name__ == "__main__":
    main()
