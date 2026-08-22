#!/usr/bin/env python3
"""Şasi + kapı muayenesi teknik raporunu (.docx) üretir.

Rapor tek kaynaktan üretilir ki bir sonraki güncellemede Word'de elle düzenleme
gerekmesin: metni buradan düzenleyip yeniden çalıştırmak yeterli.

    python3 docs/rapor_uret.py

Figürler bu dizindeki `fig_*.png` dosyalarıdır (numara YOK: şekil numaraları metin
değiştikçe kayıyordu). Nereden geldikleri:

  fig_chassis_*, fig_surface_normals, fig_target_points, fig_viewpoint_generation,
  fig_solo_viewpoints, fig_coverage_curve, fig_marginal_gain, fig_sick_collision_shell
      -- 22 Temmuz 2026 sürümünden.
  fig_multirobot_plan / _ur10e / _kawasaki
      -- paketin kendi görselleştiricisi, GÜNCEL plandan:
         python3 ../multirobot_viewpoint_planner/plan_visualizer.py \
             --output <docs>/fig_multirobot_plan.png
         (bu kutuda sistem matplotlib'i gerekir; ayrıntı için figur_uret.py'nin başı)
  fig_doors_*
      -- python3 docs/figur_uret.py
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "viewpoint_inspection_system_report.docx")
FIG_WIDTH = Inches(5.2)


# --------------------------------------------------------------------------- #
# yardımcılar
# --------------------------------------------------------------------------- #
def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def p(doc, text):
    return doc.add_paragraph(text)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for c, name in zip(t.rows[0].cells, header):
        c.text = name
    for row in rows:
        cells = t.add_row().cells
        for c, val in zip(cells, row):
            c.text = str(val)
    doc.add_paragraph()
    return t


def figure(doc, filename, caption):
    path = os.path.join(HERE, filename)
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    doc.add_picture(path, width=FIG_WIDTH)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# --------------------------------------------------------------------------- #
def build():
    doc = Document()

    doc.add_heading("Şasi ve Kapı Muayenesi için Bakış-Noktası (Viewpoint) "
                    "Planlama Sistemi", 0)
    sub = doc.add_paragraph(
        "Tek-kol UR10e ve İşbirlikli UR10e + Kawasaki RS005L Senaryoları — "
        "Teknik Rapor")
    sub.style = doc.styles["Subtitle"]
    meta = doc.add_paragraph()
    meta.add_run("Revizyon 2 — 22 Ağustos 2026    |    Hazırlayan: Cem Süha Yılmaz "
                 "   |    ROS 2 Humble / MoveIt 2").bold = True
    p(doc, "Revizyon 1 (22 Temmuz 2026) şasi senaryosunu kapsıyordu. Bu revizyon, "
           "aradaki bir aylık çalışmayı ekler: trajectory optimizasyonu (2π açma ve "
           "IK dal seçimi), tamamlanan çarpışma modeli eklemeleri ve ikinci muayene "
           "senaryosu olan kapı (doors) muayenesi. Değişen bütün sayılar mevcut plan "
           "dosyalarından yeniden okunmuştur.")

    # ---------------------------------------------------------------- 1
    h1(doc, "1. Yönetici Özeti")
    p(doc, "Bu rapor, bir üretim hücresinde kamera ile otomatik yüzey muayenesi için "
           "geliştirilen paketleri ve altlarındaki mühendislik kararlarını belgeler. "
           "Amaç, bir SICK Visionary-T Mini derinlik kamerasıyla hedef yüzeyin mümkün "
           "olan en yüksek oranda kaplanmasını (coverage) sağlayacak en az sayıda "
           "robot bakış-noktasını (viewpoint) otonom olarak seçmek ve bu noktaları "
           "gerçek robot üzerinde çarpışmadan, tekrarlanabilir biçimde gezmektir.")
    bullets(doc, [
        "viewpoint_planner — tek bir UR10e (lineer eksen üzerinde) ile çalışan baz "
        "senaryodur; sistemin çekirdek planlama hattı burada yaşar.",
        "multirobot_viewpoint_planner — UR10e ile bir Kawasaki RS005L kolunu (AGV "
        "üstünde) birlikte çalıştıran işbirlikli senaryodur; tek kolun ulaşamadığı "
        "yüzeyleri ikinci kolla kapatarak kaplama tavanını yükseltir.",
        "doors_inspection — aynı çalıştırılabilir düğümleri kapıya özgü "
        "parametrelerle koşan İKİNCİ muayene senaryosudur. Hiçbir doğrulanmış modül "
        "fork edilmemiştir.",
    ])
    p(doc, "Kritik bulgu değişmedi: tek kolun kaplaması bir ERİŞİLEBİLİRLİK "
           "(reachability) tavanına çarpar, algılama geometrisine değil. İkinci kol "
           "tam da bu tavana saldırmak için eklenmiştir. Mevcut plan dosyalarından "
           "okunan güncel değerler:")
    table(doc,
          ["Senaryo", "UR bakış-noktası", "Kawasaki", "Ulaşılan kaplama"],
          [["viewpoint_planner (şasi, solo)", "24", "—", "%54.6"],
           ["multirobot_viewpoint_planner (şasi)", "19", "9", "%66.3"],
           ["doors_inspection (kapılar)", "9", "3", "%95.6"]])
    p(doc, "Kapı senaryosunun %95.6'ya çıkması hedef geometrisinin küçüklüğü ve "
           "erişilebilirliğiyle açıklanır (birleşik kapı yüzeyi 6.82 m², şasi "
           "31.97 m²) — sistemin iyileşmesiyle değil. Şasi ile kapı sayıları aynı "
           "ölçekte karşılaştırılamaz.")
    p(doc, "Bu revizyonun eklediği ikinci ana başlık trajectory optimizasyonudur: "
           "eklem hedeflerinin 2π-eşdeğerleri arasından en yakınının seçilmesi ve IK "
           "dalları arasında arama, gerçek plan üzerinde toplam eklem yolunu %35 "
           "azaltmıştır. Pozlar fiziksel olarak aynı kaldığı için kaplama etkilenmez; "
           "yalnızca yol kısalır.")

    # ---------------------------------------------------------------- 2
    h1(doc, "2. Sistem Bağlamı ve Problem")
    p(doc, "Birincil muayene hedefi bir otomobil şasisidir. Dünya-çerçevesi ölçüleri "
           "(world, m): AABB min[-1.501, 0.431, -0.018], max[-0.736, 2.011, 1.885]; "
           "yani ~0.77 × 1.58 × 1.90 m ve en yüksek noktası 1.885 m. Bu, tek lineer "
           "eksen üstündeki ~1.3 m erişimli bir UR10e için hem yüksek hem geniştir. "
           "Hücre dünya orijinindedir; UR lineer rayı (base_to_robot_mount) dünya "
           "X = -0.158'de, şasi X ∈ [-1.5, -0.74] aralığındadır — dolayısıyla "
           "X > -0.158 olan her kamera pozunda ray, kamera ile şasi arasındadır "
           "(bkz. Workspace keep-out kutusu).")
    p(doc, "Kamera: SICK Visionary-T Mini V3S145 (her iki kol da aynı fiziksel kamera "
           "modelini taşır; Kawasaki'de kol ucuna montelidir).")
    figure(doc, "fig_chassis_mesh.png",
           "Şekil 1: Muayene hedefi olan şasinin CAD mesh'i (121.956 üçgen), dünya "
           "çerçevesinde. Tüm planlama bu geometri üzerinden yapılır.")

    # ---------------------------------------------------------------- 3
    h1(doc, "3. viewpoint_planner — Tek-kol UR10e Senaryosu")

    h2(doc, "3.1 Planlama Hattı (Pipeline)")
    p(doc, "Planlama, aşağıdaki çekirdek modüllerden oluşan bir zincirdir. Hem "
           "çok-robot paketi hem kapı paketi bu çekirdeği aynen içe aktarır "
           "(fork etmez):")
    table(doc, ["Modül", "Görev"], [
        ["mesh_analyzer.py",
         "Hedef mesh'ini yüzey örnek noktalarına ayırır; aşağı-bakan yüzeyleri "
         "(min_normal_z) eler."],
        ["viewpoint_generator.py",
         "Her hedef için aday kamera pozları üretir (standoff mesafeleri + tilt)."],
        ["reachability_checker.py",
         "Her aday pozu MoveIt IK ile erişilebilirlik açısından test eder."],
        ["set_cover_optimizer.py",
         "Açgözlü (greedy) küme-kaplama ile en az sayıda pozu seçer."],
        ["viewpoint_clusterer.py",
         "Kaplama-güvenli kümeleme ile yakın pozları birleştirir (coverage+IK "
         "yeniden doğrulanır)."],
        ["viewpoint_planner_node.py",
         "Yukarıdakileri orkestre eder; plan JSON'unu yazar."],
        ["inspection_executor_node.py",
         "Seçilen pozları robotta gezip PCD + poz kaydeder (solo baz)."],
        ["plan_visualizer.py / visualization.py",
         "RViz'de bakış-noktalarını ve okları çizer."],
    ])
    figure(doc, "fig_surface_normals.png",
           "Şekil 2: Adım 1 — Yüzey normalleri. Her mesh üçgeninin dışa-bakan normali "
           "hesaplanır; bu, kamera görüş yönünün nereye bakması gerektiğini belirler.")
    figure(doc, "fig_target_points.png",
           "Şekil 3: Adım 2 — Hedef örnek noktaları. Mesh yüzeyi noktalara ayrılır; "
           "normal_z < -0.95 olan (aşağı bakan, kameranın altından göremeyeceği) "
           "yüzeyler elenir.")
    figure(doc, "fig_viewpoint_generation.png",
           "Şekil 4: Adım 3 — Aday bakış-noktası üretimi. Her hedef için normal "
           "boyunca 4 standoff mesafesi × 3 tilt (-15/0/+15°) aday kamera pozu "
           "üretilir.")
    figure(doc, "fig_solo_viewpoints.png",
           "Şekil 5: Adım 4 — Seçilen bakış-noktaları (solo UR10e). Renk ziyaret "
           "sırasını, oklar kamera görüş yönünü gösterir. Küme-kaplama, hedeflerin "
           "çoğunu en az pozla kapatır. Güncel solo plan 24 bakış-noktası ve %54.6 "
           "kaplama üretmektedir.")

    h2(doc, "3.2 Kaplama Tavanı: Kök-neden Analizi")
    p(doc, "Azaltılmak istenen ~160-170 bakış-noktası aslında tam kaplama için "
           "gerekmiyordu; bu sayı, azalan-getiri (diminishing returns) kuyruğundan "
           "geliyordu. Kapaksız (uncapped) bir koşu 170 viewpoint üretti ama yalnızca "
           "%72.2 kaplamaya ulaştı (0.98 eşiğine hiç varmadı, greedy tükendi). 170'in "
           "~91'i her biri ≤3 yeni nokta ekliyordu.")
    table(doc, ["Bakış-noktası sayısı", "Kaplama", "Not"],
          [["25", "%58.4", "Ulaşılabilir %72'nin %81'i"],
           ["40", "%63.3", ""],
           ["170 (kapaksız)", "%72.2", "Son ~130 vp toplam ~14 nokta ekliyor"],
           ["Geometrik tavan", "%91.3", "3921/4294 hedef ≥1 adaydan görülebilir"]])
    p(doc, "Kanıt: erişilebilirlik-farkındalıklı greedy 1444 IK sorgusundan 1419'unu "
           "NO_IK_SOLUTION(-31) ile reddetti (~%98 red oranı). Yani hedef GEOMETRİK "
           "olarak kaplanabilir; tek Y-rayı üstündeki UR10e bilgilendirici pozların "
           "çoğuna ULAŞAMIYOR. Collision'ı kapatma testi (ik_avoid_collisions=false) "
           "red oranını yalnızca %98.3→%98.1 düşürdü — yani darboğaz gerçek KİNEMATİK "
           "erişim, collision değil.")
    p(doc, "Denenip geri alındı: kamera optik ekseni etrafında roll varyasyonları "
           "([0,45,90,135]°). Kaplamayı iyileştirmedi (%57.9), IK reddini kötüleştirdi "
           "(%99.5) ve aday sayısını 4×'ledi. Sonuç: yönelim çeşitliliği bu erişim "
           "sınırını açmıyor. Roll/tilt sayımını tekrar denememek gerekir.")
    figure(doc, "fig_coverage_curve.png",
           "Şekil 6: Kaplama büyüme eğrisi. Çok-robot koşu, solo erişim tavanını "
           "(~%72) aşan bir bölgeye ulaşır. İki kesikli çizgi belgelenen tavanları "
           "(geometrik ~%91.3, solo erişim ~%72.2) gösterir; eğrinin dizi (~25 vp) "
           "belirgindir.")
    figure(doc, "fig_marginal_gain.png",
           "Şekil 7: Bakış-noktası başına eklenen YENİ nokta sayısı. Baştaki birkaç "
           "poz kaplamanın çoğunu getirir; uzun kuyruk her biri çok az ekler — "
           "marjinal-kazanç tabanı tam bu kuyruğu keser.")

    h2(doc, "3.3 Otonom Bakış-Noktası Sayısı — Marjinal-Kazanç Tabanı")
    p(doc, "Elle seçilen max_viewpoints kapağı, birincil kaldıraç olmaktan çıkarıldı "
           "ve yerine MARJİNAL-KAZANÇ TABANI kondu. SetCoverOptimizer."
           "min_marginal_coverage (hedeflerin kesri), en iyi sonraki viewpoint "
           "min_marginal_coverage × num_targets'tan daha az YENİ nokta kapatacaksa "
           "greedy'yi durdurur — böylece viewpoint SAYISI, azalan-getiri eğrisinin "
           "dizinden (knee) doğal olarak ortaya çıkar.")
    bullets(doc, [
        "min_marginal_coverage: yaml varsayılanı 0.005 (~%0.5 → şasi mesh'inde "
        "~25 viewpoint). 0.015-0.02'ye çıkarınca ~10-12.",
        "max_viewpoints: artık opsiyonel güvenlik kapağı (yaml 0 = kapalı; node "
        "0→None eşler).",
        "Sentetik uzun-kuyruk verisinde doğrulandı (taban 0/0.005/0.01/0.02 → "
        "60/10/8/5 viewpoint).",
    ])

    h2(doc, "3.4 Erişim Kaldıraçları (Reach Levers)")
    bullets(doc, [
        "pick_ik geri yüklendi: ifarlab_*_moveit_config kinematics.yaml dosyaları "
        "stok KDLKinematicsPlugin@5ms'e gerilemişti (bu dosyalar git-tracked DEĞİL; "
        "muhtemelen bir upload commit'i ezmiş) ve sonuçları baz çizginin ALTINA "
        "çekmişti. pick_ik/PickIkPlugin geri getirildi (global, timeout 0.05, "
        "attempts 3). Bugün yedi moveit_config'de pick_ik yüklüdür "
        "(harmony_moveit_config sonradan eklendi).",
        "minimal_displacement_weight 0.0'da BIRAKILMALIDIR. IK dal sürekliliğini "
        "solver tarafında zorlamak için 0.005 denendi ve erişilebilirliği çökertti "
        "(bkz. Bölüm 6.3).",
        "Yönelim eşiği gevşetildi: real_ur10e grubunda orientation_threshold "
        "0.01→0.05 (~0.57°→~2.9° kamera-yönü esnekliği; sensör modelinin 3° FOV "
        "payınca massedilir) ve kinematics_solver_attempts 3→5.",
        "Padding şeması: collision_padding = 0.03 m TÜM ur10e_ linklerine uygulanır "
        "(kol + mobilya + hedef). Daha geniş koridor için hedefi padding'den "
        "çıkarmak yerine collision_padding'i düşürün.",
        "Plan/exec sahne eşleme bug'ı: plan-zamanında /get_planning_scene boş bir ACM "
        "döndürüyordu, IK padding'siz koşuyor sonra motion planning'te patlıyordu; "
        "get_planning_scene çağrısı entry_names dolana kadar (10×, 0.5s) YENİDEN "
        "denenir hale getirildi.",
    ])

    h2(doc, "3.5 Diğer Solo İyileştirmeleri")
    bullets(doc, [
        "Rota optimizasyonu: seçimden sonra her kolun listesi kısa bir kartezyen "
        "ziyaret turuna (nearest-neighbour + 2-opt + or-opt) yeniden sıralanır; "
        "gerçek planlarda ~%60 yol-uzunluğu azalması. order_by_proximity (varsayılan "
        "true); SEÇİM değişmez, yalnızca ziyaret sırası.",
        "Workspace keep-out kutusu: mesh-tabanlı occlusion ray-trace, rayın "
        "arkasındaki (kör) pozları anlayamıyordu; workspace_filter_enabled + "
        "workspace_bounds_min/max ile dünya-çerçevesinde bir kutu tanımlanır.",
        "Gerçek UR bulut dönüklüğü düzeltildi: gerçek UR SICK bulutu görüş ekseni "
        "etrafında ~60° dönük geliyordu. ur_macro.xacro'ya ur10e_sick_optical_frame "
        "(rpy 0 0 1.04720 = 60° yaw) eklendi; gerçek SICK node frame_id + transformer "
        "target + executor real_ur_sensor_frame buna yönlendirildi.",
        "2π açma ve dal seçimi solo baz senaryoya da ulaştı: inspection_executor_node "
        "kendi açma mantığına sahip değildi, 2026-08-14'te bu iş pymoveit2_real'e "
        "taşınınca hiçbir çağrı-yeri değişikliği olmadan devraldı (bkz. Bölüm 6.5).",
    ])

    # ---------------------------------------------------------------- 4
    h1(doc, "4. multirobot_viewpoint_planner — İşbirlikli UR10e + Kawasaki")

    h2(doc, "4.1 Neden Var ve Tasarım Kararları")
    bullets(doc, [
        "Doğrudan solo'nun erişim tavanına saldırır: ikinci kol, tek Y-rayı üstündeki "
        "UR'nin ulaşamadığı yüzeyleri kapatır.",
        "Kawasaki VP stratejisi: 'paylaşılan küme + tahsis' — TEK bir tam-kaplama "
        "aday havuzu üretilir (solo ile aynı), sonra her viewpoint IK ile hangi kol "
        "ulaşabiliyorsa ona atanır. Eski 'ayna' (mirror) yaklaşımı değil.",
        "Amaç: maksimum kaplama + dengeli yük (paralel koşu süresi ≈ max(kol-başı "
        "sayılar)).",
        "Paket ilişkisi: 'depend + extend' — viewpoint_planner çekirdeğini "
        "(mesh_analyzer, viewpoint_generator, reachability_checker) aynen içe aktarır.",
        "İki-senaryo ayrımı: viewpoint_planner artık saf UR10e'dir (tüm Kawasaki kodu "
        "çıkarıldı); işbirlikli senaryo bu ayrı pakette yaşar.",
    ])

    h2(doc, "4.2 Kawasaki Yönelim Sorunu — Kök Neden")
    p(doc, "İki robotun kamera optik çerçeveleri FARKLI görüş-ekseni konvansiyonu "
           "kullanır. Generator pozları görüşü +Z boyunca üretir. UR çalışır çünkü IK "
           "gerçek ur10e_depth_optical_frame'i (rpy 0 0 0, +Z = görüş ekseni) "
           "hedefler. Kawasaki tek bir macro (rs005l_macro.xacro) kullandığından "
           "telafi edici fark yoktur ve camera_rgb_optic_frame görüşü +X boyuncadır. "
           "Bir .pcd'yi kendi optik çerçevesine geri dönüştürerek ampirik "
           "kanıtlanmıştır.")
    bullets(doc, [
        "Düzeltme: R_cam = R_view @ C, C = rpy(-π/2, -π/2, 0) — generator +Z görüşünü "
        "kamera +X görüşüne eşler. Param kawasaki_view_axis_correction_rpy.",
        "İkinci tuzak: kamera→tip ofseti TF ile değil ANALİTİK hesaplanır "
        "(_analytic_kawasaki_offset): sim koşusunda gerçek (öneksiz) frame'ler TF'te "
        "olmadığından TF-lookup sessizce başarısız olup yanlış linke düzeltme "
        "uyguluyordu. URDF zincirinden (joint7 rpy 0,0,-1.96; camera_joint "
        "xyz .0145 .0125 .185 rpy 0,-π/2,π/2) link6-in-camera doğrudan hesaplanır; "
        "rastgele viewpoint'lerde 0.00° hata ile doğrulandı.",
        "Yürütücü tarafında bu zincirin İKİZİ vardır (_ik_target / _camera_to_tip) ve "
        "İKİSİ AYNI KALMAK ZORUNDADIR; ayrılırlarsa kamera yanlış yöne bakar. "
        "p_off = [-0.185, 0.0145, 0.0125] olarak çevrimdışı doğrulanmıştır.",
    ])

    h2(doc, "4.3 Bakış-Noktası Sayısı ve Rota")
    bullets(doc, [
        "min_new_points (kol-başı kuyruk kesme, varsayılan 8) + "
        "max_viewpoints_per_robot (sert kap, 30) + fractional min_marginal_coverage "
        "(0.005). Etkin taban = max(8, ceil(0.005×num_targets)).",
        "Karşıt-uçtan başlama: ur_order_anchor='max_y' (hedefin ÖNÜ), "
        "kawasaki_order_anchor='min_y' (ARKA) — iki kol zıt uçlardan tarar.",
        "UR workspace keep-out kutusu (yalnız UR): kamera pozu kutu dışındaysa UR için "
        "ULAŞILMAZ raporlanır, viewpoint Kawasaki'ye gider (Kawasaki asla "
        "gate'lenmez). bounds_min z 0.20 UR'nin yere yakın pozlarını da eler.",
        "Güncel şasi planı: 19 UR + 9 Kawasaki = 28 bakış-noktası, %66.3 kaplama.",
    ])
    p(doc, "Aşağıdaki üç şekil, paketin kendi görselleştiricisi (plan_visualizer.py) "
           "tarafından doğrudan GÜNCEL plan dosyasından üretilmiştir; renk skalası her "
           "bakış-noktasının kaplamaya kattığı YENİ nokta sayısıdır (koyu = az katkı), "
           "yani rota üzerindeki azalan getiri figürün kendisinde okunabilir.")
    figure(doc, "fig_multirobot_plan.png",
           "Şekil 8: Çok-robot çözümü dört görünümde (izometrik, ön, yan, tepe): "
           "19 UR (turuncu-yeşil skala) + 9 Kawasaki (mavi skala) bakış-noktası, oklar "
           "görüş ekseni. Tepe görünümü işin özünü verir: iki kümenin X aralıkları hiç "
           "ÖRTÜŞMEZ — UR -1.22 … -0.22, Kawasaki -2.42 … -1.41 arasındadır ve şasi "
           "(-1.50 … -0.74) tam ortada kalır. Kaplama tavanının nasıl aşıldığının "
           "uzaysal kanıtı.")
    figure(doc, "fig_multirobot_ur10e.png",
           "Şekil 9: Aynı plan, yalnız UR10e'nin 19 bakış-noktası (X -1.22 … -0.22, "
           "yükseklik 0.51 … 2.14 m). Kol şasinin +X yüzünü ve üst bölgesini tarar; "
           "rayın kendisi kamera ile şasi arasında kaldığı için karşı yüz bu kola "
           "tümüyle kapalıdır.")
    figure(doc, "fig_multirobot_kawasaki.png",
           "Şekil 10: Aynı plan, yalnız Kawasaki'nin 9 bakış-noktası — hepsi "
           "X ≤ -1.41'de, yani şasinin uzak yüzünde. İkinci kolun varlık sebebi bu "
           "figürdür: bu dokuz poz sırasıyla 339, 149, 123, 76, 71, 51, 43, 29 ve 24 "
           "YENİ nokta getirir; toplamı, UR'nin tek başına ulaşabildiği tavanın "
           "üstüne çıkan farktır.")

    h2(doc, "4.4 Gerçek-Robot Yürütücüsü (Executor)")
    bullets(doc, [
        "Ölçülen-varış gating (Kawasaki): JTC action sonucuna güvenmek yerine, komut "
        "edilen her eklem birleşik /joint_states cache'inde tolerans içine girene "
        "kadar (veya arrival_timeout_sec=90) beklenir. Gerekli, çünkü kol "
        "broadcaster'ı ile AGV bridge'i world_to_agv vs joint1..6'yı AYRI mesajlarda "
        "yayınlar. arrival_joint_tol 0.10 rad / arrival_linear_tol 0.08 m.",
        "AGV async: AGV ayrı bir async platformdur (JTC world_to_agv → bridge → "
        "/agv/goal_position → agv_controller → rosbridge → ROS1 action, 0.05 m/s). "
        "Kawasaki controller path tolerance'ları 0.0'a, goal_time 3→120'ye çekildi — "
        "JTC, AGV fiziksel olarak varana dek hedefi TUTAR.",
        "Sabit başlangıç pozu: _go_to_start() döngüden önce her kolu bilinen bir "
        "başlangıç pozuna götürür, böylece ilk cached trajectory tam kaydedildiği "
        "state'ten başlar.",
        "UR pozisyon-hedefi (pose goal): planın per-viewpoint joint_positions'ı "
        "GÜVENİLMEZ tek bir IK dalıdır; bazıları eklem limitinde oturur (ör. "
        "shoulder_pan = 2π) → OMPL goal bölgesi çöker → çıplak FAILURE. UR artık "
        "planlayıcının IK çözdüğü KARTEZYEN POZ HEDEFİNE planlar (target_link "
        "ur10e_depth_optical_frame). Joint-space (açı-normalizasyonlu) fallback'tir. "
        "KAWASAKI DEĞİŞMEDİ (joint-space + ölçülen-varış korunur).",
        "UR ışınlanma / CONTROL_FAILED (-4) düzeltmesi: (a) _dispatch, göndermeden "
        "önce /joint_states'in max_age içinde taze olmasını bekler — move_group, en "
        "yeni joint state 1 sn'den eskiyse goal'u sessizce iptal ediyordu; "
        "(b) _rebase_traj_to_current koşullu yapıldı, herhangi bir eklem kayıtlı "
        "başlangıçtan uzaksa rebase etmez, önce kaydedilmiş başlangıca yürür.",
    ])

    h2(doc, "4.5 Düğüm Ayrımı (Node-Split) Mimarisi")
    p(doc, "Yürütücü, tek ~1300 satırlık dosyadan bir şablon-metod (template method) "
           "mimarisine ayrıldı; iki kol AYRI SÜREÇLER olarak çalışır:")
    table(doc, ["Dosya", "Rol"], [
        ["inspection_base.py",
         "InspectionNodeBase(Node): tek-kol ortak mantık (plan / cache / capture / "
         "padding / arrival / 2π açma / dal seçimi)."],
        ["ur_inspection_node.py",
         "URInspectionNode: pose-goal, kontrolcü-tabanlı tamamlanma, ur10e_ padding."],
        ["kawasaki_inspection_node.py",
         "KawasakiInspectionNode: joint-space, ölçülen-varış gating, async AGV, "
         "kamera→tip IK zinciri."],
    ])
    p(doc, "only_ur / only_kawasaki gate'leri her düğüm kendi içinde çözer; her ikisi "
           "de koşulsuz başlar. Cross-arm collision farkındalığı korunur çünkü "
           "paylaşılan move_group planning scene'inde (her iki kolun /joint_states'i "
           "üzerinden) yaşar.")
    p(doc, "UYARI — launch ASLA kendiliğinden bitmez: multirobot_inspection.launch.py "
           "iki yürütücüyü ve KALICI bir visualizer'ı başlatır, on_exit handler'ı "
           "yoktur. Tur bittiğinde iki yürütücü süreci ölür ama ros2 launch süresiz "
           "ayakta kalır. Tamamlanmayı launch sürecinden anlamaya çalışmak sonsuza "
           "kadar bekler; doğru yöntem yürütücü PID'lerini izlemektir "
           "(doors_mission_node bunu pgrep ile yapar).")

    # ---------------------------------------------------------------- 5
    h1(doc, "5. Çarpışma Modeli ve SRDF")
    p(doc, "Bu bölüm, gerçek robotun fiziksel olarak çarptığı her durumun kök nedenini "
           "ve eklenen hacimleri toplar. Ortak ders: MoveIt Setup Assistant tarafından "
           "üretilen SRDF'ler, rastgele örneklemede çarpışmadığı için kritik çiftleri "
           "'Never' olarak DEVRE DIŞI bırakır. DEVRE DIŞI bir çift hiç kontrol "
           "edilmez; dolayısıyla PADDING O ÇİFT İÇİN HİÇBİR ŞEY YAPMAZ — padding "
           "yalnızca kontrol edilen çiftleri şişirir.")

    h2(doc, "5.1 SRDF Düzeltmeleri")
    bullets(doc, [
        "Kablo kanalı: gerçek UR10e, bilek kablo kanalı (cable_channel) ile şasiye "
        "fiziksel olarak çarpıp protective-stop yapıyordu — 10 cm padding'e rağmen. "
        "Altı SRDF'ten cable_channel[_flange]↔chassis_part 'Never' satırları silindi "
        "(gerçek adjacent/self disable'lar korundu).",
        "Kawasaki şasi: çok-robot koşusunda Kawasaki kolu şasiye çarptı. Üç gerçek "
        "SRDF'ten TÜM link1..6↔chassis_part çiftleri (config'e göre 206-214 çift) "
        "yeniden etkinleştirildi. Padding artık gerçek bir knob; MODEST tutun "
        "(~0.02-0.03), yükseği yakın viewpoint'leri erişilmez yapar.",
        "Doğrulama (22 Ağustos 2026): real_ifarlab_moveit_config/whole_cell_hw.srdf "
        "içinde cable_channel↔chassis 0 çift, link1..6↔chassis 0 çift — her iki "
        "düzeltme de yerinde duruyor.",
    ])

    h2(doc, "5.2 UR SICK Kamera Kabuğu (6 Kutu)")
    p(doc, "sick_camera collision MESH'i gerçek gövdeyi eksik gösteriyor (optik "
           "yüzünde gerçek muhafazadan ~1.9 cm kısa kalıyor); MoveIt kamerayı UR'nin "
           "kendi eklemlerine ve kablo kanalına sokan yollar planlıyordu. Kamera, "
           "altı kutudan oluşan bir kabukla sarıldı: ön (optik yüz), arka ve dört yan. "
           "Kutular görüş ekseni etrafında 60° yaw'lı (URDF'teki ur_sick_optical_joint "
           "ile aynı açı).")
    table(doc, ["Kutu", "Boyut (cm)", "Konum — sick_camera çerçevesi (m)"], [
        ["front (optik)", "8.7 × 9.0 × 3.9", "(-0.0301, 0.0174, 0.0537)"],
        ["back", "8.7 × 9.0 × 3.9", "(-0.0301, 0.0174, -0.0628)"],
        ["+u yan", "3.9 × 9.0 × 7.75", "(0.0014, 0.0720, -0.0045)"],
        ["-u yan", "3.9 × 9.0 × 7.75", "(-0.0616, -0.0372, -0.0045)"],
        ["+v yan", "8.7 × 3.9 × 7.75", "(-0.0860, 0.0497, -0.0045)"],
        ["-v yan", "8.7 × 3.9 × 7.75", "(0.0258, -0.0149, -0.0045)"],
    ])
    p(doc, "Kabuk önce çalışma zamanında planning scene'e attach ediliyordu; "
           "sonra URDF'e gömüldü (yalnız collision, visual yok) — hem gerçek "
           "(ur_macro.xacro) hem sim (simrobot_ur_macro.xacro). Böylece move_group "
           "kutuları HEM kayıt HEM oynatmada otomatik görür. En dar boşluk "
           "wrist_3_link'e ~0.53 cm'dir; yan kutuların kalınlığını 3.9 cm'nin üstüne "
           "çıkarmayın. SRDF'e dokunulmadı: dört komşu link zaten her SRDF'te "
           "sick_camera'ya karşı kapalıdır.")
    figure(doc, "fig_sick_collision_shell.png",
           "Şekil 11: SICK kamera etrafına eklenen 6-kutu çarpışma kabuğu "
           "(ur10e_sick_camera çerçevesi, görüş ekseni etrafında 60° yaw). Eksik "
           "collision mesh'inin bıraktığı boşluğu doldurur.")

    h2(doc, "5.3 Kawasaki Kamera Kabuğu (7 Kutu) — TAMAMLANDI")
    p(doc, "Revizyon 1'de 'ertelendi' olarak geçen iş 3 Ağustos 2026'da yapıldı. "
           "camera_link (mobile_manipulator_description/urdf/khi_rs/rs005l_macro.xacro) "
           "yalnızca <inertial> taşıyordu — ne visual ne collision. MoveIt, "
           "Kawasaki'nin kamerası YOKMUŞ gibi planlıyordu.")
    p(doc, "UR kutularının kopyası DEĞİLDİR:")
    bullets(doc, [
        "camera_link'in KENDİSİ optik çerçevedir (camera_rgb_joint ve "
        "camera_rgb_optic_joint ikisi de birimdir), dolayısıyla görüş ekseni +X'tir, "
        "UR'daki gibi +Z değil; muhafaza -X boyunca uzanır.",
        "Hiç mesh olmadığı için gövde kutusu da eklendi → altı etek yerine 7 KUTU.",
        "60° yaw taklit EDİLMEDİ (o, UR braketinin montaj açısıydı, burada bilinmiyor); "
        "kesit 0.090'da kare tutuldu, böylece her roll açısında doğru kalır.",
        "Xacro property'leri: cam_body_depth 0.0775 / cam_body_cross 0.090 / "
        "cam_shell_skirt 0.039. Zarf 15.6 × 16.8 × 16.8 cm (UR'ninkiyle eşleşir), "
        "optiğin önünde 3.9 cm çıkıntı. Kawasaki erişimine fazla mal olursa "
        "cam_shell_skirt küçültülür.",
    ])
    p(doc, "SRDF işi ZORUNLUYDU: hiçbir SRDF camera_link'ten söz etmiyordu, yani yeni "
           "collision geometrisi kalıcı self-collision üretecekti. camera_link↔link7 "
           "ve camera_link↔link6 YEDİ whole_cell SRDF'in hepsine eklendi "
           "(harmony + real/sim × plain/vacuum/gripper). link5 ve daha içerideki "
           "linkler BİLEREK ETKİN BIRAKILDI — kabuk link7'yi aşıyor ve önkola "
           "katlanması gerçek bir çarpışmadır; sıfır pozda kabuğun en yakın yüzü "
           "link5'ten 6.8 cm açıktır, yani başlangıç-durumu çarpışması olmaz.")
    p(doc, "İkinci bir kopya tuzağı: rs005l_description/urdf/rs005l_macro.xacro "
           "hiç kamerası olmayan ESKİ bir kopyadır; whole_cell dosyaları "
           "mobile_manipulator_description'dakini include eder.")

    h2(doc, "5.4 AGV Güverte Keep-out Kutusu")
    p(doc, "AGV güvertesi hiçbir CAD modelinin kapsamadığı kablo ve donanım taşıyor, "
           "Kawasaki de içlerine doğru yol planlıyordu. ota_base.xacro içindeki "
           "ota_base_link'e ikinci bir <collision> kutusu eklendi.")
    bullets(doc, [
        "ota_base_link çerçevesinde z +0.165 .. +0.263; taban izi AGV gövdesiyle aynı "
        "(1.02605 × 0.728), böylece keep-out platformdan taşmaz.",
        "0.165 = AGV gövde kutusunun üstü; 0.263 = Kawasaki'nin en alçak noktası "
        "(base2manipulatorbase 0.212 + base_mount 0.051). Aradaki boşluk 9.8 cm — "
        "kullanıcının 'yaklaşık 10 cm' tahminiyle örtüşür.",
        "SRDF DEĞİŞİKLİĞİ GEREKMEDİ: ota_base_link'e karşı zaten kapalı olanlar "
        "manipulator_base_link, base_link, link1, link2; hâlâ KONTROL EDİLENLER "
        "link3..link7 ve camera_link — yani güverteye uzanabilen tam da o linkler. "
        "57 chassis_part da zaten kapalıydı.",
        "ur10e_* linkleri ota_base_link'e karşı kapalı DEĞİL, yani UR de artık "
        "güverte hacminden uzak tutuluyor; UR viewpoint'leri beklenmedik şekilde "
        "düşmeye başlarsa aday sebep budur.",
    ])

    h2(doc, "5.5 link7 Gripper Pin Bloğu — Bilinen Boşluklarla")
    p(doc, "Gripper pinlerini temsil eden bir kutu 6 Ağustos 2026'da Kawasaki link7'ye "
           "eklendi (wrist_block_* property'leri). Yeni bir link değil, link7 üstünde "
           "ek geometri olduğu için link7'nin mevcut SRDF satırlarını miras alır ve "
           "SRDF düzenlemesi gerekmedi. Son ofsetler RViz'de "
           "scripts/wrist_block_tuner.py ile ayarlandı, hesaplanmadı.")
    p(doc, "Ölçülen maliyet: link4'e karşı (bu çift kapalı DEĞİL) joint5 × joint6 "
           "taraması (61×73) blok yüzünden 94 pozu (%2.1) yeni engelliyor, hepsi "
           "|joint5| ≥ 116° bölgesinde; home pozu temiz. Konuma çok duyarlı — dışarı "
           "ayardan önce 22, 24 mm daha alçak bir denemede 249 poz.")
    p(doc, "ÜÇ BOŞLUK HÂLÂ AÇIK (kullanıcı 6 Ağustos'ta 'şimdilik böyle kalsın' dedi):")
    bullets(doc, [
        "link7 hiç collision padding almıyor: kawasaki_inspection_node."
        "_arm_link_names '^link[1-6]$' kullanıyor, yani pinler tam modellenmiş "
        "boyutta kontrol edilirken link1-6'ya 4 cm veriliyor.",
        "Kontrol edilen 28 linkin yalnızca 9'unda herhangi bir marj var ve o da "
        "yalnızca engel tarafından geliyor; kalan 19'unda (AGV tabanı, 8 teker, "
        "2 tekerlek, barrier_kaw1..3, ifarlab_ray_link, kawa_cable_channel) SIFIR.",
        "only_kawasaki:=true bunu da götürüyor: UR düğümü should_run()'da erkenden "
        "dönüyor ve _apply_collision_padding() hiç çalışmıyor, yani ur10e_* hiç "
        "padding almıyor ve pinler kapılara sıfır marjla yaklaşıyor.",
        "Ayrıca link7'nin 25 kapalı çifti var — pinler UR'nin wrist_2/wrist_3/flange/"
        "tool0/omuz/üst kol/kablo kanalına, masalara, konveyöre, vida besleyiciye ve "
        "kabinete karşı HİÇ kontrol edilmiyor. O 'Never' kararları ESKİ, daha küçük "
        "link7 ile üretilmişti.",
    ])

    # ---------------------------------------------------------------- 6
    h1(doc, "6. Trajectory Optimizasyonu — 2π Açma ve IK Dal Seçimi")
    p(doc, "Bu bölümün tamamı Revizyon 1'den sonradır (3-14 Ağustos 2026). Gerçek "
           "robot koşularında kullanıcı, kolun ardışık viewpoint'ler arasında gereksiz "
           "tam turlar attığını bildirdi. İki AYRI kök neden bulundu ve ikisi de ayrı "
           "ayrı çözüldü.")

    h2(doc, "6.1 Birinci Neden: IK'nın Sakladığı Rastgele 2π Dalı")
    p(doc, "IK, bir çözüm ailesinin TEK bir keyfî dalını döndürür ve bu her viewpoint "
           "için bağımsız seçilir; aşağı akışta bunları uzlaştıran hiçbir şey yoktu. "
           "Plan dosyasında doğrulandı: ur_vp_007'nin shoulder_lift'i, komşuları "
           "+89.8 / +78.3 iken -279.1°; ur_vp_014'te -359.9° ve wrist_2 270°. Aynı "
           "fiziksel açı, ama arada bir tam tur.")
    p(doc, "Çözüm joint_wrap.py'dir: parse_joint_limits(urdf) + wrap_to_reference(), "
           "bir referans açıya en yakın 2π-eşdeğerini seçer ve bunu YALNIZCA eklemin "
           "URDF limitleri eksi wrap_limit_margin içinde kalırsa kabul eder.")
    p(doc, "LİMİT KONTROLÜ İŞİN TAM KENDİSİDİR. Limit-kör bir 'en yakın mevcut açı' "
           "kuralı önce denendi ve geri alındı (Kawasaki joint1'i 1.928 → -4.355'e, "
           "yani menzil dışına itti ve kolu yanlış yönelimli pozlara gönderdi). "
           "Limit-farkındalıklı olduğunda, gerekmeyen yerde otomatik olarak hiçbir şey "
           "yapmaz:")
    table(doc, ["Eklem", "Menzil", "Açılabilir mi?"], [
        ["UR shoulder_pan / shoulder_lift / wrist_1 / wrist_2 / wrist_3",
         "±360°", "EVET (span > 2π)"],
        ["Kawasaki joint4 / joint6", "±360°", "EVET"],
        ["UR elbow", "±180° (yaml bilerek yarıya indirir)", "HAYIR"],
        ["Kawasaki joint1", "±180° (span TAM 2π)",
         "HAYIR — tek alternatif limit kenarında oturur, marj reddeder"],
        ["Kawasaki joint2 / joint3 / joint5", "-80..135 / -172..118 / ±145", "HAYIR"],
        ["world_to_agv, ur10e_base_to_robot_mount", "prizmatik", "HAYIR"],
    ])
    p(doc, "Bu tablo neden Kawasaki'nin 'hâlâ çok dönüyor' göründüğünü de açıklar: "
           "Kawasaki 6 ekleminin yalnızca 2'sini açabilir, UR ise 5'ini. Kayıtlı "
           "veride en çok yol yapan eklem olan joint1 (146° ve 248°'lik hoplar) tam "
           "olarak 2π açıklığındadır ve kalıcı olarak açılamaz.")
    p(doc, "ÖLÇÜLEN KAZANÇ (gerçek plan, gerçek URDF limitleri): UR turu "
           "10860° → 6969° (%-36, 17 yeniden yazma), Kawasaki 4495° → 3009° (%-33, "
           "6 yeniden yazma), toplam %-35. Sim koşusu doğruladı (18/18 UR + "
           "9/9 Kawasaki): ur_vp_014 1406° → 433° (%-69), ur_vp_016 1015° → 187° "
           "(%-82). Pozlar fiziksel olarak aynı olduğu için kaplama/erişilebilirlik "
           "değişemez — yalnızca yol.")

    h2(doc, "6.2 İkinci Neden: IK Dal Sıçramaları")
    p(doc, "Açmadan sonra tek hoplar HÂLÂ 574/828/853° tutuyordu. Tekillik DEĞİL: "
           "her UR viewpoint'inin wrist_2'si tekil değerlerden (0 / ±180) en az 30° "
           "uzak ölçüldü, en yakını 30.6°. Gerçek sebep: ur_vp_000 → ur_vp_001 "
           "kartezyende 10.5 cm / 30° uzakken eklem uzayında 530°, ve bunun büyük "
           "kısmı shoulder_pan -187° + bileğin telafisi — yani FARKLI BİR IK DALI "
           "(her pozun ~8 dalı var: omuz/dirsek/bilek flip'leri). Açma bunlara "
           "dokunamaz; bunlar tam-tur ofseti değil, gerçekten farklı "
           "konfigürasyonlardır.")
    p(doc, "Üç parçalı çözümün iki parçası kullanımdadır:")
    bullets(doc, [
        "Planlayıcı _rechain_ik(): sıralamadan sonra her viewpoint, bir önceki "
        "viewpoint'in çözümüyle TOHUMLANARAK yeniden çözülür ve yalnızca daha yakınsa "
        "tutulur (rechain_ik_attempts 4). _unwind_tour'dan ÖNCE çalışır (önce dal, "
        "sonra tam turlar). _joint_distance mod-2π puanlar, böylece saf bir sarma dal "
        "değişikliğiyle karıştırılmaz.",
        "Yürütücü _branch_candidates() + _solve_ik(): UR için /compute_ik'i KENDİMİZ, "
        "MEVCUT ölçülen poza tohumlanmış olarak çözer, dalları örneklemek için "
        "tekrarlar (ik_seed_attempts 4), her birini açar, yola göre sıralar ve yol "
        "üreten en yakın dala planlar; çıplak pose goal artık yalnızca FALLBACK'tir. "
        "Belirleyici olan budur, çünkü MoveIt'in kendi goal sampler'ı IK'yı rastgele "
        "tohumlar.",
        "_plan_best_branch(): 'plan üreten ilk adayı al' yerine "
        "branch_plan_candidates (varsayılan 3) kadar aday planlanır ve GERÇEK "
        "_traj_travel'ı en kısa olan tutulur — hedef-uzayı sıralaması yalnızca bir "
        "vekildi. Plan üretemeyen adaylar bütçeyi tüketmez.",
    ])
    p(doc, "NOT: pymoveit2'nin compute_ik'i burada kullanılamaz — ik_link_name'i hiç "
           "set etmez (grup ucu için çözer, kamera çerçevesi için değil) ve bloklayan "
           "biçimi içeride rclpy.spin_once çağırarak bu düğümün executor wait set'ini "
           "bozar.")

    h2(doc, "6.3 Denenip Geri Alınan: pick_ik minimal_displacement_weight")
    p(doc, "Dal sürekliliğini solver tarafında zorlamak için ur10e gruplarında "
           "minimal_displacement_weight 0.0 → 0.005 yapıldı. ERİŞİLEBİLİRLİĞİ "
           "ÇÖKERTTİ:")
    table(doc, ["Metrik", "weight = 0.0", "weight = 0.005"], [
        ["UR NO_IK_SOLUTION", "538", "1780"],
        ["UR viewpoint sayısı", "18", "1"],
        ["Kawasaki viewpoint sayısı", "6", "21"],
        ["Kaplama", "%65.7", "%50.7"],
        ["Planlama süresi", "41 s", "128 s"],
    ])
    p(doc, "Sebep pick_ik'in kendi parametre dokümanındadır: "
           "minimal_displacement_weight '>0.0 maliyet fonksiyonunu ETKİNLEŞTİRİR' ve "
           "cost_threshold 'tüm maliyet fonksiyonları bu değerin altında dönerse' "
           "koşulunu ekler. 0.0'da maliyet fonksiyonu KAPALI olduğu için buradaki çok "
           "sıkı cost_threshold (0.001) hiç uygulanmıyordu; etkinleştirince tohuma "
           "zaten çok yakın olmayan her çözüm eşiği geçemedi. Altı dosyada da 0.0'a "
           "geri alındı, her birine uyarı yorumu eklendi. Denenmemiş alternatif: "
           "ağırlıkla birlikte cost_threshold'u da yükseltmek (~0.01) — deney olarak "
           "ele alınmalı ve NO_IK_SOLUTION + viewpoint sayıları yeniden ölçülmelidir.")
    p(doc, "Önemli: kod tarafındaki düzeltmeler bu knob'a BAĞLI DEĞİLDİR. Yürütücü "
           "birkaç IK çözümü arasından kendisi seçer ve _rechain_ik yalnızca kesinlikle "
           "daha yakınsa tutar, dolayısıyla ikisi de ağırlık 0.0 iken çalışır.")

    h2(doc, "6.4 Kawasaki'ye Özgü Ek İşler ve Park Edilen Kısım")
    bullets(doc, [
        "_joint_distance prizmatik eklemleri tamamen ATLIYORDU, dolayısıyla "
        "Kawasaki'nin FAZLALIKLI (ray + 6) grubunda AGV'yi bir metre süren aday "
        "bedava puanlanıyordu. Artık prizmatik yol rechain_rail_weight (2.0 rad/m) "
        "ile ücretlendiriliyor.",
        "_seed_variants(): prev, mevcut çözüm, prev + ray ±rechain_rail_probe "
        "(0.25 m) ve prev + joint1 ±π (o 248°'lik hopu hedefleyen taban-yaw dal "
        "probu). Hepsi URDF limitlerine kırpılır; sonuç yalnızca kesinlikle daha "
        "yakınsa tutulur, yani bir turu kötüleştiremez.",
        "_rail_shift_ok() koruması: _goal_travel prizmatik eklem ücretlendirmediği "
        "için bu olmadan bir aday sessizce AGV'yi metrelerce sürebilirdi. "
        "branch_max_rail_shift varsayılanı 0.0 = ray planlanan değere kilitli; "
        "planlayıcı o AGV konumunu kaplama için seçmişti.",
        "PARK EDİLDİ (4 Ağustos): Kawasaki yürütücüsü her viewpoint'te 'nearest-branch "
        "IK: 1 distinct branch(es)' logluyor. Önce KDL suçlandı; çalışan move_group'a "
        "canlı /compute_ik probu bunu ÇÜRÜTTÜ — 24 rastgele tohumla viewpoint başına "
        "16-20 AYRI çözüm dönüyor. Gerçek iki sebep: (a) _ik_seeds yalnızca 2-4 tohum "
        "üretiyor ve hepsi planlanan hedefin çevresinde kümeleniyor, KDL ise YEREL bir "
        "çözücü olduğundan hepsi aynı çözüme yakınsıyor; (b) branch_max_rail_shift=0.0 "
        "ile _rail_shift_ok neredeyse her alternatifi reddediyor (kawa_vp_005'te 19 "
        "çözümden yalnız 1'i rayı koruyor). Uygulanacak çözüm: tohum listesine N "
        "rastgele arm tohumu eklemek ve rayı yapısal olarak sabitleyen YENİ bir 6-DOF "
        "SRDF grubu (chain base_link → link7) kullanmak.",
    ])

    h2(doc, "6.5 Kayıt/Oynatma Sözleşmesi ve pymoveit2_real'e Taşıma")
    p(doc, "Trajectory cache: deterministik, önceden-doğrulanmış yollar için her "
           "viewpoint'in yolu <plans>/trajectories/<arm>_<vp_id>.json'a kaydedilir ve "
           "sonradan REPLAY edilir (use_trajectory_cache true, force_replan false). "
           "Bu revizyonda üç şey değişti:")
    bullets(doc, [
        "GEÇERSİZLEŞTİRME ETİKETİ: eski boolean 'unwound' alanı, "
        "goal_policy = 'unwind=N,branch=N,jbranch=N' oldu. Hedef-seçim politikasının "
        "HERHANGİ bir değişikliği kayıtlı yolları tam BİR KEZ geçersizleştirir. "
        "Bu gerekliydi, çünkü bir UR cache girdisi pose_goal'a göre anahtarlanır ve "
        "açma pose_goal'ı DEĞİŞTİRMEZ — eski uzun yollar sonsuza dek oynatılırdı.",
        "KAYITLI YOLUN KENDİSİ AÇILIR: _unwind_traj_to_current(), karşılaştırmayı "
        "değil KAYITLI YOLU tam turlardan arındırır. _at_pose'u mod-2π "
        "karşılaştıracak şekilde değiştirmek tuzak olurdu: 'zaten oradasın' der ama "
        "trajectory hâlâ uzak sayıları taşır ve onu göndermek tek kontrol cycle'ında "
        "360°'lik bir adımdır. Kaydırma yalnızca eklemin span'ı > 2π ise VE kaydırılmış "
        "yol TÜM noktalarda limitlerde kalıyorsa kabul edilir.",
        "SON VİEWPOİNT → HOME BACAĞI DA KAYDEDİLİR (Ağustos 2026): _home_arm artık "
        "'return_home' adlı bir geçiş olarak _run_cached_move'dan geçer ve "
        "<arm>_return_home.json altında saklanır. Önceden turun tek plansız-kaydedilmemiş "
        "hareketi buydu; gerçek robotta canlı yeniden planlanıyordu.",
    ])
    p(doc, "TAŞIMA (14 Ağustos 2026): hücrede bu senaryo dışında üç kullanım senaryosu "
           "daha var ve açmayı her birinde yeniden yazmak kabul edilmezdi. Kullanıcının "
           "çizdiği sınır: yalnızca YENİ TRAJECTORY işi taşınsın; cache/kayıt tarafı "
           "(goal_policy etiketi, _unwind_traj_to_current) ve viewpoint-planlama tarafı "
           "(_rechain_ik, _unwind_tour) yerinde kalsın. Bu sınır sağlamdır: taze bir "
           "hedefi açmak poz-özdeş ve limit-korumalıdır, yani senaryo başına donanım "
           "doğrulaması gerektirmez; oysa cache etiketi donanımda doğrulanmış kayıtları "
           "GEÇERSİZLEŞTİRİR — bu bir kütüphane varsayılanı değil, senaryo başına "
           "karardır.")
    bullets(doc, [
        "Nereye gitti: pymoveit2_real/moveit2.py içindeki MoveIt2.set_joint_goal(). "
        "Tek darboğaz noktası budur — move_to_configuration() ve plan_async() ikisi de "
        "oraya iner — dolayısıyla cleaning_mission_runner, sensing_robot, "
        "pick_and_place, HRC, ur_inspection_scenario ve solo şasi baz senaryosu "
        "(viewpoint_planner/inspection_executor_node.py) sıfır çağrı-yeri "
        "değişikliğiyle bu özelliği kazandı.",
        "joint_wrap.py, viewpoint_planner'dan pymoveit2_real'e taşındı; eski konum "
        "yeniden-dışa-aktarım (re-export) shim'i olarak durduğu için eski importlar "
        "çalışmaya devam ediyor.",
        "Yeni ctor parametreleri: unwind_joint_goals (varsayılan True), "
        "unwind_limit_margin (0.05), unwind_min_gain (0.35). Yeni public "
        "unwind_joint_goal() ve joint_limits.",
        "ÜÇ YERDE BİLEREK KAPALI: ur_inspection_node ve kawasaki_inspection_node "
        "(yürütücü açmayı daha fazla bağlamla kendisi sahipleniyor; ikisi birden açık "
        "olsaydı yürütücünün bilerek HAM tuttuğu son-çare denemesi açılmış olanın "
        "kopyası olurdu ve wrap_goals_to_current:=false hiçbir şeyi devre dışı "
        "bırakmaz hale gelirdi) ve MoveIt2Gripper (gripper eklemleri asla sarılamaz).",
        "Doğrulama: gerçek whole_cell_hw.urdf.xacro'ya karşı gerçek MoveIt2'yi süren "
        "11/11 test (doğru açılabilir kümeler UR 5/7, Kawasaki 2/7; ray dokunulmuyor; "
        "idempotent; min_gain tabanı; reference= onurlandırılıyor; opt-out çalışıyor; "
        "hiçbir sarılmış değer URDF limitlerinden çıkmıyor) ve shim üzerinden 27/27 "
        "test_joint_wrap.py.",
        "BİLİNEN BOŞLUK: pymoveit2_sim kendi moveit2.py'si olan ayrı bir fork'tur ve "
        "bunu ALMADI. HIL ihtiyacına göre karar verilecek.",
    ])

    # ---------------------------------------------------------------- 7
    h1(doc, "7. Gerçek Muayene Çıktısı — Şasi Octomap Rekonstrüksiyonu")
    p(doc, "Sistemin uçtan-uca sonucu: gerçek robot planlanan bakış-noktalarını gezer, "
           "her pozda SICK kamerasından bir nokta bulutu (PCD) + poz kaydeder; bu "
           "bulutlar dünya çerçevesinde birleştirilir ve bir octomap'e (voksel "
           "ızgarası) dönüştürülür.")
    figure(doc, "fig_chassis_octomap.png",
           "Şekil 12: GERÇEK octomap çıktısı: fiziksel hücrede yakalanan SICK "
           "bulutlarından oluşturulan renkli voksel rekonstrüksiyonu. Şasi iskeleti, "
           "rafları ve ayakları belirgin biçimde geri kazanılmıştır — planlama "
           "hattının fiziksel doğrulaması.")
    p(doc, "Metrik notu: planlayıcının kaplama tahmini ile octomap kaplaması FARKLI "
           "şeyler ölçer ve eşleşmeleri BEKLENMEZ. Planlayıcı, katı bir sensör modeli "
           "altında (FOV-3°, geliş açısı ≤80°, menzil, occlusion) görülebilir mesh "
           "örnek noktalarının kesridir — muhafazakâr bir TAHMİNdir. Octomap ise "
           "referans PCD'lerin vokselleştirilmiş halinde gerçekte isabet alan "
           "vokselleri sayar (farklı payda, voksel başına any-hit).")

    # ---------------------------------------------------------------- 8
    h1(doc, "8. doors_inspection — İkinci Senaryo: Kapı Muayenesi")

    h2(doc, "8.1 Ne Yeni, Ne Değil")
    p(doc, "Doğrulanmış hattan hiçbir şey fork EDİLMEDİ. multirobot_planner_node, "
           "ur_inspection_node ve kawasaki_inspection_node zaten tam "
           "parametrelendirilmiş olduğu için bu paket AYNI çalıştırılabilirleri kapıya "
           "özgü parametrelerle koşar. Dolayısıyla şasi işi için yapılan düzeltmeler "
           "(trajectory cache doğrulaması, 2π açma, en-yakın-dal IK) kapı işine "
           "kendiliğinden uygulanır. Yalnızca kapıya özgü kod yenidir:")
    table(doc, ["Dosya", "Neden var"], [
        ["merge_door_meshes.py",
         "Planlayıcı TEK bir mesh analiz eder ve kapılar tek bir tur olarak "
         "istendiğinden iki STL birleştirilir. Bu aynı zamanda occlusion'ı doğru "
         "yapar: tek mesh'te her kapı diğerine giden ışınları engeller, tıpkı gerçek "
         "çift gibi."],
        ["door_parts.py",
         "Her kapıyı bir yama (patch) PCD ızgarasına böler: octomap aşamasının "
         "belief-map referansı."],
        ["src/build_doors_octomap.cpp",
         "Belief + occupancy octomap'i ve kaplama raporunu üretir; pcd2octomap_builder'ın "
         "kapıya uyarlanmış hâli."],
    ])
    p(doc, "Paket bu yüzden HİBRİTTİR (ament_cmake + ament_cmake_python): octomap "
           "builder gerçek octomap kütüphanesine link olmak zorunda ve burada "
           "ColorOcTree serileştirmesi için kullanılabilir Python binding'i yok. "
           "setup.py kaldırıldı, iki konsol giriş noktası scripts/ altına taşındı. "
           "Bu paket hiçbir URDF, mesh veya SRDF'e DOKUNMAZ.")

    h2(doc, "8.2 Kapı Mesh'leri ve Birim Tuzağı")
    table(doc, ["", "windowdoor", "windowlessdoor"], [
        ["üçgen", "773.109", "129.462"],
        ["alan", "2.880 m²", "3.941 m²"],
        ["x", "-1.211 … -0.794", "-1.210 … -0.806"],
        ["y", "+1.301 … +2.478", "+0.088 … +1.319"],
        ["z", "+0.204 … +1.708", "+0.201 … +1.706"],
    ])
    p(doc, "Birleşik: 902.571 üçgen, 6.821 m² (şasi 31.97 m², yani ~1/5).")
    p(doc, "BİRİMLER: kapı STL'leri zaten METREdir (URDF onları scale=\"1 1 1\" ile "
           "yükler), millimetre olan chassis.stl'in aksine. Planlama launch'ı "
           "mesh_scale:=1.0 verir; 0.001 geçmek kapıları 2 mm'lik bir zerreye "
           "küçültür ve her viewpoint erişilmez döner. Aynı tuzağın octomap "
           "tarafındaki ikizi için bkz. 8.4.")
    p(doc, "HER İKİ YÜZ de muayene edilir. Her kapı neredeyse eşit +X ve -X alanı "
           "taşır (windowless'ta %48.8 / %48.8) ve iki kol karşılıklı taraflarda "
           "oturur — UR rayı x ≈ -0.16, Kawasaki AGV rayı x ≈ -2.2, kapılar arada "
           "x ≈ -1.2 … -0.79. Yani +X yüzleri UR'ye, -X yüzleri Kawasaki'ye geometri "
           "gereği düşer; tahsisi elle yazılmış bir taraf kuralı değil, "
           "erişilebilirlik çözer.")

    h2(doc, "8.3 Plan")
    p(doc, "Güncel kapı planı 9 UR + 3 Kawasaki bakış-noktası ve %95.6 kaplamadır. "
           "Kamera konfigürasyonu şasidekinden farklıdır: viewpoint_distances "
           "[1.0, 1.2, 1.4, 1.6] m (kapı için 7 Ağustos'ta 1.0 m'ye yükseltildi), "
           "FOV 70° × 60°, güvenlik payı 3°, menzil 0.2-1.8 m, maksimum geliş açısı "
           "80°.")
    figure(doc, "fig_doors_viewpoints.png",
           "Şekil 13: Kapı planı dört bakış açısından — 9 UR (mavi) + 3 Kawasaki "
           "(kırmızı) bakış-noktası ve görüş yönleri, birleştirilmiş kapı mesh'i "
           "üzerinde. Üst sıra iki kolun kendi yüzünü gösterir; sol alttaki tepe "
           "görüntüsü iki kolun kapıların ZIT taraflarında durduğunu, sağ alttaki yan "
           "görüntü ise bakış-noktalarının yükseklik dağılımını ortaya koyar. "
           "İşaretler derinlik testi olmadan mesh'in üzerine çizilir, yani bir "
           "işaretin mesh üstünde görünmesi onun ön tarafta olduğu anlamına gelmez.")

    h2(doc, "8.4 Octomap Aşaması")
    p(doc, "Şasiyle aynı iki aşamalı reçete: görülmesi GEREKEN yüzeyin belief "
           "haritası, sonra yakalanan bulutların onun üstüne damgalanması ve parça "
           "başına occupied/belief voksel oranlarının kaplama raporu olması. "
           "insertRay KULLANILMAZ — ışın boyunca boş alan oymak belief haritasını yok "
           "ederdi; yalnızca uç noktalar occupied işaretlenir.")
    p(doc, "Üç şey kapıya özgü bir builder'ı zorunlu kıldı:")
    bullets(doc, [
        "BİRİMLER: şasi builder'ı parça bulutlarına /1000'i SABİT KODLAR, çünkü şasi "
        "CAD'i millimetredir. Kapılar metredir, bu yüzden ölçek --parts_scale "
        "parametresidir ve burada varsayılanı 1.0'dır. Şasi builder'ını kapılarda "
        "çalıştırmak onları 1.2 mm'lik bir zerreye küçültür ve rapor %0 okur.",
        "VERİ KÖKÜ: ~/colcon_ws/src/pcds/doors/{sim,real}_pcds; asla şasi ağacı değil.",
        "RAPOR ÇÖZÜNÜRLÜĞÜ: şasi CAD'inden 57 parça PCD'si bedava geliyordu. Kapılar "
        "iki STL'dir ve 'windowdoor %94' ifadesi eksiğin NEREDE olduğunu söylemez — "
        "bu yüzden door_parts.py her kapıyı 3×4 yama ızgarasına böler (patches_y, "
        "patches_z; 1×1 kapı başına tek parça verir). Builder yamaları '__' "
        "öncesindeki isme göre gruplar, kapı başına ve genel toplam basar; octovis o "
        "zaman görülmemiş yamanın belli olduğu renkli bir yamalı desen gösterir.",
    ])
    p(doc, "windowlessdoor.stl SIFIR KALINLIKTA çift yüzlü bir kabuktur (trimesh "
           "hacim 0.0, +X alanı 1.951 m² = -X alanı). İki yüzü herhangi bir "
           "çözünürlükte AYNI voksele düşer, dolayısıyla bir voksel iki yüzden "
           "HERHANGİ biri görüldüğünde kaplı sayılır. windowdoor ise gerçek bir katıdır "
           "(hacim 0.084 m³, ≈6 cm kalınlık) ve iki yüzü ayrı vokselleri işgal eder. Bu "
           "yüzden octomap kaplaması planlayıcının yüz-seviyesi kaplamasıyla aynı "
           "büyüklük değildir ve windowless kapıda daha yüksek okur.")
    p(doc, "Aşağıdaki rekonstrüksiyon, hâlen ~/colcon_ws/src/pcds/doors altında duran "
           "yakalamalardan üretilmiştir: her mod için 9 UR + 3 Kawasaki bulutu. Renkler "
           "belief haritasının yama ızgarasından gelir; boşluk, hiç görülmemiş yüzeydir. "
           "İki satır arasındaki fark bir sonraki alt bölümün konusudur — sim'de kapı "
           "neredeyse eksiksiz kapanırken gerçekte boşluklar ALT bölgede toplanır.")
    figure(doc, "fig_doors_octomap.png",
           "Şekil 14: Kapı octomap rekonstrüksiyonu (2 cm voksel), her iki yüzden. "
           "Üst sıra simülasyon (%93.0 kaplama, 13.436 voksel), alt sıra gerçek robot "
           "(%62.9, snap 2 + cloud_fix ile, 8.686 voksel). Gerçek verideki eksik, "
           "kapı YÜZEYİNE dağılmış değil, ALTA yığılmıştır — kamera oraya bakmadığı "
           "için değil, kapılar orada modelden ~8 cm sapmış olduğu için.")

    h2(doc, "8.5 Gerçek ile Sim Arasındaki Fark — Kapılar Modelin Yerinde Değil")
    p(doc, "İlk olarak 7 Ağustos 2026'da ölçüldü, 22 Ağustos'ta diskteki güncel "
           "yakalamalarla yeniden ölçüldü. Düzeltmesiz gerçek kapı octomap'i %36.0 "
           "okurken sim %93.0 okuyor ve kapı altları taranmamış görünüyordu. "
           "TARANMAMIŞ DEĞİLLERDİ: gerçek veride 0.20-0.35 m arasında sim'den DAHA ÇOK "
           "ayrık 1 cm voksel vardı (10.718'e karşı 7.080).")
    p(doc, "Mekanizma: build_doors_octomap.cpp bir belief vokselini yalnızca TAM "
           "isabette kredilendirir (2 cm ızgara). Yakalanan buluttan referans mesh'e "
           "medyan mesafe: SİM 3.0 mm, GERÇEK 37.8 mm. Sim bedava geçer; gerçek "
           "vokseli ıskalar ve 'kapı dışı' diye atılır.")
    p(doc, "KÖK NEDEN ROBOTLAR DEĞİL, KAPILARDIR. Hata neredeyse tamamen dünya X'inde "
           "(dy≈0) ve yükseklikle doğrusal değişiyor — ~5°'lik bir eğim. Altta İKİ KOL "
           "DA BAĞIMSIZ OLARAK AYNI OFSETİ ölçüyor: UR -81.4 mm, Kawasaki -81.5 mm. "
           "Ayrı kinematik zincirler, karşıt yüzler, aynı cevap ⇒ kapılar URDF'in "
           "söylediği yerde monte değil. İkincil ±3-5 cm viewpoint'ler arası dağılım "
           "hand-eye kalibrasyonudur ve tek başına 2 cm vokseli zaten aşar.")
    p(doc, "Bunun sonucu olarak viewpoint_distances'ı 1.0 m'ye çıkarmak bu sorunu "
           "ÇÖZMEZ — aynı yanlış yerde daha çok nokta yine vokseli ıskalar. Bunun "
           "yerine iki telafi eklendi:")
    bullets(doc, [
        "--snap_radius (voksel; metre = yarıçap × çözünürlük): bir voksel yanına "
        "düşen noktanın da sayılmasına izin verir (26-komşu arama) ve EN YAKIN belief "
        "vokselini seçer. Eski snap dx,dy,dz'yi -1'den tarayıp ilk isabeti alıyordu, "
        "yani noktaları -x,-y,-z'ye yanlıyordu; bu düzeltildi.",
        "--cloud_fix \"roll,pitch,yaw,x,y,z\": YAKALANAN bulutlara vokselleştirmeden "
        "önce katı bir dönüşüm uygular. Bilerek belief haritasına uygulanmaz — belief "
        "haritası voksel ızgarasını ve kaplama paydasını tanımlar; modeli oynatmak "
        "paydayı her koşuda değiştirirdi. Nokta-düzlem ICP ile (%20 kırpılmış, "
        "183.740 nokta) fit edildi: rpy = -0.740, -5.128, +0.347 derece; "
        "xyz = +0.109, +0.002, +0.061 m. Modele medyan mesafe 41.5 mm → 18.6 mm. "
        "İki kapı ayrı ayrı fit edildiğinde AYNI pitch çıkıyor (-5.04°, -4.80°) ve "
        "ayrı fitler tek küresel fiti yenmiyor, bu yüzden tek dönüşüm kullanılıyor "
        "(launch sabiti DOORS_FIT).",
    ])
    p(doc, "Sekiz kombinasyonun tamamı 22 Ağustos 2026'da diskteki güncel yakalamalarla "
           "yeniden ölçülmüştür (builder'ın kendi raporu, 14.945 belief vokseli):")
    table(doc, ["Ayar", "Gerçek robot", "Simülasyon"],
          [["ham (düzeltmesiz)", "%36.0  (5387)", "%93.0  (13905)"],
           ["snap r=2", "%52.1  (7790)", "%93.1  (13915)"],
           ["cloud_fix", "%49.6  (7414)", "%23.6  (3530)"],
           ["snap r=2 + cloud_fix", "%62.9  (9401)", "%50.5  (7544)"]])
    figure(doc, "fig_doors_coverage.png",
           "Şekil 15: Kapı octomap kaplamasının ölçülen değerleri. snap ve cloud_fix "
           "birlikte gerçek veriyi %36.0'dan %62.9'a taşır; sim'in hatası zaten 3 mm "
           "olduğu için snap sim'i neredeyse hiç oynatmaz (%93.0 → %93.1) — bu da "
           "düzeltmenin gerçek dünyaya özgü bir ofseti kodladığının kanıtıdır. "
           "Gri bölge cloud_fix'in yalnız gerçek veri için olduğunu işaretler.")
    p(doc, "Kapı başına gerçek veride: windowdoor %28.4 → %51.5, windowlessdoor "
           "%46.5 → %78.3. Snap'lenen nokta oranı %84'ten %52.7'ye, ortalama snap "
           "mesafesi 33.1 → 29.0 mm'ye düştü — yani cloud_fix gerçek hatayı yeniden "
           "etiketlemek yerine gerçekten soğurdu.")
    p(doc, "YALNIZ GERÇEK İÇİN. Aynı düzeltmeyi sim'e uygulamak onu %93.0'dan %23.6'ya "
           "düşürüyor (ölçüldü); builder --mode sim ile --cloud_fix birlikte "
           "kullanılırsa yüksek sesle uyarır. Kalan 18.6 mm ≈ bir voksel, "
           "viewpoint'ler arası hand-eye dağılımıdır ve katı bir dönüşümle "
           "giderilemez. SNAP BİR RAPORLAMA TELAFİSİDİR; hizasızlık hâlâ oradadır ve "
           "builder snap'lenen oranı ile ortalama snap mesafesini basar ki bu asla "
           "temiz kaplama diye okunmasın.")

    h2(doc, "8.6 Görev Düğümü (doors_mission_node)")
    p(doc, "Kapı işi tek başına bir tur değil, bir görevdir: sensing turu → tespit "
           "edilen defect'lerin yayınlanması. Launch'ın asla bitmemesi (bkz. 4.5) bu "
           "düğümün ilk sürümünü bozdu. Çalışan tespit yöntemi: pgrep -g <pgid> -f "
           "inspection_node ile yürütücü süreçlerini izlemek; bir kez ortaya çıkıp "
           "sonra hepsi kaybolduğunda tur bitmiştir, ardından launch açıkça "
           "kapatılmalıdır — yoksa ikinci koşu aynı topic üzerinde ikinci bir "
           "visualizer ile çakışır. Kollar başarısızlıkta da aynı şekilde çıktığı için "
           "'süreçler gitti' tek başına başarı demek değildir; doors_mission_node "
           "defect yayınlamadan önce sıfırdan farklı capture sayısıyla en az bir "
           "'sequence complete' şartı arar.")

    # ---------------------------------------------------------------- 9
    h1(doc, "9. Donanım Katmanı Bulguları (Lineer Eksen / Festo)")
    p(doc, "Fiziksel lineer eksende (192.168.3.1) ÖLÇÜLEN — çıkarsanan değil — "
           "davranışlar:")
    bullets(doc, [
        "position_scaling=1000000 (µm), velocity_scaling=1000, base_velocity=5.0.",
        "Drive mdi_velocity'yi YOK SAYAR: 50/200/500 mm/s komutları hep ~324 mm/s "
        "üretti (sabit drive-tarafı maksimum). Düzeltmesi CMMT parametrelemesi ister, "
        "ROS kodu değil.",
        "Her kontrol cycle'ında değişen bir hedef, profil üreticisini takar: her cycle "
        "yazılınca 147.6 mm takip hatası; her 10 mm'de yazılınca 14.2 mm. Bu yüzden "
        "linear_axis_adapter.py mdi_tarpos'u BATCH'ler (setpoint_batch_m=0.01).",
        "TopicBasedSystem state'i /joint_states'ten OKUMAMALI: broadcaster kendi "
        "çıktısını tükettiği için state URDF initial_value'ya (1.0 m) kilitlendi. Ayrı "
        "/linear_axis/joint_states verilerek çözüldü.",
        "edcon: ComModbus arka-plan I/O thread'i; send_io/recv_io varsayılan bloklar "
        "(~10 ms). position_task 'Traversing task aborted' loglar ama aslında komut "
        "hiç gönderilmedi (drive operational değil). Servo takip için "
        "pos_stw1.continuous_update=True.",
    ])

    # ---------------------------------------------------------------- 10
    h1(doc, "10. Operasyonel Notlar (Sık Tuzaklar)")
    bullets(doc, [
        "BUILD ŞART: Config/SRDF/URDF install-share'den DÜZ DOSYA olarak yüklenir "
        "(workspace --symlink-install değil). Bunları düzenledikten sonra ilgili "
        "paketi colcon build ile yeniden derleyin. Aynısı xacro'nun "
        "$(find mobile_manipulator_description) çözümlemesi için de geçerlidir — "
        "src'deki değişiklik o paket derlenene kadar hiçbir şeyi değiştirmez.",
        "CACHE GEÇERSİZLEME: cache, hedefe ve goal_policy'ye göre anahtarlanır; "
        "collision_padding, SRDF veya URDF geometri değişikliğini FARK ETMEZ. Kamera "
        "kabuğu, AGV güverte kutusu, link7 pinleri ve her SRDF düzenlemesinden sonra "
        "bir kez force_replan:=true ile koşun — yoksa eski clearance ile planlanmış "
        "yollar güvensizce oynatılır.",
        "PLANI DA YENİDEN ÜRETİN: yeni collision hacmi IK'yı da etkiler, bazı "
        "viewpoint'ler düşebilir. Kabuk/kutu eklendikten sonra planlama da yeniden "
        "koşulmalıdır.",
        "İŞ AKIŞI: fake_hardware ile trajectory kaydedin, sonra cache'i gerçek robotta "
        "oynatın. Gerçek oynatmada force_replan KAPALI olmalı.",
        "LAUNCH BİTMEZ: tur tamamlandığında ros2 launch ayakta kalır (kalıcı "
        "visualizer, on_exit yok). Tamamlanmayı yürütücü PID'lerinden anlayın.",
        "DİSKTEKİ .ot HANGİ AYARLA ÜRETİLDİĞİNİ TAŞIMAZ. 22 Ağustos 2026'da "
        "pcds/doors/sim_pcds/occupancyMap_doors_sim.ot'un cloud_fix UYGULANMIŞ olarak "
        "üretildiği tespit edildi (%23.6 okuyor, temiz üretimde %93.0). Dosyanın "
        "kendisinde bunu belli eden bir alan yok. Bir .ot'u yorumlamadan önce onu "
        "hangi bayraklarla ürettiğinizi bilin ya da yeniden üretin — 0.4 saniye sürer.",
        "AĞ ÇEKİŞMESİ: her iki SICK kamera + UR external control aynı NIC'i "
        "(192.168.3.0/24) paylaşınca UR'nin 500 Hz reverse channel'ı açlıktan flap "
        "yapar. Çözüm: kamera ethernet'ini 100 Mbps switch'e alarak flood'u "
        "sınırlayın (kesin çözüm: kameralara ayrı NIC).",
        "kinematics.yaml DOSYALARI GIT-TRACKED DEĞİL ve daha önce bir upload commit'i "
        "tarafından bir kez ezildi. pick_ik, orientation_threshold 0.05, attempts 5 ve "
        "minimal_displacement_weight 0.0 değerlerini periyodik olarak doğrulayın.",
    ])

    # ---------------------------------------------------------------- 11
    h1(doc, "11. Açık Konular ve Sonraki Adımlar")
    bullets(doc, [
        "AGV-önce sıralama (kol topla → AGV yalnız git + varış bekle → kol uzat) hâlâ "
        "UYGULANMADI — gerçek robottaki kol/AGV desync çarpışmasının asıl çözümü "
        "budur. Kodda izi yoktur.",
        "Kawasaki en-yakın-dal IK park edilmiş durumda: rastgele arm tohumları + "
        "rayı hariç tutan 6-DOF SRDF grubu gerekiyor (bkz. 6.4). Yeni grup her "
        "kinematics.yaml'da kendi kinematics_solver girdisini de ister.",
        "OMPL DETOURU DOKUNULMADI: kapı cache'inde ölçüldü — Kawasaki yolları doğrudan "
        "eklem mesafesinin 1.54 katı (en kötü 3.50×: doğrudan 182° → yol 638°), UR "
        "1.05 katı. Kawasaki'nin 3053°'sinin ~1070°'si saf OMPL dolambacıdır, hedef "
        "seçimi değil. Çözümü aynı hedefe N kez plan yapıp en kısasını tutmaktır.",
        "pymoveit2_sim 2π açmayı ALMADI (ayrı fork, kendi moveit2.py'si). HIL "
        "ihtiyacına göre karar verilecek.",
        "pcds/doors/sim_pcds altındaki sim octomap'i yanlış bayrakla üretilmiş halde "
        "duruyor (cloud_fix uygulanmış, %23.6). Temiz sürümü %93.0 okur; dosya "
        "yeniden üretilmelidir.",
        "Kapıların fiziksel hizasızlığı sürüyor: cloud_fix ve snap raporlama "
        "telafileridir. Kalıcı çözüm ya kapıları URDF'in söylediği yere monte etmek ya "
        "da URDF'i ölçülen konuma güncellemektir (fit edilmiş dönüşüm elde mevcut).",
        "kawa_vp_003 israf: gerçek koşuda kapı kutusu içine 81.674 noktadan 6'sını "
        "yakaladı (%0.01), sıfır voksel katkısı; sim'de bile yalnızca %13.4. Bu "
        "viewpoint düşürülmeli ya da yeniden planlanmalı.",
        "link7 pin bloğunun üç padding boşluğu açık (bkz. 5.5) ve link7'nin 25 kapalı "
        "SRDF çifti eski, daha küçük link7 ile üretilmişti.",
        "Erişim tavanı yazılım kaldıraçları: kol her yere ulaşabiliyor, sorun bölge "
        "başına dar yönelim seti — poz başına daha çok seçenek. Donanım (Z-lift, ek "
        "eksen) en son çare.",
    ])

    p(doc, "Bu rapor, viewpoint_planner, multirobot_viewpoint_planner ve "
           "doors_inspection paketlerinde 22 Ağustos 2026 itibarıyla yürütülen "
           "mühendislik çalışmasının bir özetidir.")

    doc.save(OUT)
    print(f"yazıldı: {OUT}")


if __name__ == "__main__":
    build()
